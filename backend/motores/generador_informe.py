"""Generador de Informe Técnico — Word (.docx) y PDF descargables.

Empaqueta, para UNA entidad, un informe técnico completo que el docente
puede entregar a sus estudiantes como evidencia de ejecución del sistema
y como fuente de trabajo para la asignatura. El informe integra:

  1. Portada institucional (SIIEAP, entidad, fecha)
  2. Contextualización de la Administración Pública Contemporánea:
     un bloque FIJO (igual en todos los informes) que recorre el arco
     teórico completo del microcurrículo — desde los antecedentes griegos
     de lo público/privado, pasando por el Nuevo Institucionalismo, la
     Nueva Gestión Pública (NGP) y la post-NGP, hasta el esquema
     integrador de la Administración Pública contemporánea (Estado
     Digital, Transformación Digital, Gobernanza Inteligente, IA,
     Gobierno Abierto, Gobernanza de Datos, Valor Público, Administración
     Pública Basada en Evidencia, Capacidades Estatales, Resiliencia
     Institucional, Agenda 2030/ODS, OCDE, CEPAL y Naciones Unidas).
  3. Resultado real del diagnóstico IDI-MIPG de la entidad (dimensiones,
     brechas).
  4. El análisis integral generado por IA (motor_analisis_ia.py): lectura
     desde las tres teorías del curso, recomendaciones técnicas/jurídicas/
     financieras, valoración de riesgo y prospectiva.
  5. Nota de trazabilidad y disclaimer académico.

Este módulo NO llama a la API de Claude — solo da formato a un texto de
análisis que ya fue generado antes por motor_analisis_ia.py.

Dependencias nuevas a agregar en requirements.txt:
    python-docx
    reportlab
"""
from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image,
)

from backend.motores.graficas_informe import generar_grafica_dimensiones, generar_grafica_brechas


# ---------------------------------------------------------------------------
# Identidad visual institucional compartida por los 3 informes (Técnico,
# Estudio de Caso, Alcaldes/Gobernadores): logos, banner de portada,
# divisores de sección con ícono, franjas alternas en tablas y el esquema
# de color por quintil oficial del MIPG. Todo lo que vive aquí se importa
# desde los otros dos módulos para no duplicar código ni diseño.
# ---------------------------------------------------------------------------

COLOR_INSTITUCIONAL = "1F3864"  # azul institucional ESAP, usado en banners y divisores

_CARPETA_ASSETS = Path(__file__).resolve().parent.parent.parent / "assets"
_RUTA_LOGO_ESAP = _CARPETA_ASSETS / "logo_esap.png"
_RUTA_LOGO_CERTIFICACIONES = _CARPETA_ASSETS / "certificaciones.png"


def _agregar_logos_docx(doc) -> None:
    """Inserta, si existen los archivos, el logo de la ESAP (izquierda) y el
    de certificaciones ICONTEC/IQNET (derecha) lado a lado en la parte
    superior de la portada. Si la carpeta assets/ no está presente (por
    ejemplo, no se subió al repo de despliegue), el informe se genera igual,
    simplemente sin logos — nunca se rompe la app por un archivo faltante."""
    try:
        if not (_RUTA_LOGO_ESAP.exists() and _RUTA_LOGO_CERTIFICACIONES.exists()):
            return
        tabla_logos = doc.add_table(rows=1, cols=2)
        tabla_logos.autofit = False
        celda_izq, celda_der = tabla_logos.rows[0].cells
        celda_izq.width = Cm(8.5)
        celda_der.width = Cm(8.5)
        p_izq = celda_izq.paragraphs[0]
        p_izq.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_izq.add_run().add_picture(str(_RUTA_LOGO_ESAP), height=Cm(1.6))
        p_der = celda_der.paragraphs[0]
        p_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_der.add_run().add_picture(str(_RUTA_LOGO_CERTIFICACIONES), height=Cm(1.9))
        doc.add_paragraph()
    except Exception:
        # Nunca romper la generación del informe por un problema con los logos
        pass


def _logos_pdf_flowables():
    """Devuelve una lista de flowables (posiblemente vacía) con los dos
    logos institucionales lado a lado, para insertar al inicio de la
    portada PDF. Igual que en Word: si faltan los archivos, no falla nada."""
    try:
        if not (_RUTA_LOGO_ESAP.exists() and _RUTA_LOGO_CERTIFICACIONES.exists()):
            return []
        img_esap = Image(str(_RUTA_LOGO_ESAP), width=5.2 * cm, height=1.4 * cm)
        img_cert = Image(str(_RUTA_LOGO_CERTIFICACIONES), width=2.9 * cm, height=2.4 * cm)
        tabla_logos = Table([[img_esap, img_cert]], colWidths=[10 * cm, 6.5 * cm])
        tabla_logos.setStyle(TableStyle([
            ("ALIGN", (0, 0), (0, 0), "LEFT"),
            ("ALIGN", (1, 0), (1, 0), "RIGHT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        return [tabla_logos, Spacer(1, 10)]
    except Exception:
        return []


def _agregar_banner_docx(doc, titulo: str, subtitulo: str | None = None, color_hex: str = COLOR_INSTITUCIONAL) -> None:
    """Banda institucional de ancho completo (fondo de color, texto blanco)
    para usar como encabezado de portada en vez de un título plano
    centrado. Reutilizable por los 3 informes."""
    banner = doc.add_table(rows=1, cols=1)
    banner.autofit = False
    celda = banner.rows[0].cells[0]
    celda.width = Cm(17)
    _sombrear_celda(celda, color_hex)
    p_titulo = celda.paragraphs[0]
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)
    run_titulo.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if subtitulo:
        p_sub = celda.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(subtitulo)
        run_sub.italic = True
        run_sub.font.size = Pt(11)
        run_sub.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def _dibujar_banda_portada_pdf(titulo: str, subtitulo: str | None = None, color_hex: str = COLOR_INSTITUCIONAL, alto_cm: float = 3.2):
    """Fábrica de función onFirstPage para reportlab: dibuja una banda de
    color a todo el ancho en la parte superior de la primera página, con el
    título y subtítulo en blanco, y delega el pie de página a
    `funcion_pie` si se le pasa (ver `_combinar_callbacks_primera_pagina`)."""
    def _dibujar(canvas_pdf, doc_pdf):
        canvas_pdf.saveState()
        ancho_pagina, alto_pagina = LETTER
        alto_banda = alto_cm * cm
        canvas_pdf.setFillColor(colors.HexColor(f"#{color_hex}"))
        canvas_pdf.rect(0, alto_pagina - alto_banda, ancho_pagina, alto_banda, fill=1, stroke=0)
        canvas_pdf.setFillColor(colors.white)
        canvas_pdf.setFont("Helvetica-Bold", 15)
        canvas_pdf.drawCentredString(ancho_pagina / 2, alto_pagina - alto_banda * 0.45, titulo)
        if subtitulo:
            canvas_pdf.setFont("Helvetica-Oblique", 10)
            canvas_pdf.drawCentredString(ancho_pagina / 2, alto_pagina - alto_banda * 0.72, subtitulo)
        canvas_pdf.restoreState()
    return _dibujar


def _banner_portada_pdf_flowables(titulo: str, subtitulo: str | None = None, color_hex: str = COLOR_INSTITUCIONAL):
    """Versión del banner de portada como flowables (Table de una celda con
    fondo de color), para insertar directamente en el flujo del documento
    igual que en Word, sin depender de un callback de canvas."""
    filas = [[titulo]]
    if subtitulo:
        filas.append([subtitulo])
    tabla = Table(filas, colWidths=[17 * cm])
    estilo = [
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{color_hex}")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.white),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 15),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]
    if subtitulo:
        estilo.append(("FONTNAME", (0, 1), (-1, 1), "Helvetica-Oblique"))
        estilo.append(("FONTSIZE", (0, 1), (-1, 1), 10.5))
    tabla.setStyle(TableStyle(estilo))
    return [tabla, Spacer(1, 14)]


def _combinar_callbacks_primera_pagina(*funciones):
    """Combina varias funciones onFirstPage/onLaterPages de reportlab
    (por ejemplo la banda de portada + el pie de página) en una sola,
    ya que SimpleDocTemplate solo acepta un callback por evento."""
    def _combinado(canvas_pdf, doc_pdf):
        for funcion in funciones:
            if funcion:
                funcion(canvas_pdf, doc_pdf)
    return _combinado


def _agregar_divisor_seccion_docx(doc, titulo: str, icono: str = "▪", color_hex: str = COLOR_INSTITUCIONAL):
    """Divisor de sección de ancho completo (barra de color + ícono +
    título) para usar en vez de un doc.add_heading plano. Devuelve el
    párrafo del título por si se necesita seguir estilizando."""
    tabla = doc.add_table(rows=1, cols=1)
    tabla.autofit = False
    celda = tabla.rows[0].cells[0]
    celda.width = Cm(17)
    _sombrear_celda(celda, color_hex)
    p = celda.paragraphs[0]
    run = p.add_run(f"{icono}  {titulo}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()
    return p


def _divisor_seccion_pdf(titulo: str, icono: str = "▪", color_hex: str = COLOR_INSTITUCIONAL):
    """Versión PDF (lista de flowables) del divisor de sección: una tabla
    de una celda con fondo de color, texto blanco e ícono, seguida de un
    espaciador. Se antepone a cada Heading1 de los informes."""
    tabla = Table([[f"{icono}  {titulo}"]], colWidths=[17 * cm])
    tabla.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{color_hex}")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 12),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return [tabla, Spacer(1, 8)]


def _franjas_alternas_docx(tabla, color_hex_par: str = "F2F4F8") -> None:
    """Aplica sombreado en filas alternas (excluyendo el encabezado) a una
    tabla de python-docx, para tablas más legibles en documentos largos."""
    for indice_fila, fila in enumerate(tabla.rows[1:], start=1):
        if indice_fila % 2 == 0:
            for celda in fila.cells:
                _sombrear_celda(celda, color_hex_par)


# Esquema OFICIAL de 5 quintiles del MIPG (rangos de puntaje 0-100), usado
# para colorear de forma consistente cualquier tabla que muestre puntajes
# en los 3 informes — no debe confundirse con el esquema de 3 niveles de
# riesgo (alta/media/baja) de las brechas, que es una convención distinta
# y ya existente en este mismo módulo (_COLOR_HEX_POR_RIESGO).
_QUINTILES_MIPG = [
    (20, "E6534A"),   # 0-20   crítico
    (40, "F0A02E"),   # 21-40  bajo
    (60, "F5D65A"),   # 41-60  medio
    (80, "8CC152"),   # 61-80  satisfactorio
    (100, "2E8B57"),  # 81-100 sobresaliente
]


def _color_hex_quintil_mipg(puntaje) -> str | None:
    """Devuelve el color hex del quintil oficial MIPG correspondiente a un
    puntaje de 0 a 100, o None si el puntaje no es un número válido."""
    try:
        puntaje = float(puntaje)
    except (TypeError, ValueError):
        return None
    for limite, color_hex in _QUINTILES_MIPG:
        if puntaje <= limite:
            return color_hex
    return _QUINTILES_MIPG[-1][1]


def _color_hex_quintil_mipg(puntaje) -> str | None:
    """Devuelve el color hex del quintil oficial MIPG correspondiente a un
    puntaje de 0 a 100, o None si el puntaje no es un número válido."""
    try:
        puntaje = float(puntaje)
    except (TypeError, ValueError):
        return None
    for limite, color_hex in _QUINTILES_MIPG:
        if puntaje <= limite:
            return color_hex
    return _QUINTILES_MIPG[-1][1]


_QUINTILES_MIPG_NOMBRE_EMOJI = [
    (20, "Crítico", "🔴"),
    (40, "Bajo", "🟠"),
    (60, "Medio", "🟡"),
    (80, "Medio-alto", "🟢"),
    (100, "Consolidación", "✅"),
]


def _quintil_mipg(puntaje):
    """Devuelve (color_hex, emoji, nombre_categoria) para un puntaje de 0 a
    100, según el esquema oficial de 5 quintiles del MIPG. Usado por el
    Informe Ejecutivo para Alcaldes/Gobernadores (semáforo en lenguaje
    llano) y disponible para los otros informes si lo necesitan."""
    try:
        puntaje_num = float(puntaje)
    except (TypeError, ValueError):
        return ("999999", "⚪", "Sin dato")
    for limite, nombre, emoji in _QUINTILES_MIPG_NOMBRE_EMOJI:
        if puntaje_num <= limite:
            color_hex = _color_hex_quintil_mipg(puntaje_num)
            return (color_hex, emoji, nombre)
    color_hex = _QUINTILES_MIPG[-1][1]
    return (color_hex, "✅", "Consolidación")


DESCARGO_RESPONSABILIDAD_AMPLIADO = (
    "Este informe fue generado por el Sistema Integral de Diagnóstico del Desempeño "
    "Institucional (SIIEAP) a partir de datos reales del Índice de Desempeño "
    "Institucional (IDI-MIPG) publicado por el Departamento Administrativo de la "
    "Función Pública, complementado con un análisis generado por inteligencia "
    "artificial (Claude, Anthropic) como punto de partida académico y metodológico. "
    "El semáforo de quintiles, las brechas priorizadas y las lecturas de riesgo "
    "legal, financiero, administrativo y disciplinario que aquí se presentan son "
    "un ejercicio de apoyo a la toma de decisiones de la Alta Dirección — NO "
    "constituyen un dictamen jurídico, fiscal ni disciplinario oficial, y no "
    "sustituyen la valoración de la Oficina de Control Interno, la Oficina "
    "Asesora Jurídica, la Secretaría de Hacienda ni de los organismos de control "
    "(Procuraduría General de la Nación, Contraloría General de la República, "
    "Contralorías territoriales) de la entidad. Se entrega con fines académicos, "
    "como evidencia de ejecución del sistema y como insumo de trabajo para la "
    "asignatura."
)


# ---------------------------------------------------------------------------
# Bloque fijo: Contextualización de la Administración Pública Contemporánea
# ---------------------------------------------------------------------------

CONTEXTUALIZACION_TITULO = "Contextualización de la Administración Pública Contemporánea"

CONTEXTUALIZACION_INTRO = (
    "Este informe se enmarca en el recorrido teórico de la asignatura Enfoques y "
    "Teorías de la Administración Pública, que inicia en los antecedentes de la "
    "cultura griega para la distinción entre lo público y lo privado, y llega hasta "
    "los debates más recientes sobre la transformación digital del Estado. A "
    "continuación se sintetiza ese arco completo, como marco de lectura para el "
    "diagnóstico institucional que se presenta más adelante."
)

CONTEXTUALIZACION_ARCO_HISTORICO = [
    (
        "Antecedentes griegos de lo público y lo privado",
        "La distinción entre la esfera pública (la polis, los asuntos comunes, el "
        "koinon) y la esfera privada (el oikos, el hogar, el idion) es el punto de "
        "partida clásico para pensar qué es \"lo público\" y cómo se organiza su "
        "gobierno. De esa distinción original —lo que compete a todos frente a lo "
        "que compete a cada uno— se deriva la pregunta que atraviesa toda la "
        "disciplina: ¿quién decide sobre los asuntos comunes y con qué reglas?",
    ),
    (
        "Administración Pública clásica y el modelo burocrático weberiano",
        "La Administración Pública clásica (Wilson, Weber) responde a esa pregunta "
        "con un modelo de burocracia racional-legal: jerarquía, especialización, "
        "reglas escritas, impersonalidad y separación entre política y "
        "administración. Este modelo dota al Estado de previsibilidad y control, "
        "pero es también el que después será criticado por su rigidez, su lentitud "
        "y su desconexión de los resultados que efectivamente produce para el "
        "ciudadano.",
    ),
    (
        "Nuevo Institucionalismo en la Administración Pública",
        "March y Olsen (1989) proponen \"redescubrir las instituciones\": las reglas, "
        "normas y rutinas no son un telón de fondo neutro, sino que moldean "
        "directamente el comportamiento de las organizaciones públicas y explican "
        "por qué entidades con recursos similares logran resultados distintos. El "
        "institucionalismo distingue al menos tres variantes que se complementan: "
        "el institucionalismo racional (las reglas como incentivos que estructuran "
        "decisiones), el histórico (la dependencia de la trayectoria: lo que una "
        "entidad hizo antes condiciona lo que puede hacer ahora) y el sociológico "
        "(DiMaggio y Powell, 1983), que explica el isomorfismo institucional: las "
        "organizaciones se parecen entre sí no solo por eficiencia, sino por "
        "presión coercitiva (la norma obliga), mimética (se copia al que parece "
        "legítimo) y normativa (las profesiones y gremios difunden un estándar). "
        "Esta última variante es clave para leer el MIPG: explica por qué una "
        "entidad puede tener todos los formatos y comités que la norma exige "
        "(cumplimiento formal, isomorfismo) sin que eso se traduzca aún en una "
        "práctica real institucionalizada — la brecha entre \"reglas en el papel\" "
        "y \"reglas en uso\".",
    ),
    (
        "Nueva Gestión Pública (NGP)",
        "Surge en los años 80 y 90 (Hood, 1991; Osborne y Gaebler, 1992) como "
        "respuesta a la crisis fiscal del Estado de bienestar y a la percepción de "
        "ineficiencia burocrática. Retoma técnicas de gestión privada: "
        "descentralización de la autoridad gerencial, orientación a resultados e "
        "indicadores de desempeño, introducción de mecanismos de mercado y "
        "competencia, orientación al \"cliente-ciudadano\", y medición explícita del "
        "desempeño (los \"siete doctrinas\" de Hood). Su aporte fue instalar la "
        "cultura de la medición y la rendición de cuentas por resultados —la base "
        "misma de instrumentos como el FURAG—, pero ha sido criticada por "
        "fragmentar el aparato estatal en unidades autónomas difíciles de "
        "coordinar, por debilitar el ethos de lo público al importar lógicas de "
        "mercado, y por reducir el éxito de la gestión a indicadores que no "
        "siempre capturan el bienestar colectivo generado.",
    ),
    (
        "post-Nueva Gestión Pública (post-NGP) y Gobernanza Digital",
        "Ante la fragmentación que dejó la NGP, autores como Dunleavy et al. (2006, "
        "\"Digital Era Governance\") y Osborne (2006, \"Nueva Gobernanza Pública\") "
        "proponen reintegrar lo que la NGP separó: menos unidades autónomas y más "
        "coordinación interinstitucional, menos competencia interna y más redes de "
        "colaboración, menos cliente y más ciudadano co-productor del servicio. La "
        "post-NGP no abandona la orientación a resultados de la NGP, pero la "
        "combina con una lógica de red (gobernanza) y con las posibilidades que "
        "abre la digitalización del Estado para integrar de nuevo lo que estaba "
        "fragmentado. Desde la ESAP, Chica-Vélez y Salazar-Ortiz (2021) profundizan "
        "en esta transición y muestran cómo la gobernanza y la innovación pública se "
        "resignifican precisamente en el tránsito de la NGP hacia la post-NGP, "
        "dejando de ser un simple discurso de modernización para convertirse en una "
        "forma concreta de organizar y gestionar lo público.",
    ),
    (
        "Gobernanza Pública y Valor Público",
        "La gobernanza (Kooiman, Rhodes) reconoce que el Estado ya no gobierna en "
        "solitario: coordina una red de actores —mercado, sociedad civil, "
        "cooperación internacional— para producir resultados que ninguno lograría "
        "por separado. En paralelo, Mark Moore (1995, \"Creating Public Value\") "
        "propone el valor público como el criterio último de éxito de la gestión: "
        "no basta con ser eficiente (NGP) ni con coordinar bien una red "
        "(gobernanza) si el resultado no se traduce en bienestar colectivo "
        "legítimo y sostenible. El valor público exige, además, capacidad "
        "operativa real y legitimidad/respaldo político-social — el \"triángulo "
        "estratégico\" de Moore.",
    ),
    (
        "MIPG — Modelo Integrado de Planeación y Gestión",
        "El MIPG (Decreto 1499 de 2017) es la traducción institucional colombiana "
        "de estas tres corrientes: conserva del institucionalismo la atención a "
        "reglas y rutinas (dimensiones y políticas normadas), toma de la NGP la "
        "medición sistemática del desempeño (el FURAG y el IDI), y adopta de la "
        "gobernanza y el valor público la idea de que la meta final no es cumplir "
        "un formato sino generar resultados y valor para el ciudadano. Articula "
        "planeación, gestión del riesgo, talento humano, control interno y "
        "evaluación de resultados en un solo modelo de gestión pública.",
    ),
]

CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_TITULO = (
    "Ampliación: las teorías y enfoques de la Administración Pública Contemporánea "
    "(no son solo tres, y se incluyen aquí las corrientes de vanguardia más recientes)"
)

CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_INTRO = (
    "Las tres corrientes anteriores (Nuevo Institucionalismo, NGP, post-NGP) explican "
    "la base teórica del MIPG, pero la Administración Pública contemporánea ha seguido "
    "evolucionando en la última década con nuevos enfoques, cada uno con su propio "
    "cuerpo de literatura. Estos no son solo marcos conceptuales importados: cada uno "
    "se desarrolla aquí junto con la norma colombiana vigente que lo valida, citando "
    "el artículo específico cuando la fuente es una ley o decreto. Se incluyen, "
    "además de los diez enfoques del esquema integrador original, cinco corrientes "
    "adicionales de vanguardia administrativa (Gobierno como Plataforma, "
    "interoperabilidad de datos, co-creación y Design Thinking, Administración "
    "Pública Conductual y Gemelos Digitales de Territorio). Aclaración metodológica "
    "importante: los documentos CONPES (Consejo Nacional de Política Económica y "
    "Social) son documentos de POLÍTICA PÚBLICA, no leyes, y por tanto no tienen "
    "\"artículos\" — se citan aquí por su número y objetivo, no deben confundirse con "
    "normas de rango legal. Y dos salvedades honestas, señaladas explícitamente donde "
    "corresponde: la Inteligencia Artificial cuenta en Colombia con política pública "
    "(CONPES), pero todavía no con una ley integral que la regule; y la "
    "Administración Pública Conductual y los Gemelos Digitales de Territorio son, a "
    "la fecha de este informe, corrientes de vanguardia SIN desarrollo normativo "
    "propio verificado en Colombia — se incluyen como oportunidad de política "
    "pública, no como hecho normativo consumado."
)

CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS = [
    (
        "Estado Digital (e-Gobierno)",
        "El Estado Digital, o e-Gobierno, es la etapa en la que las tecnologías de la "
        "información dejan de ser un simple apoyo administrativo y se integran "
        "directamente en la prestación de servicios y en la toma de decisiones "
        "(Fountain, 2001, \"Building the Virtual State\"; Dunleavy et al., 2006). Su "
        "aporte clave es mostrar que la tecnología no se limita a automatizar un "
        "trámite existente: lo redefine, porque las reglas y la estructura de la "
        "organización terminan moldeadas por lo que el sistema de información permite "
        "o restringe (la \"tecnología en uso\" de Fountain). Colombia lo valida "
        "normativamente en el artículo 2, numeral 8, de la Ley 1341 de 2009 (que ordena "
        "al Gobierno fijar los mecanismos para la masificación del Gobierno en Línea, "
        "hoy Gobierno Digital), desarrollado por el Decreto 1008 de 2018 y actualizado "
        "por el Decreto 767 de 2022, cuyo artículo 2.2.9.1.1.3 (del Decreto 1078 de "
        "2015) fija los principios de la Política de Gobierno Digital vigente, una de "
        "las políticas de gestión y desempeño del MIPG.",
    ),
    (
        "Transformación Digital del Estado",
        "La Transformación Digital va un paso más allá del Estado Digital: no basta "
        "con digitalizar un procedimiento, sino que exige repensar de fondo la "
        "arquitectura institucional, los flujos de datos entre entidades y la cultura "
        "organizacional para operar bajo una lógica \"digital por defecto\" (OCDE, 2014, "
        "\"Recommendation on Digital Government Strategies\"). Colombia la desarrolla "
        "mediante el Decreto 1263 de 2022, que adiciona el Título 23 a la Parte 2 del "
        "Libro 2 del Decreto 1078 de 2015 con los lineamientos de transformación "
        "digital pública, y mediante la Estrategia Nacional Digital 2023-2026 del DNP "
        "(instrumento de política, no de rango legal), que se traducen en la Política "
        "de Gobierno Digital del MIPG.",
    ),
    (
        "Gobernanza Inteligente (Smart Governance)",
        "La Gobernanza Inteligente (Meijer y Bolívar, 2016) describe cómo los "
        "gobiernos usan datos, sensores y tecnologías de ciudad inteligente para "
        "coordinar la acción colectiva de forma más ágil y basada en evidencia casi "
        "en tiempo real. Extiende la lógica de red de la gobernanza pública (post-NGP) "
        "hacia un entorno digital instrumentado. En Colombia se apoya en el modelo de "
        "territorios y ciudades inteligentes que exige el Decreto 767 de 2022 y en el "
        "documento CONPES 3995 de 2020 (Política Nacional de Confianza y Seguridad "
        "Digital, sin artículos por ser documento de política, no ley), que da el "
        "marco de confianza sobre el que puede operar esa coordinación.",
    ),
    (
        "Inteligencia Artificial en el sector público",
        "La incorporación de Inteligencia Artificial en la Administración Pública "
        "(OCDE, 2019, \"Recomendación sobre Inteligencia Artificial\") ofrece "
        "oportunidades reales —automatización de análisis, diagnóstico predictivo, "
        "como el que realiza este mismo sistema— pero exige gestionar riesgos "
        "conocidos: sesgos algorítmicos, opacidad (\"caja negra\") y la necesidad de "
        "supervisión humana significativa sobre cualquier resultado generado por IA. "
        "En Colombia, esta corriente se ha traducido en política pública —los "
        "documentos CONPES 3975 de 2019 y CONPES 4144 de 2025 (Política Nacional de "
        "Inteligencia Artificial, con hoja de ruta a 2030), ninguno de los cuales tiene "
        "artículos por no ser leyes— pero, a la fecha de este informe, el país AÚN NO "
        "cuenta con una ley integral que regule la IA: existen varios proyectos de ley "
        "en trámite en el Congreso sobre la materia, aún sin sanción. Por eso este "
        "informe insiste en que el análisis generado con IA es siempre un punto de "
        "partida metodológico, nunca un dictamen definitivo.",
    ),
    (
        "Gobierno Abierto",
        "El Gobierno Abierto (Open Government Partnership, 2011) descansa en tres "
        "pilares que se refuerzan entre sí: transparencia (acceso a la información "
        "pública), participación ciudadana (involucrar a la ciudadanía en el diseño "
        "de la política) y colaboración (co-producción de soluciones entre Estado, "
        "sociedad civil y sector privado). Colombia es miembro de la Open Government "
        "Partnership desde 2011, y su pilar de transparencia está anclado en la Ley "
        "1712 de 2014 (Ley de Transparencia y del Derecho de Acceso a la Información "
        "Pública): el artículo 1 fija su objeto, el artículo 2 establece el principio "
        "de máxima publicidad (toda información en poder de una entidad pública es "
        "pública salvo excepción constitucional o legal expresa) y el artículo 3 fija "
        "los demás principios de la transparencia y el acceso a la información.",
    ),
    (
        "Gobernanza de Datos",
        "La Gobernanza de Datos trata el dato público como un activo estratégico que "
        "debe gestionarse con reglas claras de calidad, interoperabilidad, seguridad "
        "y uso responsable. En Colombia esta corriente tiene tres soportes legales "
        "concretos que operan en tensión productiva: el Decreto 1389 de 2022 "
        "(compilado como artículo 2.2.24.3.4 del Decreto 1078 de 2015), que crea el "
        "Comité Nacional de Datos para impulsar su gobernanza, uso y reutilización; el "
        "artículo 2 de la Ley 1712 de 2014, que exige apertura y máxima publicidad; y "
        "los artículos 1 y 2 de la Ley 1581 de 2012 (Habeas Data), que protegen los "
        "datos personales y ponen el límite a esa apertura cuando el dato es de una "
        "persona natural.",
    ),
    (
        "Administración Pública Basada en Evidencia",
        "La Administración Pública Basada en Evidencia (Evidence-Based Policy Making; "
        "Head, 2008; Nutley et al., 2007) exige que las decisiones de política "
        "pública se apoyen en datos y evaluaciones rigurosas, no solo en la intuición "
        "o en la presión política coyuntural. En Colombia, el propio Decreto 1499 de "
        "2017 que crea el MIPG es la traducción normativa de esta corriente: el "
        "artículo 2.2.22.3.2 lo define como marco de referencia para dirigir, "
        "planear, ejecutar, hacer seguimiento, evaluar y controlar la gestión "
        "pública, y el artículo 2.2.22.3.3 fija sus objetivos, obligando a medir el "
        "desempeño institucional (FURAG, IDI) como base de la toma de decisiones.",
    ),
    (
        "Capacidades Estatales",
        "Las Capacidades Estatales (Fukuyama, 2004; Grindle, 1997) son la aptitud "
        "real de un Estado o entidad para formular e implementar política pública, "
        "y suelen distinguirse al menos tres tipos: capacidad administrativa (talento "
        "humano, sistemas de información, procesos), capacidad fiscal (recursos "
        "disponibles y su ejecución) y capacidad política (legitimidad y respaldo "
        "para decidir y sostener una decisión en el tiempo). Colombia reconoce "
        "explícitamente estas asimetrías de capacidad entre entidades mediante el "
        "artículo 1 de la Ley 617 de 2000 (categorización presupuestal de "
        "departamentos, en desarrollo del artículo 302 de la Constitución) y su "
        "artículo 2 (categorización de distritos y municipios), junto con el "
        "artículo 3 de la Ley 1454 de 2011 (LOOT), que fija la autonomía, la "
        "descentralización y la asociatividad como principios rectores del "
        "ordenamiento territorial — la base normativa de la tipología y el 'Grupo "
        "par' que este sistema usa para no comparar entidades de capacidad muy "
        "distinta entre sí.",
    ),
    (
        "Resiliencia Institucional",
        "La Resiliencia Institucional (OCDE, 2020, en el marco de la respuesta a la "
        "pandemia) es la capacidad de una entidad para anticipar, absorber y "
        "recuperarse de choques —crisis sanitarias, desastres naturales, choques "
        "fiscales— sin perder su capacidad de generar valor público. Es, de todos los "
        "enfoques contemporáneos, el que tiene el respaldo legal colombiano más "
        "directo y antiguo: el artículo 8 de la Ley 1523 de 2012 crea el Sistema "
        "Nacional de Gestión del Riesgo de Desastres, y su artículo 42 obliga "
        "expresamente a toda entidad pública o privada que preste servicios públicos "
        "o desarrolle actividades de riesgo a realizar un análisis específico de "
        "riesgo, mientras el artículo 43 exige el respectivo plan de gestión del "
        "riesgo — desarrollado en detalle por el Decreto 2157 de 2017, que reglamenta "
        "justamente ese artículo 42. Esto incorpora la pregunta de resiliencia "
        "directamente al análisis institucional que hace este informe.",
    ),
    (
        "Agenda 2030, ODS y marcos multilaterales (OCDE, CEPAL, Naciones Unidas)",
        "La Agenda 2030 y los Objetivos de Desarrollo Sostenible (Naciones Unidas, "
        "2015) sitúan la gestión pública territorial dentro de compromisos globales, "
        "en particular el ODS 16 (paz, justicia e instituciones sólidas), que exige "
        "instituciones eficaces, responsables, transparentes y con participación "
        "ciudadana real. Colombia adoptó formalmente esta agenda mediante el "
        "documento CONPES 3918 de 2018 (estrategia de implementación de los ODS en el "
        "país, sin artículos por ser documento de política, no ley), y es miembro "
        "pleno de la OCDE desde 2020, lo que la obliga a alinear su gestión pública "
        "con los estándares técnicos de esa organización. La CEPAL adapta esos "
        "estándares al contexto latinoamericano y sus asimetrías estructurales, y "
        "Naciones Unidas los articula en metas verificables — el marco de referencia "
        "final frente al cual se mide, en último término, el desempeño institucional "
        "que reporta este sistema.",
    ),
    (
        "Gobierno como Plataforma",
        "El Gobierno como Plataforma (O'Reilly, 2010) propone que el Estado deje de "
        "construir un sistema aislado por cada trámite y en su lugar ofrezca "
        "infraestructura digital compartida y reutilizable (identidad, autenticación, "
        "interoperabilidad) sobre la cual cualquier entidad pueda montar sus "
        "servicios, evitando duplicar esfuerzos. Colombia lo implementa de forma "
        "concreta mediante el Decreto 620 de 2020: el artículo 2.2.17.2.2.1 garantiza "
        "el acceso a los servicios ciudadanos digitales base a través del "
        "'Articulador' (la Agencia Nacional Digital), y el artículo 2.2.17.2.2.2 "
        "establece que el servicio de interoperabilidad del Estado se presta de forma "
        "exclusiva a través de ese Articulador — es decir, la plataforma común que "
        "describe la teoría ya es, en Colombia, una obligación normativa concreta, "
        "fundamentada además en el artículo 147 de la Ley 1955 de 2019 (Plan Nacional "
        "de Desarrollo 2018-2022).",
    ),
    (
        "Interoperabilidad de datos entre entidades",
        "La interoperabilidad —la capacidad de que distintos sistemas de información "
        "intercambien datos sin fricción, siguiendo el modelo de referencia de países "
        "como Estonia (X-Road)— es la condición técnica que hace posible el Gobierno "
        "como Plataforma y la Gobernanza de Datos. En Colombia queda regulada por el "
        "mismo Decreto 620 de 2020 (artículo 2.2.17.2.2.2, servicio de "
        "interoperabilidad exclusivo del Articulador) y por el Decreto 1078 de 2015, "
        "que define el 'marco de interoperabilidad' como la estructura común de "
        "principios, recomendaciones y directrices —políticas, legales, "
        "organizacionales, semánticas y técnicas— que orientan el intercambio de "
        "información entre entidades del Estado.",
    ),
    (
        "Co-creación y Design Thinking en política pública",
        "La co-creación y el Design Thinking proponen que la ciudadanía deje de ser "
        "receptora pasiva de un servicio público y pase a codiseñarlo junto con la "
        "entidad, mediante prototipado rápido y validación directa con el usuario "
        "final. En Colombia esta corriente está parcialmente respaldada: el Decreto "
        "767 de 2022 incorpora la 'Innovación Pública Digital' como uno de los "
        "elementos estructurales de la Política de Gobierno Digital, y bajo ese "
        "marco el Centro de Innovación Pública Digital de MinTIC opera la "
        "metodología CoCrearE. Es importante precisar, con honestidad metodológica, "
        "que CoCrearE es una metodología operativa del Centro, no una norma con "
        "artículos propios: su respaldo legal es el Decreto 767 de 2022 que crea el "
        "marco dentro del cual esa metodología funciona.",
    ),
    (
        "Administración Pública Conductual (Nudge)",
        "La Administración Pública Conductual, o enfoque Nudge (Thaler y Sunstein, "
        "2008, \"Nudge: Improving Decisions About Health, Wealth, and Happiness\"), "
        "propone rediseñar la forma en que se presentan las opciones al ciudadano "
        "(el 'arquitecto de decisiones') para facilitar mejores decisiones sin "
        "restringir la libertad de elegir, apoyándose en evidencia de las ciencias "
        "del comportamiento. Es el enfoque contemporáneo que ha tenido más desarrollo "
        "internacional (la Behavioural Insights Team del Reino Unido, la oficina de "
        "Cass Sunstein en la administración Obama en Estados Unidos) pero, siendo "
        "estrictos con la evidencia disponible, no se identificó una ley o CONPES "
        "colombiano que institucionalice de forma específica una unidad o política "
        "de ciencias del comportamiento a nivel nacional. Se incluye aquí como "
        "corriente teórica vigente y de vanguardia, no como corriente ya "
        "normativizada en Colombia — una oportunidad de política pública más que un "
        "hecho normativo consumado.",
    ),
    (
        "Gemelos Digitales de Territorio (Digital Twins)",
        "Un gemelo digital de territorio es una réplica virtual, alimentada con "
        "datos reales y actualizados, de una ciudad o región, que permite simular el "
        "efecto de una decisión de política pública (una obra, una reubicación de "
        "población, un cambio de uso del suelo) antes de ejecutarla. Es, de los cinco "
        "enfoques añadidos en esta ampliación, el más incipiente en Colombia: no se "
        "identificó un decreto, ley o CONPES que regule específicamente los gemelos "
        "digitales de territorio a la fecha de este informe; el CONPES 4144 de "
        "2025 menciona tecnologías emergentes de forma general dentro de la Política "
        "Nacional de Inteligencia Artificial, pero sin un desarrollo propio para "
        "esta técnica en particular. Se incluye por su relevancia para el debate "
        "actual de vanguardia en gestión territorial, dejando explícito que aún es "
        "una oportunidad de política pública pendiente de desarrollo normativo en el "
        "país.",
    ),
]

CONTEXTUALIZACION_ESQUEMA_INTEGRADOR = [
    "Estado Digital",
    "Transformación Digital",
    "Gobernanza Inteligente",
    "Inteligencia Artificial",
    "Gobierno Abierto",
    "Gobernanza de Datos",
    "Valor Público",
    "Administración Pública Basada en Evidencia",
    "Capacidades Estatales",
    "Resiliencia Institucional",
    "Agenda 2030 y Objetivos de Desarrollo Sostenible (ODS)",
]

CONTEXTUALIZACION_CIERRE = (
    "Este esquema integrador evidencia que la Administración Pública contemporánea "
    "no es una ruptura frente a las teorías clásicas, sino su evolución: el Estado "
    "Digital y la Transformación Digital habilitan una Gobernanza Inteligente "
    "apoyada en Inteligencia Artificial; el Gobierno Abierto y la Gobernanza de "
    "Datos fortalecen el Valor Público y una Administración Pública Basada en "
    "Evidencia; y todo ello construye Capacidades Estatales y Resiliencia "
    "Institucional, en línea con la Agenda 2030 y los ODS y los marcos de "
    "referencia de organismos internacionales como la OCDE, la CEPAL y las "
    "Naciones Unidas, que orientan a los países en la modernización de su gestión "
    "pública hacia la innovación, la transparencia y la generación de valor "
    "público."
)

CADENA_INTERPRETACION_TITULO = "Cadena de interpretación institucional del SIIEAP"

CADENA_INTERPRETACION_INTRO = (
    "El SIIEAP no se limita a contar recomendaciones de Función Pública: cada "
    "brecha detectada se interpreta siguiendo una cadena de análisis que va desde "
    "el dato oficial hasta el plan de tratamiento y su seguimiento posterior. Esta "
    "es la ruta que sigue el sistema para cada brecha priorizada de la entidad:"
)

CADENA_INTERPRETACION_PASOS = [
    "Pregunta FURAG",
    "Dimensión – Política – Índice MIPG",
    "Estándar metodológico esperado",
    "Resultado oficial MIPG",
    "Recomendación oficial de Función Pública",
    "Brecha de implementación identificada",
    "Interpretación desde las teorías y enfoques de la Administración Pública",
    "Paradigma administrativo predominante",
    "Modelo de gestión pública relacionado",
    "Capacidad estatal comprometida",
    "Valor público afectado",
    "Gobernanza comprometida",
    "Transformación digital involucrada",
    "Normatividad aplicable",
    "Riesgo institucional",
    "Impacto en: índice, política, dimensión, IDI, capacidades institucionales y valor público",
    "Plan de tratamiento",
    "Indicadores KPI e indicadores KRI",
    "Seguimiento",
    "Nueva medición FURAG",
]

CADENA_INTERPRETACION_CIERRE = (
    "Esta cadena evita el error de tratar todas las recomendaciones como "
    "equivalentes: lo que importa no es cuántas brechas tiene una entidad, sino "
    "qué tan lejos está cada una del estándar exigido, qué capacidad estatal y "
    "qué valor público compromete, y qué tan bien se cierra el ciclo hasta la "
    "siguiente medición FURAG. El análisis de riesgo que se presenta más "
    "adelante en este informe sigue esta misma lógica de distancia al estándar, "
    "no de simple conteo."
)

DISCLAIMER_INFORME = (
    "Este informe fue generado por el Sistema Integral de Diagnóstico del "
    "Desempeño Institucional (SIIEAP), a partir de datos reales del Índice de "
    "Desempeño Institucional (IDI-MIPG) de Función Pública, complementado con un "
    "análisis generado por inteligencia artificial (Claude, Anthropic) como punto "
    "de partida académico y metodológico. No constituye un dictamen oficial de "
    "Función Pública ni sustituye la validación técnica, jurídica y del líder de "
    "proceso de la entidad analizada. Se entrega con fines académicos, como "
    "evidencia de ejecución del sistema y como insumo de trabajo para la "
    "asignatura."
)

# ---------------------------------------------------------------------------
# Glosario y convenciones del informe
#
# Fuente de los términos OFICIALES: Glosario MIPG v7 (octubre de 2021),
# Departamento Administrativo de la Función Pública — definiciones
# parafraseadas aquí, no citadas textualmente. Se listan aparte los términos
# de CONVENCIÓN PROPIA de SIIEAP, que NO existen en el glosario oficial y son
# exclusivos de este sistema, para que el lector nunca confunda unos con
# otros.
# ---------------------------------------------------------------------------

GLOSARIO_OFICIAL_MIPG = [
    (
        "MIPG (Modelo Integrado de Planeación y Gestión)",
        "Marco de referencia oficial del Estado colombiano (Decreto 1499 de "
        "2017) para dirigir, planear, ejecutar, hacer seguimiento, evaluar y "
        "controlar la gestión de las entidades públicas, con el fin de generar "
        "resultados que atiendan los planes de desarrollo y las necesidades de "
        "la ciudadanía.",
    ),
    (
        "Índice de Desempeño Institucional (IDI)",
        "Cifra oficial, publicada por Función Pública, que resume qué tan "
        "orientada está una entidad hacia la eficacia, la eficiencia y la "
        "calidad de su gestión. Es siempre la cifra de referencia principal en "
        "este informe.",
    ),
    (
        "FURAG (Formulario Único de Reporte y Avance de Gestión)",
        "Cuestionario oficial en línea, aplicado anualmente por Función "
        "Pública, con el que se recolecta la información que luego produce el "
        "IDI y mide el avance de las políticas de MIPG y del Modelo Estándar "
        "de Control Interno (MECI).",
    ),
    (
        "MECI (Modelo Estándar de Control Interno)",
        "Estructura oficial para evaluar la estrategia, la gestión y los "
        "mecanismos de autoevaluación de una entidad, adaptable a las "
        "necesidades propias de cada una.",
    ),
    (
        "Dimensión",
        "Cada uno de los grandes bloques de políticas y prácticas de gestión "
        "en que se organiza MIPG (por ejemplo, Talento Humano, Direccionamiento "
        "Estratégico, Evaluación de Resultados). El IDI se calcula agregando "
        "el desempeño de la entidad en estas dimensiones.",
    ),
    (
        "Grupo par",
        "Agrupación oficial de entidades con características homogéneas (por "
        "ejemplo, mismo tipo y tamaño de alcaldía) que Función Pública usa para "
        "que los resultados del FURAG sean comparables entre entidades "
        "similares.",
    ),
    (
        "Valor público",
        "Los cambios sociales reales y medibles que una entidad pública "
        "produce para responder a las necesidades de la ciudadanía; es el "
        "criterio último con el que se juzga si la gestión institucional tuvo "
        "sentido, más allá del cumplimiento formal de indicadores.",
    ),
    (
        "ODS (Objetivos de Desarrollo Sostenible)",
        "Los 17 objetivos globales adoptados por la ONU en 2015 para poner fin "
        "a la pobreza, proteger el planeta y garantizar la prosperidad para "
        "2030; se usan aquí como marco de referencia para la prospectiva.",
    ),
]

GLOSARIO_CONVENCIONES_SIIEAP = [
    (
        "IDI-MIPG (como se usa en este informe)",
        "Forma abreviada, propia de este sistema, para referirse al Índice de "
        "Desempeño Institucional dentro del marco de MIPG. No es una sigla "
        "oficial de Función Pública, aunque el IDI que reporta sí lo es.",
    ),
    (
        "Brecha (detectada por SIIEAP)",
        "Convención EXCLUSIVA de este sistema: un índice MIPG de la entidad "
        "con puntaje por debajo de 60 puntos, umbral de alerta metodológico "
        "interno que SIIEAP usa para priorizar qué revisar primero. NO es una "
        "cifra que publique Función Pública, y no debe confundirse con el IDI "
        "oficial ni con las recomendaciones oficiales del banco consolidado.",
    ),
    (
        "ISVPT (Índice Sintético de Valor Público Territorial)",
        "Indicador propio de este informe, construido con metodología "
        "académica (normalización min-max y agregación simple, siguiendo "
        "directrices de la OCDE de 2008 para indicadores compuestos), que "
        "compara las 7 dimensiones del IDI-MIPG de la entidad dentro de su "
        "grupo par. No es un índice oficial de Función Pública; es un "
        "ejercicio complementario de este sistema.",
    ),
    (
        "Umbral de alerta metodológico interno (60 puntos)",
        "El punto de corte que SIIEAP usa internamente para marcar un índice "
        "como brecha. Se insiste en todo el informe en que la meta real y "
        "oficial de la gestión pública es siempre el 100%, nunca el 60%.",
    ),
    (
        "Recomendaciones oficiales FP",
        "A diferencia de las brechas, este dato SÍ es oficial: es el total de "
        "recomendaciones del banco consolidado del Departamento Administrativo "
        "de la Función Pública, emitidas específicamente para la entidad "
        "analizada.",
    ),
]


def _agregar_glosario_docx(doc, estilo_normal=None):
    """Inserta la sección 'Glosario y convenciones' en el documento Word,
    separando siempre los términos OFICIALES (Glosario MIPG v7, Función
    Pública) de las CONVENCIONES PROPIAS de SIIEAP, para que el lector nunca
    confunda un dato exclusivo de este informe con un dato oficial.
    """
    doc.add_heading("Glosario y convenciones de este informe", level=1)
    doc.add_paragraph(
        "Los términos técnicos usados a lo largo de este informe se explican "
        "aquí en dos grupos: primero los términos OFICIALES del Modelo "
        "Integrado de Planeación y Gestión, tal como los define el Glosario "
        "MIPG versión 7 (octubre de 2021) del Departamento Administrativo de "
        "la Función Pública; y luego las CONVENCIONES PROPIAS de este sistema "
        "(SIIEAP), que no existen en el glosario oficial y no deben "
        "confundirse con cifras que publique Función Pública."
    )
    doc.add_heading("Términos oficiales (Glosario MIPG v7, Función Pública)", level=2)
    for termino, definicion in GLOSARIO_OFICIAL_MIPG:
        p = doc.add_paragraph()
        run_t = p.add_run(f"{termino}: ")
        run_t.bold = True
        p.add_run(definicion)
    doc.add_heading("Convenciones propias de SIIEAP (NO oficiales)", level=2)
    for termino, definicion in GLOSARIO_CONVENCIONES_SIIEAP:
        p = doc.add_paragraph()
        run_t = p.add_run(f"{termino}: ")
        run_t.bold = True
        run_t.font.color.rgb = RGBColor(0xB8, 0x5C, 0x00)
        p.add_run(definicion)
    doc.add_page_break()


def _agregar_glosario_pdf(elementos, estilos, estilo_normal, estilo_h2):
    """Versión PDF (reportlab) de _agregar_glosario_docx: misma separación
    estricta entre términos oficiales (Glosario MIPG v7, Función Pública) y
    convenciones propias de SIIEAP."""
    elementos.append(Paragraph("Glosario y convenciones de este informe", estilos["Heading1"]))
    elementos.append(Paragraph(
        "Los términos técnicos usados a lo largo de este informe se explican aquí en dos "
        "grupos: primero los términos OFICIALES del Modelo Integrado de Planeación y "
        "Gestión, tal como los define el Glosario MIPG versión 7 (octubre de 2021) del "
        "Departamento Administrativo de la Función Pública; y luego las CONVENCIONES "
        "PROPIAS de este sistema (SIIEAP), que no existen en el glosario oficial y no deben "
        "confundirse con cifras que publique Función Pública.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph("Términos oficiales (Glosario MIPG v7, Función Pública)", estilo_h2))
    for termino, definicion in GLOSARIO_OFICIAL_MIPG:
        elementos.append(Paragraph(f"<b>{termino}:</b> {definicion}", estilo_normal))
        elementos.append(Spacer(1, 4))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph("Convenciones propias de SIIEAP (NO oficiales)", estilo_h2))
    for termino, definicion in GLOSARIO_CONVENCIONES_SIIEAP:
        elementos.append(Paragraph(
            f'<font color="#B85C00"><b>{termino}:</b></font> {definicion}',
            estilo_normal,
        ))
        elementos.append(Spacer(1, 4))
    elementos.append(PageBreak())


def _ajustar_tabla_docx(tabla, anchos_cm, tamano_fuente_pt=9.5):
    """Fija anchos de columna reales (no solo sugeridos) y un tamaño de letra
    más pequeño para toda la tabla, evitando que el texto se desborde de las
    celdas — python-docx requiere fijar el ancho en CADA celda de CADA fila,
    no solo en la tabla, para que Word lo respete de forma consistente."""
    tabla.autofit = False
    tabla.allow_autofit = False
    for fila in tabla.rows:
        for celda, ancho in zip(fila.cells, anchos_cm):
            celda.width = Cm(ancho)
            for parrafo in celda.paragraphs:
                parrafo.paragraph_format.space_after = Pt(2)
                for run in parrafo.runs:
                    run.font.size = Pt(tamano_fuente_pt)


def _fecha_hoy_es():
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
        "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    hoy = datetime.now()
    return f"{hoy.day} de {meses[hoy.month - 1]} de {hoy.year}"


_COLOR_HEX_POR_RIESGO = {
    "alta": "F5B7B1",
    "media": "FAD7A0",
    "baja": "A9DFBF",
}


# Mapeo determinístico política MIPG -> enfoque(s) contemporáneo(s) + norma.
# Esto NO depende de que la IA lo mencione: se calcula siempre por código a
# partir de las brechas reales de la entidad, garantizando que la conexión
# con los 15 enfoques contemporáneos sea visible en TODOS los informes.
POLITICA_A_ENFOQUE_Y_NORMA = {
    "gestión estratégica del talento humano": ("Capacidades Estatales", "Ley 617 de 2000, art. 1-2; Ley 1454 de 2011, art. 3"),
    "integridad": ("Gobierno Abierto", "Ley 1712 de 2014; Ley 2013 de 2019 y Decreto 830 de 2021 (conflictos de interés)"),
    "planeación institucional": ("Administración Pública Basada en Evidencia", "Decreto 1499 de 2017, art. 2.2.22.3.2-3.3"),
    "gestión presupuestal y eficiencia del gasto": ("Capacidades Estatales", "Ley 610 de 2000, art. 1 (responsabilidad fiscal)"),
    "compras y contratación pública": ("Gobierno como Plataforma", "Decreto 620 de 2020, art. 2.2.17.2.2.1-2.2.2 (TVEC/Articulador)"),
    "fortalecimiento organizacional y simplificación de procesos": ("Capacidades Estatales", "Ley 617 de 2000; Ley 1454 de 2011"),
    "gobierno digital": ("Estado Digital / Transformación Digital del Estado", "Decreto 767 de 2022; Decreto 1263 de 2022"),
    "seguridad digital": ("Gobernanza de Datos", "CONPES 3995 de 2020 (Confianza y Seguridad Digital)"),
    "defensa jurídica": ("Capacidades Estatales", "Ley 87 de 1993 (control jurídico interno)"),
    "mejora normativa": ("Administración Pública Basada en Evidencia", "Decreto 1499 de 2017"),
    "servicio a las ciudadanías": ("Gobierno Abierto / Co-creación y Design Thinking", "Ley 1712 de 2014; Decreto 767 de 2022 (Innovación Pública Digital)"),
    "racionalización de trámites": ("Estado Digital", "Decreto 019 de 2012 (antitrámites); Decreto 767 de 2022"),
    "participación ciudadana en la gestión pública": ("Gobierno Abierto", "Ley 1712 de 2014; Colombia miembro OGP desde 2011"),
    "seguimiento y evaluación del desempeño institucional": ("Administración Pública Basada en Evidencia / Resiliencia Institucional", "Decreto 1499 de 2017; Ley 1523 de 2012"),
    "transparencia, acceso a la información y lucha contra la corrupción": ("Gobierno Abierto", "Ley 1712 de 2014, art. 1-3; ODS 16 (marco externo, ONU 2015)"),
    "gestión documental": ("Gobernanza de Datos", "Ley 594 de 2000 (Ley General de Archivos); Decreto 1389 de 2022"),
    "gestión de la información estadística": ("Gobernanza de Datos / Administración Pública Basada en Evidencia", "Decreto 1389 de 2022"),
    "gestión del conocimiento": ("Administración Pública Basada en Evidencia / Capacidades Estatales", "Decreto 1499 de 2017"),
    "control interno": ("Resiliencia Institucional / Capacidades Estatales", "Ley 87 de 1993, art. 1-2-9; Ley 1523 de 2012, art. 8"),
}


def _enfoque_y_norma_de_politica(nombre_politica: str):
    """Busca el enfoque contemporáneo y la norma asociada a una política MIPG,
    normalizando el nombre para tolerar variantes de mayúsculas/redacción."""
    clave = str(nombre_politica).strip().lower()
    clave = clave.replace("política de ", "").replace("política ", "")
    for politica_conocida, valor in POLITICA_A_ENFOQUE_Y_NORMA.items():
        if politica_conocida in clave or clave in politica_conocida:
            return valor
    return ("Capacidades Estatales", "Decreto 1499 de 2017 (MIPG)")


# ---------------------------------------------------------------------------
# Marco de responsabilidad de la Alta Dirección: artículo 6 de la
# Constitución y cruce política MIPG -> tipo de riesgo (jurídico, fiscal,
# administrativo o disciplinario) -> norma -> consecuencia posible.
# Uso exclusivo del Informe Ejecutivo para Alcaldes/Gobernadores, pero vive
# aquí (módulo compartido) por si otro informe necesita la misma lectura.
# ---------------------------------------------------------------------------

ARTICULO_6_CONSTITUCION_TITULO = "Artículo 6 de la Constitución Política de Colombia (1991)"

ARTICULO_6_CONSTITUCION_TEXTO = (
    "Los particulares sólo son responsables ante las autoridades por infringir "
    "la Constitución y las leyes. Los servidores públicos lo son por la misma "
    "causa y por omisión o extralimitación en el ejercicio de sus funciones."
)

ARTICULO_6_CONSTITUCION_EXPLICACION = (
    "Esta es la razón de fondo por la que cada brecha de este informe importa más "
    "allá del puntaje: a diferencia de un particular, que solo responde por infringir "
    "la ley, el servidor público de esta entidad responde también por OMISIÓN — es "
    "decir, por no hacer lo que la ley le exige — y por EXTRALIMITACIÓN — por hacer "
    "más de lo que la ley le permite. Una brecha del MIPG sostenida en el tiempo, sin "
    "plan de mejoramiento ni evidencia de gestión, es precisamente ese escenario de "
    "omisión que el artículo 6 sanciona, y puede derivar en responsabilidad jurídica, "
    "fiscal, administrativa y/o disciplinaria de quien tenía el deber legal de actuar."
)

# (tipo_de_riesgo, norma_principal, consecuencia_posible_breve)
POLITICA_A_RIESGO_ALTA_DIRECCION = {
    "gestión estratégica del talento humano": (
        "Administrativo y disciplinario",
        "Ley 909 de 2004 (empleo público); Ley 1952 de 2019 modif. Ley 2094 de 2021 (Código General Disciplinario)",
        "Faltas en la carrera administrativa y en la evaluación del desempeño pueden derivar en nulidad de actos de personal y en investigación disciplinaria contra el nominador.",
    ),
    "integridad": (
        "Disciplinario",
        "Ley 1952 de 2019 modif. Ley 2094 de 2021; Ley 2013 de 2019 y Decreto 830 de 2021 (conflictos de interés)",
        "Los conflictos de interés no declarados son falta disciplinaria autónoma, independiente de si hubo o no un daño patrimonial.",
    ),
    "planeación institucional": (
        "Administrativo",
        "Decreto 1499 de 2017, art. 2.2.22.3.2-3.3",
        "La ausencia de planeación institucional debilita la defensa de la entidad ante entes de control al no poder demostrar la debida diligencia de gestión.",
    ),
    "gestión presupuestal y eficiencia del gasto": (
        "Fiscal",
        "Ley 610 de 2000 (proceso de responsabilidad fiscal); Ley 617 de 2000 (límites de gasto)",
        "El detrimento patrimonial derivado de una gestión presupuestal deficiente puede abrir proceso de responsabilidad fiscal ante la Contraloría, con obligación de resarcir el daño con el patrimonio del responsable.",
    ),
    "compras y contratación pública": (
        "Fiscal y disciplinario",
        "Ley 80 de 1993 y Ley 1150 de 2007 (contratación estatal); Ley 1474 de 2011 (Estatuto Anticorrupción); Ley 610 de 2000",
        "Irregularidades contractuales son la causa más frecuente de procesos de responsabilidad fiscal y disciplinaria contra alcaldes y gobernadores en Colombia.",
    ),
    "fortalecimiento organizacional y simplificación de procesos": (
        "Administrativo",
        "Ley 617 de 2000; Ley 1454 de 2011 (ordenamiento territorial)",
        "Una estructura organizacional no ajustada a la capacidad fiscal real de la entidad es causal de observaciones de viabilidad fiscal ante el Ministerio de Hacienda.",
    ),
    "gobierno digital": (
        "Administrativo",
        "Decreto 767 de 2022; Decreto 1263 de 2022",
        "El incumplimiento de la política de Gobierno Digital es hoy objeto de seguimiento directo por parte del MinTIC y de Función Pública dentro del FURAG.",
    ),
    "seguridad digital": (
        "Administrativo y disciplinario",
        "CONPES 3995 de 2020; Ley 1952 de 2019 (deber de custodia de la información)",
        "La pérdida o filtración de información institucional por fallas de seguridad digital compromete disciplinariamente a quien tenía el deber de custodia.",
    ),
    "defensa jurídica": (
        "Jurídico y fiscal",
        "Ley 2213 de 2022 (procesos judiciales); Ley 678 de 2001 (acción de repetición); Ley 610 de 2000",
        "Una defensa jurídica débil incrementa el riesgo de condenas judiciales contra la entidad, que a su vez pueden derivar en acción de repetición contra el servidor responsable.",
    ),
    "mejora normativa": (
        "Administrativo",
        "Decreto 1499 de 2017",
        "Normas internas desactualizadas o contradictorias con la ley superior son causal de nulidad de actos administrativos ante lo contencioso administrativo.",
    ),
    "servicio a las ciudadanías": (
        "Disciplinario",
        "Ley 1755 de 2015 (derecho de petición); Ley 1952 de 2019",
        "No responder oportunamente las peticiones ciudadanas es falta disciplinaria expresa y puede derivar en acción de tutela contra la entidad.",
    ),
    "racionalización de trámites": (
        "Administrativo",
        "Decreto 019 de 2012 (antitrámites); Decreto 767 de 2022",
        "Exigir requisitos no autorizados por la ley es causal de responsabilidad disciplinaria del funcionario que los exige.",
    ),
    "participación ciudadana en la gestión pública": (
        "Disciplinario y administrativo",
        "Ley 1757 de 2015 (participación democrática)",
        "El incumplimiento de los espacios de participación obligatorios (rendición de cuentas, presupuesto participativo) es objeto de seguimiento por los entes de control.",
    ),
    "seguimiento y evaluación del desempeño institucional": (
        "Administrativo",
        "Decreto 1499 de 2017; Ley 1523 de 2012 (gestión del riesgo)",
        "La falta de seguimiento y evaluación impide demostrar la debida diligencia de la Alta Dirección ante un eventual proceso de responsabilidad.",
    ),
    "transparencia, acceso a la información y lucha contra la corrupción": (
        "Disciplinario y penal",
        "Ley 1712 de 2014 (transparencia); Ley 1474 de 2011 (Estatuto Anticorrupción); Código Penal, Título XV",
        "Las faltas en transparencia son, junto con la contratación, la puerta de entrada más común a investigaciones disciplinarias y, en los casos más graves, penales.",
    ),
    "gestión documental": (
        "Disciplinario",
        "Ley 594 de 2000 (Ley General de Archivos); Decreto 1389 de 2022",
        "La pérdida, alteración o destrucción de documentos públicos es falta disciplinaria gravísima según el Código General Disciplinario.",
    ),
    "gestión de la información estadística": (
        "Administrativo",
        "Decreto 1389 de 2022",
        "Reportar información estadística inexacta a Función Pública compromete la validez del FURAG y de los indicadores oficiales de la entidad.",
    ),
    "gestión del conocimiento": (
        "Administrativo",
        "Decreto 1499 de 2017",
        "La pérdida de conocimiento institucional por rotación de personal debilita la capacidad de respuesta de la entidad ante requerimientos de entes de control.",
    ),
    "control interno": (
        "Disciplinario y fiscal",
        "Ley 87 de 1993, art. 1-2-9; Ley 1523 de 2012, art. 8",
        "Un sistema de control interno débil o inexistente agrava la responsabilidad de la Alta Dirección, porque elimina la primera línea de defensa que debía advertir el riesgo a tiempo.",
    ),
}


def _riesgo_alta_direccion_de_politica(nombre_politica: str):
    """Análogo a _enfoque_y_norma_de_politica, pero para la lectura de
    riesgo legal/fiscal/administrativo/disciplinario dirigida a la Alta
    Dirección (alcalde/gobernador)."""
    clave = str(nombre_politica).strip().lower()
    clave = clave.replace("política de ", "").replace("política ", "")
    for politica_conocida, valor in POLITICA_A_RIESGO_ALTA_DIRECCION.items():
        if politica_conocida in clave or clave in politica_conocida:
            return valor
    return (
        "Administrativo",
        "Decreto 1499 de 2017 (MIPG)",
        "Toda brecha sostenida en el MIPG debilita la capacidad de la entidad para demostrar debida diligencia ante los entes de control.",
    )


def _color_hex_riesgo(nivel_riesgo) -> str | None:
    if not nivel_riesgo:
        return None
    return _COLOR_HEX_POR_RIESGO.get(str(nivel_riesgo).strip().lower())


def _sombrear_celda(celda, color_hex: str) -> None:
    """Colorea el fondo de una celda de tabla de python-docx (no hay API
    pública directa, se manipula el XML de la celda)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    propiedades_celda = celda._tc.get_or_add_tcPr()
    sombreado = OxmlElement("w:shd")
    sombreado.set(qn("w:val"), "clear")
    sombreado.set(qn("w:color"), "auto")
    sombreado.set(qn("w:fill"), color_hex)
    propiedades_celda.append(sombreado)


# ---------------------------------------------------------------------------
# Generación del .docx
# ---------------------------------------------------------------------------

def generar_reporte_docx(nombre_entidad, diag, analisis_ia_texto, resultado_isvpt=None, resultado_360=None, idi_oficial=None, cruce_recomendaciones=None, total_recomendaciones_entidad=None):
    """Devuelve un BytesIO con el informe técnico en formato Word (.docx).

    resultado_isvpt (opcional): un motor_isvpt.ResultadoISVPT ya calculado
    para el grupo de comparación de la entidad. Si se pasa, se agrega una
    sección con el Índice Sintético de Valor Público Territorial (novedad
    metodológica inspirada en el ISDEL de Vélez Tamayo et al., 2026).
    resultado_360 (opcional): un motor_analisis_360.ResultadoAnalisis360 ya
    calculado, usado para el Resumen Ejecutivo (comparación real contra el
    grupo par, igual que en un informe profesional de auditoría).
    idi_oficial (opcional pero MUY recomendado): el IDI oficial publicado
    por Función Pública para esta entidad, leído directamente del archivo
    oficial. Si se proporciona, este informe SIEMPRE lo usa como la cifra
    protagonista (portada, resumen ejecutivo); diag.idi_estimado (el
    cálculo interno de SIIEAP) se muestra únicamente como nota metodológica
    de verificación, nunca como reemplazo del dato oficial.
    """
    doc = Document()

    # Tipografía base más grande y legible (el informe se lee en pantalla y se imprime)
    estilo_normal_doc = doc.styles["Normal"]
    estilo_normal_doc.font.size = Pt(11.5)
    estilo_normal_doc.font.name = "Calibri"
    estilo_normal_doc.paragraph_format.space_after = Pt(8)
    estilo_normal_doc.paragraph_format.line_spacing = 1.15

    # El IDI OFICIAL de Función Pública es siempre la cifra protagonista.
    # Si por alguna razón no se cargó (ej. captura manual), se usa el cálculo
    # interno como única cifra disponible, dejándolo claro en el texto.
    idi_protagonista = idi_oficial if idi_oficial is not None else diag.idi_estimado
    hay_diferencia_idi = (
        idi_oficial is not None and diag.idi_estimado is not None
        and round(idi_oficial, 2) != round(diag.idi_estimado, 2)
    )
    nivel_riesgo_global = "ALTO" if (idi_protagonista or 0) < 40 else ("MEDIO" if (idi_protagonista or 0) < 70 else "BAJO")

    # Portada — logos institucionales + banner de color en vez de un título plano
    _agregar_logos_docx(doc)
    _agregar_banner_docx(
        doc,
        "Modelo de Conocimiento Institucional del Sistema de Inteligencia Artificial "
        "para la Evaluación Integral del Desempeño Institucional en Entidades Públicas (SIIEAP)",
        "Informe de Diagnóstico Institucional y Plan de Mejoramiento Prospectivo",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nombre_entidad)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"Generado el {_fecha_hoy_es()} · Índice de Desempeño Institucional (IDI-MIPG), Decreto 1499 de 2017")

    p_autoria = doc.add_paragraph()
    p_autoria.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_autoria = p_autoria.add_run(
        "Docente: Norma Elizabeth Álvarez Grajales · Área de conocimiento: Entidades "
        "Públicas y del Desarrollo · Escuela Superior de Administración Pública (ESAP)"
    )
    run_autoria.italic = True
    run_autoria.font.size = Pt(9)

    # Total de recomendaciones OFICIALES de Función Pública para esta entidad.
    # Prioridad: el total real de recomendaciones cargadas para la entidad
    # (todas las que trae el consolidado/archivo, sin filtrar por brecha).
    # Solo si no se recibió ese dato, se cae al conteo derivado del cruce
    # (que subestima el total real, porque solo cuenta recomendaciones
    # vinculadas a brechas detectadas).
    total_recomendaciones = (
        total_recomendaciones_entidad
        if total_recomendaciones_entidad is not None
        else (sum(len(lista) for lista in cruce_recomendaciones.values()) if cruce_recomendaciones else None)
    )

    doc.add_paragraph()
    n_columnas_portada = 4 if total_recomendaciones is not None else 3
    tabla_portada = doc.add_table(rows=1, cols=n_columnas_portada)
    tabla_portada.style = "Light Grid Accent 1"
    tabla_portada.autofit = False
    tabla_portada.allow_autofit = False
    celdas_portada = tabla_portada.rows[0].cells
    celdas_portada[0].text = f"IDI oficial (Función Pública)\n{idi_protagonista}"
    celdas_portada[1].text = f"Nivel de riesgo global\n{nivel_riesgo_global}"
    celdas_portada[2].text = f"Brechas detectadas (dato exclusivo de este informe)\n{len(diag.brechas)}"
    ancho_celda = Cm(4.2) if n_columnas_portada == 4 else Cm(5.5)
    if total_recomendaciones is not None:
        celdas_portada[3].text = f"Recomendaciones oficiales FP\n{total_recomendaciones}"
    for celda in celdas_portada:
        celda.width = ancho_celda
        for parrafo in celda.paragraphs:
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run_c in parrafo.runs:
                run_c.bold = True
                run_c.font.size = Pt(13)

    doc.add_paragraph()
    p_explica = doc.add_paragraph()
    run_explica = p_explica.add_run(
        "⚠ Brechas detectadas: esta cifra es un cálculo EXCLUSIVO de este informe (SIIEAP) — "
        "NO es un dato que publique el Departamento Administrativo de la Función Pública. "
        "Corresponde al número de índices MIPG de esta entidad con puntaje por debajo de 60 "
        "puntos, umbral de alerta metodológico interno del sistema (la meta plena de la gestión "
        "pública es siempre el 100%, no 60)."
        + (
            f" ✔ Recomendaciones oficiales FP: este dato SÍ es oficial — corresponde al total de "
            f"recomendaciones del banco consolidado del Departamento Administrativo de la Función "
            f"Pública emitidas para esta entidad (no se limita a las brechas detectadas por este "
            f"informe)."
            if total_recomendaciones is not None else ""
        )
    )
    run_explica.italic = True
    run_explica.font.size = Pt(9)
    run_explica.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        p_grupo = doc.add_paragraph()
        p_grupo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_grupo.add_run(
            f"Grupo de comparación: {resultado_360.filtro_descripcion} "
            f"({resultado_360.n_entidades} entidades) · IDI promedio del grupo: {resultado_360.promedio_idi}"
        )
    if hay_diferencia_idi:
        p_nota_idi = doc.add_paragraph()
        p_nota_idi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nota_idi = p_nota_idi.add_run(
            f"(Nota: el cálculo interno de verificación de SIIEAP, revisando la suma/ponderación de los "
            f"índices reportados, arroja {diag.idi_estimado}. Esta es una cifra de validación metodológica "
            f"interna que NO reemplaza ni desvirtúa el IDI oficial de Función Pública, que es siempre el "
            f"que prevalece.)"
        )
        run_nota_idi.italic = True
        run_nota_idi.font.size = Pt(9)
        run_nota_idi.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    _agregar_glosario_docx(doc)

    # 1. Resumen ejecutivo — cifras reales de ESTA entidad, arriba de todo
    _agregar_divisor_seccion_docx(doc, "1. Resumen ejecutivo", icono="📊")
    parrafos_resumen = [
        f"{nombre_entidad} obtuvo un Índice de Desempeño Institucional (IDI) oficial de {idi_protagonista} "
        f"sobre 100 en la vigencia analizada, con un nivel de riesgo global {nivel_riesgo_global}."
    ]
    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        brecha_grupo = round((idi_protagonista or 0) - resultado_360.promedio_idi, 2)
        comparativo = "por debajo" if brecha_grupo < 0 else "por encima"
        parrafos_resumen.append(
            f"Frente a su grupo de comparación ({resultado_360.filtro_descripcion}, "
            f"{resultado_360.n_entidades} entidades), la entidad se ubica {abs(brecha_grupo)} puntos "
            f"{comparativo} del promedio del grupo ({resultado_360.promedio_idi})."
        )
        if resultado_360.percentil_entidad_referencia is not None:
            parrafos_resumen.append(
                f"Esto la ubica en el percentil {resultado_360.percentil_entidad_referencia}% de su grupo."
            )
    dims_criticas = sorted(diag.resultados_por_dimension, key=lambda r: (r.promedio if r.promedio is not None else 0))[:3]
    if dims_criticas:
        nombres_criticas = ", ".join(f"{r.codigo} {r.nombre} ({r.promedio})" for r in dims_criticas)
        parrafos_resumen.append(f"Las dimensiones más críticas son: {nombres_criticas}.")
    dims_fuertes = [r for r in diag.resultados_por_dimension if (r.promedio or 0) >= 60]
    if dims_fuertes:
        nombres_fuertes = ", ".join(f"{r.codigo} {r.nombre} ({r.promedio})" for r in dims_fuertes)
        parrafos_resumen.append(f"Como fortaleza relativa, se destaca(n): {nombres_fuertes}.")
    else:
        parrafos_resumen.append(
            "Ninguna dimensión alcanza el umbral de 60 puntos: el patrón de brechas simultáneas en "
            "múltiples dimensiones sugiere una debilidad estructural de capacidad institucional "
            "(Oszlak), más que fallas puntuales y aisladas de gestión."
        )
    if hay_diferencia_idi:
        parrafos_resumen.append(
            f"Nota metodológica: Función Pública reporta un IDI oficial de {idi_oficial} para esta "
            f"entidad, cifra que este informe usa siempre como referencia principal. El sistema SIIEAP, "
            f"al revisar internamente la suma y ponderación de los índices del archivo oficial, calcula "
            f"además un valor de verificación de {diag.idi_estimado}; esta diferencia es un insumo para "
            f"depurar la metodología interna del sistema, y en ningún caso reemplaza, desvirtúa ni "
            f"contradice la cifra oficial de Función Pública."
        )
    parrafos_resumen.append(
        f"Este informe (SIIEAP) detectó, con metodología propia y exclusiva del sistema — no con "
        f"una cifra publicada por la Función Pública —, {len(diag.brechas)} brechas de "
        "implementación por debajo de 60 puntos (el umbral de alerta metodológico interno que usa "
        "este sistema para priorizar el análisis; la meta plena de la gestión pública, se insiste, "
        "es el 100% de cumplimiento, no 60 puntos), todas desarrolladas en detalle en este informe "
        "— sin selección ni recorte."
    )
    for parrafo_r in parrafos_resumen:
        doc.add_paragraph(parrafo_r)

    # Contextualización académica (bloque fijo)
    _agregar_divisor_seccion_docx(doc, CONTEXTUALIZACION_TITULO, icono="📚")
    doc.add_paragraph(CONTEXTUALIZACION_INTRO)

    doc.add_heading("Arco teórico: de los griegos a la Administración Pública contemporánea", level=2)
    for titulo_hito, texto_hito in CONTEXTUALIZACION_ARCO_HISTORICO:
        p = doc.add_paragraph()
        run_hito = p.add_run(f"{titulo_hito}. ")
        run_hito.bold = True
        p.add_run(texto_hito)

    doc.add_heading(CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_TITULO, level=2)
    doc.add_paragraph(CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_INTRO)
    for titulo_teoria, texto_teoria in CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS:
        p_teoria = doc.add_paragraph()
        run_titulo_teoria = p_teoria.add_run(f"{titulo_teoria}. ")
        run_titulo_teoria.bold = True
        p_teoria.add_run(texto_teoria)

    doc.add_heading("Esquema integrador de la Administración Pública contemporánea", level=2)
    p_esquema = doc.add_paragraph()
    run_esquema = p_esquema.add_run(" → ".join(CONTEXTUALIZACION_ESQUEMA_INTEGRADOR))
    run_esquema.italic = True
    doc.add_paragraph(CONTEXTUALIZACION_CIERRE)

    doc.add_page_break()

    # Cadena de interpretación institucional (bloque fijo)
    _agregar_divisor_seccion_docx(doc, CADENA_INTERPRETACION_TITULO, icono="🔗")
    doc.add_paragraph(CADENA_INTERPRETACION_INTRO)
    for paso in CADENA_INTERPRETACION_PASOS:
        doc.add_paragraph(paso, style="List Bullet")
    doc.add_paragraph()
    p_cierre_cadena = doc.add_paragraph()
    run_cierre_cadena = p_cierre_cadena.add_run(CADENA_INTERPRETACION_CIERRE)
    run_cierre_cadena.italic = True

    doc.add_page_break()

    # Resultado real del diagnóstico
    _agregar_divisor_seccion_docx(doc, "Resultado del diagnóstico institucional (datos reales)", icono="📈")
    p_idi = doc.add_paragraph()
    run_idi = p_idi.add_run(f"IDI estimado: {diag.idi_estimado}")
    run_idi.bold = True
    run_idi.font.size = Pt(14)

    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Light Grid Accent 1"
    encabezados = tabla.rows[0].cells
    for celda, texto in zip(encabezados, ["Dimensión", "Promedio", "Riesgo", "Índices"]):
        celda.text = texto
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True

    for r in diag.resultados_por_dimension:
        fila = tabla.add_row().cells
        fila[0].text = f"{r.codigo} {r.nombre}"
        fila[1].text = str(r.promedio)
        fila[2].text = str(r.nivel_riesgo)
        fila[3].text = f"{r.n_indices_evaluados}/{r.n_indices_esperados}"
        color_fondo = _color_hex_riesgo(r.nivel_riesgo)
        if color_fondo:
            _sombrear_celda(fila[2], color_fondo)
        color_quintil = _color_hex_quintil_mipg(r.promedio)
        if color_quintil:
            _sombrear_celda(fila[1], color_quintil)
    _ajustar_tabla_docx(tabla, anchos_cm=[9.5, 2.3, 2.3, 2.0])

    doc.add_paragraph()
    try:
        buffer_grafica_dim = generar_grafica_dimensiones(diag)
        doc.add_picture(buffer_grafica_dim, width=Inches(6.2))
        parrafo_imagen = doc.paragraphs[-1]
        parrafo_imagen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass  # si falla la gráfica, el informe sigue sin ella

    doc.add_heading("Brechas priorizadas (todas las detectadas)", level=2)
    if diag.brechas:
        try:
            buffer_grafica_brechas = generar_grafica_brechas(diag)
            if buffer_grafica_brechas:
                doc.add_picture(buffer_grafica_brechas, width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
        except Exception:
            pass
        for b in diag.brechas:
            doc.add_paragraph(
                f"{b.codigo_indice} ({b.puntaje}) — {b.nombre_indice} · {b.politica}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No se detectaron brechas por debajo del umbral con los datos disponibles.")

    # Tabla de conexión política -> enfoque contemporáneo -> norma (garantizada por código)
    if diag.brechas:
        doc.add_page_break()
        doc.add_heading("Conexión de las brechas con los enfoques contemporáneos de la Administración Pública", level=2)
        doc.add_paragraph(
            "Cada política con brechas se conecta aquí, de forma sistemática, con el enfoque "
            "contemporáneo que mejor la explica (de los 15 desarrollados en la contextualización "
            "de este informe) y con la norma colombiana que lo respalda."
        )
        politicas_con_brecha = []
        vistas = set()
        for b in diag.brechas:
            if b.politica not in vistas:
                vistas.add(b.politica)
                politicas_con_brecha.append(b.politica)
        tabla_enfoques = doc.add_table(rows=1, cols=3)
        tabla_enfoques.style = "Light Grid Accent 1"
        enc_ef = tabla_enfoques.rows[0].cells
        for celda, texto in zip(enc_ef, ["Política con brecha", "Enfoque contemporáneo", "Norma"]):
            celda.text = texto
            for parrafo in celda.paragraphs:
                for run_enc in parrafo.runs:
                    run_enc.bold = True
        for politica in politicas_con_brecha:
            enfoque, norma = _enfoque_y_norma_de_politica(politica)
            fila = tabla_enfoques.add_row().cells
            fila[0].text = politica
            fila[1].text = enfoque
            fila[2].text = norma
        _ajustar_tabla_docx(tabla_enfoques, anchos_cm=[5.0, 4.5, 6.0], tamano_fuente_pt=8.5)
        _franjas_alternas_docx(tabla_enfoques)

    if resultado_isvpt is not None and resultado_isvpt.isvpt_entidad_referencia is not None:
        doc.add_page_break()
        _agregar_divisor_seccion_docx(doc, "El Termómetro del Valor Público: Índice Sintético de Valor Público Territorial (ISVPT)", icono="🎯")
        doc.add_paragraph(
            "Como complemento al IDI oficial, este informe incorpora un ejercicio de índice "
            "sintético construido con la metodología académica validada por Vélez Tamayo, "
            "Ortiz-Muñoz y Cardona Montoya (2026) para el Índice Sintético de Desarrollo "
            "Económico Local (ISDEL) — normalización min-max y agregación aritmética simple, "
            "siguiendo las directrices de la OCDE (2008) para indicadores compuestos — "
            "aplicada aquí a las 7 dimensiones reales del IDI-MIPG dentro del grupo de "
            "comparación de esta entidad."
        )
        p_isvpt = doc.add_paragraph()
        run_isvpt = p_isvpt.add_run(
            f"ISVPT de {nombre_entidad}: {resultado_isvpt.isvpt_entidad_referencia} "
            f"(posición {resultado_isvpt.posicion_entidad_referencia} de "
            f"{resultado_isvpt.n_entidades} en su grupo de comparación)."
        )
        run_isvpt.bold = True

        if resultado_isvpt.subindices_por_dimension_entidad:
            doc.add_heading("Subíndices normalizados por dimensión (0 = el más rezagado del grupo, 1 = el mejor)", level=2)
            tabla_isvpt = doc.add_table(rows=1, cols=2)
            tabla_isvpt.style = "Light Grid Accent 1"
            enc = tabla_isvpt.rows[0].cells
            enc[0].text, enc[1].text = "Dimensión", "Subíndice normalizado"
            for celda in enc:
                for parrafo in celda.paragraphs:
                    for run_enc in parrafo.runs:
                        run_enc.bold = True
            for dim, valor in resultado_isvpt.subindices_por_dimension_entidad.items():
                fila = tabla_isvpt.add_row().cells
                fila[0].text = dim
                fila[1].text = str(valor)

        doc.add_paragraph()
        p_nota_isvpt = doc.add_paragraph()
        run_nota_isvpt = p_nota_isvpt.add_run(resultado_isvpt.nota_metodologica)
        run_nota_isvpt.italic = True
        run_nota_isvpt.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # Análisis integral IA
    _agregar_divisor_seccion_docx(doc, "Análisis integral, señales de riesgo multinivel y Plan de Mejoramiento Prospectivo (generado por IA a partir de los datos reales)", icono="⚠️")
    p_cita_riesgo = doc.add_paragraph()
    run_cita_riesgo = p_cita_riesgo.add_run(
        "La lectura de riesgo institucional que sigue retoma la metodología de gestión integrada "
        "de riesgos (ISO 31000) documentada por Jurado-Zambrano y Villanueva (2021, ESAP) para el "
        "sector público colombiano, aplicada aquí a los datos reales de esta entidad."
    )
    run_cita_riesgo.italic = True
    run_cita_riesgo.font.size = Pt(9.5)
    for parrafo in analisis_ia_texto.split("\n"):
        if parrafo.strip():
            doc.add_paragraph(parrafo)

    doc.add_page_break()

    # Disclaimer
    doc.add_heading("Nota metodológica", level=2)
    p_disc = doc.add_paragraph()
    run_disc = p_disc.add_run(DISCLAIMER_INFORME)
    run_disc.italic = True
    run_disc.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Generación del PDF (reportlab)
# ---------------------------------------------------------------------------

def generar_reporte_pdf(nombre_entidad, diag, analisis_ia_texto, resultado_isvpt=None, resultado_360=None, idi_oficial=None, cruce_recomendaciones=None, total_recomendaciones_entidad=None):
    """Devuelve un BytesIO con el informe técnico en formato PDF.

    resultado_isvpt (opcional): ver docstring de generar_reporte_docx.
    resultado_360 (opcional): ver docstring de generar_reporte_docx.
    idi_oficial (opcional pero MUY recomendado): ver docstring de generar_reporte_docx.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=LETTER,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )
    estilos = getSampleStyleSheet()
    estilo_titulo = ParagraphStyle("TituloSIIEAP", parent=estilos["Title"], fontSize=20)
    estilo_h2 = ParagraphStyle("H2SIIEAP", parent=estilos["Heading2"], spaceBefore=14, fontSize=14)
    estilo_normal = ParagraphStyle("NormalSIIEAP", parent=estilos["Normal"], fontSize=11, leading=15)
    estilo_cursiva = ParagraphStyle("Cursiva", parent=estilo_normal, fontName="Helvetica-Oblique", textColor=colors.grey)

    idi_protagonista = idi_oficial if idi_oficial is not None else diag.idi_estimado
    hay_diferencia_idi = (
        idi_oficial is not None and diag.idi_estimado is not None
        and round(idi_oficial, 2) != round(diag.idi_estimado, 2)
    )
    nivel_riesgo_global = "ALTO" if (idi_protagonista or 0) < 40 else ("MEDIO" if (idi_protagonista or 0) < 70 else "BAJO")

    elementos = []

    # Portada — logos institucionales + banner de color en vez de un título plano
    elementos.extend(_logos_pdf_flowables())
    elementos.extend(_banner_portada_pdf_flowables(
        "Modelo de Conocimiento Institucional del Sistema de Inteligencia Artificial "
        "para la Evaluación Integral del Desempeño Institucional en Entidades Públicas (SIIEAP)",
        "Informe de Diagnóstico Institucional y Plan de Mejoramiento Prospectivo",
    ))
    elementos.append(Paragraph(f"<b>{nombre_entidad}</b>", ParagraphStyle("EntidadSIIEAP", parent=estilos["Heading2"], textColor=colors.HexColor("#1F3864"))))
    elementos.append(Paragraph(f"Generado el {_fecha_hoy_es()} · Índice de Desempeño Institucional (IDI-MIPG), Decreto 1499 de 2017", estilo_normal))
    elementos.append(Paragraph(
        "Docente: Norma Elizabeth Álvarez Grajales · Área de conocimiento: Entidades "
        "Públicas y del Desarrollo · Escuela Superior de Administración Pública (ESAP)",
        ParagraphStyle("AutoriaSIIEAP", parent=estilo_normal, fontSize=9, fontName="Helvetica-Oblique"),
    ))
    elementos.append(Spacer(1, 14))

    # Igual que en la versión docx: se prioriza el total real de recomendaciones
    # de la entidad (todas las cargadas, sin filtrar por brecha) sobre el conteo
    # derivado del cruce, que subestima el total oficial.
    total_recomendaciones = (
        total_recomendaciones_entidad
        if total_recomendaciones_entidad is not None
        else (sum(len(lista) for lista in cruce_recomendaciones.values()) if cruce_recomendaciones else None)
    )

    if total_recomendaciones is not None:
        datos_tabla_portada = [
            ["IDI oficial (Función Pública)", "Nivel de riesgo global", "Brechas detectadas\n(dato exclusivo de este informe)", "Recomendaciones oficiales FP"],
            [str(idi_protagonista), nivel_riesgo_global, str(len(diag.brechas)), str(total_recomendaciones)],
        ]
        anchos_portada = [4 * cm, 4 * cm, 4 * cm, 4.5 * cm]
    else:
        datos_tabla_portada = [
            ["IDI oficial (Función Pública)", "Nivel de riesgo global", "Brechas detectadas\n(dato exclusivo de este informe)"],
            [str(idi_protagonista), nivel_riesgo_global, str(len(diag.brechas))],
        ]
        anchos_portada = [5 * cm, 5 * cm, 5 * cm]
    tabla_portada = Table(datos_tabla_portada, hAlign="CENTER", colWidths=anchos_portada)
    tabla_portada.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("FONTSIZE", (0, 1), (-1, 1), 14),
        ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elementos.append(tabla_portada)
    elementos.append(Spacer(1, 6))
    texto_explica_portada = (
        "⚠ Brechas detectadas: cifra EXCLUSIVA de este informe (SIIEAP) — NO es un dato que "
        "publique el Departamento Administrativo de la Función Pública. Corresponde al número de "
        "índices MIPG de esta entidad con puntaje por debajo de 60 puntos, umbral de alerta "
        "metodológico interno del sistema (la meta plena de la gestión pública es siempre el "
        "100%, no 60)."
    )
    if total_recomendaciones is not None:
        texto_explica_portada += (
            " ✔ Recomendaciones oficiales FP: este dato SÍ es oficial — corresponde al total de "
            "recomendaciones del banco consolidado del Departamento Administrativo de la Función "
            "Pública emitidas para esta entidad (no se limita a las brechas detectadas por este "
            "informe)."
        )
    elementos.append(Paragraph(texto_explica_portada, estilo_cursiva))
    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        elementos.append(Spacer(1, 10))
        elementos.append(Paragraph(
            f"Grupo de comparación: {resultado_360.filtro_descripcion} "
            f"({resultado_360.n_entidades} entidades) · IDI promedio del grupo: {resultado_360.promedio_idi}",
            estilo_cursiva,
        ))
    if hay_diferencia_idi:
        elementos.append(Spacer(1, 8))
        elementos.append(Paragraph(
            f"(Nota: el cálculo interno de verificación de SIIEAP, revisando la suma/ponderación de los "
            f"índices reportados, arroja {diag.idi_estimado}. Esta es una cifra de validación metodológica "
            f"interna que NO reemplaza ni desvirtúa el IDI oficial de Función Pública, que es siempre el "
            f"que prevalece.)",
            ParagraphStyle("NotaIDI", parent=estilo_cursiva, fontSize=8),
        ))
    elementos.append(PageBreak())

    _agregar_glosario_pdf(elementos, estilos, estilo_normal, estilo_h2)

    # 1. Resumen ejecutivo — cifras reales de ESTA entidad, arriba de todo
    elementos.extend(_divisor_seccion_pdf("1. Resumen ejecutivo", icono="📊"))
    parrafos_resumen = [
        f"{nombre_entidad} obtuvo un Índice de Desempeño Institucional (IDI) oficial de {idi_protagonista} "
        f"sobre 100 en la vigencia analizada, con un nivel de riesgo global {nivel_riesgo_global}."
    ]
    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        brecha_grupo = round((idi_protagonista or 0) - resultado_360.promedio_idi, 2)
        comparativo = "por debajo" if brecha_grupo < 0 else "por encima"
        parrafos_resumen.append(
            f"Frente a su grupo de comparación ({resultado_360.filtro_descripcion}, "
            f"{resultado_360.n_entidades} entidades), la entidad se ubica {abs(brecha_grupo)} puntos "
            f"{comparativo} del promedio del grupo ({resultado_360.promedio_idi})."
        )
        if resultado_360.percentil_entidad_referencia is not None:
            parrafos_resumen.append(f"Esto la ubica en el percentil {resultado_360.percentil_entidad_referencia}% de su grupo.")
    dims_criticas = sorted(diag.resultados_por_dimension, key=lambda r: (r.promedio if r.promedio is not None else 0))[:3]
    if dims_criticas:
        nombres_criticas = ", ".join(f"{r.codigo} {r.nombre} ({r.promedio})" for r in dims_criticas)
        parrafos_resumen.append(f"Las dimensiones más críticas son: {nombres_criticas}.")
    dims_fuertes = [r for r in diag.resultados_por_dimension if (r.promedio or 0) >= 60]
    if dims_fuertes:
        nombres_fuertes = ", ".join(f"{r.codigo} {r.nombre} ({r.promedio})" for r in dims_fuertes)
        parrafos_resumen.append(f"Como fortaleza relativa, se destaca(n): {nombres_fuertes}.")
    else:
        parrafos_resumen.append(
            "Ninguna dimensión alcanza el umbral de 60 puntos: el patrón de brechas simultáneas en "
            "múltiples dimensiones sugiere una debilidad estructural de capacidad institucional "
            "(Oszlak), más que fallas puntuales y aisladas de gestión."
        )
    if hay_diferencia_idi:
        parrafos_resumen.append(
            f"Nota metodológica: Función Pública reporta un IDI oficial de {idi_oficial} para esta "
            f"entidad, cifra que este informe usa siempre como referencia principal. El sistema SIIEAP, "
            f"al revisar internamente la suma y ponderación de los índices del archivo oficial, calcula "
            f"además un valor de verificación de {diag.idi_estimado}; esta diferencia es un insumo para "
            f"depurar la metodología interna del sistema, y en ningún caso reemplaza, desvirtúa ni "
            f"contradice la cifra oficial de Función Pública."
        )
    parrafos_resumen.append(
        f"Este informe (SIIEAP) detectó, con metodología propia y exclusiva del sistema — no con "
        f"una cifra publicada por la Función Pública —, {len(diag.brechas)} brechas de "
        "implementación por debajo de 60 puntos (el umbral de alerta metodológico interno que usa "
        "este sistema para priorizar el análisis; la meta plena de la gestión pública, se insiste, "
        "es el 100% de cumplimiento, no 60 puntos), todas desarrolladas en detalle en este informe "
        "— sin selección ni recorte."
    )
    for parrafo_r in parrafos_resumen:
        elementos.append(Paragraph(parrafo_r, estilo_normal))
        elementos.append(Spacer(1, 4))
    elementos.append(PageBreak())

    # Contextualización
    elementos.extend(_divisor_seccion_pdf(CONTEXTUALIZACION_TITULO, icono="📚"))
    elementos.append(Paragraph(CONTEXTUALIZACION_INTRO, estilo_normal))
    elementos.append(Paragraph("Arco teórico: de los griegos a la Administración Pública contemporánea", estilo_h2))
    for titulo_hito, texto_hito in CONTEXTUALIZACION_ARCO_HISTORICO:
        elementos.append(Paragraph(f"<b>{titulo_hito}.</b> {texto_hito}", estilo_normal))
        elementos.append(Spacer(1, 6))

    elementos.append(Paragraph(CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_TITULO, estilo_h2))
    elementos.append(Paragraph(CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_INTRO, estilo_normal))
    elementos.append(Spacer(1, 4))
    for titulo_teoria, texto_teoria in CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS:
        elementos.append(Paragraph(f"<b>{titulo_teoria}.</b> {texto_teoria}", estilo_normal))
        elementos.append(Spacer(1, 6))

    elementos.append(Paragraph("Esquema integrador de la Administración Pública contemporánea", estilo_h2))
    elementos.append(Paragraph("<i>" + " → ".join(CONTEXTUALIZACION_ESQUEMA_INTEGRADOR) + "</i>", estilo_normal))
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph(CONTEXTUALIZACION_CIERRE, estilo_normal))
    elementos.append(PageBreak())

    # Cadena de interpretación institucional (bloque fijo)
    elementos.extend(_divisor_seccion_pdf(CADENA_INTERPRETACION_TITULO, icono="🔗"))
    elementos.append(Paragraph(CADENA_INTERPRETACION_INTRO, estilo_normal))
    elementos.append(Spacer(1, 6))
    for paso in CADENA_INTERPRETACION_PASOS:
        elementos.append(Paragraph(f"• {paso}", estilo_normal))
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph(CADENA_INTERPRETACION_CIERRE, estilo_cursiva))
    elementos.append(PageBreak())

    # Resultado real del diagnóstico
    elementos.extend(_divisor_seccion_pdf("Resultado del diagnóstico institucional (datos reales)", icono="📈"))
    elementos.append(Paragraph(f"<b>IDI estimado: {diag.idi_estimado}</b>", ParagraphStyle("IDIGrande", parent=estilo_normal, fontSize=15)))
    elementos.append(Spacer(1, 8))

    datos_tabla = [["Dimensión", "Promedio", "Riesgo", "Índices"]]
    for r in diag.resultados_por_dimension:
        datos_tabla.append([
            f"{r.codigo} {r.nombre}", str(r.promedio), str(r.nivel_riesgo),
            f"{r.n_indices_evaluados}/{r.n_indices_esperados}",
        ])
    tabla = Table(datos_tabla, hAlign="LEFT", colWidths=[7 * cm, 2.5 * cm, 2.5 * cm, 2.5 * cm])
    estilo_tabla_dim = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    _COLOR_PDF_POR_RIESGO = {
        "alta": colors.HexColor("#F5B7B1"),
        "media": colors.HexColor("#FAD7A0"),
        "baja": colors.HexColor("#A9DFBF"),
    }
    for indice_fila, r in enumerate(diag.resultados_por_dimension, start=1):
        color_riesgo = _COLOR_PDF_POR_RIESGO.get(str(r.nivel_riesgo).strip().lower())
        if color_riesgo:
            estilo_tabla_dim.append(("BACKGROUND", (2, indice_fila), (2, indice_fila), color_riesgo))
        color_quintil_hex = _color_hex_quintil_mipg(r.promedio)
        if color_quintil_hex:
            estilo_tabla_dim.append(("BACKGROUND", (1, indice_fila), (1, indice_fila), colors.HexColor(f"#{color_quintil_hex}")))
    tabla.setStyle(TableStyle(estilo_tabla_dim))
    elementos.append(tabla)
    elementos.append(Spacer(1, 10))

    try:
        buffer_grafica_dim = generar_grafica_dimensiones(diag)
        elementos.append(Image(buffer_grafica_dim, width=16 * cm, height=16 * cm * 0.5))
        elementos.append(Spacer(1, 10))
    except Exception:
        pass

    elementos.append(Paragraph("Brechas priorizadas (todas las detectadas)", estilo_h2))
    if diag.brechas:
        try:
            buffer_grafica_brechas = generar_grafica_brechas(diag)
            if buffer_grafica_brechas:
                elementos.append(Image(buffer_grafica_brechas, width=16 * cm, height=16 * cm * 0.55))
                elementos.append(Spacer(1, 10))
        except Exception:
            pass
        for b in diag.brechas:
            elementos.append(Paragraph(
                f"• {b.codigo_indice} ({b.puntaje}) — {b.nombre_indice} · {b.politica}", estilo_normal,
            ))
    else:
        elementos.append(Paragraph("No se detectaron brechas por debajo del umbral con los datos disponibles.", estilo_normal))
    elementos.append(PageBreak())

    # Tabla de conexión política -> enfoque contemporáneo -> norma (garantizada por código)
    if diag.brechas:
        elementos.append(Paragraph("Conexión de las brechas con los enfoques contemporáneos de la Administración Pública", estilo_h2))
        elementos.append(Paragraph(
            "Cada política con brechas se conecta aquí, de forma sistemática, con el enfoque "
            "contemporáneo que mejor la explica y con la norma colombiana que lo respalda.",
            estilo_normal,
        ))
        elementos.append(Spacer(1, 6))
        politicas_con_brecha = []
        vistas = set()
        for b in diag.brechas:
            if b.politica not in vistas:
                vistas.add(b.politica)
                politicas_con_brecha.append(b.politica)
        datos_tabla_enfoques = [["Política con brecha", "Enfoque contemporáneo", "Norma"]]
        for politica in politicas_con_brecha:
            enfoque, norma = _enfoque_y_norma_de_politica(politica)
            datos_tabla_enfoques.append([politica, enfoque, norma])
        tabla_enfoques = Table(datos_tabla_enfoques, hAlign="LEFT", colWidths=[5 * cm, 5 * cm, 6 * cm])
        tabla_enfoques.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F4F8")]),
        ]))
        elementos.append(tabla_enfoques)
        elementos.append(PageBreak())

    if resultado_isvpt is not None and resultado_isvpt.isvpt_entidad_referencia is not None:
        elementos.extend(_divisor_seccion_pdf("El Termómetro del Valor Público: Índice Sintético de Valor Público Territorial (ISVPT)", icono="🎯"))
        elementos.append(Paragraph(
            "Como complemento al IDI oficial, este informe incorpora un ejercicio de índice "
            "sintético construido con la metodología académica validada por Vélez Tamayo, "
            "Ortiz-Muñoz y Cardona Montoya (2026) para el ISDEL — normalización min-max y "
            "agregación aritmética simple, siguiendo las directrices de la OCDE (2008) para "
            "indicadores compuestos — aplicada aquí a las 7 dimensiones reales del IDI-MIPG "
            "dentro del grupo de comparación de esta entidad.",
            estilo_normal,
        ))
        elementos.append(Spacer(1, 6))
        elementos.append(Paragraph(
            f"<b>ISVPT de {nombre_entidad}: {resultado_isvpt.isvpt_entidad_referencia} "
            f"(posición {resultado_isvpt.posicion_entidad_referencia} de "
            f"{resultado_isvpt.n_entidades} en su grupo de comparación).</b>",
            estilo_normal,
        ))
        elementos.append(Spacer(1, 8))

        if resultado_isvpt.subindices_por_dimension_entidad:
            elementos.append(Paragraph("Subíndices normalizados por dimensión (0 = más rezagado del grupo, 1 = mejor del grupo)", estilo_h2))
            datos_tabla_isvpt = [["Dimensión", "Subíndice normalizado"]]
            for dim, valor in resultado_isvpt.subindices_por_dimension_entidad.items():
                datos_tabla_isvpt.append([dim, str(valor)])
            tabla_isvpt = Table(datos_tabla_isvpt, hAlign="LEFT", colWidths=[10 * cm, 4 * cm])
            tabla_isvpt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            elementos.append(tabla_isvpt)
            elementos.append(Spacer(1, 8))

        elementos.append(Paragraph(resultado_isvpt.nota_metodologica, estilo_cursiva))
        elementos.append(PageBreak())

    # Análisis integral IA
    elementos.extend(_divisor_seccion_pdf("Análisis integral, señales de riesgo multinivel y Plan de Mejoramiento Prospectivo (generado por IA a partir de los datos reales)", icono="⚠️"))
    elementos.append(Paragraph(
        "La lectura de riesgo institucional que sigue retoma la metodología de gestión integrada de "
        "riesgos (ISO 31000) documentada por Jurado-Zambrano y Villanueva (2021, ESAP) para el sector "
        "público colombiano, aplicada aquí a los datos reales de esta entidad.",
        estilo_cursiva,
    ))
    elementos.append(Spacer(1, 6))
    for parrafo in analisis_ia_texto.split("\n"):
        if parrafo.strip():
            texto_escapado = parrafo.replace("&amp;", "&amp;amp;").replace("&lt;", "&amp;lt;").replace("&gt;", "&amp;gt;")
            elementos.append(Paragraph(texto_escapado, estilo_normal))
            elementos.append(Spacer(1, 4))
    elementos.append(PageBreak())

    # Disclaimer
    elementos.append(Paragraph("Nota metodológica", estilos["Heading2"]))
    elementos.append(Paragraph(DISCLAIMER_INFORME, estilo_cursiva))

    def _pie_de_pagina(canvas_pdf, doc_pdf):
        canvas_pdf.saveState()
        canvas_pdf.setFont("Helvetica", 8)
        canvas_pdf.setFillColor(colors.grey)
        canvas_pdf.drawString(2 * cm, 1.3 * cm, f"SIIEAP — {nombre_entidad}")
        canvas_pdf.drawRightString(LETTER[0] - 2 * cm, 1.3 * cm, f"Página {doc_pdf.page}")
        canvas_pdf.restoreState()

    doc.build(elementos, onFirstPage=_pie_de_pagina, onLaterPages=_pie_de_pagina)
    buffer.seek(0)
    return buffer
