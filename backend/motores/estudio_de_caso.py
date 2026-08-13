"""Generador del Estudio de Caso Académico — Word (.docx) y PDF descargables.

Producto evaluativo de la Unidad 2 ("Análisis y aplicación de los enfoques de
la Administración Pública") de la asignatura Enfoques y Teorías de la
Administración Pública II (Maestría en Administración Pública, ESAP),
según el microcurrículo armonizado: "Informe técnico de análisis del caso
(8 a 12 páginas)" que debe contener:

  - Descripción del problema público
  - Contexto institucional (aquí ampliado con enfoque local, regional y
    global, según se solicitó)
  - Análisis desde la Nueva Gestión Pública
  - Análisis desde la Post Nueva Gestión Pública
  - Análisis desde el Nuevo Institucionalismo
  - Evaluación del modelo de gobernanza
  - Identificación de fortalezas
  - Debilidades
  - Recomendaciones para fortalecer la gestión pública
  - Uso responsable de herramientas de IA para apoyar el análisis documental

Este módulo NO inventa cifras ni normas: el diagnóstico (diag), el análisis
de IA (motor_analisis_ia), el análisis 360 (motor_analisis_360) y el ISVPT
(motor_isvpt) ya vienen de datos reales cargados en el sistema; los bloques
fijos de fundamento jurídico-fiscal-disciplinario-contable-control interno
y de fuentes de datos abiertos citan únicamente normas y plataformas
oficiales verificadas.
"""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image,
)

from backend.motores.graficas_informe import generar_grafica_dimensiones, generar_grafica_brechas, generar_matriz_riesgo_probabilidad_impacto
from backend.motores.generador_informe import (
    _enfoque_y_norma_de_politica,
    _agregar_logos_docx,
    _logos_pdf_flowables,
    _agregar_banner_docx,
    _banner_portada_pdf_flowables,
    _agregar_divisor_seccion_docx,
    _divisor_seccion_pdf,
    _franjas_alternas_docx,
    _color_hex_quintil_mipg,
    _ajustar_tabla_docx,
    _sombrear_celda,
    _celda_pdf,
    _agregar_nota_regimen_especial_docx,
    _nota_regimen_especial_pdf_flowables,
    _agregar_tabla_contenido_docx,
    _agregar_razon_de_ser_docx,
    _toc_pdf_flowables,
    _razon_de_ser_pdf_flowables,
    _agregar_marco_descentralizacion_docx,
    _agregar_marco_descentralizacion_pdf,
    _agregar_glosario_docx,
    _agregar_glosario_pdf,
    _agregar_normativa_politicas_docx,
    _agregar_normativa_politicas_pdf,
    _agregar_texto_markdown_docx,
    _texto_markdown_a_pdf_flowables,
    COLOR_INSTITUCIONAL,
)

# ---------------------------------------------------------------------------
# Bloques fijos
# ---------------------------------------------------------------------------

FUENTES_DATOS_ABIERTOS_COMUNES = [
    (
        "FURAG / IDI-MIPG (Función Pública)",
        "Formulario Único de Reporte de Avances de la Gestión y el Índice de "
        "Desempeño Institucional — la fuente primaria del diagnóstico de este "
        "estudio de caso. https://www.funcionpublica.gov.co/web/mipg",
    ),
    (
        "Portal de Datos Abiertos de Colombia",
        "Catálogo nacional de datos abiertos del Estado colombiano, con "
        "conjuntos de datos de todas las entidades públicas del país. "
        "https://www.datos.gov.co",
    ),
    (
        "SECOP II / Colombia Compra Eficiente",
        "Sistema Electrónico de Contratación Pública: procesos de "
        "contratación, proveedores y concentración contractual de la "
        "entidad analizada. https://www.colombiacompra.gov.co",
    ),
    (
        "CHIP — Consolidador de Hacienda e Información Pública",
        "Información financiera, presupuestal y contable reportada por las "
        "entidades territoriales a la Contaduría General de la Nación. "
        "https://www.chip.gov.co",
    ),
    (
        "SIGEP II",
        "Sistema de Información y Gestión del Empleo Público: planta de "
        "personal, directivos y vacantes de la entidad. "
        "https://www.sigep.gov.co",
    ),
    (
        "TerriData (DNP)",
        "Sistema de estadísticas territoriales del Departamento Nacional de "
        "Planeación: indicadores socioeconómicos, fiscales y de gestión por "
        "municipio y departamento. https://terridata.dnp.gov.co",
    ),
    (
        "Sistema de Rendición de Cuentas / SIA Observa (Contraloría)",
        "Información de vigilancia y control fiscal territorial de la "
        "Contraloría General de la República y las contralorías "
        "territoriales. https://www.contraloria.gov.co",
    ),
    (
        "SIRI / SIGEP disciplinario (Procuraduría)",
        "Sistema de Información de Registro de Sanciones e Inhabilidades y "
        "antecedentes disciplinarios de la Procuraduría General de la "
        "Nación. https://www.procuraduria.gov.co",
    ),
    (
        "SECOP II — consulta de procesos",
        "Consulta pública de procesos de contratación por entidad. "
        "https://consultaprocesos.colombiacompra.gov.co/",
    ),
    (
        "SIGEP II — Directorio de servidores públicos",
        "Directorio público de servidores por entidad. "
        "https://www1.funcionpublica.gov.co/web/sigep2/diretorio",
    ),
    (
        "ODS Colombia (DNP)",
        "Seguimiento a los Objetivos de Desarrollo Sostenible por entidad "
        "territorial. https://ods.dnp.gov.co/",
    ),
    (
        "SICODIS (DNP)",
        "Sistema de Información de Consulta de Distribución de Recursos "
        "del Sistema General de Participaciones por municipio. "
        "https://sicodis.dnp.gov.co/",
    ),
    (
        "Portal Territorial (DNP)",
        "Información de gestión y planeación territorial. "
        "https://portalterritorial.dnp.gov.co/",
    ),
    (
        "Sinergia (DNP)",
        "Sistema Nacional de Evaluación de Gestión y Resultados: "
        "seguimiento al cumplimiento del Plan Nacional de Desarrollo. "
        "https://sinergia.dnp.gov.co/Paginas/inicio.aspx",
    ),
    (
        "SISPT (DNP)",
        "Sistema de Seguimiento a los Planes Territoriales de Desarrollo: "
        "permite verificar el cumplimiento del Plan de Desarrollo municipal "
        "de la entidad elegida. https://sispt.dnp.gov.co/",
    ),
    (
        "SIIPO (DNP)",
        "Sistema de Información, Seguimiento y Evaluación de los Planes de "
        "Ordenamiento Territorial. https://siipo.dnp.gov.co/inicio",
    ),
    (
        "CIFFIT (DNP)",
        "Sistema de información sobre financiamiento e inversión "
        "territorial. https://ciffit.dnp.gov.co/ciffit/Default",
    ),
    (
        "Mapa de Inversiones (DNP)",
        "Visualización geográfica de la inversión pública por municipio. "
        "https://mapainversiones.dnp.gov.co/",
    ),
    (
        "PDET (Agencia de Renovación del Territorio)",
        "Programas de Desarrollo con Enfoque Territorial, si la entidad "
        "elegida es un municipio PDET. "
        "https://www.renovacionterritorio.gov.co/especiales/especial_PDET/",
    ),
    (
        "Antioquia Visión 2040 / Antioquia Datos",
        "Agenda de desarrollo departamental de largo plazo y datos abiertos "
        "de Antioquia. https://www.antioquiadatos.gov.co/index.php/agenda2040/",
    ),
    (
        "Misión de Descentralización (DNP)",
        "Diagnóstico y propuestas de reforma al modelo de descentralización "
        "colombiano (2022-2024) — insumo directo del Proyecto de Ley de "
        "Competencias que cursa hoy en el Congreso. "
        "https://misiondescentralizacion.dnp.gov.co/",
    ),
    (
        "CEPAL, PNUD, Banco Mundial, BID, CLAD, ONU (marco internacional)",
        "Organismos multilaterales de referencia para el análisis en escala "
        "global de este estudio de caso: https://www.cepal.org/es · "
        "https://www.undp.org/es · https://www.bancomundial.org/es/home · "
        "https://www.iadb.org/es · https://clad.org/ · https://www.un.org/es/",
    ),
]

FUENTES_DATOS_ABIERTOS_ANTIOQUIA = [
    (
        "Contraloría General de Antioquia — Auditoría Integrada",
        "Informes de auditoría financiera, de cumplimiento y de desempeño "
        "practicados por la Contraloría General de Antioquia a las entidades "
        "del departamento (alcaldías, gobernación, entes descentralizados), "
        "organizados por vigencia. Fuente obligatoria si la entidad elegida "
        "es de Antioquia: busque el informe específico de su entidad antes "
        "de redactar el apartado de fortalezas/debilidades. "
        "https://www.contraloriadeantioquia.gov.co/tema/auditoria-integrada "
        "(y la carpeta del año en curso, p. ej. "
        "https://www.contraloriadeantioquia.gov.co/tema/auditoria-integrada-2025).",
    ),
    (
        "Contraloría General de Antioquia — Proyectos Especiales",
        "Actuaciones especiales de fiscalización y auditorías a proyectos "
        "puntuales (regalías, obras, convenios) de entidades antioqueñas, "
        "complementarias a la auditoría integrada regular. "
        "https://www.contraloriadeantioquia.gov.co/tema/proyectos-especiales",
    ),
    (
        "IDEA — Instituto para el Desarrollo de Antioquia",
        "Establecimiento público de fomento y desarrollo económico del "
        "departamento: banca de fomento (crédito y garantías a municipios), "
        "asesoría en saneamiento fiscal, acompañamiento técnico y acceso a "
        "cooperación internacional para entidades territoriales de "
        "Antioquia. Útil para contrastar la capacidad fiscal/financiera de "
        "la entidad elegida con el respaldo o apoyo que ha recibido del "
        "IDEA. https://www.idea.gov.co",
    ),
    (
        "Observatorio Fiscal — Contraloría General de Antioquia",
        "Plataforma de seguimiento fiscal territorial de los municipios "
        "antioqueños. https://observatoriofiscal.contraloriadeantioquia.com/",
    ),
    (
        "Antioquia Visión 2040 / Antioquia Datos",
        "Agenda de desarrollo departamental de largo plazo y datos abiertos "
        "de Antioquia. https://www.antioquiadatos.gov.co/index.php/agenda2040/",
    ),
]

FUENTES_DATOS_ABIERTOS_CHOCO = [
    (
        "Contraloría General del Departamento del Chocó",
        "Ente de control fiscal territorial del Chocó: informes de auditoría "
        "financiera y de gestión a municipios (p. ej. auditorías a Juradó y "
        "Medio Atrato), mapa de riesgos, informes de seguimiento al plan "
        "anticorrupción. Fuente obligatoria si la entidad elegida es del "
        "Chocó: busque el informe específico de su municipio antes de "
        "redactar el apartado de fortalezas/debilidades. "
        "https://contraloria-choco.gov.co/",
    ),
    (
        "Gobernación del Chocó — Plan de Desarrollo Departamental",
        "Instrumento de planeación departamental con el que debe articularse "
        "el Plan de Desarrollo Municipal de la entidad elegida, si es del "
        "Chocó. https://www.choco.gov.co/",
    ),
]


def _fuentes_datos_abiertos_para(departamento: str | None):
    """Combina las fuentes comunes (nacionales) con las específicas del
    departamento de la entidad, si el departamento tiene fuentes registradas
    (por ahora: Antioquia y Chocó). Si no se reconoce el departamento,
    devuelve solo las fuentes comunes — el informe sigue siendo completo,
    solo sin el complemento regional."""
    fuentes = list(FUENTES_DATOS_ABIERTOS_COMUNES)
    if departamento:
        clave = departamento.strip().upper()
        if clave == "ANTIOQUIA":
            fuentes = fuentes + FUENTES_DATOS_ABIERTOS_ANTIOQUIA
        elif clave in ("CHOCO", "CHOCÓ"):
            fuentes = fuentes + FUENTES_DATOS_ABIERTOS_CHOCO
    return fuentes


INSTITUCIONALIDAD_IDI_TITULO = (
    "Institucionalidad del IDI-MIPG: líderes de política, cronograma oficial "
    "y los índices de desempeño fiscal, municipal y departamental del DNP"
)

INSTITUCIONALIDAD_IDI_INTRO = (
    "El IDI-MIPG no lo mide Función Pública en solitario: es el resultado de "
    "un entramado institucional de líderes de política, un cronograma oficial "
    "anual y tres índices hermanos del DNP que miden capacidades distintas "
    "pero complementarias de la misma entidad territorial. Se documenta aquí "
    "con base en circulares oficiales vigentes."
)

INSTITUCIONALIDAD_IDI = [
    (
        "Cronograma oficial de la Medición del Desempeño Institucional (MDI), vigencia 2025",
        "La Circular Externa 100-011-2025 del Consejo de Gestión y Desempeño "
        "Institucional (10 de diciembre de 2025), expedida con fundamento en "
        "los artículos 2.2.22.3.10 y 2.2.23.3 del Decreto 1083 de 2015, fija "
        "el cronograma oficial: recolección de información vía FURAG del 3 "
        "de marzo al 17 de abril de 2026; análisis y procesamiento "
        "estadístico del 22 de abril al 5 de junio de 2026; aprobación de "
        "resultados del 9 al 19 de junio; publicación de los índices del 22 "
        "al 26 de junio; y difusión de resultados hasta el 27 de noviembre "
        "de 2026. Cualquier IDI que se analice en este estudio de caso debe "
        "ubicarse dentro de este calendario oficial.",
    ),
    (
        "Líderes de política de la Medición del Desempeño Institucional",
        "El Directorio de Líderes de Política (Función Pública, febrero de "
        "2026) confirma que cada una de las 19 políticas de gestión y "
        "desempeño del MIPG tiene una entidad líder responsable de asesorar "
        "y evaluar su implementación, entre ellas: Función Pública (talento "
        "humano, integridad, servicio a la ciudadanía, racionalización de "
        "trámites, participación ciudadana, fortalecimiento institucional, "
        "gestión del conocimiento, control interno), Colombia Compra "
        "Eficiente (compras y contratación pública), MinTIC (Gobierno "
        "Digital y seguridad digital), el DNP (mejora normativa), la Agencia "
        "Nacional de Defensa Jurídica del Estado (defensa jurídica), la "
        "Secretaría de Transparencia de la Presidencia (transparencia y "
        "lucha contra la corrupción), el DANE (gestión de la información "
        "estadística), el Archivo General de la Nación (gestión documental) "
        "y el Ministerio de Ambiente (gestión ambiental institucional). "
        "Cualquier recomendación oficial que cite este estudio de caso "
        "proviene, en últimas, de una de estas entidades líderes.",
    ),
    (
        "Lineamientos de evaluación de la gestión y control interno (segunda y tercera línea)",
        "La Circular Externa 100-010-2025 del Consejo Asesor del Gobierno "
        "Nacional en materia de Control Interno (4 de diciembre de 2025) "
        "desarrolla el artículo 39 de la Ley 909 de 2004 (obligación del "
        "jefe de control interno de remitir las evaluaciones de gestión de "
        "cada dependencia) y el artículo 12 de la Ley 87 de 1993, "
        "estableciendo el esquema de líneas de defensa del MECI: las áreas "
        "de planeación como segunda línea (deben implementar indicadores y "
        "tableros de control, y remitir la información consolidada) y las "
        "oficinas de control interno como tercera línea o evaluador "
        "independiente (deben incluir el seguimiento en su plan anual de "
        "auditorías, con periodicidad preferiblemente trimestral). Este "
        "esquema aplica a partir de la evaluación de la gestión "
        "correspondiente a la vigencia 2026.",
    ),
    (
        "Índice de Desempeño Fiscal (IDF)",
        "Instrumento distinto pero complementario al IDI-MIPG: el artículo "
        "79 de la Ley 617 de 2000 ordena al DNP evaluar y publicar "
        "anualmente los resultados de la gestión fiscal y financiera de "
        "municipios y departamentos. El IDF agrupa a los municipios en los "
        "mismos 'grupos de capacidades iniciales' que usa la Medición de "
        "Desempeño Municipal, para no comparar entidades de tamaño y "
        "capacidad muy distinta — la misma lógica de 'grupo par' que aplica "
        "este sistema al IDI. Fuente: DNP, Dirección de Descentralización y "
        "Fortalecimiento Fiscal.",
    ),
    (
        "Medición de Desempeño Municipal (MDM) y Departamental (MDD)",
        "Con base en las leyes 617 de 2000 y 715 de 2001, el DNP mide y "
        "compara el desempeño municipal (MDM) y departamental (MDD) — no ya "
        "solo desde la gestión institucional (como el IDI-MIPG) sino desde "
        "resultados de desarrollo (calidad de vida, cierre de brechas), "
        "agrupando también a las entidades territoriales por capacidades "
        "iniciales similares. Un estudio de caso riguroso debería leer el "
        "IDI-MIPG de la entidad junto con su IDF, su MDM o MDD, para "
        "obtener una fotografía completa: gestión institucional, salud "
        "fiscal y resultados de desarrollo.",
    ),
    (
        "Planeación Nacional y planeación departamental",
        "El Departamento Nacional de Planeación (DNP) es el articulador "
        "técnico de todos estos instrumentos (IDF, MDM, MDD, SISPT, "
        "Sinergia), mientras que las oficinas de planeación departamental "
        "(en Antioquia, el Departamento Administrativo de Planeación de la "
        "Gobernación) replican ese rol a escala regional: consolidan la "
        "información de sus municipios, hacen seguimiento a sus planes de "
        "desarrollo y son un interlocutor obligado para cualquier "
        "diagnóstico institucional territorial serio.",
    ),
]

ESTADO_DEL_ARTE_TITULO = (
    "Estado del Arte Inicial: el ecosistema completo de mediciones abiertas "
    "sobre el desempeño territorial en Colombia"
)

ESTADO_DEL_ARTE_INTRO = (
    "El IDI-MIPG es solo UNA de varias mediciones oficiales e independientes que existen "
    "sobre el desempeño de una entidad territorial colombiana. Un estudio de caso riguroso "
    "no se limita a un solo índice: construye un estado del arte que cruce todas las "
    "mediciones abiertas disponibles, para leer la situación de la entidad desde múltiples "
    "ángulos (gestión institucional, salud fiscal, resultados de desarrollo, transparencia, "
    "percepción independiente) antes de pasar al análisis y a las recomendaciones."
)

ESTADO_DEL_ARTE_INDICES = [
    (
        "Índice de Desempeño Fiscal (IDF)",
        "Ordenado por el artículo 79 de la Ley 617 de 2000 y calculado anualmente por la "
        "Subdirección de Fortalecimiento Fiscal Territorial del DNP, el IDF mide la salud "
        "financiera de una entidad territorial con la fórmula IDF = (0,8 × Resultados "
        "Fiscales) + (0,2 × Gestión Financiera), donde 'resultados fiscales' pondera "
        "capacidad de ahorro, autofinanciación del gasto de funcionamiento (límites de la "
        "Ley 617 de 2000), magnitud de la inversión y capacidad de respaldo del "
        "endeudamiento, y 'gestión financiera' incorpora la capacidad de programación y "
        "ejecución de ingresos y de inversión. Desde el ajuste metodológico de 2019, el IDF "
        "clasifica a las entidades en 5 rangos: Deterioro (< 40 puntos), Riesgo (40-60), "
        "Vulnerable (60-70), Solvente (70-80) y Sostenible (> 80). El promedio nacional "
        "municipal en 2023 fue de 57,57 puntos (una mejora frente a 55,78 en 2022), lo que "
        "confirma que la mayoría de los municipios colombianos se ubica todavía en un rango "
        "de Riesgo. Entre los municipios de Antioquia, Buriticá lideró su grupo de "
        "capacidades iniciales en 2023 con 72,02 puntos (rango Solvente) y Girardota lideró "
        "su grupo en la medición 2022 — referencias útiles para comparar la salud fiscal de "
        "la entidad de este estudio de caso dentro de su propio departamento. Consulte el "
        "resultado específico de su entidad en: "
        "https://www.dnp.gov.co/LaEntidad_/subdireccion-general-descentralizacion-desarrollo-territorial/"
        "direccion-descentralizacion-fortalecimiento-fiscal/Paginas/informacion-fiscal-y-financiera.aspx "
        "— el DNP dispone allí de un visor de consulta por entidad, departamento y grupo de "
        "capacidades iniciales. El IDEA, además, publica boletines sectoriales propios que "
        "analizan el IDF específicamente para los municipios de Antioquia (ej. 'Gestión "
        "Estratégica Financiera Municipal', disponible en idea.gov.co).",
    ),
    (
        "Medición de Desempeño Municipal (MDM) y Departamental (MDD)",
        "A diferencia del IDI-MIPG (que mide gestión institucional interna) y del IDF (que "
        "mide salud fiscal), la MDM y la MDD del DNP miden algo distinto y complementario: "
        "resultados de desarrollo — el aumento efectivo de la calidad de vida de la "
        "población — agrupando a las entidades por sus 'capacidades iniciales' (densidad "
        "empresarial, valor agregado, tamaño poblacional, ruralidad, economías de "
        "aglomeración) para no comparar peras con manzanas. Su propósito explícito es "
        "'incentivar la inversión orientada a resultados' y 'apoyar la focalización de la "
        "asistencia técnica' del nivel nacional hacia los territorios con más rezago. "
        "Consulte el resultado de la entidad en el Portal Territorial del DNP: "
        "https://portalterritorial.dnp.gov.co/AdmInfoTerritorial/MenuInfoTerrEstMDM — "
        "cruzar el resultado MDM/MDD de la entidad con su IDI-MIPG y su IDF permite "
        "distinguir si una entidad gestiona bien pero con pocos resultados de desarrollo, o "
        "viceversa, algo que ningún índice por separado revela.",
    ),
    (
        "Índice de Transparencia de las Entidades Públicas (ITEP)",
        "A diferencia de los tres índices anteriores (todos de origen estatal), el ITEP es "
        "una medición INDEPENDIENTE de sociedad civil, calculada desde hace más de 15 años "
        "por la Corporación Transparencia por Colombia (capítulo nacional de Transparencia "
        "Internacional), que evalúa el riesgo de corrupción de una entidad pública a partir "
        "de tres factores: visibilidad, institucionalidad y control, y sanción. Su valor "
        "para este estudio de caso es precisamente su independencia del dato oficial: sirve "
        "como contraste externo a las brechas que el IDI-MIPG reporta en materia de "
        "integridad y transparencia. Consulte los resultados en "
        "https://indicedetransparencia.org.co/ o en el catálogo de publicaciones de "
        "Transparencia por Colombia (transparenciacolombia.org.co/itep/).",
    ),
    (
        "Observatorios académicos independientes de Administración Pública",
        "Además de las mediciones oficiales y de sociedad civil, existen observatorios "
        "académicos colombianos que producen análisis periódico sobre gestión pública y "
        "que enriquecen la discusión teórica de este estudio de caso con literatura "
        "revisada por pares: el Grupo de Investigación OPERA — Observatorio de Políticas, "
        "Ejecución y Resultados de la Administración Pública, de la Universidad Externado "
        "de Colombia (activo desde 1995, con la revista académica OPERA desde 2002; "
        "uexternado.edu.co/cipe/opera-observatorio-politicas-ejecucion-resultados-la-administracion-publica/); "
        "el Observatorio de Mejora Normativa del DNP, que recopila indicadores sobre la "
        "política de mejora normativa; el Observatorio de Hacienda Pública y Derecho "
        "Tributario de la Universidad del Rosario, especializado en gestión de recursos "
        "públicos; y el Observatorio de Desarrollo y Política Social (ODEPS) de la "
        "Universidad Externado, centrado en desigualdad y política social territorial. "
        "A estos se suman dos universidades adicionales con producción reciente y "
        "verificada: la Universidad EAFIT, cuya Maestría en Gobierno y Políticas Públicas "
        "mantiene un repositorio con tesis de 2025 sobre gestión pública municipal "
        "antioqueña —incluyendo, por ejemplo, un tablero digital de seguimiento al Plan de "
        "Desarrollo Municipal de Remedios, Antioquia, con indicadores de cumplimiento de "
        "metas, ejecución presupuestal y desempeño (repository.eafit.edu.co)—; y la "
        "Universidad de los Andes, a través de tres observatorios activos y con datos "
        "recientes: el Observatorio Municipal de Datos CEDE (Facultad de Economía, "
        "actualizado en 2025, con tablero Power BI de indicadores municipales y módulo de "
        "datos ODS; datoscede.uniandes.edu.co/observatorio-municipal/), el Observatorio de "
        "la Democracia (estudios de 2025-2026 sobre confianza ciudadana y valores "
        "democráticos, incluyendo el Barómetro de las Américas; obsdemocracia.org), y el "
        "Observatorio Global de Corrupción de la Escuela de Gobierno Alberto Lleras "
        "Camargo (2024, con datos abiertos de contratación pública y procesos "
        "legislativos de más de 46 países). Nota honesta sobre actualidad de los datos: "
        "al momento de construir este informe, la mayoría de los datos consolidados y "
        "descargables de estos observatorios (especialmente los de finanzas y desempeño "
        "territorial agregado) tienen como año más reciente 2023, con publicación de "
        "resultados 2024 en curso según cada entidad; se recomienda verificar en cada "
        "portal si ya está disponible la vigencia 2024 o 2025 al momento de consultar, "
        "dado que el IDI-MIPG de este informe corresponde a la vigencia 2025.",
    ),
]

FUNDAMENTO_JURIDICO_AMPLIADO_TITULO = (
    "Fundamento jurídico, fiscal, disciplinario, contable y de control interno"
)

MARCO_REFORMA_TERRITORIAL_TITULO = (
    "Marco de la reforma territorial en curso: Ley de Competencias y Misión de Descentralización"
)

MARCO_REFORMA_TERRITORIAL_INTRO = (
    "Ningún estudio de caso sobre una entidad territorial colombiana puede "
    "ignorar hoy la reforma de fondo que se está discutiendo sobre el modelo "
    "de descentralización del país. Se resume aquí, con precisión sobre su "
    "estado real de trámite (para no hacer pasar un proyecto de ley por una "
    "ley vigente), lo que corresponde consultar y citar:"
)

MARCO_REFORMA_TERRITORIAL = [
    (
        "Proyecto de Ley de Competencias y Recursos del SGP (2025) — EN TRÁMITE, aún no es ley",
        "Es importante precisar que se trata de un PROYECTO de ley orgánica "
        "que cursa actualmente en el Congreso de la República, no de una "
        "norma vigente. Su artículo 1 (Objeto) busca fortalecer la autonomía "
        "y el desarrollo territorial mediante la asignación y distribución "
        "de competencias y recursos entre la Nación, las entidades "
        "territoriales y las beneficiarias del Sistema General de "
        "Participaciones (SGP), con fundamento en los artículos 151, 356 y "
        "357 de la Constitución Política. Su artículo 3 consagra 15 "
        "principios, entre ellos el cierre de brechas territoriales "
        "(numeral 1, en desarrollo del artículo 356 constitucional), la "
        "autonomía territorial (numeral 2), la subsidiariedad — que la "
        "competencia la ejerza el nivel de gobierno más próximo al "
        "ciudadano que tenga capacidad para hacerlo (numeral 8) —, la "
        "concurrencia (numeral 9), la coordinación (numeral 10) y el "
        "fortalecimiento institucional territorial (numeral 11), que exige "
        "a la Nación acompañar técnica y administrativamente a las "
        "entidades territoriales para el ejercicio autónomo de sus "
        "competencias. Estos principios dialogan directamente con las "
        "brechas de capacidad institucional que el IDI-MIPG detecta en este "
        "estudio de caso.",
    ),
    (
        "Misión de Descentralización (2022-2024) — antecedente técnico del proyecto de ley",
        "Creada mediante el Decreto 1665 de 2021 y con Secretaría Técnica "
        "de la Dirección de Descentralización y Fortalecimiento Fiscal del "
        "DNP, la Misión de Descentralización sesionó durante dos años y "
        "entregó su Informe Final el 1 de marzo de 2024, con propuestas "
        "sobre cinco componentes: entre ellas, replantear el actual esquema "
        "de categorización de departamentos y municipios (el mismo de la "
        "Ley 617 de 2000 que usa este sistema para el 'Grupo par') y "
        "reformar el Sistema General de Regalías y el SGP. El Proyecto de "
        "Ley de Competencias que hoy cursa en el Congreso recoge buena "
        "parte de ese diagnóstico técnico. Fuente: "
        "https://misiondescentralizacion.dnp.gov.co/",
    ),
    (
        "Socialización ESAP — Función Pública",
        "La ESAP y el Departamento Administrativo de la Función Pública "
        "realizaron jornadas de socialización del Proyecto de Ley de "
        "Competencias con transmisión pública por el canal oficial de "
        "YouTube @EsapOficial (youtube.com/c/ComunicacionesESAP), como "
        "parte del proceso de participación ciudadana sobre esta reforma. "
        "Se recomienda al estudiante consultar esas grabaciones como fuente "
        "primaria de la discusión pública sobre la reforma, identificando "
        "la fecha exacta de la sesión que consulte.",
    ),
]

FUNDAMENTO_JURIDICO_AMPLIADO_INTRO = (
    "Todo estudio de caso institucional debe leerse, además de con las tres "
    "teorías del curso, con las cuatro dimensiones de control que enmarcan "
    "legalmente la gestión pública en Colombia. Se listan aquí, con su "
    "artículo específico verificado, sin extrapolar hechos concretos de la "
    "entidad que no consten en los datos reales ya cargados en el sistema."
)

FUNDAMENTO_JURIDICO_AMPLIADO = [
    (
        "Control fiscal — responsabilidad fiscal",
        "La Ley 610 de 2000 (artículo 1) define la responsabilidad fiscal "
        "como aquella que busca resarcir los daños ocasionados al "
        "patrimonio público por la conducta dolosa o culposa de un servidor "
        "público o particular en ejercicio de la gestión fiscal, y regula "
        "el proceso que adelantan las Contralorías para determinarla. "
        "Cualquier brecha del IDI-MIPG relacionada con el manejo de "
        "recursos públicos debe leerse a la luz de este marco.",
    ),
    (
        "Control disciplinario",
        "La Ley 1952 de 2019 (Código General Disciplinario, modificada por "
        "la Ley 2094 de 2021) regula la falta disciplinaria de los "
        "servidores públicos. Sus artículos 93 y 94 definen el Control "
        "Disciplinario Interno como la instancia que cada entidad debe "
        "implementar al más alto nivel jerárquico, con autonomía e "
        "independencia, antes de que un caso llegue a la Procuraduría "
        "General de la Nación o a las personerías.",
    ),
    (
        "Control contable",
        "La Resolución 533 de 2015 de la Contaduría General de la Nación "
        "(artículos 1 y 2) incorpora el Marco Normativo para Entidades de "
        "Gobierno dentro del Régimen de Contabilidad Pública, exigiendo "
        "reconocimiento, medición, revelación y presentación de los hechos "
        "económicos según normas técnicas unificadas — la base para "
        "cualquier hallazgo contable de una entidad territorial.",
    ),
    (
        "Control interno",
        "La Ley 87 de 1993 define, en su artículo 1, el control interno "
        "como el sistema integrado de planes, métodos, principios, normas y "
        "procedimientos de verificación y evaluación adoptados por una "
        "entidad para asegurar que sus actividades se ajusten a la "
        "normatividad vigente; el artículo 2 fija sus objetivos y el "
        "artículo 9 crea la Unidad u Oficina de Coordinación del Control "
        "Interno como su instancia responsable — precisamente la Dimensión "
        "7 (Control Interno) que evalúa el IDI-MIPG.",
    ),
]


def _fecha_hoy_es():
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
        "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    hoy = datetime.now()
    return f"{hoy.day} de {meses[hoy.month - 1]} de {hoy.year}"


def _construir_contexto_local_regional_global(nombre_entidad, resultado_360, resultado_isvpt):
    """Arma el texto del contexto institucional con enfoque local/regional/global,
    usando SOLO los resultados reales ya calculados por motor_analisis_360 y
    motor_isvpt (si se proporcionaron). No inventa cifras nuevas."""
    lineas = []
    lineas.append(
        f"El caso analizado — {nombre_entidad} — se sitúa simultáneamente en "
        "tres escalas de lectura que exige el enfoque contemporáneo de la "
        "Administración Pública: la escala LOCAL (la entidad y su propia "
        "capacidad instalada), la escala REGIONAL (su comparación frente a "
        "un grupo de entidades similares en Antioquia u otro departamento) "
        "y la escala GLOBAL (los marcos de referencia de la OCDE, la CEPAL, "
        "Naciones Unidas y la Agenda 2030/ODS, particularmente el ODS 16)."
    )
    if resultado_360 is not None:
        lineas.append(
            f"En la escala regional, el análisis 360 del sistema ubica a la entidad "
            f"dentro del grupo de comparación '{resultado_360.filtro_descripcion}', "
            f"compuesto por {resultado_360.n_entidades} entidades, con un IDI "
            f"promedio de grupo de {resultado_360.promedio_idi}."
        )
        if resultado_360.percentil_entidad_referencia is not None:
            lineas.append(
                f"La entidad se ubica en el percentil {resultado_360.percentil_entidad_referencia}% "
                "de ese grupo, con un IDI oficial de "
                f"{resultado_360.idi_entidad_referencia}."
            )
    if resultado_isvpt is not None and resultado_isvpt.isvpt_entidad_referencia is not None:
        lineas.append(
            "El Índice Sintético de Valor Público Territorial (ISVPT) de esta "
            "entidad dentro de su grupo de comparación se desarrolla en la "
            "sección siguiente, como novedad metodológica de este sistema."
        )
    lineas.append(
        "En la escala global, el ODS 16 (paz, justicia e instituciones "
        "sólidas) exige instituciones eficaces, responsables y transparentes; "
        "el desempeño de la entidad en el IDI-MIPG es, en esa medida, un dato "
        "que también dialoga con los compromisos internacionales de Colombia "
        "como miembro pleno de la OCDE desde 2020 y con los estándares que "
        "para la región promueve la CEPAL."
    )
    return lineas


ANALISIS_FINANZAS_PUBLICAS_TITULO = (
    "Análisis de Finanzas Públicas y Control Interno Financiero — matriz de riesgos y controles"
)

ANALISIS_FINANZAS_PUBLICAS_INTRO = (
    "Ningún estudio de caso institucional está completo sin mirar las finanzas públicas "
    "de la entidad, más allá del IDI-MIPG. Se incorpora aquí la metodología oficial de "
    "auditoría al proceso financiero por auditores no financieros, usada por las "
    "Contralorías territoriales, con tres instrumentos reales: la Matriz de Riesgos y "
    "Controles, la Prueba de Recorrido a procesos clave, y la Evaluación del Control "
    "Interno Financiero basada en los 5 componentes del modelo COSO."
)

ANALISIS_FINANZAS_PUBLICAS_COMPONENTES = [
    (
        "Matriz de Riesgos y Controles (probabilidad × impacto)",
        "Para cada riesgo del proceso financiero (Gestión Financiera y Contable, Gestión "
        "Presupuestal y Contractual, entre otros macroprocesos), el auditor califica la "
        "PROBABILIDAD (1 a 5) y el IMPACTO (1 a 5) del riesgo inherente; su producto ubica "
        "el riesgo en una de cuatro zonas: Baja (< 3), Moderada (3 a <6), Alta (6 a <15) o "
        "Extrema (≥ 15). Sobre ese riesgo inherente se evalúa después el DISEÑO del control "
        "(ponderando: control apropiado 30%, tipo automático/manual 25%, frecuencia 10%, "
        "segregación de funciones 25%, documentación 5%, clase preventivo/detectivo 5%) y "
        "la EFECTIVIDAD de su operación (evidencia de uso 20%, incorrecciones materiales "
        "60%, hallazgos repetidos en auditorías anteriores 20%), para llegar a un RIESGO "
        "RESIDUAL o combinado final.",
    ),
    (
        "Prueba de Recorrido a procesos clave",
        "Complementaria a la matriz, consiste en que el auditor 'camine' un proceso "
        "financiero clave de principio a fin (por ejemplo, el ciclo completo de un pago o "
        "de un ingreso), verificando en la práctica —no solo en el papel— que cada control "
        "documentado realmente ocurre como está descrito, con evidencia física o digital "
        "de cada paso.",
    ),
    (
        "Evaluación del Control Interno Financiero (5 componentes COSO)",
        "Un cuestionario estructurado en los 5 componentes del modelo COSO —(A) Ambiente "
        "de Control, (B) Evaluación de Riesgos, (C) Actividades de Control, (D) "
        "Información y Comunicación, (E) Actividades de Supervisión (Monitoreo)— con "
        "preguntas cerradas (Sí/Parcial/No) que se califican y consolidan en un resultado "
        "por componente. La pregunta tipo del componente Ambiente de Control ilustra bien "
        "el nivel de detalle exigido: '¿Se cuenta con un área contable debidamente "
        "conformada y con la segregación de funciones necesaria?', calificada como Sí (1), "
        "Parcial (2) o No (3).",
    ),
]

ANALISIS_FINANZAS_PUBLICAS_CIERRE = (
    "Aplicando esta metodología a las brechas ya identificadas en este estudio de caso "
    "(particularmente las de Control Interno, Gestión para Resultados con Valores y "
    "Evaluación de Resultados), el equipo directivo de la entidad puede construir su "
    "propia matriz de riesgos y controles del proceso financiero, usando exactamente estos "
    "tres instrumentos, disponibles como plantillas oficiales de auditoría."
)

# Dimensiones IDI-MIPG con vínculo directo al proceso financiero/de control, tal como
# las cita ANALISIS_FINANZAS_PUBLICAS_CIERRE. El impacto fijo por dimensión refleja el
# tipo de consecuencia que cada una tiene sobre el proceso financiero: Control Interno
# compromete directamente la responsabilidad disciplinaria/fiscal (impacto máximo),
# mientras que Gestión para Resultados con Valores y Evaluación de Resultados
# comprometen la calidad del gasto y la rendición de cuentas (impacto alto).
_DIMENSIONES_MATRIZ_FINANCIERA = {
    "Dimensión de Control Interno": 5,
    "Dimensión Gestión para Resultados con Valores": 4,
    "Dimensión de Evaluación de Resultados": 4,
}

# La matriz de riesgos y controles usa su propio vocabulario de 4 zonas
# (Extrema/Alta/Moderada/Baja, según probabilidad × impacto), distinto del
# vocabulario de 3 niveles (Alta/Media/Baja) que usa _color_hex_riesgo en el
# resto del sistema — por eso tiene su propio mapa de color aquí.
_COLOR_HEX_POR_ZONA_MATRIZ = {
    "Extrema": "E6B0AA",
    "Alta": "F5B7B1",
    "Moderada": "FAD7A0",
    "Baja": "A9DFBF",
}


def _filas_matriz_riesgos_financieros(diag, top_n: int = 6):
    """Aplica la metodología descrita en ANALISIS_FINANZAS_PUBLICAS_COMPONENTES a las
    brechas REALES del diagnóstico (diag.brechas), filtradas a las tres dimensiones que
    el propio texto de cierre menciona (Control Interno, Gestión para Resultados con
    Valores, Evaluación de Resultados). Convierte cada brecha en una fila de matriz de
    riesgos y controles: probabilidad × impacto → zona de riesgo inherente, diseño del
    control, efectividad de operación y riesgo residual — exactamente las cuatro
    calificaciones que describe la Matriz de Riesgos y Controles del componente 1.

    Devuelve una lista de tuplas:
      (riesgo, probabilidad, impacto, zona_inherente, diseño, efectividad, residual)
    """
    if diag is None:
        return []
    brechas_financieras = [
        b for b in diag.brechas if b.dimension in _DIMENSIONES_MATRIZ_FINANCIERA
    ]
    if not brechas_financieras:
        return []
    brechas_financieras = sorted(brechas_financieras, key=lambda b: b.puntaje)[:top_n]

    filas = []
    for b in brechas_financieras:
        impacto = _DIMENSIONES_MATRIZ_FINANCIERA[b.dimension]
        if b.puntaje < 15:
            probabilidad = 5
        elif b.puntaje < 30:
            probabilidad = 4
        elif b.puntaje < 45:
            probabilidad = 3
        else:
            probabilidad = 2
        producto = probabilidad * impacto
        if producto >= 15:
            zona_inherente = "Extrema"
        elif producto >= 6:
            zona_inherente = "Alta"
        elif producto >= 3:
            zona_inherente = "Moderada"
        else:
            zona_inherente = "Baja"

        if b.puntaje < 20:
            diseno = "Débil — control no documentado o inexistente"
            efectividad = "No efectivo"
        elif b.puntaje < 40:
            diseno = "Parcial — existe pero sin segregación de funciones ni evidencia suficiente"
            efectividad = "Parcialmente efectivo"
        else:
            diseno = "Adecuado en el diseño, débil en la operación"
            efectividad = "Efectivo con hallazgos menores"

        residual_por_zona = {
            "Extrema": "Alta (persiste tras controles compensatorios)",
            "Alta": "Moderada (persiste tras controles compensatorios)",
            "Moderada": "Baja (persiste tras controles compensatorios)",
            "Baja": "Baja",
        }
        riesgo_residual = residual_por_zona[zona_inherente]

        nombre_riesgo = f"{b.codigo_indice} ({b.puntaje}) — {b.nombre_indice}"
        filas.append((
            nombre_riesgo, str(probabilidad), str(impacto), zona_inherente,
            diseno, efectividad, riesgo_residual,
        ))
    return filas


ESCENARIOS_PROSPECTIVOS_TITULO = (
    "Prospectiva: 5 escenarios de desarrollo territorial a partir del cierre de brechas "
    "(2026-2030)"
)

ESCENARIOS_PROSPECTIVOS_INTRO = (
    "A partir del panorama inicial construido en las secciones anteriores (diagnóstico "
    "IDI-MIPG, ISVPT, IDF, MDM/MDD y el Plan de Desarrollo Municipal vigente), se proyectan "
    "aquí 5 escenarios prospectivos —siguiendo la lógica de planeación por escenarios de "
    "Peter Schwartz (1991), adaptada a la gestión pública territorial— que muestran cómo la "
    "entidad puede convertir el cierre de sus brechas en valor público real: cambios "
    "sociales observables y medibles en la calidad de vida de quienes habitan el "
    "territorio. Cada escenario se lee con el triángulo estratégico de Mark Moore (1995): "
    "legitimidad y respaldo político-social, valor sustantivo para la ciudadanía, y "
    "capacidad operativa real para sostenerlo."
)


def _construir_escenarios_prospectivos(nombre_entidad, diag, resultado_360=None):
    """Construye los 5 escenarios prospectivos usando las brechas y políticas REALES
    de la entidad (diag.brechas), sin inventar cifras de desarrollo social — los
    indicadores de medición que se citan son siempre instrumentos oficiales
    existentes (MDM/MDD, IDF, Encuesta de Calidad de Vida, FURAG), no proyecciones
    numéricas propias."""
    politicas_brecha = sorted({b.politica for b in diag.brechas}) if diag.brechas else []
    resumen_politicas = "; ".join(politicas_brecha[:6]) if politicas_brecha else "sus políticas de gestión y desempeño"

    escenarios = [
        (
            "Escenario 1 — Tendencial (si no se interviene)",
            f"Si {nombre_entidad} no ejecuta ningún plan de mejoramiento sobre las brechas "
            f"detectadas ({resumen_politicas}), el patrón esperado —según la evidencia "
            "de capacidad institucional de Oszlak— es la persistencia o profundización "
            "de la brecha frente al grupo par en la siguiente medición FURAG, con riesgo "
            "adicional de deterioro en su Índice de Desempeño Fiscal (IDF) si las "
            "debilidades de control interno y gestión no se corrigen (los rangos IDF caen "
            "de 'Riesgo' a 'Deterioro' cuando la gestión financiera se descuida de forma "
            "sostenida). No genera valor público nuevo: solo mantiene el statu quo, y en el "
            "triángulo de Moore pierde legitimidad política a mediano plazo por ausencia de "
            "resultados visibles para la ciudadanía.",
        ),
        (
            "Escenario 2 — Cierre de brechas institucionales (MIPG)",
            f"Ejecutando un plan de mejoramiento completo sobre {len(diag.brechas)} brechas "
            "detectadas —siguiendo el formato oficial de Función Pública (hallazgo, causa "
            "raíz, acción correctiva, responsable, fecha)— la entidad puede razonablemente "
            "proyectar una mejora medible en su siguiente IDI-MIPG (vigencia siguiente), "
            "verificable de forma pública y comparable en el mismo tablero de Función "
            "Pública consultado para este informe. El valor sustantivo para la ciudadanía "
            "es indirecto pero real: una entidad con mejor control interno, gestión "
            "documental y gestión del conocimiento comete menos errores administrativos que "
            "afectan directamente al ciudadano (trámites perdidos, respuestas tardías, "
            "información no disponible).",
        ),
        (
            "Escenario 3 — Fortalecimiento fiscal y territorial",
            "Complementando el cierre de brechas MIPG con una gestión fiscal más sólida "
            "(mejor programación y ejecución de ingresos, mayor autofinanciación del gasto "
            "de funcionamiento según los límites de la Ley 617 de 2000), la entidad puede "
            "aspirar a mejorar su rango en el Índice de Desempeño Fiscal (IDF) del DNP en la "
            "siguiente vigencia — medible y público, comparando el resultado de este año "
            "contra el siguiente en el mismo grupo de capacidades iniciales. Este escenario "
            "cobra más fuerza si el Proyecto de Ley de Competencias (en trámite) se "
            "convierte en ley: sus principios de subsidiariedad y fortalecimiento "
            "institucional territorial podrían traer nuevas fuentes de acompañamiento "
            "técnico y financiero desde la Nación hacia entidades en esta situación.",
        ),
        (
            "Escenario 4 — Desarrollo social con equidad territorial",
            "Este escenario conecta el cierre de brechas institucionales con las líneas "
            "estratégicas propias del Plan de Desarrollo Municipal vigente de la entidad "
            "(típicamente equidad y desarrollo social, atención a grupos vulnerables, "
            "deporte, cultura e infraestructura social, según la estructura estándar de los "
            "PDM municipales colombianos 2024-2027). Una entidad con mejor gestión del "
            "servicio a la ciudadanía y mejor información y comunicación puede identificar "
            "con más precisión a su población vulnerable y medir mejor el impacto real de "
            "sus programas sociales. La medición de este escenario NO depende de una cifra "
            "inventada por este sistema: se verifica con instrumentos oficiales existentes "
            "— la Medición de Desempeño Municipal (MDM) del DNP, que evalúa resultados de "
            "desarrollo (no solo gestión), y la Encuesta de Calidad de Vida departamental "
            "(en Antioquia, la Encuesta de Calidad de Vida de la Gobernación, edición "
            "25/26 en curso), que permiten observar cambios reales y medibles en la calidad "
            "de vida de la población entre una medición y la siguiente.",
        ),
        (
            "Escenario 5 — Territorio inteligente, abierto y resiliente",
            "El escenario más ambicioso conecta el cierre de brechas con los enfoques "
            "contemporáneos de vanguardia desarrollados en este informe: Estado Digital y "
            "Transformación Digital (cerrando la brecha en Gobierno Digital, si existe), "
            "Gobierno Abierto (mejorando transparencia y participación), y Resiliencia "
            "Institucional (fortaleciendo la capacidad de la entidad para sostener su "
            "desempeño ante choques fiscales, climáticos o sociales, en línea con la Ley "
            "1523 de 2012). Su indicador de éxito es doble y verificable: de un lado, la "
            "mejora en las políticas digitales del IDI-MIPG en la siguiente vigencia; del "
            "otro, un mejor resultado en mediciones independientes de transparencia como el "
            "ITEP de Transparencia por Colombia. En el triángulo de Moore, este escenario es "
            "el que más fortalece la legitimidad política de largo plazo, porque hace "
            "visible y verificable para la ciudadanía —no solo para el auditor— el avance "
            "institucional de la entidad.",
        ),
    ]
    return escenarios


# ---------------------------------------------------------------------------
# Generación del .docx
# ---------------------------------------------------------------------------

def generar_estudio_de_caso_docx(
    nombre_entidad, diag, analisis_ia_texto, cruce_recomendaciones=None,
    resultado_360=None, resultado_isvpt=None, idi_oficial=None, departamento=None,
    tipo_regimen_especial=None,
):
    """Devuelve un BytesIO con el Estudio de Caso Académico en Word (.docx),
    siguiendo la estructura exigida por el microcurrículo para la Unidad 2.

    idi_oficial (opcional pero MUY recomendado): el IDI oficial publicado
    por Función Pública. Este documento SIEMPRE lo usa como cifra
    protagonista; diag.idi_estimado (cálculo interno de SIIEAP) se muestra
    solo como nota metodológica de verificación, nunca como reemplazo.
    departamento (opcional): departamento de la entidad (ej. "ANTIOQUIA" o
    "CHOCÓ"), usado para incluir las fuentes de datos abiertos regionales
    correctas (Contraloría/IDEA para Antioquia, Contraloría/Gobernación
    para Chocó). Si se omite o no está registrado, se usan solo las
    fuentes de alcance nacional.
    """
    doc = Document()
    idi_protagonista = idi_oficial if idi_oficial is not None else diag.idi_estimado
    hay_diferencia_idi = (
        idi_oficial is not None and diag.idi_estimado is not None
        and round(idi_oficial, 2) != round(diag.idi_estimado, 2)
    )

    _agregar_logos_docx(doc)
    _agregar_banner_docx(
        doc,
        "Informe Académico y de Investigación — Estudio de Caso Institucional "
        "(Enfoques y Teorías de la Administración Pública II)",
    )
    _agregar_nota_regimen_especial_docx(doc, tipo_regimen_especial)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nombre_entidad)
    run.bold = True
    run.font.size = Pt(16)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(
        f"Maestría en Administración Pública · ESAP Territorial Antioquia · "
        f"Generado el {_fecha_hoy_es()}"
    )
    p_autoria_ec = doc.add_paragraph()
    p_autoria_ec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_autoria_ec = p_autoria_ec.add_run(
        "Docente: Norma Elizabeth Álvarez Grajales · Área del conocimiento "
        "Organizaciones Públicas y Gestión · ESAP"
    )
    run_autoria_ec.italic = True
    run_autoria_ec.font.size = Pt(9)
    doc.add_page_break()

    _agregar_tabla_contenido_docx(doc)
    _agregar_razon_de_ser_docx(doc, "estudio_caso")

    _agregar_glosario_docx(doc)
    _agregar_normativa_politicas_docx(doc, diag=diag)

    # 1. Descripción del problema público
    _agregar_divisor_seccion_docx(doc, "1. Descripción del problema público", icono="📝")
    try:
        buffer_grafica_dim_ec = generar_grafica_dimensiones(diag)
        doc.add_picture(buffer_grafica_dim_ec, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    if diag.brechas:
        doc.add_paragraph(
            f"A partir del diagnóstico real IDI-MIPG de {nombre_entidad} "
            f"(IDI oficial de Función Pública: {idi_protagonista}), este estudio de caso "
            f"identifica —con metodología propia y exclusiva de SIIEAP, no con una cifra "
            f"publicada por la Función Pública— {len(diag.brechas)} brechas de "
            "implementación por debajo del umbral de alerta metodológico interno del "
            "sistema (60 puntos; la meta plena de la gestión pública es siempre el 100%, "
            "no 60). El problema público objeto de este estudio de "
            "caso es, en términos generales, la brecha entre el estándar de "
            "desempeño institucional exigido por el Modelo Integrado de "
            "Planeación y Gestión (MIPG) y la capacidad real de la entidad "
            "para cumplirlo, particularmente en las siguientes políticas:"
        )
        try:
            buffer_grafica_brechas_ec = generar_grafica_brechas(diag)
            if buffer_grafica_brechas_ec:
                doc.add_picture(buffer_grafica_brechas_ec, width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
        for b in diag.brechas:
            doc.add_paragraph(
                f"{b.codigo_indice} ({b.puntaje}) — {b.nombre_indice} · {b.politica}",
                style="List Bullet",
            )
        doc.add_paragraph()
        doc.add_heading("Conexión de las brechas con los enfoques contemporáneos de la Administración Pública", level=2)
        politicas_con_brecha_ec = []
        vistas_ec = set()
        for b in diag.brechas:
            if b.politica not in vistas_ec:
                vistas_ec.add(b.politica)
                politicas_con_brecha_ec.append(b.politica)
        tabla_enfoques_ec = doc.add_table(rows=1, cols=3)
        tabla_enfoques_ec.style = "Light Grid Accent 1"
        enc_ef_ec = tabla_enfoques_ec.rows[0].cells
        for celda, texto in zip(enc_ef_ec, ["Política con brecha", "Enfoque contemporáneo", "Norma"]):
            celda.text = texto
            for parrafo in celda.paragraphs:
                for run_enc in parrafo.runs:
                    run_enc.bold = True
        for politica in politicas_con_brecha_ec:
            enfoque, norma = _enfoque_y_norma_de_politica(politica)
            fila = tabla_enfoques_ec.add_row().cells
            fila[0].text = politica
            fila[1].text = enfoque
            fila[2].text = norma
        _franjas_alternas_docx(tabla_enfoques_ec)
        doc.add_paragraph(
            f"El diagnóstico IDI-MIPG de {nombre_entidad} no registra brechas "
            "por debajo del umbral con los datos disponibles; el problema "
            "público de este estudio de caso puede orientarse entonces hacia "
            "el sostenimiento y profundización del desempeño ya alcanzado."
        )
    if hay_diferencia_idi:
        p_nota_idi_ec = doc.add_paragraph()
        run_nota_idi_ec = p_nota_idi_ec.add_run(
            f"Nota metodológica: el IDI oficial de Función Pública para esta entidad es "
            f"{idi_oficial}, cifra que este estudio de caso usa siempre como referencia principal "
            f"y que en ningún caso se desvirtúa. El sistema SIIEAP calcula, además, un valor interno "
            f"de verificación de {diag.idi_estimado}, revisando la suma y ponderación de los índices "
            f"del archivo oficial — un ejercicio de validación metodológica interna, no un dato "
            f"alternativo al oficial."
        )
        run_nota_idi_ec.italic = True
        run_nota_idi_ec.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 2. Contexto institucional (local, regional, global)
    _agregar_divisor_seccion_docx(doc, "2. Contexto institucional — enfoque local, regional y global", icono="🌍")
    for linea in _construir_contexto_local_regional_global(nombre_entidad, resultado_360, resultado_isvpt):
        doc.add_paragraph(linea)

    # 2bis. Índice Sintético de Valor Público Territorial (ISVPT) — novedad metodológica
    if resultado_isvpt is not None and resultado_isvpt.isvpt_entidad_referencia is not None:
        doc.add_heading("2.1 🎯 El Termómetro del Valor Público: Índice Sintético de Valor Público Territorial (ISVPT)", level=2)
        doc.add_paragraph(
            "Complementario al IDI oficial, este índice normaliza (min-max) las 7 "
            "dimensiones del IDI-MIPG DENTRO del grupo de comparación de la "
            "entidad y las agrega en un solo valor relativo entre 0 y 1, "
            "siguiendo la metodología académica validada por Vélez Tamayo, "
            "Ortiz-Muñoz y Cardona Montoya (2026) para el Índice Sintético de "
            "Desarrollo Económico Local (ISDEL, Administración & Desarrollo, "
            "56(1), e-1352) y las directrices de la OCDE (2008, 'Handbook on "
            "Constructing Composite Indicators') para indicadores compuestos."
        )
        p_isvpt_ec = doc.add_paragraph()
        run_isvpt_ec = p_isvpt_ec.add_run(
            f"ISVPT de {nombre_entidad}: {resultado_isvpt.isvpt_entidad_referencia} "
            f"(posición {resultado_isvpt.posicion_entidad_referencia} de "
            f"{resultado_isvpt.n_entidades} en su grupo de comparación)."
        )
        run_isvpt_ec.bold = True
        if resultado_isvpt.subindices_por_dimension_entidad:
            tabla_isvpt_ec = doc.add_table(rows=1, cols=2)
            tabla_isvpt_ec.style = "Light Grid Accent 1"
            enc = tabla_isvpt_ec.rows[0].cells
            enc[0].text, enc[1].text = "Dimensión", "Subíndice normalizado (0 a 1)"
            for celda in enc:
                for parrafo in celda.paragraphs:
                    for run_enc in parrafo.runs:
                        run_enc.bold = True
            for dim, valor in resultado_isvpt.subindices_por_dimension_entidad.items():
                fila = tabla_isvpt_ec.add_row().cells
                fila[0].text = dim
                fila[1].text = str(valor)

        doc.add_paragraph()
        p_nota_isvpt_ec = doc.add_paragraph()
        run_nota_isvpt_ec = p_nota_isvpt_ec.add_run(resultado_isvpt.nota_metodologica)
        run_nota_isvpt_ec.italic = True
        run_nota_isvpt_ec.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if resultado_isvpt.nota_valor_extremo:
            p_nota_extremo_ec = doc.add_paragraph()
            p_nota_extremo_ec.paragraph_format.space_before = Pt(6)
            run_nota_extremo_ec = p_nota_extremo_ec.add_run(resultado_isvpt.nota_valor_extremo)
            run_nota_extremo_ec.bold = True
            run_nota_extremo_ec.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

    # 2.1bis Estado del Arte Inicial: ecosistema de mediciones abiertas
    doc.add_page_break()
    doc.add_heading(ESTADO_DEL_ARTE_TITULO, level=2)
    doc.add_paragraph(ESTADO_DEL_ARTE_INTRO)
    for titulo_indice, texto_indice in ESTADO_DEL_ARTE_INDICES:
        p_ei = doc.add_paragraph()
        run_ei = p_ei.add_run(f"{titulo_indice}. ")
        run_ei.bold = True
        p_ei.add_run(texto_indice)

    # 2.2 Institucionalidad del IDI-MIPG
    doc.add_heading(INSTITUCIONALIDAD_IDI_TITULO, level=2)
    doc.add_paragraph(INSTITUCIONALIDAD_IDI_INTRO)
    for titulo_i, texto_i in INSTITUCIONALIDAD_IDI:
        p_i = doc.add_paragraph()
        run_i = p_i.add_run(f"{titulo_i}. ")
        run_i.bold = True
        p_i.add_run(texto_i)

    # 2.3 Marco de la reforma territorial en curso
    doc.add_heading(MARCO_REFORMA_TERRITORIAL_TITULO, level=2)
    doc.add_paragraph(MARCO_REFORMA_TERRITORIAL_INTRO)
    for titulo_r, texto_r in MARCO_REFORMA_TERRITORIAL:
        p_r = doc.add_paragraph()
        run_r = p_r.add_run(f"{titulo_r}. ")
        run_r.bold = True
        p_r.add_run(texto_r)

    # 3. Fuentes de datos abiertos
    _agregar_divisor_seccion_docx(doc, "3. Fuentes de datos abiertos utilizadas y recomendadas", icono="🗂️")
    doc.add_paragraph(
        "Este estudio de caso se fundamenta en la fuente oficial IDI-MIPG "
        "cargada en el sistema; se listan además las plataformas oficiales "
        "de datos abiertos del Estado colombiano que el estudiante puede "
        "consultar para profundizar el análisis exógeno y endógeno de la "
        "entidad (contratación, finanzas, talento humano, control fiscal y "
        "disciplinario):"
    )
    for nombre_fuente, descripcion_fuente in _fuentes_datos_abiertos_para(departamento):
        p_fuente = doc.add_paragraph(style="List Bullet")
        run_fuente = p_fuente.add_run(f"{nombre_fuente}: ")
        run_fuente.bold = True
        p_fuente.add_run(descripcion_fuente)

    p_indep = doc.add_paragraph()
    run_indep = p_indep.add_run(
        "Nota de independencia metodológica: el diagnóstico, los cálculos (IDI interno de "
        "verificación, ISVPT, análisis 360) y este informe completo se generan a partir del "
        "archivo oficial de resultados ya cargado en el sistema, sin depender de que estos "
        "portales de consulta estén disponibles en el momento de la entrega. Las plataformas "
        "listadas son un complemento para profundizar el análisis, no un requisito para "
        "generar el informe: ningún estudiante depende de la disponibilidad de un sitio web "
        "del Gobierno para completar su trabajo académico."
    )
    run_indep.italic = True
    run_indep.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # 4. Análisis desde las tres teorías + gobernanza + riesgo + plan de mejoramiento (contenido de motor_analisis_ia)
    _agregar_divisor_seccion_docx(
        doc,
        "4. Análisis desde la Nueva Gestión Pública, la Post-Nueva Gestión "
        "Pública y el Nuevo Institucionalismo, señales de riesgo multinivel "
        "y Plan de Mejoramiento Prospectivo orientado a valor público",
        icono="🧭",
    )
    p_fundamentacion_ec = doc.add_paragraph()
    run_fundamentacion_ec = p_fundamentacion_ec.add_run(
        "Fundamentación teórica de referencia para esta sección: desde la ESAP, "
        "Chica-Vélez y Salazar-Ortiz (2021) documentan el tránsito de la Nueva "
        "Gestión Pública hacia la Post-Nueva Gestión Pública y la gobernanza como "
        "forma de organización de lo público; y Jurado-Zambrano y Villanueva (2021) "
        "aportan la metodología de gestión integrada de riesgos (ISO 31000) aplicada "
        "al sector público colombiano que orienta la lectura de riesgo institucional "
        "que sigue."
    )
    run_fundamentacion_ec.italic = True
    run_fundamentacion_ec.font.size = Pt(9.5)
    doc.add_paragraph(
        "El siguiente análisis fue generado con apoyo de inteligencia "
        "artificial (Claude, Anthropic), a partir EXCLUSIVAMENTE de los "
        "datos reales del diagnóstico de la entidad y de las recomendaciones "
        "oficiales de Función Pública ya cargadas en el sistema. Es un punto "
        "de partida metodológico para la discusión académica, no un dictamen "
        "definitivo (ver nota de uso responsable de IA al final)."
    )
    _agregar_texto_markdown_docx(doc, analisis_ia_texto)

    doc.add_page_break()

    # 5. Fundamento jurídico, fiscal, disciplinario, contable y de control interno
    _agregar_divisor_seccion_docx(doc, "5. " + FUNDAMENTO_JURIDICO_AMPLIADO_TITULO, icono="⚖️")
    doc.add_paragraph(FUNDAMENTO_JURIDICO_AMPLIADO_INTRO)
    for titulo_f, texto_f in FUNDAMENTO_JURIDICO_AMPLIADO:
        p_f = doc.add_paragraph()
        run_f = p_f.add_run(f"{titulo_f}. ")
        run_f.bold = True
        p_f.add_run(texto_f)

    # 5.1 Análisis de Finanzas Públicas y Control Interno Financiero
    doc.add_page_break()
    doc.add_heading("5.1 " + ANALISIS_FINANZAS_PUBLICAS_TITULO, level=2)
    doc.add_paragraph(ANALISIS_FINANZAS_PUBLICAS_INTRO)
    try:
        buffer_matriz_riesgo = generar_matriz_riesgo_probabilidad_impacto()
        doc.add_picture(buffer_matriz_riesgo, width=Inches(4.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    for titulo_fin, texto_fin in ANALISIS_FINANZAS_PUBLICAS_COMPONENTES:
        p_fin = doc.add_paragraph()
        run_fin = p_fin.add_run(f"{titulo_fin}. ")
        run_fin.bold = True
        p_fin.add_run(texto_fin)

    # Matriz de riesgos y controles APLICADA a las brechas reales de esta entidad
    # (no solo la metodología en abstracto) — ver docstring de
    # _filas_matriz_riesgos_financieros.
    filas_matriz_fin = _filas_matriz_riesgos_financieros(diag)
    if filas_matriz_fin:
        doc.add_heading("Matriz de riesgos y controles aplicada a las brechas de esta entidad", level=3)
        doc.add_paragraph(
            "A manera de ejemplo de aplicación concreta —no como sustituto del ejercicio "
            "completo que debe hacer el equipo directivo—, se califican aquí, con la "
            "metodología de probabilidad × impacto descrita arriba, las brechas detectadas "
            "por SIIEAP en las tres dimensiones con vínculo directo al proceso financiero "
            "(Control Interno, Gestión para Resultados con Valores y Evaluación de "
            "Resultados)."
        )
        tabla_matriz_fin = doc.add_table(rows=1, cols=7)
        tabla_matriz_fin.style = "Light Grid Accent 1"
        encabezados_matriz_fin = [
            "Riesgo (brecha SIIEAP)", "Prob.", "Imp.", "Zona (inherente)",
            "Diseño del control", "Efectividad", "Riesgo residual",
        ]
        for celda, texto in zip(tabla_matriz_fin.rows[0].cells, encabezados_matriz_fin):
            celda.text = texto
            _sombrear_celda(celda, COLOR_INSTITUCIONAL)
            for parrafo in celda.paragraphs:
                for run_enc in parrafo.runs:
                    run_enc.bold = True
                    run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for fila_datos in filas_matriz_fin:
            fila = tabla_matriz_fin.add_row().cells
            for idx_col, valor in enumerate(fila_datos):
                fila[idx_col].text = valor
            _sombrear_celda(fila[3], _COLOR_HEX_POR_ZONA_MATRIZ.get(fila_datos[3], "FFFFFF"))
        _ajustar_tabla_docx(
            tabla_matriz_fin,
            anchos_cm=[4.0, 1.0, 1.0, 2.2, 3.4, 2.5, 2.9],
            tamano_fuente_pt=8,
        )
        doc.add_paragraph()

    doc.add_paragraph(ANALISIS_FINANZAS_PUBLICAS_CIERRE)

    # 6. Recomendaciones oficiales de Función Pública vinculadas a las brechas
    _agregar_divisor_seccion_docx(doc, "6. Recomendaciones oficiales de Función Pública", icono="📋")
    if cruce_recomendaciones:
        for codigo, lista_recos in cruce_recomendaciones.items():
            doc.add_paragraph(f"Brecha {codigo}:", style="List Bullet")
            for texto_reco in lista_recos:
                doc.add_paragraph(texto_reco, style="List Bullet 2")
    else:
        doc.add_paragraph(
            "No se cargó un consolidado de recomendaciones oficiales para "
            "esta entidad en esta sesión; se recomienda cargarlo en la "
            "barra lateral del sistema para enriquecer este apartado con el "
            "texto oficial de Función Pública."
        )

    doc.add_page_break()

    # 6.1 Plan de Mejoramiento (formato oficial de Función Pública)
    doc.add_heading("6.1 Plan de Mejoramiento — formato oficial de Función Pública", level=2)
    doc.add_paragraph(
        "Se estructura a continuación en el formato exacto que usan las entidades "
        "públicas colombianas para su Plan de Mejoramiento institucional (columnas: "
        "Ítem, Observación/hallazgo, Riesgo materializado, Causa raíz, Acción correctiva, "
        "Fecha inicio, Fecha fin, Soporte, Responsable), consolidando TODAS las brechas "
        "detectadas — no una selección — como exige un ejercicio de auditoría real."
    )
    tabla_pm = doc.add_table(rows=1, cols=6)
    tabla_pm.style = "Light Grid Accent 1"
    enc_pm = tabla_pm.rows[0].cells
    for celda, texto in zip(enc_pm, ["Ítem", "Observación/Hallazgo", "Causa raíz probable", "Acción correctiva a implementar", "Plazo sugerido", "Responsable sugerido"]):
        celda.text = texto
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
    for i, b in enumerate(diag.brechas, start=1):
        enfoque_pm, norma_pm = _enfoque_y_norma_de_politica(b.politica)
        fila = tabla_pm.add_row().cells
        fila[0].text = str(i)
        fila[1].text = f"{b.codigo_indice} ({b.puntaje}) — {b.nombre_indice} · {b.politica}"
        fila[2].text = "Debilidad de capacidad institucional (Oszlak) en el proceso asociado; ver análisis de riesgo de la sección 4."
        fila[3].text = f"Implementar acción correctiva alineada con {enfoque_pm} ({norma_pm}); ver detalle en recomendaciones de la sección 4 y 6."
        fila[4].text = "0-6 meses" if b.puntaje < 20 else ("6-18 meses" if b.puntaje < 45 else "18-36 meses")
        fila[5].text = "Líder de proceso / Jefe de Control Interno"
    # Sin un ancho de columna explícito, Word puede repartir el espacio de
    # forma pareja entre las 6 columnas sin importar cuánto texto tenga cada
    # una — con columnas de texto largo (Observación, Acción correctiva) eso
    # las deja demasiado angostas y el contenido se ve apretado o se sale de
    # los márgenes de la página. _ajustar_tabla_docx fija anchos razonables
    # y sincroniza también el tblGrid (ver su docstring).
    _ajustar_tabla_docx(
        tabla_pm,
        anchos_cm=[1.0, 4.0, 3.3, 4.5, 2.2, 2.5],
        tamano_fuente_pt=8.5,
    )

    doc.add_paragraph()
    doc.add_heading("Indicadores de efectividad del Plan de Mejoramiento (desde resultados)", level=2)
    doc.add_paragraph(
        "Siguiendo el formato oficial de Seguimiento al Plan de Mejoramiento de Función "
        "Pública (columnas Estado, Fecha de seguimiento, Observación), se sugieren los "
        "siguientes indicadores de efectividad — medidos desde el RESULTADO, no desde el "
        "cumplimiento formal de la tarea:"
    )
    for texto_ind in [
        "% de brechas cerradas efectivamente (puntaje ≥ 60 en la siguiente medición FURAG) sobre el total de brechas del plan.",
        "Variación real del IDI-MIPG oficial entre esta vigencia y la siguiente (puntos ganados/perdidos).",
        "Variación del Índice de Desempeño Fiscal (IDF) entre vigencias, como indicador de resultado financiero.",
        "Número de hallazgos de auditoría (interna o de la Contraloría) reincidentes de una vigencia a otra (debe tender a cero).",
        "Tiempo promedio real de cierre de cada acción correctiva frente al plazo sugerido (eficacia en la ejecución, no solo en el registro).",
    ]:
        doc.add_paragraph(texto_ind, style="List Bullet")

    doc.add_page_break()

    # 7. Fortalezas y debilidades
    _agregar_divisor_seccion_docx(doc, "7. Fortalezas y debilidades identificadas", icono="💪")
    fortalezas = [r for r in diag.resultados_por_dimension if (r.promedio or 0) >= 60]
    debilidades = [r for r in diag.resultados_por_dimension if (r.promedio or 0) < 60]
    doc.add_heading("Fortalezas", level=2)
    if fortalezas:
        for r in fortalezas:
            doc.add_paragraph(f"{r.codigo} {r.nombre}: {r.promedio} puntos", style="List Bullet")
    else:
        doc.add_paragraph("Con los datos disponibles, no se identifican dimensiones por encima de 60 puntos.")
    doc.add_heading("Debilidades", level=2)
    if debilidades:
        for r in debilidades:
            doc.add_paragraph(f"{r.codigo} {r.nombre}: {r.promedio} puntos", style="List Bullet")
    else:
        doc.add_paragraph("Con los datos disponibles, no se identifican dimensiones por debajo de 60 puntos.")

    doc.add_page_break()

    # 8. Prospectiva: 5 escenarios de desarrollo territorial
    _agregar_divisor_seccion_docx(doc, "8. " + ESCENARIOS_PROSPECTIVOS_TITULO, icono="🔮")
    doc.add_paragraph(ESCENARIOS_PROSPECTIVOS_INTRO)
    for titulo_esc, texto_esc in _construir_escenarios_prospectivos(nombre_entidad, diag, resultado_360):
        p_esc = doc.add_paragraph()
        run_esc = p_esc.add_run(f"{titulo_esc}. ")
        run_esc.bold = True
        p_esc.add_run(texto_esc)

    # 9. Nota de uso responsable de IA
    _agregar_divisor_seccion_docx(doc, "9. Nota de uso responsable de inteligencia artificial", icono="🤖")
    p_ia = doc.add_paragraph()
    run_ia = p_ia.add_run(
        "Este estudio de caso usó inteligencia artificial (Claude, Anthropic) "
        "como herramienta de apoyo para estructurar el análisis teórico a "
        "partir de datos reales, siguiendo el criterio transversal de uso "
        "ético de la IA exigido por el microcurrículo. La IA no sustituye el "
        "juicio académico del estudiante: los datos institucionales, "
        "normativos y de recomendaciones son reales y verificables; la "
        "interpretación teórica generada debe ser revisada, discutida y "
        "complementada por el estudiante antes de su entrega."
    )
    run_ia.italic = True
    run_ia.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _agregar_marco_descentralizacion_docx(doc, diag=diag, nombre_entidad=nombre_entidad)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Generación del PDF (reportlab)
# ---------------------------------------------------------------------------

def generar_estudio_de_caso_pdf(
    nombre_entidad, diag, analisis_ia_texto, cruce_recomendaciones=None,
    resultado_360=None, resultado_isvpt=None, idi_oficial=None, departamento=None,
    tipo_regimen_especial=None,
):
    """Devuelve un BytesIO con el Estudio de Caso Académico en PDF.

    idi_oficial (opcional pero MUY recomendado): ver docstring de
    generar_estudio_de_caso_docx.
    departamento (opcional): ver docstring de generar_estudio_de_caso_docx.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=LETTER,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )
    estilos = getSampleStyleSheet()
    estilo_titulo = ParagraphStyle("TituloEC", parent=estilos["Title"], fontSize=16)
    estilo_h2 = ParagraphStyle("H2EC", parent=estilos["Heading2"], spaceBefore=12)
    estilo_normal = estilos["Normal"]
    estilo_cursiva = ParagraphStyle("CursivaEC", parent=estilos["Normal"], fontName="Helvetica-Oblique", textColor=colors.grey)

    idi_protagonista = idi_oficial if idi_oficial is not None else diag.idi_estimado
    hay_diferencia_idi = (
        idi_oficial is not None and diag.idi_estimado is not None
        and round(idi_oficial, 2) != round(diag.idi_estimado, 2)
    )

    elementos = []
    elementos.extend(_logos_pdf_flowables())
    elementos.extend(_banner_portada_pdf_flowables(
        "Informe Académico y de Investigación — Estudio de Caso Institucional "
        "(Enfoques y Teorías de la Administración Pública II)",
    ))
    elementos.extend(_nota_regimen_especial_pdf_flowables(tipo_regimen_especial))
    elementos.append(Paragraph(f"<b>{nombre_entidad}</b>", estilos["Heading2"]))
    elementos.append(Paragraph(f"Maestría en Administración Pública · ESAP Territorial Antioquia · Generado el {_fecha_hoy_es()}", estilo_normal))
    elementos.append(Paragraph(
        "Docente: Norma Elizabeth Álvarez Grajales · Área del conocimiento "
        "Organizaciones Públicas y Gestión · ESAP",
        estilo_cursiva,
    ))
    elementos.append(PageBreak())

    elementos.extend(_toc_pdf_flowables([
        "Descripción del problema público",
        "Contexto institucional y fuentes de datos abiertos",
        "Análisis teórico (NGP, Post-NGP, Nuevo Institucionalismo)",
        "Fundamento jurídico ampliado",
        "Fortalezas, debilidades y escenarios prospectivos",
        "Nota de uso responsable de IA",
    ]))
    elementos.extend(_razon_de_ser_pdf_flowables("estudio_caso"))

    _agregar_glosario_pdf(elementos, estilos, estilo_normal, estilo_h2)
    _agregar_normativa_politicas_pdf(elementos, estilos, estilo_normal, estilo_h2, diag=diag)

    elementos.extend(_divisor_seccion_pdf("1. Descripción del problema público", icono="📝"))
    try:
        buffer_grafica_dim_ec = generar_grafica_dimensiones(diag)
        elementos.append(Image(buffer_grafica_dim_ec, width=16 * cm, height=16 * cm * 0.5))
        elementos.append(Spacer(1, 8))
    except Exception:
        pass
    if diag.brechas:
        elementos.append(Paragraph(
            f"A partir del diagnóstico real IDI-MIPG de {nombre_entidad} (IDI oficial de Función Pública: "
            f"{idi_protagonista}), se identifican {len(diag.brechas)} brechas de "
            "implementación por debajo del umbral esperado, entre ellas:",
            estilo_normal,
        ))
        try:
            buffer_grafica_brechas_ec = generar_grafica_brechas(diag)
            if buffer_grafica_brechas_ec:
                elementos.append(Image(buffer_grafica_brechas_ec, width=16 * cm, height=16 * cm * 0.55))
                elementos.append(Spacer(1, 8))
        except Exception:
            pass
        for b in diag.brechas:
            elementos.append(Paragraph(f"• {b.codigo_indice} ({b.puntaje}) — {b.nombre_indice} · {b.politica}", estilo_normal))
        elementos.append(Spacer(1, 8))
        elementos.append(Paragraph("Conexión de las brechas con los enfoques contemporáneos de la Administración Pública", estilo_h2))
        politicas_con_brecha_ec = []
        vistas_ec = set()
        for b in diag.brechas:
            if b.politica not in vistas_ec:
                vistas_ec.add(b.politica)
                politicas_con_brecha_ec.append(b.politica)
        datos_tabla_enfoques_ec = [[_celda_pdf("Política con brecha", encabezado=True), _celda_pdf("Enfoque contemporáneo", encabezado=True), _celda_pdf("Norma", encabezado=True)]]
        for politica in politicas_con_brecha_ec:
            enfoque, norma = _enfoque_y_norma_de_politica(politica)
            datos_tabla_enfoques_ec.append([_celda_pdf(politica), _celda_pdf(enfoque), _celda_pdf(norma)])
        tabla_enfoques_ec = Table(datos_tabla_enfoques_ec, hAlign="LEFT", colWidths=[5 * cm, 5 * cm, 6 * cm])
        tabla_enfoques_ec.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        elementos.append(tabla_enfoques_ec)
    else:
        elementos.append(Paragraph(
            f"El diagnóstico IDI-MIPG de {nombre_entidad} no registra brechas por "
            "debajo del umbral con los datos disponibles.", estilo_normal,
        ))
    if hay_diferencia_idi:
        elementos.append(Paragraph(
            f"Nota metodológica: el IDI oficial de Función Pública para esta entidad es {idi_oficial}, "
            f"cifra que este estudio de caso usa siempre como referencia principal. El sistema SIIEAP "
            f"calcula además un valor interno de verificación de {diag.idi_estimado}, revisando la suma "
            f"y ponderación de los índices del archivo oficial — un ejercicio de validación metodológica "
            f"interna, que en ningún caso reemplaza ni desvirtúa el dato oficial.",
            estilo_cursiva,
        ))
    elementos.append(Spacer(1, 8))

    elementos.extend(_divisor_seccion_pdf("2. Contexto institucional — enfoque local, regional y global", icono="🌍"))
    for linea in _construir_contexto_local_regional_global(nombre_entidad, resultado_360, resultado_isvpt):
        elementos.append(Paragraph(linea, estilo_normal))
        elementos.append(Spacer(1, 4))

    if resultado_isvpt is not None and resultado_isvpt.isvpt_entidad_referencia is not None:
        elementos.append(Paragraph("2.1 🎯 El Termómetro del Valor Público: Índice Sintético de Valor Público Territorial (ISVPT)", estilo_h2))
        elementos.append(Paragraph(
            "Complementario al IDI oficial, este índice normaliza (min-max) las 7 "
            "dimensiones del IDI-MIPG dentro del grupo de comparación de la entidad, "
            "siguiendo la metodología académica de Vélez Tamayo, Ortiz-Muñoz y "
            "Cardona Montoya (2026, ISDEL) y las directrices de la OCDE (2008) para "
            "indicadores compuestos.",
            estilo_normal,
        ))
        elementos.append(Spacer(1, 4))
        elementos.append(Paragraph(
            f"<b>ISVPT de {nombre_entidad}: {resultado_isvpt.isvpt_entidad_referencia} "
            f"(posición {resultado_isvpt.posicion_entidad_referencia} de "
            f"{resultado_isvpt.n_entidades} en su grupo de comparación).</b>",
            estilo_normal,
        ))
        if resultado_isvpt.subindices_por_dimension_entidad:
            elementos.append(Spacer(1, 6))
            datos_tabla_isvpt_ec = [["Dimensión", "Subíndice normalizado"]]
            for dim, valor in resultado_isvpt.subindices_por_dimension_entidad.items():
                datos_tabla_isvpt_ec.append([dim, str(valor)])
            tabla_isvpt_ec = Table(datos_tabla_isvpt_ec, hAlign="LEFT", colWidths=[10 * cm, 4 * cm])
            tabla_isvpt_ec.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            elementos.append(tabla_isvpt_ec)
        elementos.append(Spacer(1, 8))

        elementos.append(Paragraph(resultado_isvpt.nota_metodologica, estilo_cursiva))

        if resultado_isvpt.nota_valor_extremo:
            estilo_nota_extremo_ec = ParagraphStyle(
                "NotaExtremoISVPT_EC", parent=estilo_normal,
                textColor=colors.HexColor("#C05000"), fontName="Helvetica-Bold",
                spaceBefore=6,
            )
            elementos.append(Paragraph(resultado_isvpt.nota_valor_extremo, estilo_nota_extremo_ec))
            elementos.append(Spacer(1, 6))

    elementos.append(PageBreak())
    elementos.append(Paragraph(ESTADO_DEL_ARTE_TITULO, estilo_h2))
    elementos.append(Paragraph(ESTADO_DEL_ARTE_INTRO, estilo_normal))
    elementos.append(Spacer(1, 6))
    for titulo_indice, texto_indice in ESTADO_DEL_ARTE_INDICES:
        elementos.append(Paragraph(f"<b>{titulo_indice}.</b> {texto_indice}", estilo_normal))
        elementos.append(Spacer(1, 8))

    elementos.append(Paragraph(INSTITUCIONALIDAD_IDI_TITULO, estilo_h2))
    elementos.append(Paragraph(INSTITUCIONALIDAD_IDI_INTRO, estilo_normal))
    elementos.append(Spacer(1, 4))
    for titulo_i, texto_i in INSTITUCIONALIDAD_IDI:
        elementos.append(Paragraph(f"<b>{titulo_i}.</b> {texto_i}", estilo_normal))
        elementos.append(Spacer(1, 6))

    elementos.append(Paragraph(MARCO_REFORMA_TERRITORIAL_TITULO, estilo_h2))
    elementos.append(Paragraph(MARCO_REFORMA_TERRITORIAL_INTRO, estilo_normal))
    elementos.append(Spacer(1, 4))
    for titulo_r, texto_r in MARCO_REFORMA_TERRITORIAL:
        elementos.append(Paragraph(f"<b>{titulo_r}.</b> {texto_r}", estilo_normal))
        elementos.append(Spacer(1, 6))

    elementos.append(Paragraph("3. Fuentes de datos abiertos utilizadas y recomendadas", estilo_h2))
    for nombre_fuente, descripcion_fuente in _fuentes_datos_abiertos_para(departamento):
        elementos.append(Paragraph(f"<b>{nombre_fuente}:</b> {descripcion_fuente}", estilo_normal))
        elementos.append(Spacer(1, 3))
    elementos.append(Paragraph(
        "Nota de independencia metodológica: el diagnóstico, los cálculos (IDI interno de "
        "verificación, ISVPT, análisis 360) y este informe completo se generan a partir del "
        "archivo oficial de resultados ya cargado en el sistema, sin depender de que estos "
        "portales de consulta estén disponibles en el momento de la entrega. Las plataformas "
        "listadas son un complemento para profundizar el análisis, no un requisito para "
        "generar el informe: ningún estudiante depende de la disponibilidad de un sitio web "
        "del Gobierno para completar su trabajo académico.",
        estilo_cursiva,
    ))
    elementos.append(PageBreak())

    elementos.extend(_divisor_seccion_pdf(
        "4. Análisis desde la Nueva Gestión Pública, la Post-Nueva Gestión Pública "
        "y el Nuevo Institucionalismo, señales de riesgo multinivel y Plan de "
        "Mejoramiento Prospectivo orientado a valor público",
        icono="🧭",
    ))
    elementos.append(Paragraph(
        "Fundamentación teórica de referencia para esta sección: desde la ESAP, "
        "Chica-Vélez y Salazar-Ortiz (2021) documentan el tránsito de la Nueva "
        "Gestión Pública hacia la Post-Nueva Gestión Pública y la gobernanza como "
        "forma de organización de lo público; y Jurado-Zambrano y Villanueva (2021) "
        "aportan la metodología de gestión integrada de riesgos (ISO 31000) aplicada "
        "al sector público colombiano que orienta la lectura de riesgo institucional "
        "que sigue.",
        estilo_cursiva,
    ))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph(
        "Análisis generado con apoyo de IA a partir de datos reales del diagnóstico "
        "y las recomendaciones oficiales de Función Pública. Punto de partida "
        "metodológico, no dictamen definitivo.", estilo_cursiva,
    ))
    elementos.append(Spacer(1, 6))
    elementos.extend(_texto_markdown_a_pdf_flowables(analisis_ia_texto, estilo_normal, estilo_h2))
    elementos.append(PageBreak())

    elementos.extend(_divisor_seccion_pdf("5. " + FUNDAMENTO_JURIDICO_AMPLIADO_TITULO, icono="⚖️"))
    elementos.append(Paragraph(FUNDAMENTO_JURIDICO_AMPLIADO_INTRO, estilo_normal))
    elementos.append(Spacer(1, 6))
    for titulo_f, texto_f in FUNDAMENTO_JURIDICO_AMPLIADO:
        elementos.append(Paragraph(f"<b>{titulo_f}.</b> {texto_f}", estilo_normal))
        elementos.append(Spacer(1, 6))
    elementos.append(PageBreak())

    elementos.append(Paragraph("5.1 " + ANALISIS_FINANZAS_PUBLICAS_TITULO, estilo_h2))
    elementos.append(Paragraph(ANALISIS_FINANZAS_PUBLICAS_INTRO, estilo_normal))
    elementos.append(Spacer(1, 6))
    try:
        buffer_matriz_riesgo = generar_matriz_riesgo_probabilidad_impacto()
        elementos.append(Image(buffer_matriz_riesgo, width=11 * cm, height=11 * cm))
        elementos.append(Spacer(1, 8))
    except Exception:
        pass
    for titulo_fin, texto_fin in ANALISIS_FINANZAS_PUBLICAS_COMPONENTES:
        elementos.append(Paragraph(f"<b>{titulo_fin}.</b> {texto_fin}", estilo_normal))
        elementos.append(Spacer(1, 6))

    # Matriz de riesgos y controles APLICADA a las brechas reales de esta entidad
    # (ver docstring de _filas_matriz_riesgos_financieros).
    filas_matriz_fin = _filas_matriz_riesgos_financieros(diag)
    if filas_matriz_fin:
        elementos.append(Paragraph("Matriz de riesgos y controles aplicada a las brechas de esta entidad", estilo_h2))
        elementos.append(Paragraph(
            "A manera de ejemplo de aplicación concreta —no como sustituto del ejercicio "
            "completo que debe hacer el equipo directivo—, se califican aquí, con la "
            "metodología de probabilidad × impacto descrita arriba, las brechas detectadas "
            "por SIIEAP en las tres dimensiones con vínculo directo al proceso financiero "
            "(Control Interno, Gestión para Resultados con Valores y Evaluación de Resultados).",
            estilo_normal,
        ))
        elementos.append(Spacer(1, 6))
        datos_matriz_fin = [[
            _celda_pdf("Riesgo (brecha SIIEAP)", encabezado=True),
            _celda_pdf("Prob.", encabezado=True),
            _celda_pdf("Imp.", encabezado=True),
            _celda_pdf("Zona (inherente)", encabezado=True),
            _celda_pdf("Diseño del control", encabezado=True),
            _celda_pdf("Efectividad", encabezado=True),
            _celda_pdf("Riesgo residual", encabezado=True),
        ]]
        colores_fila_matriz_fin = []
        for fila_datos in filas_matriz_fin:
            datos_matriz_fin.append([_celda_pdf(v) for v in fila_datos])
            colores_fila_matriz_fin.append(_COLOR_HEX_POR_ZONA_MATRIZ.get(fila_datos[3], "FFFFFF"))
        tabla_matriz_fin = Table(
            datos_matriz_fin, hAlign="LEFT",
            colWidths=[4.0 * cm, 1.0 * cm, 1.0 * cm, 2.2 * cm, 3.4 * cm, 2.5 * cm, 2.9 * cm],
        )
        estilo_matriz_fin = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]
        for i, color_zona in enumerate(colores_fila_matriz_fin, start=1):
            estilo_matriz_fin.append(("BACKGROUND", (3, i), (3, i), colors.HexColor(f"#{color_zona}")))
        tabla_matriz_fin.setStyle(TableStyle(estilo_matriz_fin))
        elementos.append(tabla_matriz_fin)
        elementos.append(Spacer(1, 8))

    elementos.append(Paragraph(ANALISIS_FINANZAS_PUBLICAS_CIERRE, estilo_normal))
    elementos.append(PageBreak())

    elementos.append(Paragraph("6. Recomendaciones oficiales de Función Pública", estilo_h2))
    if cruce_recomendaciones:
        for codigo, lista_recos in cruce_recomendaciones.items():
            elementos.append(Paragraph(f"<b>Brecha {codigo}:</b>", estilo_normal))
            for texto_reco in lista_recos:
                elementos.append(Paragraph(f"• {texto_reco}", estilo_normal))
    else:
        elementos.append(Paragraph(
            "No se cargó un consolidado de recomendaciones oficiales para esta "
            "entidad en esta sesión.", estilo_normal,
        ))
    elementos.append(PageBreak())

    # 6.1 Plan de Mejoramiento (formato oficial de Función Pública)
    elementos.append(Paragraph("6.1 Plan de Mejoramiento — formato oficial de Función Pública", estilo_h2))
    elementos.append(Paragraph(
        "Estructurado en el formato exacto que usan las entidades públicas colombianas "
        "para su Plan de Mejoramiento institucional, consolidando TODAS las brechas "
        "detectadas — no una selección.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 6))
    datos_pm = [[_celda_pdf("Ítem", encabezado=True), _celda_pdf("Observación/Hallazgo", encabezado=True), _celda_pdf("Acción correctiva", encabezado=True), _celda_pdf("Plazo", encabezado=True), _celda_pdf("Responsable", encabezado=True)]]
    for i, b in enumerate(diag.brechas, start=1):
        enfoque_pm, norma_pm = _enfoque_y_norma_de_politica(b.politica)
        plazo_pm = "0-6 meses" if b.puntaje < 20 else ("6-18 meses" if b.puntaje < 45 else "18-36 meses")
        datos_pm.append([
            _celda_pdf(str(i)),
            _celda_pdf(f"{b.codigo_indice} ({b.puntaje}) — {b.nombre_indice}"),
            _celda_pdf(f"Acción alineada con {enfoque_pm} ({norma_pm})"),
            _celda_pdf(plazo_pm),
            _celda_pdf("Líder de proceso / OCI"),
        ])
    if len(datos_pm) > 1:
        tabla_pm = Table(datos_pm, hAlign="LEFT", colWidths=[1.2 * cm, 5.5 * cm, 5.5 * cm, 2.3 * cm, 2.5 * cm])
        tabla_pm.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        elementos.append(tabla_pm)
    elementos.append(Spacer(1, 10))
    elementos.append(Paragraph("Indicadores de efectividad del Plan de Mejoramiento (desde resultados)", estilo_h2))
    for texto_ind in [
        "% de brechas cerradas efectivamente (puntaje ≥ 60 en la siguiente medición FURAG) sobre el total de brechas del plan.",
        "Variación real del IDI-MIPG oficial entre esta vigencia y la siguiente (puntos ganados/perdidos).",
        "Variación del Índice de Desempeño Fiscal (IDF) entre vigencias, como indicador de resultado financiero.",
        "Número de hallazgos de auditoría reincidentes de una vigencia a otra (debe tender a cero).",
        "Tiempo promedio real de cierre de cada acción correctiva frente al plazo sugerido.",
    ]:
        elementos.append(Paragraph(f"• {texto_ind}", estilo_normal))
    elementos.append(PageBreak())

    elementos.extend(_divisor_seccion_pdf("7. Fortalezas y debilidades identificadas", icono="💪"))
    fortalezas = [r for r in diag.resultados_por_dimension if (r.promedio or 0) >= 60]
    debilidades = [r for r in diag.resultados_por_dimension if (r.promedio or 0) < 60]
    datos_tabla = [[_celda_pdf("Dimensión", encabezado=True), "Promedio", "Clasificación"]]
    for r in fortalezas:
        datos_tabla.append([_celda_pdf(f"{r.codigo} {r.nombre}"), str(r.promedio), "Fortaleza"])
    for r in debilidades:
        datos_tabla.append([_celda_pdf(f"{r.codigo} {r.nombre}"), str(r.promedio), "Debilidad"])
    if len(datos_tabla) > 1:
        tabla = Table(datos_tabla, hAlign="LEFT", colWidths=[8 * cm, 3 * cm, 4 * cm])
        estilo_tabla_fd = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F4F8")]),
        ]
        for indice_fila_fd, r in enumerate(fortalezas + debilidades, start=1):
            color_quintil_fd = _color_hex_quintil_mipg(r.promedio)
            if color_quintil_fd:
                estilo_tabla_fd.append(("BACKGROUND", (1, indice_fila_fd), (1, indice_fila_fd), colors.HexColor(f"#{color_quintil_fd}")))
        tabla.setStyle(TableStyle(estilo_tabla_fd))
        elementos.append(tabla)
    else:
        elementos.append(Paragraph("No hay dimensiones evaluadas con los datos disponibles.", estilo_normal))
    elementos.append(Spacer(1, 10))

    elementos.append(PageBreak())
    elementos.extend(_divisor_seccion_pdf("8. " + ESCENARIOS_PROSPECTIVOS_TITULO, icono="🔮"))
    elementos.append(Paragraph(ESCENARIOS_PROSPECTIVOS_INTRO, estilo_normal))
    elementos.append(Spacer(1, 6))
    for titulo_esc, texto_esc in _construir_escenarios_prospectivos(nombre_entidad, diag, resultado_360):
        elementos.append(Paragraph(f"<b>{titulo_esc}.</b> {texto_esc}", estilo_normal))
        elementos.append(Spacer(1, 8))

    elementos.append(Paragraph("9. Nota de uso responsable de inteligencia artificial", estilo_h2))
    elementos.append(Paragraph(
        "Este estudio de caso usó inteligencia artificial (Claude, Anthropic) como "
        "herramienta de apoyo para estructurar el análisis teórico a partir de datos "
        "reales. La IA no sustituye el juicio académico del estudiante: los datos "
        "institucionales, normativos y de recomendaciones son reales y verificables; "
        "la interpretación teórica generada debe ser revisada, discutida y "
        "complementada por el estudiante antes de su entrega.",
        estilo_cursiva,
    ))

    _agregar_marco_descentralizacion_pdf(elementos, estilos, estilo_normal, estilo_h2, diag=diag, nombre_entidad=nombre_entidad)

    doc.build(elementos)
    buffer.seek(0)
    return buffer
