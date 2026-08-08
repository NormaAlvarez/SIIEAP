"""Generador del Informe Ejecutivo para Representantes Legales — Word (.docx)
y PDF descargables.

Tercer tipo de informe de SIIEAP, pensado para el representante legal de la
entidad (alcalde, gobernador, contralor, personero, gerente de ESE, gerente
de entidad descentralizada, rector de universidad pública, gerente de
empresa industrial y comercial del Estado) — NO para el estudiante ni para
el técnico de la oficina de planeación. Por eso:

  - Usa lenguaje llano, directo y sin jerga académica (nada de "Nuevo
    Institucionalismo", "post-NGP", etc. — eso vive en el Informe Técnico
    y en el Estudio de Caso, no aquí).
  - Es tan extenso como el Informe Técnico, pero prioriza SIEMPRE que el
    contenido se entienda de un vistazo: semáforos de color, gráficos,
    cifras grandes, antes que párrafos largos.
  - Muestra las brechas MÁS CRÍTICAS (top, priorizadas por gravedad) para
    que la Alta Dirección sepa dónde mirar primero, pero muestra la
    TOTALIDAD de las recomendaciones oficiales de la Función Pública, no
    solo las de las brechas prioritarias — la Alta Dirección tiene derecho
    a ver el banco completo, no un recorte.
  - Mantiene la misma disciplina de trazabilidad de los otros dos informes:
    dejar clarísimo qué es dato oficial de Función Pública y qué es cálculo
    exclusivo de SIIEAP.

Reutiliza utilidades ya construidas en generador_informe.py (colores de
riesgo, formato de tablas, fecha en español) para mantener consistencia
visual entre los tres informes.
"""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image,
    HRFlowable, KeepTogether,
)

from backend.motores.graficas_informe import generar_grafica_dimensiones, generar_grafica_brechas, generar_matriz_riesgo_probabilidad_impacto
from backend.motores.generador_informe import (
    _fecha_hoy_es,
    _color_hex_riesgo,
    _sombrear_celda,
    _ajustar_tabla_docx,
    _enfoque_y_norma_de_politica,
    _quintil_mipg,
    DESCARGO_RESPONSABILIDAD_AMPLIADO,
    _agregar_logos_docx,
    _logos_pdf_flowables,
    ARTICULO_6_CONSTITUCION_TITULO,
    ARTICULO_6_CONSTITUCION_TEXTO,
    ARTICULO_6_CONSTITUCION_EXPLICACION,
    _riesgo_alta_direccion_de_politica,
    _celda_pdf,
    _agregar_nota_regimen_especial_docx,
    _nota_regimen_especial_pdf_flowables,
    _agregar_tabla_contenido_docx,
    _agregar_razon_de_ser_docx,
    _toc_pdf_flowables,
    _razon_de_ser_pdf_flowables,
    _agregar_marco_descentralizacion_docx,
    _agregar_marco_descentralizacion_pdf,
    _agregar_glosario_docx,
    _agregar_glosario_pdf,
    _agregar_normativa_politicas_docx,
    _agregar_normativa_politicas_pdf,
)


# ---------------------------------------------------------------------------
# Textos fijos en lenguaje llano (sin jerga académica)
# ---------------------------------------------------------------------------

MENSAJE_APERTURA = (
    "Este informe resume, en un lenguaje directo y sin tecnicismos, cómo está "
    "funcionando la gestión de la entidad según la medición oficial del "
    "Gobierno Nacional (el Índice de Desempeño Institucional — IDI). El "
    "objetivo es que usted, como máxima autoridad de la entidad, pueda tomar "
    "decisiones informadas sin necesidad de leer el informe técnico completo. "
    "Al final encontrará la lista completa de recomendaciones oficiales que "
    "la Función Pública ya le entregó a la entidad, para que su equipo las "
    "convierta en acciones concretas."
)

NOTA_SEMAFORO = (
    "El semáforo de color que usa este informe sigue el esquema oficial de 5 quintiles del "
    "MIPG (0-20 crítico, 21-40 bajo, 41-60 medio, 61-80 medio-alto, 81-100 consolidación — "
    "este último es el proceso esperado). El dato oficial en sentido estricto es siempre el "
    "número (el IDI y los puntajes de cada índice); el color es una ayuda visual para leerlo "
    "más rápido."
)

NOTA_FINAL_LLANA = (
    "Este informe fue generado por SIIEAP a partir de datos reales del Índice "
    "de Desempeño Institucional (IDI-MIPG) que publica la Función Pública. Las "
    "cifras oficiales son siempre las que reporta la Función Pública; el "
    "semáforo de 5 quintiles y el listado de \"brechas más críticas\" son una "
    "lectura propia de este sistema para facilitar la toma de decisiones, y "
    "no reemplazan el criterio técnico, jurídico ni la validación del líder "
    "de proceso de la entidad."
)


def _top_brechas(diag, top_n: int = 15):
    """Las N brechas más críticas (menor puntaje primero), para la sección
    de lectura rápida. La sección de recomendaciones, en cambio, usa TODAS
    las brechas — esta función solo decide qué se destaca visualmente."""
    return sorted(diag.brechas, key=lambda b: b.puntaje)[:top_n]


def _mensaje_ejecutivo_riesgo(nombre_quintil_global: str, nombre_entidad: str) -> str:
    """Mensaje de una línea, en lenguaje llano, según el quintil MIPG global
    de la entidad (CRÍTICO / BAJO / MEDIO / MEDIO-ALTO / CONSOLIDACIÓN)."""
    nivel = str(nombre_quintil_global or "").strip().upper()
    if nivel in ("CRÍTICO", "BAJO"):
        return (
            f"{nombre_entidad} requiere atención prioritaria de la Alta Dirección: "
            "varias dimensiones de la gestión están por debajo de lo esperado."
        )
    if nivel in ("MEDIO", "MEDIO-ALTO"):
        return (
            f"{nombre_entidad} tiene un desempeño intermedio: hay avances reales, "
            "pero persisten brechas que conviene cerrar en el corto plazo."
        )
    return (
        f"{nombre_entidad} se ubica en el quintil de Consolidación: un desempeño "
        "favorable frente a la medición oficial, con oportunidades puntuales de mejora."
    )


# ---------------------------------------------------------------------------
# Helpers de diagramación — banda de color institucional (#1F3864), franjas
# alternas en tablas, y encabezados de sección con icono, para dar al
# Informe Ejecutivo el mismo nivel de acabado visual que un documento
# institucional de alto nivel del Estado.
# ---------------------------------------------------------------------------

COLOR_INSTITUCIONAL = "1F3864"
COLOR_INSTITUCIONAL_CLARO = "D9E2F3"
COLOR_FRANJA_ALTERNA = "F2F5FB"

ICONO_SECCION = {
    1: "📊",
    2: "🎯",
    3: "📈",
    4: "📋",
}


def _linea_divisora_docx(doc, color_hex: str = COLOR_INSTITUCIONAL, alto_pt: float = 2.2) -> None:
    """Inserta una línea horizontal de color a todo el ancho, usando el borde
    inferior de un párrafo vacío (más confiable en Word que una tabla de
    1 fila). Sirve como divisor visual bajo los títulos de sección."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._p.get_or_add_pPr()
    borde_inferior = OxmlElement("w:pBdr")
    borde = OxmlElement("w:bottom")
    borde.set(qn("w:val"), "single")
    borde.set(qn("w:sz"), str(int(alto_pt * 8)))
    borde.set(qn("w:space"), "1")
    borde.set(qn("w:color"), color_hex)
    borde_inferior.append(borde)
    p_pr.append(borde_inferior)


def _titulo_seccion_docx(doc, numero: int, texto: str) -> None:
    """Título de sección consistente: icono + texto en el color institucional,
    seguido de una línea divisora de color — reemplaza los doc.add_heading()
    sueltos para que las 4 secciones del informe se vean como partes de un
    mismo documento diseñado, no como bloques de texto pegados uno tras otro."""
    icono = ICONO_SECCION.get(numero, "▪")
    encabezado = doc.add_heading(level=1)
    run = encabezado.add_run(f"{icono}  {numero}. {texto}")
    run.font.color.rgb = RGBColor.from_string(COLOR_INSTITUCIONAL)
    _linea_divisora_docx(doc)


def _bandear_filas_docx(tabla, fila_inicio: int = 1) -> None:
    """Sombrea en franjas alternas (banded rows) las filas de datos de una
    tabla, a partir de fila_inicio (por defecto se salta el encabezado).
    Mejora sustancialmente la legibilidad de tablas largas."""
    for i in range(fila_inicio, len(tabla.rows)):
        if (i - fila_inicio) % 2 == 1:
            for celda in tabla.rows[i].cells:
                # No sobrescribir celdas que ya tengan un color de semáforo propio
                tc_pr = celda._tc.get_or_add_tcPr()
                ya_sombreada = tc_pr.find(qn("w:shd")) is not None
                if not ya_sombreada:
                    _sombrear_celda(celda, COLOR_FRANJA_ALTERNA)


def _fila_alterna_pdf(indice_datos: int) -> tuple:
    """Color de fondo para franjas alternas en tablas PDF (reportlab),
    dado el índice de fila de datos (0 = primera fila de datos)."""
    return colors.HexColor(f"#{COLOR_FRANJA_ALTERNA}") if indice_datos % 2 == 1 else colors.white


def _dibujar_banda_portada_pdf(nombre_entidad: str, aplica_mipg_integral: bool = True):
    """Devuelve una función callback compatible con onFirstPage de reportlab
    que dibuja, solo en la primera página, una banda de color institucional a
    todo el ancho con el título en blanco — el mismo recurso visual que usan
    los documentos oficiales de alto nivel (MinTIC, DNP, Función Pública)
    para anclar la identidad del documento desde el primer vistazo."""
    subtitulo_banner = (
        "Resultado de la medición oficial del Índice de Desempeño Institucional (IDI-MIPG)"
        if aplica_mipg_integral
        else "Resultado de la medición oficial del Índice de Control Interno (MECI)"
    )

    def _callback(canvas_pdf, doc_pdf):
        ancho_pagina, alto_pagina = LETTER
        canvas_pdf.saveState()
        canvas_pdf.setFillColor(colors.HexColor(f"#{COLOR_INSTITUCIONAL}"))
        canvas_pdf.rect(0, alto_pagina - 3.4 * cm, ancho_pagina, 3.4 * cm, fill=1, stroke=0)
        canvas_pdf.setFillColor(colors.HexColor("#E8B84B"))
        canvas_pdf.rect(0, alto_pagina - 3.5 * cm, ancho_pagina, 0.1 * cm, fill=1, stroke=0)
        canvas_pdf.setFillColor(colors.white)
        canvas_pdf.setFont("Helvetica-Bold", 19)
        canvas_pdf.drawCentredString(ancho_pagina / 2, alto_pagina - 1.6 * cm, "INFORME EJECUTIVO PARA LA ALTA DIRECCIÓN")
        canvas_pdf.setFont("Helvetica-Oblique", 11)
        canvas_pdf.setFillColor(colors.HexColor("#E8B84B"))
        canvas_pdf.drawCentredString(
            ancho_pagina / 2, alto_pagina - 2.3 * cm,
            subtitulo_banner,
        )
        # CORREGIDO: el nombre de la entidad es de longitud variable (algunos
        # nombres de municipios/departamentos son largos); antes se dibujaba
        # con drawCentredString a tamaño fijo y podía desbordarse. Ahora el
        # tamaño de fuente se reduce automáticamente si el nombre es largo.
        canvas_pdf.setFillColor(colors.white)
        tamano_nombre = 9 if len(nombre_entidad) <= 55 else 7.5
        canvas_pdf.setFont("Helvetica", tamano_nombre)
        canvas_pdf.drawCentredString(ancho_pagina / 2, alto_pagina - 2.95 * cm, nombre_entidad)
        canvas_pdf.restoreState()
        _pie_de_pagina_alcaldes(canvas_pdf, doc_pdf, nombre_entidad)
    return _callback


def _pie_de_pagina_alcaldes(canvas_pdf, doc_pdf, nombre_entidad: str) -> None:
    canvas_pdf.saveState()
    canvas_pdf.setFont("Helvetica", 8)
    canvas_pdf.setFillColor(colors.grey)
    canvas_pdf.drawString(2 * cm, 1.3 * cm, f"SIIEAP — Informe Ejecutivo — {nombre_entidad}")
    canvas_pdf.drawRightString(LETTER[0] - 2 * cm, 1.3 * cm, f"Página {doc_pdf.page}")
    canvas_pdf.restoreState()


def _divisor_seccion_pdf(numero: int, texto: str, estilos) -> list:
    """Título de sección + línea divisora de color institucional, análogo a
    _titulo_seccion_docx, para que las 4 secciones se vean como partes de un
    mismo documento diseñado en ambos formatos (Word y PDF)."""
    icono = ICONO_SECCION.get(numero, "▪")
    estilo_titulo_seccion = ParagraphStyle(
        f"TituloSeccion{numero}", parent=estilos["Heading1"],
        textColor=colors.HexColor(f"#{COLOR_INSTITUCIONAL}"), spaceAfter=4,
    )
    return [
        Paragraph(f"{icono}&nbsp;&nbsp;{numero}. {texto}", estilo_titulo_seccion),
        HRFlowable(width="100%", thickness=2.2, color=colors.HexColor(f"#{COLOR_INSTITUCIONAL}"), spaceAfter=10),
    ]


# ---------------------------------------------------------------------------
# Versión Word (.docx)
# ---------------------------------------------------------------------------

def generar_informe_alcaldes_docx(
    nombre_entidad,
    diag,
    resultado_isvpt=None,
    resultado_360=None,
    idi_oficial=None,
    cruce_recomendaciones=None,
    total_recomendaciones_entidad=None,
    top_n_brechas: int = 15,
    tipo_regimen_especial=None,
):
    """Genera el Informe Ejecutivo para Alcaldes/Gobernadores en Word.

    Parámetros: mismo patrón que generar_reporte_docx (generador_informe.py),
    sin analisis_ia_texto porque este informe no incluye el análisis
    académico con IA — solo cifras, semáforos, gráficos y recomendaciones.
    """
    doc = Document()

    idi_protagonista = idi_oficial if idi_oficial is not None else diag.idi_estimado
    color_global, emoji_global, texto_global = _quintil_mipg(idi_protagonista)

    total_recomendaciones = (
        total_recomendaciones_entidad
        if total_recomendaciones_entidad is not None
        else (sum(len(lista) for lista in cruce_recomendaciones.values()) if cruce_recomendaciones else None)
    )

    # --- Portada: banda de color institucional a todo el ancho, como en un
    # documento oficial de alto nivel, en vez de solo texto centrado ---
    _agregar_logos_docx(doc)
    banner = doc.add_table(rows=1, cols=1)
    banner.autofit = False
    celda_banner = banner.rows[0].cells[0]
    celda_banner.width = Cm(17)
    _sombrear_celda(celda_banner, COLOR_INSTITUCIONAL)
    celda_banner.vertical_alignment = 1  # centro
    p_banner_titulo = celda_banner.paragraphs[0]
    p_banner_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_banner_titulo.paragraph_format.space_before = Pt(14)
    p_banner_titulo.paragraph_format.space_after = Pt(2)
    run_banner_titulo = p_banner_titulo.add_run("INFORME EJECUTIVO PARA LA ALTA DIRECCIÓN")
    run_banner_titulo.bold = True
    run_banner_titulo.font.size = Pt(20)
    run_banner_titulo.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p_banner_sub = celda_banner.add_paragraph()
    p_banner_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_banner_sub.paragraph_format.space_after = Pt(14)
    run_banner_sub = p_banner_sub.add_run(
        "Resultado de la medición oficial del Índice de Desempeño Institucional (IDI-MIPG)"
        if diag.aplica_mipg_integral
        else "Resultado de la medición oficial del Índice de Control Interno (MECI)"
    )
    run_banner_sub.italic = True
    run_banner_sub.font.size = Pt(12)
    run_banner_sub.font.color.rgb = RGBColor(0xE8, 0xB8, 0x4B)
    for fila_banner in banner.rows:
        tr_pr = fila_banner._tr.get_or_add_trPr()
        # sin acción adicional: la altura se ajusta por el contenido

    doc.add_paragraph()
    _agregar_nota_regimen_especial_docx(doc, tipo_regimen_especial)

    p_ent = doc.add_paragraph()
    p_ent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ent = p_ent.add_run(nombre_entidad)
    run_ent.bold = True
    run_ent.font.size = Pt(20)
    run_ent.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fecha.add_run(f"Generado el {_fecha_hoy_es()} · Decreto 1499 de 2017")

    doc.add_paragraph()

    # Tabla grande de cifras clave (IDI + semáforo + brechas + recomendaciones)
    n_columnas = 4 if total_recomendaciones is not None else 3
    tabla_cifras = doc.add_table(rows=2, cols=n_columnas)
    tabla_cifras.style = "Light Grid Accent 1"
    tabla_cifras.alignment = WD_TABLE_ALIGNMENT.CENTER
    encabezados = [("IDI oficial" if diag.aplica_mipg_integral else "Índice Control Interno"), f"Nivel de riesgo {emoji_global}", "Brechas más críticas detectadas"]
    valores = [str(idi_protagonista), texto_global, str(len(diag.brechas))]
    if total_recomendaciones is not None:
        encabezados.append("Recomendaciones oficiales de Función Pública")
        valores.append(str(total_recomendaciones))
    for celda, texto in zip(tabla_cifras.rows[0].cells, encabezados):
        celda.text = texto
        for parrafo in celda.paragraphs:
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run_c in parrafo.runs:
                run_c.bold = True
                run_c.font.size = Pt(10)
    for celda, texto in zip(tabla_cifras.rows[1].cells, valores):
        celda.text = texto
        for parrafo in celda.paragraphs:
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run_c in parrafo.runs:
                run_c.bold = True
                run_c.font.size = Pt(18)
    ancho_celda = Cm(4.2) if n_columnas == 4 else Cm(5.5)
    _ajustar_tabla_docx(tabla_cifras, anchos_cm=[ancho_celda.cm] * n_columnas, tamano_fuente_pt=10)
    _sombrear_celda(tabla_cifras.rows[1].cells[1], color_global)

    doc.add_paragraph()
    p_msg = doc.add_paragraph()
    run_msg = p_msg.add_run(_mensaje_ejecutivo_riesgo(texto_global, nombre_entidad))
    run_msg.bold = True
    run_msg.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph(MENSAJE_APERTURA)

    p_nota_sem = doc.add_paragraph()
    run_nota_sem = p_nota_sem.add_run(NOTA_SEMAFORO)
    run_nota_sem.italic = True
    run_nota_sem.font.size = Pt(9)
    run_nota_sem.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    _agregar_tabla_contenido_docx(doc)
    _agregar_razon_de_ser_docx(doc, "ejecutivo")

    _agregar_glosario_docx(doc)
    _agregar_normativa_politicas_docx(doc, diag=diag)

    # --- 1. Cómo está la entidad, dimensión por dimensión (semáforo) ---
    _titulo_seccion_docx(doc, 1, "Cómo está la entidad, dimensión por dimensión")
    doc.add_paragraph(
        "Cada dimensión agrupa un conjunto de temas de la gestión pública (talento "
        "humano, planeación, contratación, servicio al ciudadano, entre otras). El "
        "color indica qué tan urgente es la atención de la Alta Dirección en esa "
        "dimensión."
    )
    tabla_dim = doc.add_table(rows=1, cols=3)
    tabla_dim.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_dim.rows[0].cells, ["Dimensión", "Puntaje (sobre 100)", "Semáforo"]):
        celda.text = texto
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
    for r in diag.resultados_por_dimension:
        fila = tabla_dim.add_row().cells
        fila[0].text = f"{r.nombre}"
        fila[1].text = str(r.promedio)
        color_dim, emoji_dim, texto_dim = _quintil_mipg(r.promedio)
        fila[2].text = f"{emoji_dim} {texto_dim}"
        _sombrear_celda(fila[2], color_dim)
    _bandear_filas_docx(tabla_dim)
    _ajustar_tabla_docx(tabla_dim, anchos_cm=[9.0, 4.0, 4.0], tamano_fuente_pt=10.5)

    doc.add_paragraph()
    try:
        buffer_grafica_dim = generar_grafica_dimensiones(diag)
        doc.add_picture(buffer_grafica_dim, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    doc.add_paragraph()

    # --- 2. Las brechas más críticas (top, priorizadas) ---
    _titulo_seccion_docx(doc, 2, "Los puntos que más necesitan atención ahora")
    doc.add_paragraph(
        f"De todas las brechas detectadas por SIIEAP, estas son las {min(top_n_brechas, len(diag.brechas))} "
        "más críticas — es decir, las de puntaje más bajo. Son el punto de partida "
        "recomendado para el plan de acción de la entidad."
    )
    top_brechas = _top_brechas(diag, top_n_brechas)
    if top_brechas:
        try:
            buffer_grafica_brechas = generar_grafica_brechas(diag)
            if buffer_grafica_brechas:
                doc.add_picture(buffer_grafica_brechas, width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
        except Exception:
            pass
        tabla_top = doc.add_table(rows=1, cols=3)
        tabla_top.style = "Light Grid Accent 1"
        for celda, texto in zip(tabla_top.rows[0].cells, ["Tema afectado", "Puntaje", "Área de gestión responsable"]):
            celda.text = texto
            for parrafo in celda.paragraphs:
                for run_enc in parrafo.runs:
                    run_enc.bold = True
        for b in top_brechas:
            fila = tabla_top.add_row().cells
            fila[0].text = b.nombre_indice
            fila[1].text = str(b.puntaje)
            fila[2].text = b.politica
        _bandear_filas_docx(tabla_top)
        _ajustar_tabla_docx(tabla_top, anchos_cm=[8.0, 2.5, 6.5], tamano_fuente_pt=9.5)
    else:
        doc.add_paragraph("Con los datos disponibles, no se detectaron brechas críticas.")

    doc.add_paragraph()

    # --- 3. Comparación con entidades similares (si hay datos) ---
    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        _titulo_seccion_docx(doc, 3, "Cómo se compara la entidad frente a otras similares")
        brecha_grupo = round((idi_protagonista or 0) - resultado_360.promedio_idi, 2)
        comparativo = "por debajo" if brecha_grupo < 0 else "por encima"
        doc.add_paragraph(
            f"Comparando con {resultado_360.n_entidades} entidades de características "
            f"similares ({resultado_360.filtro_descripcion}), {nombre_entidad} está "
            f"{abs(brecha_grupo)} puntos {comparativo} del promedio de ese grupo "
            f"({resultado_360.promedio_idi})."
        )
        if resultado_360.percentil_entidad_referencia is not None:
            doc.add_paragraph(
                f"Esto ubica a la entidad en el percentil {resultado_360.percentil_entidad_referencia}% "
                "de su grupo de comparación (entre más alto el percentil, mejor la posición relativa)."
            )
        doc.add_paragraph()

    # --- 4. Recomendaciones oficiales de la Función Pública (TODAS) ---
    _titulo_seccion_docx(doc, 4, "Recomendaciones oficiales de la Función Pública")
    doc.add_paragraph(
        "A diferencia de la sección anterior (que es una lectura propia de SIIEAP), "
        "esta lista SÍ es información oficial: es la totalidad de las recomendaciones "
        "que el Departamento Administrativo de la Función Pública ya entregó "
        "formalmente a la entidad, sin recortar solo a las más urgentes."
    )
    if cruce_recomendaciones:
        for codigo, lista_recos in cruce_recomendaciones.items():
            p_cod = doc.add_paragraph()
            run_cod = p_cod.add_run(f"Sobre {codigo}:")
            run_cod.bold = True
            for texto_reco in lista_recos:
                doc.add_paragraph(texto_reco, style="List Bullet")
    else:
        doc.add_paragraph(
            "No se cargó el consolidado de recomendaciones oficiales de la entidad en "
            "esta sesión; cárguelo en la barra lateral del sistema para incluir aquí "
            "el listado completo."
        )

    doc.add_paragraph()

    # --- 5. Implicaciones legales, fiscales, administrativas y disciplinarias ---
    _titulo_seccion_docx(doc, 5, "Lo que esto significa en lo legal, lo fiscal, lo administrativo y lo disciplinario")
    doc.add_paragraph(
        "Esta sección traduce cada brecha detectada a su implicación concreta de "
        "responsabilidad para la Alta Dirección — no como una amenaza, sino como el "
        "mapa de riesgo que su equipo jurídico, financiero y de control interno "
        "necesita para actuar a tiempo."
    )

    p_art6_titulo = doc.add_paragraph()
    run_art6_titulo = p_art6_titulo.add_run(ARTICULO_6_CONSTITUCION_TITULO)
    run_art6_titulo.bold = True
    p_art6_texto = doc.add_paragraph()
    run_art6_texto = p_art6_texto.add_run(f"“{ARTICULO_6_CONSTITUCION_TEXTO}”")
    run_art6_texto.italic = True
    doc.add_paragraph(ARTICULO_6_CONSTITUCION_EXPLICACION)
    doc.add_paragraph()

    try:
        buffer_matriz_riesgo_alc = generar_matriz_riesgo_probabilidad_impacto()
        doc.add_picture(buffer_matriz_riesgo_alc, width=Inches(4.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    except Exception:
        pass

    if diag.brechas:
        doc.add_paragraph(
            "Cada brecha priorizada, cruzada de forma sistemática (por código, no por "
            "interpretación de la IA) con el tipo de riesgo que compromete, la norma "
            "que lo respalda y la consecuencia posible más frecuente en entidades "
            "territoriales colombianas:"
        )
        tabla_riesgo_alta_direccion = doc.add_table(rows=1, cols=4)
        tabla_riesgo_alta_direccion.style = "Light Grid Accent 1"
        encabezados_riesgo = tabla_riesgo_alta_direccion.rows[0].cells
        for celda, texto in zip(encabezados_riesgo, ["Política con brecha", "Tipo de riesgo", "Norma principal", "Consecuencia posible"]):
            celda.text = texto
            for parrafo in celda.paragraphs:
                for run_enc in parrafo.runs:
                    run_enc.bold = True
        politicas_ya_vistas = set()
        for b in diag.brechas:
            if b.politica in politicas_ya_vistas:
                continue
            politicas_ya_vistas.add(b.politica)
            tipo_riesgo, norma_riesgo, consecuencia_riesgo = _riesgo_alta_direccion_de_politica(b.politica)
            fila_riesgo = tabla_riesgo_alta_direccion.add_row().cells
            fila_riesgo[0].text = b.politica
            fila_riesgo[1].text = tipo_riesgo
            fila_riesgo[2].text = norma_riesgo
            fila_riesgo[3].text = consecuencia_riesgo
        _ajustar_tabla_docx(tabla_riesgo_alta_direccion, anchos_cm=[4.0, 2.8, 4.5, 5.7], tamano_fuente_pt=8.5)
        _bandear_filas_docx(tabla_riesgo_alta_direccion)
    else:
        doc.add_paragraph(
            "No se detectaron brechas por debajo del umbral con los datos disponibles; "
            "de cualquier forma, se recomienda mantener vigente el sistema de control "
            "interno (Ley 87 de 1993) como primera línea de defensa preventiva."
        )

    doc.add_paragraph()

    # --- 6. Nota metodológica en lenguaje llano ---
    doc.add_heading("Nota metodológica", level=2)
    p_nota = doc.add_paragraph()
    run_nota = p_nota.add_run(NOTA_FINAL_LLANA)
    run_nota.italic = True
    run_nota.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p_descargo = doc.add_paragraph()
    run_descargo = p_descargo.add_run(DESCARGO_RESPONSABILIDAD_AMPLIADO)
    run_descargo.italic = True
    run_descargo.font.size = Pt(9)
    run_descargo.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _agregar_marco_descentralizacion_docx(doc, diag=diag, nombre_entidad=nombre_entidad)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Versión PDF (reportlab)
# ---------------------------------------------------------------------------

def generar_informe_alcaldes_pdf(
    nombre_entidad,
    diag,
    resultado_isvpt=None,
    resultado_360=None,
    idi_oficial=None,
    cruce_recomendaciones=None,
    total_recomendaciones_entidad=None,
    top_n_brechas: int = 15,
    tipo_regimen_especial=None,
):
    """Versión PDF (reportlab) de generar_informe_alcaldes_docx — mismo
    contenido y misma disciplina de trazabilidad oficial/propio."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=LETTER,
        topMargin=4.2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )
    estilos = getSampleStyleSheet()
    estilo_titulo = ParagraphStyle("TituloAlcaldes", parent=estilos["Title"], fontSize=20)
    estilo_h2 = ParagraphStyle("H2Alcaldes", parent=estilos["Heading2"], spaceBefore=14, fontSize=14)
    estilo_normal = ParagraphStyle("NormalAlcaldes", parent=estilos["Normal"], fontSize=11.5, leading=16)
    estilo_cursiva = ParagraphStyle("CursivaAlcaldes", parent=estilo_normal, fontName="Helvetica-Oblique", textColor=colors.grey, fontSize=9)
    estilo_mensaje = ParagraphStyle("MensajeAlcaldes", parent=estilo_normal, fontName="Helvetica-Bold", fontSize=13)

    idi_protagonista = idi_oficial if idi_oficial is not None else diag.idi_estimado
    color_global, emoji_global, texto_global = _quintil_mipg(idi_protagonista)

    total_recomendaciones = (
        total_recomendaciones_entidad
        if total_recomendaciones_entidad is not None
        else (sum(len(lista) for lista in cruce_recomendaciones.values()) if cruce_recomendaciones else None)
    )

    elementos = []

    # --- Portada: el título y subtítulo ya se dibujan en la banda superior
    # (ver _dibujar_banda_portada_pdf); aquí van los logos y el resto del contenido ---
    elementos.extend(_logos_pdf_flowables())
    elementos.extend(_nota_regimen_especial_pdf_flowables(tipo_regimen_especial))
    elementos.append(Paragraph(f"<b>{nombre_entidad}</b>", ParagraphStyle("EntAlcaldes", parent=estilos["Heading1"], textColor=colors.HexColor(f"#{COLOR_INSTITUCIONAL}"))))
    elementos.append(Paragraph(f"Generado el {_fecha_hoy_es()} · Decreto 1499 de 2017", estilo_normal))
    elementos.append(Spacer(1, 14))

    n_columnas = 4 if total_recomendaciones is not None else 3
    encabezados = [("IDI oficial" if diag.aplica_mipg_integral else "Índice Control Interno"), f"Nivel de riesgo {emoji_global}", "Brechas más críticas detectadas"]
    valores = [str(idi_protagonista), texto_global, str(len(diag.brechas))]
    if total_recomendaciones is not None:
        encabezados.append("Recomendaciones oficiales FP")
        valores.append(str(total_recomendaciones))
    datos_tabla = [encabezados, valores]
    anchos = [4.2 * cm] * n_columnas if n_columnas == 4 else [5.5 * cm] * n_columnas
    tabla_cifras = Table(datos_tabla, hAlign="CENTER", colWidths=anchos)
    estilo_tabla = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, 0), 10),
        ("FONTSIZE", (0, 1), (-1, 1), 16),
        ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (1, 1), (1, 1), colors.HexColor(f"#{color_global}")),
    ]
    tabla_cifras.setStyle(TableStyle(estilo_tabla))
    elementos.append(tabla_cifras)
    elementos.append(Spacer(1, 12))

    elementos.append(Paragraph(_mensaje_ejecutivo_riesgo(texto_global, nombre_entidad), estilo_mensaje))
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph(MENSAJE_APERTURA, estilo_normal))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph(NOTA_SEMAFORO, estilo_cursiva))
    elementos.append(PageBreak())

    elementos.extend(_toc_pdf_flowables([
        "Cómo está la entidad, dimensión por dimensión",
        "Brechas más críticas y su implicación legal, fiscal y disciplinaria",
        "Matriz de riesgo probabilidad-impacto",
        "Recomendaciones oficiales de la Función Pública",
        "Artículo 6 constitucional y responsabilidad de la Alta Dirección",
    ]))
    elementos.extend(_razon_de_ser_pdf_flowables("ejecutivo"))

    _agregar_glosario_pdf(elementos, estilos, estilo_normal, estilo_h2)
    _agregar_normativa_politicas_pdf(elementos, estilos, estilo_normal, estilo_h2, diag=diag)

    # --- 1. Semáforo por dimensión ---
    elementos.extend(_divisor_seccion_pdf(1, "Cómo está la entidad, dimensión por dimensión", estilos))
    elementos.append(Paragraph(
        "Cada dimensión agrupa un conjunto de temas de la gestión pública. El color "
        "indica qué tan urgente es la atención de la Alta Dirección en esa dimensión.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 6))
    datos_dim = [["Dimensión", "Puntaje (sobre 100)", "Semáforo"]]
    colores_fila_dim = []
    for r in diag.resultados_por_dimension:
        color_dim, emoji_dim, texto_dim = _quintil_mipg(r.promedio)
        datos_dim.append([r.nombre, str(r.promedio), f"{emoji_dim} {texto_dim}"])
        colores_fila_dim.append(color_dim)
    tabla_dim = Table(datos_dim, hAlign="LEFT", colWidths=[9 * cm, 3.5 * cm, 3.5 * cm])
    estilo_tabla_dim = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    for i, color_dim in enumerate(colores_fila_dim, start=1):
        if i % 2 == 0:
            estilo_tabla_dim.append(("BACKGROUND", (0, i), (1, i), colors.HexColor(f"#{COLOR_FRANJA_ALTERNA}")))
        estilo_tabla_dim.append(("BACKGROUND", (2, i), (2, i), colors.HexColor(f"#{color_dim}")))
    tabla_dim.setStyle(TableStyle(estilo_tabla_dim))
    elementos.append(tabla_dim)
    elementos.append(Spacer(1, 10))

    try:
        buffer_grafica_dim = generar_grafica_dimensiones(diag)
        elementos.append(Image(buffer_grafica_dim, width=16 * cm, height=16 * cm * 0.5))
        elementos.append(Spacer(1, 10))
    except Exception:
        pass

    # --- 2. Top brechas críticas ---
    elementos.extend(_divisor_seccion_pdf(2, "Los puntos que más necesitan atención ahora", estilos))
    top_brechas = _top_brechas(diag, top_n_brechas)
    elementos.append(Paragraph(
        f"De todas las brechas detectadas por SIIEAP, estas son las {min(top_n_brechas, len(diag.brechas))} "
        "más críticas — el punto de partida recomendado para el plan de acción.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 6))
    if top_brechas:
        try:
            buffer_grafica_brechas = generar_grafica_brechas(diag)
            if buffer_grafica_brechas:
                elementos.append(Image(buffer_grafica_brechas, width=16 * cm, height=16 * cm * 0.55))
                elementos.append(Spacer(1, 10))
        except Exception:
            pass
        datos_top = [["Tema afectado", "Puntaje", "Área de gestión responsable"]]
        for b in top_brechas:
            datos_top.append([b.nombre_indice, str(b.puntaje), b.politica])
        tabla_top = Table(datos_top, hAlign="LEFT", colWidths=[8 * cm, 2.5 * cm, 6.5 * cm])
        estilo_tabla_top = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]
        for i in range(1, len(datos_top)):
            if i % 2 == 0:
                estilo_tabla_top.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor(f"#{COLOR_FRANJA_ALTERNA}")))
        tabla_top.setStyle(TableStyle(estilo_tabla_top))
        elementos.append(tabla_top)
    else:
        elementos.append(Paragraph("Con los datos disponibles, no se detectaron brechas críticas.", estilo_normal))
    elementos.append(Spacer(1, 10))

    # --- 3. Comparación con entidades similares ---
    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        elementos.extend(_divisor_seccion_pdf(3, "Cómo se compara la entidad frente a otras similares", estilos))
        brecha_grupo = round((idi_protagonista or 0) - resultado_360.promedio_idi, 2)
        comparativo = "por debajo" if brecha_grupo < 0 else "por encima"
        elementos.append(Paragraph(
            f"Comparando con {resultado_360.n_entidades} entidades similares "
            f"({resultado_360.filtro_descripcion}), {nombre_entidad} está "
            f"{abs(brecha_grupo)} puntos {comparativo} del promedio de ese grupo "
            f"({resultado_360.promedio_idi}).",
            estilo_normal,
        ))
        if resultado_360.percentil_entidad_referencia is not None:
            elementos.append(Paragraph(
                f"Esto ubica a la entidad en el percentil {resultado_360.percentil_entidad_referencia}% "
                "de su grupo de comparación.",
                estilo_normal,
            ))
        elementos.append(Spacer(1, 10))

    # --- 4. TODAS las recomendaciones oficiales ---
    elementos.extend(_divisor_seccion_pdf(4, "Recomendaciones oficiales de la Función Pública", estilos))
    elementos.append(Paragraph(
        "A diferencia de la sección anterior (lectura propia de SIIEAP), esta lista SÍ "
        "es información oficial: la totalidad de las recomendaciones que la Función "
        "Pública ya entregó formalmente a la entidad, sin recortar solo a las más urgentes.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 6))
    if cruce_recomendaciones:
        for codigo, lista_recos in cruce_recomendaciones.items():
            elementos.append(Paragraph(f"<b>Sobre {codigo}:</b>", estilo_normal))
            for texto_reco in lista_recos:
                elementos.append(Paragraph(f"• {texto_reco}", estilo_normal))
            elementos.append(Spacer(1, 4))
    else:
        elementos.append(Paragraph(
            "No se cargó el consolidado de recomendaciones oficiales de la entidad en "
            "esta sesión; cárguelo en la barra lateral del sistema para incluir aquí "
            "el listado completo.",
            estilo_normal,
        ))
    elementos.append(Spacer(1, 10))

    # --- 5. Implicaciones legales, fiscales, administrativas y disciplinarias ---
    elementos.extend(_divisor_seccion_pdf(5, "Lo que esto significa en lo legal, lo fiscal, lo administrativo y lo disciplinario", estilos))
    elementos.append(Paragraph(
        "Esta sección traduce cada brecha detectada a su implicación concreta de "
        "responsabilidad para la Alta Dirección — no como una amenaza, sino como el "
        "mapa de riesgo que su equipo jurídico, financiero y de control interno "
        "necesita para actuar a tiempo.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph(f"<b>{ARTICULO_6_CONSTITUCION_TITULO}</b>", estilo_normal))
    elementos.append(Paragraph(f"“{ARTICULO_6_CONSTITUCION_TEXTO}”", estilo_cursiva))
    elementos.append(Spacer(1, 4))
    elementos.append(Paragraph(ARTICULO_6_CONSTITUCION_EXPLICACION, estilo_normal))
    elementos.append(Spacer(1, 8))

    try:
        buffer_matriz_riesgo_alc_pdf = generar_matriz_riesgo_probabilidad_impacto()
        elementos.append(Image(buffer_matriz_riesgo_alc_pdf, width=11 * cm, height=11 * cm))
        elementos.append(Spacer(1, 8))
    except Exception:
        pass

    if diag.brechas:
        elementos.append(Paragraph(
            "Cada brecha priorizada, cruzada de forma sistemática (por código, no por "
            "interpretación de la IA) con el tipo de riesgo que compromete, la norma que "
            "lo respalda y la consecuencia posible más frecuente en entidades "
            "territoriales colombianas:",
            estilo_normal,
        ))
        elementos.append(Spacer(1, 6))
        datos_riesgo_alc = [[_celda_pdf("Política con brecha", encabezado=True), _celda_pdf("Tipo de riesgo", encabezado=True), _celda_pdf("Norma principal", encabezado=True), _celda_pdf("Consecuencia posible", encabezado=True)]]
        politicas_ya_vistas_pdf = set()
        for b in diag.brechas:
            if b.politica in politicas_ya_vistas_pdf:
                continue
            politicas_ya_vistas_pdf.add(b.politica)
            tipo_riesgo, norma_riesgo, consecuencia_riesgo = _riesgo_alta_direccion_de_politica(b.politica)
            datos_riesgo_alc.append([_celda_pdf(b.politica), _celda_pdf(tipo_riesgo), _celda_pdf(norma_riesgo), _celda_pdf(consecuencia_riesgo)])
        tabla_riesgo_alc_pdf = Table(datos_riesgo_alc, hAlign="LEFT", colWidths=[3.8 * cm, 2.6 * cm, 4.6 * cm, 6 * cm])
        tabla_riesgo_alc_pdf.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{COLOR_INSTITUCIONAL}")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7.5),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(f"#{COLOR_FRANJA_ALTERNA}")]),
        ]))
        elementos.append(tabla_riesgo_alc_pdf)
    else:
        elementos.append(Paragraph(
            "No se detectaron brechas por debajo del umbral con los datos disponibles; de "
            "cualquier forma, se recomienda mantener vigente el sistema de control interno "
            "(Ley 87 de 1993) como primera línea de defensa preventiva.",
            estilo_normal,
        ))
    elementos.append(Spacer(1, 10))

    # --- Nota metodológica ---
    elementos.append(Paragraph("Nota metodológica", estilos["Heading2"]))
    elementos.append(Paragraph(NOTA_FINAL_LLANA, estilo_cursiva))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph(DESCARGO_RESPONSABILIDAD_AMPLIADO, ParagraphStyle("DescargoAmpliadoAlc", parent=estilo_cursiva, fontSize=8.5)))

    _agregar_marco_descentralizacion_pdf(elementos, estilos, estilo_normal, estilo_h2, diag=diag, nombre_entidad=nombre_entidad)

    doc.build(
        elementos,
        onFirstPage=_dibujar_banda_portada_pdf(nombre_entidad, diag.aplica_mipg_integral),
        onLaterPages=lambda c, d: _pie_de_pagina_alcaldes(c, d, nombre_entidad),
    )
    buffer.seek(0)
    return buffer
