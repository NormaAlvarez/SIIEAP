"""Generador de Informe Técnico — Word (.docx) y PDF descargables.

Empaqueta, para UNA entidad, un informe técnico completo que el docente
puede entregar a sus estudiantes como evidencia de ejecución del sistema
y como fuente de trabajo para la asignatura. El informe integra:

  1. Portada institucional (SIIEAP, entidad, fecha)
  2. Contextualización de la Administración Pública Contemporánea:
     un bloque FIJO (igual en todos los informes) que recorre el arco
     teórico completo del microcurrículo — desde los antecedentes griegos
     de lo público/privado, pasando por el Nuevo Institucionalismo, la
     Nueva Gestión Pública (NGP) y la post-NGP, hasta el esquema
     integrador de la Administración Pública contemporánea (Estado
     Digital, Transformación Digital, Gobernanza Inteligente, IA,
     Gobierno Abierto, Gobernanza de Datos, Valor Público, Administración
     Pública Basada en Evidencia, Capacidades Estatales, Resiliencia
     Institucional, Agenda 2030/ODS, OCDE, CEPAL y Naciones Unidas).
  3. Resultado real del diagnóstico IDI-MIPG de la entidad (dimensiones,
     brechas).
  4. El análisis integral generado por IA (motor_analisis_ia.py): lectura
     desde las tres teorías del curso, recomendaciones técnicas/jurídicas/
     financieras, valoración de riesgo y prospectiva.
  5. Nota de trazabilidad y disclaimer académico.

Este módulo NO llama a la API de Claude — solo da formato a un texto de
análisis que ya fue generado antes por motor_analisis_ia.py.

Dependencias nuevas a agregar en requirements.txt:
    python-docx
    reportlab
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image,
)

from backend.motores.graficas_informe import generar_grafica_dimensiones, generar_grafica_brechas
from backend.base_conocimiento.normativa_politicas import (
    todas_las_politicas_con_normativa,
    normativa_de_politica,
)


_ESTILO_CELDA_TABLA_PDF = ParagraphStyle(
    "CeldaTablaPDF", fontName="Helvetica", fontSize=8, leading=10, wordWrap="CJK",
)
_ESTILO_CELDA_TABLA_PDF_ENCABEZADO = ParagraphStyle(
    "CeldaTablaPDFEncabezado", fontName="Helvetica-Bold", fontSize=8, leading=10,
    textColor=colors.white, wordWrap="CJK",
)


def _celda_pdf(texto, encabezado: bool = False, tamano_fuente: int | None = None):
    """Envuelve el texto de una celda de tabla reportlab en un Paragraph con
    ajuste de línea (wordWrap) explícito. Reportlab NO ajusta de forma
    confiable strings planos dentro de una Table — cuando el texto es más
    largo que el ancho de columna, se desborda visualmente sobre las
    celdas vecinas en vez de partirse en varias líneas. Usar esta función
    en cualquier columna de texto largo/variable (política, norma,
    consecuencia, enfoque, recomendación, etc.) es la forma correcta y
    estable de evitarlo — las columnas puramente numéricas o de 1-2
    palabras cortas (puntaje, código, sí/no) pueden seguir como string."""
    estilo = _ESTILO_CELDA_TABLA_PDF_ENCABEZADO if encabezado else _ESTILO_CELDA_TABLA_PDF
    if tamano_fuente is not None:
        estilo = ParagraphStyle(
            "CeldaTablaPDFAd hoc", parent=estilo, fontSize=tamano_fuente, leading=tamano_fuente + 2,
        )
    texto_escapado = str(texto).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(texto_escapado, estilo)


# ---------------------------------------------------------------------------
# Identidad visual institucional compartida por los 3 informes (Técnico,
# Estudio de Caso, Alcaldes/Gobernadores): logos, banner de portada,
# divisores de sección con ícono, franjas alternas en tablas y el esquema
# de color por quintil oficial del MIPG. Todo lo que vive aquí se importa
# desde los otros dos módulos para no duplicar código ni diseño.
# ---------------------------------------------------------------------------

COLOR_INSTITUCIONAL = "1F3864"  # azul institucional ESAP, usado en banners y divisores

_CARPETA_ASSETS = Path(__file__).resolve().parent.parent.parent / "assets"
_RUTA_LOGO_ESAP = _CARPETA_ASSETS / "logo_esap.png"
_RUTA_LOGO_CERTIFICACIONES = _CARPETA_ASSETS / "certificaciones.png"


def _agregar_logos_docx(doc) -> None:
    """Inserta, si existen los archivos, el logo de la ESAP (izquierda) y el
    de certificaciones ICONTEC/IQNET (derecha) lado a lado en la parte
    superior de la portada. Si la carpeta assets/ no está presente (por
    ejemplo, no se subió al repo de despliegue), el informe se genera igual,
    simplemente sin logos — nunca se rompe la app por un archivo faltante."""
    try:
        if not (_RUTA_LOGO_ESAP.exists() and _RUTA_LOGO_CERTIFICACIONES.exists()):
            return
        tabla_logos = doc.add_table(rows=1, cols=2)
        tabla_logos.autofit = False
        celda_izq, celda_der = tabla_logos.rows[0].cells
        celda_izq.width = Cm(8.5)
        celda_der.width = Cm(8.5)
        p_izq = celda_izq.paragraphs[0]
        p_izq.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_izq.add_run().add_picture(str(_RUTA_LOGO_ESAP), height=Cm(1.6))
        p_der = celda_der.paragraphs[0]
        p_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_der.add_run().add_picture(str(_RUTA_LOGO_CERTIFICACIONES), height=Cm(1.9))
        doc.add_paragraph()
    except Exception:
        # Nunca romper la generación del informe por un problema con los logos
        pass


def _logos_pdf_flowables():
    """Devuelve una lista de flowables (posiblemente vacía) con los dos
    logos institucionales lado a lado, para insertar al inicio de la
    portada PDF. Igual que en Word: si faltan los archivos, no falla nada."""
    try:
        if not (_RUTA_LOGO_ESAP.exists() and _RUTA_LOGO_CERTIFICACIONES.exists()):
            return []
        img_esap = Image(str(_RUTA_LOGO_ESAP), width=5.2 * cm, height=1.4 * cm)
        img_cert = Image(str(_RUTA_LOGO_CERTIFICACIONES), width=2.9 * cm, height=2.4 * cm)
        tabla_logos = Table([[img_esap, img_cert]], colWidths=[10 * cm, 6.5 * cm])
        tabla_logos.setStyle(TableStyle([
            ("ALIGN", (0, 0), (0, 0), "LEFT"),
            ("ALIGN", (1, 0), (1, 0), "RIGHT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        return [tabla_logos, Spacer(1, 10)]
    except Exception:
        return []


def _agregar_banner_docx(doc, titulo: str, subtitulo: str | None = None, color_hex: str = COLOR_INSTITUCIONAL) -> None:
    """Banda institucional de ancho completo (fondo de color, texto blanco)
    para usar como encabezado de portada en vez de un título plano
    centrado. Reutilizable por los 3 informes."""
    banner = doc.add_table(rows=1, cols=1)
    banner.autofit = False
    celda = banner.rows[0].cells[0]
    celda.width = Cm(17)
    _sombrear_celda(celda, color_hex)
    p_titulo = celda.paragraphs[0]
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)
    run_titulo.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if subtitulo:
        p_sub = celda.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = p_sub.add_run(subtitulo)
        run_sub.italic = True
        run_sub.font.size = Pt(11)
        run_sub.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def _dibujar_banda_portada_pdf(titulo: str, subtitulo: str | None = None, color_hex: str = COLOR_INSTITUCIONAL, alto_cm: float = 3.2):
    """Fábrica de función onFirstPage para reportlab: dibuja una banda de
    color a todo el ancho en la parte superior de la primera página, con el
    título y subtítulo en blanco, y delega el pie de página a
    `funcion_pie` si se le pasa (ver `_combinar_callbacks_primera_pagina`).

    CORREGIDO: títulos largos (p. ej. el nombre completo de SIIEAP) ya no
    se dibujan con drawCentredString en una sola línea sin ajuste — eso
    los hacía desbordarse fuera del ancho de la página. Ahora se envuelven
    con reportlab.platypus.Paragraph dentro del ancho real de la banda,
    con tamaño de fuente que se reduce automáticamente si el título es muy
    largo, y la altura de la banda se ajusta al número de líneas reales."""
    def _dibujar(canvas_pdf, doc_pdf):
        from reportlab.platypus import Paragraph as _P
        ancho_pagina, alto_pagina = LETTER
        margen_lateral = 1.5 * cm
        ancho_disponible = ancho_pagina - 2 * margen_lateral

        tamano_titulo = 15 if len(titulo) <= 70 else (12.5 if len(titulo) <= 130 else 10.5)
        estilo_titulo = ParagraphStyle(
            "BandaTitulo", fontName="Helvetica-Bold", fontSize=tamano_titulo,
            leading=tamano_titulo * 1.2, textColor=colors.white, alignment=1,  # 1 = centro
        )
        p_titulo = _P(titulo, estilo_titulo)
        ancho_usado, alto_titulo = p_titulo.wrap(ancho_disponible, 10 * cm)

        alto_subtitulo = 0
        p_subtitulo = None
        if subtitulo:
            estilo_subtitulo = ParagraphStyle(
                "BandaSubtitulo", fontName="Helvetica-Oblique", fontSize=10,
                leading=12, textColor=colors.white, alignment=1,
            )
            p_subtitulo = _P(subtitulo, estilo_subtitulo)
            _, alto_subtitulo = p_subtitulo.wrap(ancho_disponible, 10 * cm)

        # Alto real de la banda: máximo entre el mínimo estético (alto_cm) y
        # lo que de verdad ocupa el texto envuelto, con margen de respiro.
        alto_contenido = alto_titulo + (alto_subtitulo + 6 if subtitulo else 0) + 24
        alto_banda = max(alto_cm * cm, alto_contenido)

        canvas_pdf.saveState()
        canvas_pdf.setFillColor(colors.HexColor(f"#{color_hex}"))
        canvas_pdf.rect(0, alto_pagina - alto_banda, ancho_pagina, alto_banda, fill=1, stroke=0)

        y_actual = alto_pagina - 14 - alto_titulo
        p_titulo.drawOn(canvas_pdf, margen_lateral, y_actual)
        if p_subtitulo:
            y_actual -= (alto_subtitulo + 6)
            p_subtitulo.drawOn(canvas_pdf, margen_lateral, y_actual)
        canvas_pdf.restoreState()
    return _dibujar


def _banner_portada_pdf_flowables(titulo: str, subtitulo: str | None = None, color_hex: str = COLOR_INSTITUCIONAL):
    """Versión del banner de portada como flowables (Table de una celda con
    fondo de color), para insertar directamente en el flujo del documento
    igual que en Word, sin depender de un callback de canvas.

    CORREGIDO: el título y subtítulo ahora van envueltos en Paragraph (no
    como string plano dentro de la Table), para que el texto largo haga
    salto de línea dentro de la celda en vez de desbordarse."""
    tamano_titulo = 15 if len(titulo) <= 70 else (12.5 if len(titulo) <= 130 else 10.5)
    estilo_titulo = ParagraphStyle(
        "BannerFlowTitulo", fontName="Helvetica-Bold", fontSize=tamano_titulo,
        leading=tamano_titulo * 1.25, textColor=colors.white, alignment=1,
    )
    filas = [[Paragraph(titulo, estilo_titulo)]]
    if subtitulo:
        estilo_subtitulo = ParagraphStyle(
            "BannerFlowSubtitulo", fontName="Helvetica-Oblique", fontSize=10.5,
            leading=13, textColor=colors.white, alignment=1,
        )
        filas.append([Paragraph(subtitulo, estilo_subtitulo)])
    tabla = Table(filas, colWidths=[17 * cm])
    estilo = [
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{color_hex}")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ("LEFTPADDING", (0, 0), (-1, -1), 14),
        ("RIGHTPADDING", (0, 0), (-1, -1), 14),
    ]
    tabla.setStyle(TableStyle(estilo))
    return [tabla, Spacer(1, 14)]


def _combinar_callbacks_primera_pagina(*funciones):
    """Combina varias funciones onFirstPage/onLaterPages de reportlab
    (por ejemplo la banda de portada + el pie de página) en una sola,
    ya que SimpleDocTemplate solo acepta un callback por evento."""
    def _combinado(canvas_pdf, doc_pdf):
        for funcion in funciones:
            if funcion:
                funcion(canvas_pdf, doc_pdf)
    return _combinado


def _agregar_divisor_seccion_docx(doc, titulo: str, icono: str = "▪", color_hex: str = COLOR_INSTITUCIONAL):
    """Divisor de sección de ancho completo (barra de color + ícono +
    título) para usar en vez de un doc.add_heading plano. Devuelve el
    párrafo del título por si se necesita seguir estilizando."""
    tabla = doc.add_table(rows=1, cols=1)
    tabla.autofit = False
    celda = tabla.rows[0].cells[0]
    celda.width = Cm(17)
    _sombrear_celda(celda, color_hex)
    p = celda.paragraphs[0]
    run = p.add_run(f"{icono}  {titulo}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()
    return p


def _divisor_seccion_pdf(titulo: str, icono: str = "▪", color_hex: str = COLOR_INSTITUCIONAL):
    """Versión PDF (lista de flowables) del divisor de sección: una tabla
    de una celda con fondo de color, texto blanco e ícono, seguida de un
    espaciador. Se antepone a cada Heading1 de los informes.

    CORREGIDO: antes el título iba como string plano dentro de la Table,
    lo que hacía que títulos largos (algunos de más de 100 caracteres, como
    el de la sección de análisis integral) se desbordaran fuera de la
    celda o se vieran amontonados. Ahora va envuelto en Paragraph, con el
    tamaño de fuente ajustado según la longitud del título."""
    tamano_fuente = 12.5 if len(titulo) <= 55 else (11 if len(titulo) <= 100 else 9.5)
    estilo_divisor = ParagraphStyle(
        "DivisorSeccionPDF", fontName="Helvetica-Bold", fontSize=tamano_fuente,
        leading=tamano_fuente * 1.25, textColor=colors.white,
    )
    tabla = Table([[Paragraph(f"{icono}&nbsp;&nbsp;{titulo}", estilo_divisor)]], colWidths=[17 * cm])
    tabla.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{color_hex}")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
    ]))
    return [tabla, Spacer(1, 10)]


def _franjas_alternas_docx(tabla, color_hex_par: str = "F2F4F8") -> None:
    """Aplica sombreado en filas alternas (excluyendo el encabezado) a una
    tabla de python-docx, para tablas más legibles en documentos largos."""
    for indice_fila, fila in enumerate(tabla.rows[1:], start=1):
        if indice_fila % 2 == 0:
            for celda in fila.cells:
                _sombrear_celda(celda, color_hex_par)


# ---------------------------------------------------------------------------
# Identidad compartida — Tabla de contenido, Razón de ser y Nota cruzada
# entre los 3 informes. Agregado a solicitud explícita de Norma: cada
# informe debe (a) traer su propia tabla de contenido, (b) explicar al
# inicio para qué sirve y a quién está dirigido, y (c) decirle a quien lo
# recibe que existen otros 2 informes y qué contienen, para que sepa que
# puede solicitarlos.
# ---------------------------------------------------------------------------

# Claves válidas: "tecnico", "estudio_caso", "ejecutivo"
RAZON_DE_SER_POR_INFORME = {
    "tecnico": {
        "titulo": "¿Qué es este Informe Técnico y para quién es?",
        "texto": (
            "Este documento está dirigido a los equipos técnicos y de planeación "
            "de la entidad. Responde a la pregunta: ¿qué dice exactamente el "
            "diagnóstico institucional y cómo se construye, a partir de él, el "
            "plan de mejoramiento? Contiene el diagnóstico completo por dimensión "
            "e índice del MIPG, el Índice Sintético de Valor Público Territorial "
            "(ISVPT) y el análisis integral generado por inteligencia artificial, "
            "con fundamentación teórica (Nueva Gestión Pública, Post-NGP y Nuevo "
            "Institucionalismo)."
        ),
        "nota_cruzada": (
            "Este es 1 de 3 informes que SIIEAP genera para esta misma entidad. "
            "Si usted es el representante legal (alcalde, gobernador, contralor, "
            "personero, gerente o rector) y necesita una versión breve, en "
            "lenguaje llano, con las consecuencias legales, fiscales y "
            "disciplinarias de cada brecha, solicite el Informe Ejecutivo para "
            "Representantes Legales. Si necesita el análisis académico completo, "
            "con fundamentación teórica ampliada y escenarios prospectivos, "
            "solicite el Estudio de Caso Académico."
        ),
    },
    "estudio_caso": {
        "titulo": "¿Qué es este Estudio de Caso y para quién es?",
        "texto": (
            "Este documento está dirigido a estudiantes e investigadores, en el "
            "marco de la Maestría en Administración Pública de la ESAP. Responde "
            "a la pregunta: ¿cómo se explica este caso institucional desde la "
            "teoría de la administración pública? Contiene la misma "
            "fundamentación teórica del Informe Técnico, generada dinámicamente "
            "en función de la entidad elegida, el Análisis 360° regional y "
            "escenarios prospectivos."
        ),
        "nota_cruzada": (
            "Este es 1 de 3 informes que SIIEAP genera para esta misma entidad. "
            "Si necesita el diagnóstico completo por dimensión e índice para "
            "construir un plan de mejoramiento, solicite el Informe Técnico. Si "
            "usted es el representante legal de la entidad y necesita una "
            "versión breve, en lenguaje llano, con las consecuencias legales, "
            "fiscales y disciplinarias, solicite el Informe Ejecutivo para "
            "Representantes Legales."
        ),
    },
    "ejecutivo": {
        "titulo": "¿Qué es este Informe Ejecutivo y para quién es?",
        "texto": (
            "Este documento está dirigido al representante legal de la entidad "
            "— alcalde, gobernador, contralor, personero, gerente de una ESE o "
            "de una entidad descentralizada, rector de universidad pública, o "
            "gerente de una empresa industrial y comercial del Estado — en su "
            "condición de presidente del Comité Institucional de Gestión y "
            "Desempeño y del Comité Institucional de Coordinación de Control "
            "Interno. Responde a la pregunta: ¿qué le falta a mi entidad, por "
            "qué me expone legalmente, y qué debo decidir ya? Sin tecnicismos: "
            "semáforo visual de 5 quintiles del MIPG, matriz de riesgo "
            "probabilidad-impacto, y el marco legal, fiscal, administrativo y "
            "disciplinario aplicable a cada brecha, incluido el artículo 6 "
            "constitucional."
        ),
        "nota_cruzada": (
            "Este es 1 de 3 informes que SIIEAP genera para esta misma entidad. "
            "Si su equipo técnico necesita el diagnóstico completo por dimensión "
            "e índice para construir el plan de mejoramiento, solicite el "
            "Informe Técnico. Si necesita el análisis académico con "
            "fundamentación teórica ampliada, solicite el Estudio de Caso "
            "Académico."
        ),
    },
}


def _agregar_tabla_contenido_docx(doc) -> None:
    """Inserta un campo de Tabla de Contenido NATIVO de Word (niveles 1-3).

    También activa la marca 'actualizar campos al abrir' en settings.xml,
    para que Microsoft Word ofrezca (o en muchas configuraciones, resuelva
    directamente) la actualización del índice apenas se abre el documento,
    en vez de depender de que quien lo reciba sepa hacer clic derecho >
    Actualizar campo. Si el documento se abre en un visor que no soporta
    campos (algunos lectores de PDF-desde-Word, apps móviles muy básicas),
    el texto de repuesto ("Haga clic derecho...") sigue siendo visible como
    respaldo."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Activar 'actualizar campos al abrir' en settings.xml (una sola vez por documento)
    settings_element = doc.settings.element
    if settings_element.find(qn("w:updateFields")) is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings_element.append(update_fields)

    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run("Tabla de contenido")
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)
    run_titulo.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Haga clic derecho sobre esta tabla y seleccione 'Actualizar campo' para generar el índice."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr_text)
    r_element.append(fld_separate)
    r_element.append(fld_text)
    r_element.append(fld_end)

    doc.add_page_break()


def _agregar_razon_de_ser_docx(doc, tipo_informe: str) -> None:
    """Inserta el bloque de 'razón de ser' (qué es este informe, para quién)
    y la nota cruzada (qué otros 2 informes existen), justo después de la
    tabla de contenido y antes del contenido propio del informe."""
    info = RAZON_DE_SER_POR_INFORME[tipo_informe]

    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run(info["titulo"])
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    run_titulo.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    p_texto = doc.add_paragraph(info["texto"])
    p_texto.paragraph_format.space_after = Pt(10)

    p_nota = doc.add_paragraph()
    run_nota = p_nota.add_run(info["nota_cruzada"])
    run_nota.italic = True
    run_nota.font.size = Pt(9.5)
    run_nota.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def _toc_pdf_flowables(secciones: list[str]):
    """Índice estático (sin números de página dinámicos) para la versión
    PDF: lista los títulos de sección en orden. A diferencia del .docx,
    reportlab no resuelve automáticamente números de página sin un
    refactor mayor del pipeline de construcción (doc.build -> multiBuild);
    se prioriza no arriesgar la generación ya estable del PDF."""
    elementos = [Paragraph("<b>Contenido</b>", ParagraphStyle(
        "TituloIndicePDF", fontName="Helvetica-Bold", fontSize=14,
        textColor=colors.HexColor(f"#{COLOR_INSTITUCIONAL}"),
    ))]
    elementos.append(Spacer(1, 6))
    for seccion in secciones:
        elementos.append(Paragraph(f"• {seccion}", _ESTILO_CELDA_TABLA_PDF))
    elementos.append(Spacer(1, 10))
    elementos.append(PageBreak())
    return elementos


def _razon_de_ser_pdf_flowables(tipo_informe: str):
    """Versión PDF (lista de flowables) de _agregar_razon_de_ser_docx."""
    info = RAZON_DE_SER_POR_INFORME[tipo_informe]
    estilo_titulo = ParagraphStyle(
        "TituloRazonSerPDF", fontName="Helvetica-Bold", fontSize=13,
        textColor=colors.HexColor("#1F4E79"), spaceAfter=6,
    )
    estilo_texto = ParagraphStyle(
        "TextoRazonSerPDF", fontName="Helvetica", fontSize=10.5, leading=14, spaceAfter=8,
    )
    estilo_nota = ParagraphStyle(
        "NotaCruzadaPDF", fontName="Helvetica-Oblique", fontSize=9,
        textColor=colors.grey, spaceAfter=10,
    )
    return [
        Paragraph(info["titulo"], estilo_titulo),
        Paragraph(info["texto"], estilo_texto),
        Paragraph(info["nota_cruzada"], estilo_nota),
        PageBreak(),
    ]


# Esquema OFICIAL de 5 quintiles del MIPG (rangos de puntaje 0-100), usado
# para colorear de forma consistente cualquier tabla que muestre puntajes
# en los 3 informes — no debe confundirse con el esquema de 3 niveles de
# riesgo (alta/media/baja) de las brechas, que es una convención distinta
# y ya existente en este mismo módulo (_COLOR_HEX_POR_RIESGO).
_QUINTILES_MIPG = [
    (20, "E6534A"),   # 0-20   crítico
    (40, "F0A02E"),   # 21-40  bajo
    (60, "F5D65A"),   # 41-60  medio
    (80, "8CC152"),   # 61-80  satisfactorio
    (100, "2E8B57"),  # 81-100 sobresaliente
]


def _color_hex_quintil_mipg(puntaje) -> str | None:
    """Devuelve el color hex del quintil oficial MIPG correspondiente a un
    puntaje de 0 a 100, o None si el puntaje no es un número válido."""
    try:
        puntaje = float(puntaje)
    except (TypeError, ValueError):
        return None
    for limite, color_hex in _QUINTILES_MIPG:
        if puntaje <= limite:
            return color_hex
    return _QUINTILES_MIPG[-1][1]


def _color_hex_quintil_mipg(puntaje) -> str | None:
    """Devuelve el color hex del quintil oficial MIPG correspondiente a un
    puntaje de 0 a 100, o None si el puntaje no es un número válido."""
    try:
        puntaje = float(puntaje)
    except (TypeError, ValueError):
        return None
    for limite, color_hex in _QUINTILES_MIPG:
        if puntaje <= limite:
            return color_hex
    return _QUINTILES_MIPG[-1][1]


_QUINTILES_MIPG_NOMBRE_EMOJI = [
    (20, "Crítico", "🔴"),
    (40, "Bajo", "🟠"),
    (60, "Medio", "🟡"),
    (80, "Medio-alto", "🟢"),
    (100, "Consolidación", "✅"),
]


def _quintil_mipg(puntaje):
    """Devuelve (color_hex, emoji, nombre_categoria) para un puntaje de 0 a
    100, según el esquema oficial de 5 quintiles del MIPG. Usado por el
    Informe Ejecutivo para Alcaldes/Gobernadores (semáforo en lenguaje
    llano) y disponible para los otros informes si lo necesitan."""
    try:
        puntaje_num = float(puntaje)
    except (TypeError, ValueError):
        return ("999999", "⚪", "Sin dato")
    for limite, nombre, emoji in _QUINTILES_MIPG_NOMBRE_EMOJI:
        if puntaje_num <= limite:
            color_hex = _color_hex_quintil_mipg(puntaje_num)
            return (color_hex, emoji, nombre)
    color_hex = _QUINTILES_MIPG[-1][1]
    return (color_hex, "✅", "Consolidación")


DESCARGO_RESPONSABILIDAD_AMPLIADO = (
    "Este informe fue generado por el Sistema Integral de Diagnóstico del Desempeño "
    "Institucional (SIIEAP) a partir de datos reales del Índice de Desempeño "
    "Institucional (IDI-MIPG) publicado por el Departamento Administrativo de la "
    "Función Pública, complementado con un análisis generado por inteligencia "
    "artificial (Claude, Anthropic) como punto de partida académico y metodológico. "
    "El semáforo de quintiles, las brechas priorizadas y las lecturas de riesgo "
    "legal, financiero, administrativo y disciplinario que aquí se presentan son "
    "un ejercicio de apoyo a la toma de decisiones de la Alta Dirección — NO "
    "constituyen un dictamen jurídico, fiscal ni disciplinario oficial, y no "
    "sustituyen la valoración de la Oficina de Control Interno, la Oficina "
    "Asesora Jurídica, la Secretaría de Hacienda ni de los organismos de control "
    "(Procuraduría General de la Nación, Contraloría General de la República, "
    "Contralorías territoriales) de la entidad. Se entrega con fines académicos, "
    "como evidencia de ejecución del sistema y como insumo de trabajo para la "
    "asignatura."
)


# ---------------------------------------------------------------------------
# Régimen especial MIPG/MECI: algunas entidades NO están obligadas a
# implementar el MIPG en su integralidad (7 dimensiones, 19 políticas) sino
# únicamente la política de Control Interno (Ley 87 de 1993 → política 19 →
# dimensión 7 → MECI), en virtud del artículo 40 de la Ley 489 de 1998 y del
# artículo 2.2.22.3.4 del Decreto 1499 de 2017. Esto es distinto de una
# entidad de la Rama Ejecutiva territorial (alcaldías, gobernaciones,
# establecimientos públicos, EICE) que sí debe reportar el MIPG completo.
#
# Catálogo verificado con conceptos oficiales de Función Pública (radicados
# 20265000131771 de 2026 sobre Contralorías territoriales, y 20265000162111
# de 2026 sobre la Universidad de Antioquia como ente universitario autónomo)
# más el análisis normativo propio de Ley 489/1998 arts. 38-40 para EICE.
# ---------------------------------------------------------------------------

REGIMEN_ESPECIAL_NINGUNO = "ninguno"  # Rama Ejecutiva: MIPG íntegro aplica (alcaldías, gobernaciones, EICE, ESE, EOSP...)
REGIMEN_ESPECIAL_UNIVERSIDAD_AUTONOMA = "universidad_autonoma"
REGIMEN_ESPECIAL_ORGANO_CONTROL = "organo_control"
REGIMEN_ESPECIAL_BANCO_REPUBLICA = "banco_republica"
REGIMEN_ESPECIAL_CORPORACION_AUTONOMA = "corporacion_autonoma_regional"
REGIMEN_ESPECIAL_CONCEJO_ASAMBLEA = "concejo_asamblea"
REGIMEN_ESPECIAL_PERSONERIA = "personeria"
REGIMEN_ESPECIAL_RAMA_LEGISLATIVA = "rama_legislativa"
REGIMEN_ESPECIAL_ORGANO_CONTROL_NACIONAL = "organo_control_nacional"
REGIMEN_ESPECIAL_RAMA_JUDICIAL = "rama_judicial"
REGIMEN_ESPECIAL_ORGANIZACION_ELECTORAL = "organizacion_electoral"

_CATALOGO_REGIMEN_ESPECIAL = {
    REGIMEN_ESPECIAL_NINGUNO: {
        "nombre": "Rama Ejecutiva (MIPG íntegro aplica)",
        "aplica_mipg_integral": True,
        "nota": None,
    },
    REGIMEN_ESPECIAL_UNIVERSIDAD_AUTONOMA: {
        "nombre": "Ente universitario autónomo",
        "aplica_mipg_integral": False,
        "nota": (
            "Esta entidad es un ente universitario autónomo, sujeto a régimen especial "
            "conforme al artículo 40 de la Ley 489 de 1998. Por esa razón, el Modelo "
            "Integrado de Planeación y Gestión (MIPG) no le aplica en su integralidad: "
            "solo está obligada a implementar la política de Control Interno (Ley 87 de "
            "1993 — política 19, dimensión 7 del MIPG, desarrollada a través del MECI), "
            "conforme lo confirmó el Departamento Administrativo de la Función Pública "
            "(concepto con radicado 20265000162111 de 2026). Las demás 18 políticas se "
            "implementan solo en la medida en que le sean aplicables por su propia "
            "normativa sectorial — no por obligatoriedad del MIPG. Como consecuencia, el "
            "resultado oficial publicado para esta entidad corresponde únicamente a la "
            "Dimensión 7, y cualquier lectura de brecha en las demás dimensiones que "
            "aparezca en este informe debe leerse con esa salvedad: no es un "
            "incumplimiento del MIPG, porque el MIPG íntegro no le es exigible."
        ),
    },
    REGIMEN_ESPECIAL_ORGANO_CONTROL: {
        "nombre": "Órgano de control (Contraloría territorial)",
        "aplica_mipg_integral": False,
        "nota": (
            "Esta entidad es una Contraloría departamental, municipal o distrital y, por "
            "tanto, no hace parte de la Rama Ejecutiva del Poder Público (arts. 117 y 119 "
            "de la Constitución Política). En consecuencia, no está sujeta al Modelo "
            "Integrado de Planeación y Gestión (MIPG) en su integralidad: únicamente debe "
            "implementar la política de Control Interno (Ley 87 de 1993 — política 19, "
            "dimensión 7, desarrollada a través del MECI), conforme lo confirmó el "
            "Departamento Administrativo de la Función Pública (concepto con radicado "
            "20265000131771 de 2026). Las demás 18 políticas se implementan solo en la "
            "medida en que le sean aplicables por su propia normativa — no por "
            "obligatoriedad del MIPG. El resultado oficial publicado para esta entidad "
            "corresponde únicamente a la Dimensión 7."
        ),
    },
    REGIMEN_ESPECIAL_PERSONERIA: {
        "nombre": "Personería municipal o distrital",
        "aplica_mipg_integral": False,
        "nota": (
            "Esta entidad es una Personería municipal o distrital — un órgano del "
            "Ministerio Público en el nivel territorial (arts. 118 y 313 de la "
            "Constitución Política), distinto del alcalde y de la Rama Ejecutiva. El "
            "Manual Operativo del MIPG (Función Pública, versión 6.1, febrero de 2026) la "
            "agrupa junto con los Concejos municipales, dentro de un esquema de "
            "implementación del Sistema de Control Interno (MECI) diferenciado y "
            "simplificado — documentado explícitamente para municipios de quinta y sexta "
            "categoría —, en vez del MIPG en su integralidad de 7 dimensiones y 19 "
            "políticas. Cualquier lectura de brecha en las demás dimensiones que aparezca "
            "en este informe debe leerse con esa salvedad."
        ),
    },
    REGIMEN_ESPECIAL_CONCEJO_ASAMBLEA: {
        "nombre": "Concejo municipal o Asamblea departamental",
        "aplica_mipg_integral": False,
        "nota": (
            "Esta entidad es un Concejo municipal o una Asamblea departamental — una "
            "corporación público-administrativa de elección popular (arts. 312 y 313 de "
            "la Constitución Política), NO un organismo de la Rama Ejecutiva (el "
            "alcalde/gobernador es quien ejerce esa función, no el Concejo/Asamblea). El "
            "Manual Operativo del MIPG (Función Pública, versión 6.1, febrero de 2026) la "
            "agrupa junto con las Personerías, dentro de un esquema de implementación del "
            "Sistema de Control Interno (MECI) diferenciado y simplificado — documentado "
            "explícitamente para municipios de quinta y sexta categoría —, en vez del "
            "MIPG en su integralidad de 7 dimensiones y 19 políticas. Cualquier lectura "
            "de brecha en las demás dimensiones que aparezca en este informe debe leerse "
            "con esa salvedad."
        ),
    },
    REGIMEN_ESPECIAL_BANCO_REPUBLICA: {
        "nombre": "Banco de la República",
        "aplica_mipg_integral": False,
        "nota": (
            "El Banco de la República es una entidad de régimen especial conforme al "
            "artículo 40 de la Ley 489 de 1998, por lo que el MIPG no le aplica en su "
            "integralidad — únicamente la política de Control Interno (Ley 87 de 1993)."
        ),
    },
    REGIMEN_ESPECIAL_CORPORACION_AUTONOMA: {
        "nombre": "Corporación Autónoma Regional",
        "aplica_mipg_integral": False,
        "nota": (
            "Esta entidad es una Corporación Autónoma Regional, sujeta a régimen especial "
            "conforme al artículo 40 de la Ley 489 de 1998. El MIPG no le aplica en su "
            "integralidad — únicamente la política de Control Interno (Ley 87 de 1993 — "
            "política 19, dimensión 7, desarrollada a través del MECI). Las demás 18 "
            "políticas se implementan solo en la medida en que le sean aplicables por su "
            "propia normativa sectorial."
        ),
    },
    REGIMEN_ESPECIAL_RAMA_LEGISLATIVA: {
        "nombre": "Rama Legislativa (Congreso de la República)",
        "aplica_mipg_integral": False,
        "nota": (
            "Esta entidad hace parte de la Rama Legislativa del Poder Público (Senado de "
            "la República o Cámara de Representantes, arts. 113 y 114 de la Constitución "
            "Política) y, por tanto, no hace parte de la Rama Ejecutiva. El Modelo "
            "Integrado de Planeación y Gestión (MIPG), como sistema de gestión propio de "
            "la administración pública ejecutiva (Decreto 1499 de 2017), no le aplica en "
            "su integralidad: en el catálogo oficial de Función Pública, esta entidad "
            "reporta únicamente la política de Control Interno (Ley 87 de 1993 — política "
            "19, dimensión 7 del MIPG, desarrollada a través del MECI, formulario "
            "'MECI', no 'MIPG'). Cualquier lectura de brecha en las demás dimensiones que "
            "aparezca en este informe debe leerse con esa salvedad: no es un "
            "incumplimiento del MIPG, porque el MIPG íntegro no le es exigible por su "
            "naturaleza constitucional."
        ),
    },
    REGIMEN_ESPECIAL_ORGANO_CONTROL_NACIONAL: {
        "nombre": "Órgano de control del nivel nacional",
        "aplica_mipg_integral": False,
        "nota": (
            "Esta entidad es un órgano de control del nivel nacional (Procuraduría "
            "General de la Nación, Contraloría General de la República, Defensoría del "
            "Pueblo o Auditoría General de la República — arts. 117, 118, 119, 267 y 268 "
            "de la Constitución Política), con autonomía administrativa y presupuestal "
            "propia, distinta de la Rama Ejecutiva. En el catálogo oficial de Función "
            "Pública, esta entidad reporta únicamente la política de Control Interno (Ley "
            "87 de 1993 — política 19, dimensión 7 del MIPG, desarrollada a través del "
            "MECI, formulario 'MECI', no 'MIPG'). Cualquier lectura de brecha en las "
            "demás dimensiones que aparezca en este informe debe leerse con esa salvedad."
        ),
    },
    REGIMEN_ESPECIAL_RAMA_JUDICIAL: {
        "nombre": "Rama Judicial",
        "aplica_mipg_integral": False,
        "nota": (
            "Esta entidad hace parte de la Rama Judicial del Poder Público (p. ej. "
            "Fiscalía General de la Nación, Consejo Superior de la Judicatura o Instituto "
            "Nacional de Medicina Legal y Ciencias Forenses — art. 116 de la Constitución "
            "Política), con autonomía frente a la Rama Ejecutiva. En el catálogo oficial "
            "de Función Pública, esta entidad reporta únicamente la política de Control "
            "Interno (Ley 87 de 1993 — política 19, dimensión 7 del MIPG, desarrollada a "
            "través del MECI, formulario 'MECI', no 'MIPG'). Cualquier lectura de brecha "
            "en las demás dimensiones que aparezca en este informe debe leerse con esa "
            "salvedad."
        ),
    },
    REGIMEN_ESPECIAL_ORGANIZACION_ELECTORAL: {
        "nombre": "Organización Electoral",
        "aplica_mipg_integral": False,
        "nota": (
            "Esta entidad hace parte de la Organización Electoral (Registraduría "
            "Nacional del Estado Civil o Consejo Nacional Electoral — art. 120 de la "
            "Constitución Política), un órgano autónomo e independiente de las tres "
            "ramas del poder público. En el catálogo oficial de Función Pública, esta "
            "entidad reporta únicamente la política de Control Interno (Ley 87 de 1993 — "
            "política 19, dimensión 7 del MIPG, desarrollada a través del MECI, "
            "formulario 'MECI', no 'MIPG'). Cualquier lectura de brecha en las demás "
            "dimensiones que aparezca en este informe debe leerse con esa salvedad."
        ),
    },
}


def _info_regimen_especial(tipo_regimen_especial: str | None) -> dict | None:
    """Devuelve el registro del catálogo de régimen especial para el tipo dado,
    o None si el tipo no existe o corresponde a 'ninguno' (Rama Ejecutiva)."""
    if not tipo_regimen_especial or tipo_regimen_especial == REGIMEN_ESPECIAL_NINGUNO:
        return None
    return _CATALOGO_REGIMEN_ESPECIAL.get(tipo_regimen_especial)


def _agregar_nota_regimen_especial_docx(doc, tipo_regimen_especial: str | None) -> None:
    """Inserta, si aplica, un recuadro con la nota fija de régimen especial
    MIPG/MECI en la portada del informe (docx). No hace nada si la entidad
    no tiene régimen especial (Rama Ejecutiva ordinaria)."""
    info = _info_regimen_especial(tipo_regimen_especial)
    if not info:
        return
    tabla_nota = doc.add_table(rows=1, cols=1)
    tabla_nota.autofit = False
    celda = tabla_nota.rows[0].cells[0]
    celda.width = Cm(17)
    _sombrear_celda(celda, "FDF2D0")
    p_titulo = celda.paragraphs[0]
    run_titulo = p_titulo.add_run(f"⚠️ Régimen especial MIPG/MECI — {info['nombre']}")
    run_titulo.bold = True
    run_titulo.font.size = Pt(10)
    p_nota = celda.add_paragraph()
    run_nota = p_nota.add_run(info["nota"])
    run_nota.italic = True
    run_nota.font.size = Pt(9)
    doc.add_paragraph()


def _nota_regimen_especial_pdf_flowables(tipo_regimen_especial: str | None):
    """Versión PDF (lista de flowables) de la nota de régimen especial.
    Devuelve una lista vacía si no aplica."""
    info = _info_regimen_especial(tipo_regimen_especial)
    if not info:
        return []
    estilo_titulo_nota = ParagraphStyle(
        "TituloNotaRegimen", fontName="Helvetica-Bold", fontSize=9.5, leading=12,
    )
    estilo_cuerpo_nota = ParagraphStyle(
        "CuerpoNotaRegimen", fontName="Helvetica-Oblique", fontSize=8.5, leading=11,
        spaceBefore=4,
    )
    texto_escapado = info["nota"].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    tabla = Table(
        [[[
            Paragraph(f"⚠️ Régimen especial MIPG/MECI — {info['nombre']}", estilo_titulo_nota),
            Paragraph(texto_escapado, estilo_cuerpo_nota),
        ]]],
        colWidths=[17 * cm],
    )
    tabla.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FDF2D0")),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    return [tabla, Spacer(1, 10)]


# ---------------------------------------------------------------------------
# Bloque fijo: Contextualización de la Administración Pública Contemporánea
# ---------------------------------------------------------------------------

CONTEXTUALIZACION_TITULO = "Contextualización de la Administración Pública Contemporánea"

CONTEXTUALIZACION_INTRO = (
    "Este informe se enmarca en el recorrido teórico de la asignatura Enfoques y "
    "Teorías de la Administración Pública, que inicia en los antecedentes de la "
    "cultura griega para la distinción entre lo público y lo privado, y llega hasta "
    "los debates más recientes sobre la transformación digital del Estado. A "
    "continuación se sintetiza ese arco completo, como marco de lectura para el "
    "diagnóstico institucional que se presenta más adelante."
)

CONTEXTUALIZACION_ARCO_HISTORICO = [
    (
        "Antecedentes griegos de lo público y lo privado",
        "La distinción entre la esfera pública (la polis, los asuntos comunes, el "
        "koinon) y la esfera privada (el oikos, el hogar, el idion) es el punto de "
        "partida clásico para pensar qué es \"lo público\" y cómo se organiza su "
        "gobierno. De esa distinción original —lo que compete a todos frente a lo "
        "que compete a cada uno— se deriva la pregunta que atraviesa toda la "
        "disciplina: ¿quién decide sobre los asuntos comunes y con qué reglas?",
    ),
    (
        "Administración Pública clásica y el modelo burocrático weberiano",
        "La Administración Pública clásica (Wilson, Weber) responde a esa pregunta "
        "con un modelo de burocracia racional-legal: jerarquía, especialización, "
        "reglas escritas, impersonalidad y separación entre política y "
        "administración. Este modelo dota al Estado de previsibilidad y control, "
        "pero es también el que después será criticado por su rigidez, su lentitud "
        "y su desconexión de los resultados que efectivamente produce para el "
        "ciudadano.",
    ),
    (
        "Nuevo Institucionalismo en la Administración Pública",
        "March y Olsen (1989) proponen \"redescubrir las instituciones\": las reglas, "
        "normas y rutinas no son un telón de fondo neutro, sino que moldean "
        "directamente el comportamiento de las organizaciones públicas y explican "
        "por qué entidades con recursos similares logran resultados distintos. El "
        "institucionalismo distingue al menos tres variantes que se complementan: "
        "el institucionalismo racional (las reglas como incentivos que estructuran "
        "decisiones), el histórico (la dependencia de la trayectoria: lo que una "
        "entidad hizo antes condiciona lo que puede hacer ahora) y el sociológico "
        "(DiMaggio y Powell, 1983), que explica el isomorfismo institucional: las "
        "organizaciones se parecen entre sí no solo por eficiencia, sino por "
        "presión coercitiva (la norma obliga), mimética (se copia al que parece "
        "legítimo) y normativa (las profesiones y gremios difunden un estándar). "
        "Esta última variante es clave para leer el MIPG: explica por qué una "
        "entidad puede tener todos los formatos y comités que la norma exige "
        "(cumplimiento formal, isomorfismo) sin que eso se traduzca aún en una "
        "práctica real institucionalizada — la brecha entre \"reglas en el papel\" "
        "y \"reglas en uso\".",
    ),
    (
        "Nueva Gestión Pública (NGP)",
        "Surge en los años 80 y 90 (Hood, 1991; Osborne y Gaebler, 1992) como "
        "respuesta a la crisis fiscal del Estado de bienestar y a la percepción de "
        "ineficiencia burocrática. Retoma técnicas de gestión privada: "
        "descentralización de la autoridad gerencial, orientación a resultados e "
        "indicadores de desempeño, introducción de mecanismos de mercado y "
        "competencia, orientación al \"cliente-ciudadano\", y medición explícita del "
        "desempeño (los \"siete doctrinas\" de Hood). Su aporte fue instalar la "
        "cultura de la medición y la rendición de cuentas por resultados —la base "
        "misma de instrumentos como el FURAG—, pero ha sido criticada por "
        "fragmentar el aparato estatal en unidades autónomas difíciles de "
        "coordinar, por debilitar el ethos de lo público al importar lógicas de "
        "mercado, y por reducir el éxito de la gestión a indicadores que no "
        "siempre capturan el bienestar colectivo generado.",
    ),
    (
        "post-Nueva Gestión Pública (post-NGP) y Gobernanza Digital",
        "Ante la fragmentación que dejó la NGP, autores como Dunleavy et al. (2006, "
        "\"Digital Era Governance\") y Osborne (2006, \"Nueva Gobernanza Pública\") "
        "proponen reintegrar lo que la NGP separó: menos unidades autónomas y más "
        "coordinación interinstitucional, menos competencia interna y más redes de "
        "colaboración, menos cliente y más ciudadano co-productor del servicio. La "
        "post-NGP no abandona la orientación a resultados de la NGP, pero la "
        "combina con una lógica de red (gobernanza) y con las posibilidades que "
        "abre la digitalización del Estado para integrar de nuevo lo que estaba "
        "fragmentado. Desde la ESAP, Chica-Vélez y Salazar-Ortiz (2021) profundizan "
        "en esta transición y muestran cómo la gobernanza y la innovación pública se "
        "resignifican precisamente en el tránsito de la NGP hacia la post-NGP, "
        "dejando de ser un simple discurso de modernización para convertirse en una "
        "forma concreta de organizar y gestionar lo público.",
    ),
    (
        "Gobernanza Pública y Valor Público",
        "La gobernanza (Kooiman, Rhodes) reconoce que el Estado ya no gobierna en "
        "solitario: coordina una red de actores —mercado, sociedad civil, "
        "cooperación internacional— para producir resultados que ninguno lograría "
        "por separado. En paralelo, Mark Moore (1995, \"Creating Public Value\") "
        "propone el valor público como el criterio último de éxito de la gestión: "
        "no basta con ser eficiente (NGP) ni con coordinar bien una red "
        "(gobernanza) si el resultado no se traduce en bienestar colectivo "
        "legítimo y sostenible. El valor público exige, además, capacidad "
        "operativa real y legitimidad/respaldo político-social — el \"triángulo "
        "estratégico\" de Moore.",
    ),
    (
        "MIPG — Modelo Integrado de Planeación y Gestión",
        "El MIPG (Decreto 1499 de 2017) es la traducción institucional colombiana "
        "de estas tres corrientes: conserva del institucionalismo la atención a "
        "reglas y rutinas (dimensiones y políticas normadas), toma de la NGP la "
        "medición sistemática del desempeño (el FURAG y el IDI), y adopta de la "
        "gobernanza y el valor público la idea de que la meta final no es cumplir "
        "un formato sino generar resultados y valor para el ciudadano. Articula "
        "planeación, gestión del riesgo, talento humano, control interno y "
        "evaluación de resultados en un solo modelo de gestión pública.",
    ),
]

CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_TITULO = (
    "Ampliación: las teorías y enfoques de la Administración Pública Contemporánea "
    "(no son solo tres, y se incluyen aquí las corrientes de vanguardia más recientes)"
)

CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_INTRO = (
    "Las tres corrientes anteriores (Nuevo Institucionalismo, NGP, post-NGP) explican "
    "la base teórica del MIPG, pero la Administración Pública contemporánea ha seguido "
    "evolucionando en la última década con nuevos enfoques, cada uno con su propio "
    "cuerpo de literatura. Estos no son solo marcos conceptuales importados: cada uno "
    "se desarrolla aquí junto con la norma colombiana vigente que lo valida, citando "
    "el artículo específico cuando la fuente es una ley o decreto. Se incluyen, "
    "además de los diez enfoques del esquema integrador original, cinco corrientes "
    "adicionales de vanguardia administrativa (Gobierno como Plataforma, "
    "interoperabilidad de datos, co-creación y Design Thinking, Administración "
    "Pública Conductual y Gemelos Digitales de Territorio). Aclaración metodológica "
    "importante: los documentos CONPES (Consejo Nacional de Política Económica y "
    "Social) son documentos de POLÍTICA PÚBLICA, no leyes, y por tanto no tienen "
    "\"artículos\" — se citan aquí por su número y objetivo, no deben confundirse con "
    "normas de rango legal. Y dos salvedades honestas, señaladas explícitamente donde "
    "corresponde: la Inteligencia Artificial cuenta en Colombia con política pública "
    "(CONPES), pero todavía no con una ley integral que la regule; y la "
    "Administración Pública Conductual y los Gemelos Digitales de Territorio son, a "
    "la fecha de este informe, corrientes de vanguardia SIN desarrollo normativo "
    "propio verificado en Colombia — se incluyen como oportunidad de política "
    "pública, no como hecho normativo consumado."
)

CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS = [
    (
        "Estado Digital (e-Gobierno)",
        "El Estado Digital, o e-Gobierno, es la etapa en la que las tecnologías de la "
        "información dejan de ser un simple apoyo administrativo y se integran "
        "directamente en la prestación de servicios y en la toma de decisiones "
        "(Fountain, 2001, \"Building the Virtual State\"; Dunleavy et al., 2006). Su "
        "aporte clave es mostrar que la tecnología no se limita a automatizar un "
        "trámite existente: lo redefine, porque las reglas y la estructura de la "
        "organización terminan moldeadas por lo que el sistema de información permite "
        "o restringe (la \"tecnología en uso\" de Fountain). Colombia lo valida "
        "normativamente en el artículo 2, numeral 8, de la Ley 1341 de 2009 (que ordena "
        "al Gobierno fijar los mecanismos para la masificación del Gobierno en Línea, "
        "hoy Gobierno Digital), desarrollado por el Decreto 1008 de 2018 y actualizado "
        "por el Decreto 767 de 2022, cuyo artículo 2.2.9.1.1.3 (del Decreto 1078 de "
        "2015) fija los principios de la Política de Gobierno Digital vigente, una de "
        "las políticas de gestión y desempeño del MIPG.",
    ),
    (
        "Transformación Digital del Estado",
        "La Transformación Digital va un paso más allá del Estado Digital: no basta "
        "con digitalizar un procedimiento, sino que exige repensar de fondo la "
        "arquitectura institucional, los flujos de datos entre entidades y la cultura "
        "organizacional para operar bajo una lógica \"digital por defecto\" (OCDE, 2014, "
        "\"Recommendation on Digital Government Strategies\"). Colombia la desarrolla "
        "mediante el Decreto 1263 de 2022, que adiciona el Título 23 a la Parte 2 del "
        "Libro 2 del Decreto 1078 de 2015 con los lineamientos de transformación "
        "digital pública, y mediante la Estrategia Nacional Digital 2023-2026 del DNP "
        "(instrumento de política, no de rango legal), que se traducen en la Política "
        "de Gobierno Digital del MIPG.",
    ),
    (
        "Gobernanza Inteligente (Smart Governance)",
        "La Gobernanza Inteligente (Meijer y Bolívar, 2016) describe cómo los "
        "gobiernos usan datos, sensores y tecnologías de ciudad inteligente para "
        "coordinar la acción colectiva de forma más ágil y basada en evidencia casi "
        "en tiempo real. Extiende la lógica de red de la gobernanza pública (post-NGP) "
        "hacia un entorno digital instrumentado. En Colombia se apoya en el modelo de "
        "territorios y ciudades inteligentes que exige el Decreto 767 de 2022 y en el "
        "documento CONPES 3995 de 2020 (Política Nacional de Confianza y Seguridad "
        "Digital, sin artículos por ser documento de política, no ley), que da el "
        "marco de confianza sobre el que puede operar esa coordinación.",
    ),
    (
        "Inteligencia Artificial en el sector público",
        "La incorporación de Inteligencia Artificial en la Administración Pública "
        "(OCDE, 2019, \"Recomendación sobre Inteligencia Artificial\") ofrece "
        "oportunidades reales —automatización de análisis, diagnóstico predictivo, "
        "como el que realiza este mismo sistema— pero exige gestionar riesgos "
        "conocidos: sesgos algorítmicos, opacidad (\"caja negra\") y la necesidad de "
        "supervisión humana significativa sobre cualquier resultado generado por IA. "
        "En Colombia, esta corriente se ha traducido en política pública —los "
        "documentos CONPES 3975 de 2019 y CONPES 4144 de 2025 (Política Nacional de "
        "Inteligencia Artificial, con hoja de ruta a 2030), ninguno de los cuales tiene "
        "artículos por no ser leyes— pero, a la fecha de este informe, el país AÚN NO "
        "cuenta con una ley integral que regule la IA: existen varios proyectos de ley "
        "en trámite en el Congreso sobre la materia, aún sin sanción. Por eso este "
        "informe insiste en que el análisis generado con IA es siempre un punto de "
        "partida metodológico, nunca un dictamen definitivo.",
    ),
    (
        "Gobierno Abierto",
        "El Gobierno Abierto (Open Government Partnership, 2011) descansa en tres "
        "pilares que se refuerzan entre sí: transparencia (acceso a la información "
        "pública), participación ciudadana (involucrar a la ciudadanía en el diseño "
        "de la política) y colaboración (co-producción de soluciones entre Estado, "
        "sociedad civil y sector privado). Colombia es miembro de la Open Government "
        "Partnership desde 2011, y su pilar de transparencia está anclado en la Ley "
        "1712 de 2014 (Ley de Transparencia y del Derecho de Acceso a la Información "
        "Pública): el artículo 1 fija su objeto, el artículo 2 establece el principio "
        "de máxima publicidad (toda información en poder de una entidad pública es "
        "pública salvo excepción constitucional o legal expresa) y el artículo 3 fija "
        "los demás principios de la transparencia y el acceso a la información.",
    ),
    (
        "Gobernanza de Datos",
        "La Gobernanza de Datos trata el dato público como un activo estratégico que "
        "debe gestionarse con reglas claras de calidad, interoperabilidad, seguridad "
        "y uso responsable. En Colombia esta corriente tiene tres soportes legales "
        "concretos que operan en tensión productiva: el Decreto 1389 de 2022 "
        "(compilado como artículo 2.2.24.3.4 del Decreto 1078 de 2015), que crea el "
        "Comité Nacional de Datos para impulsar su gobernanza, uso y reutilización; el "
        "artículo 2 de la Ley 1712 de 2014, que exige apertura y máxima publicidad; y "
        "los artículos 1 y 2 de la Ley 1581 de 2012 (Habeas Data), que protegen los "
        "datos personales y ponen el límite a esa apertura cuando el dato es de una "
        "persona natural.",
    ),
    (
        "Administración Pública Basada en Evidencia",
        "La Administración Pública Basada en Evidencia (Evidence-Based Policy Making; "
        "Head, 2008; Nutley et al., 2007) exige que las decisiones de política "
        "pública se apoyen en datos y evaluaciones rigurosas, no solo en la intuición "
        "o en la presión política coyuntural. En Colombia, el propio Decreto 1499 de "
        "2017 que crea el MIPG es la traducción normativa de esta corriente: el "
        "artículo 2.2.22.3.2 lo define como marco de referencia para dirigir, "
        "planear, ejecutar, hacer seguimiento, evaluar y controlar la gestión "
        "pública, y el artículo 2.2.22.3.3 fija sus objetivos, obligando a medir el "
        "desempeño institucional (FURAG, IDI) como base de la toma de decisiones.",
    ),
    (
        "Capacidades Estatales",
        "Las Capacidades Estatales (Fukuyama, 2004; Grindle, 1997) son la aptitud "
        "real de un Estado o entidad para formular e implementar política pública, "
        "y suelen distinguirse al menos tres tipos: capacidad administrativa (talento "
        "humano, sistemas de información, procesos), capacidad fiscal (recursos "
        "disponibles y su ejecución) y capacidad política (legitimidad y respaldo "
        "para decidir y sostener una decisión en el tiempo). Colombia reconoce "
        "explícitamente estas asimetrías de capacidad entre entidades mediante el "
        "artículo 1 de la Ley 617 de 2000 (categorización presupuestal de "
        "departamentos, en desarrollo del artículo 302 de la Constitución) y su "
        "artículo 2 (categorización de distritos y municipios), junto con el "
        "artículo 3 de la Ley 1454 de 2011 (LOOT), que fija la autonomía, la "
        "descentralización y la asociatividad como principios rectores del "
        "ordenamiento territorial — la base normativa de la tipología y el 'Grupo "
        "par' que este sistema usa para no comparar entidades de capacidad muy "
        "distinta entre sí.",
    ),
    (
        "Resiliencia Institucional",
        "La Resiliencia Institucional (OCDE, 2020, en el marco de la respuesta a la "
        "pandemia) es la capacidad de una entidad para anticipar, absorber y "
        "recuperarse de choques —crisis sanitarias, desastres naturales, choques "
        "fiscales— sin perder su capacidad de generar valor público. Es, de todos los "
        "enfoques contemporáneos, el que tiene el respaldo legal colombiano más "
        "directo y antiguo: el artículo 8 de la Ley 1523 de 2012 crea el Sistema "
        "Nacional de Gestión del Riesgo de Desastres, y su artículo 42 obliga "
        "expresamente a toda entidad pública o privada que preste servicios públicos "
        "o desarrolle actividades de riesgo a realizar un análisis específico de "
        "riesgo, mientras el artículo 43 exige el respectivo plan de gestión del "
        "riesgo — desarrollado en detalle por el Decreto 2157 de 2017, que reglamenta "
        "justamente ese artículo 42. Esto incorpora la pregunta de resiliencia "
        "directamente al análisis institucional que hace este informe.",
    ),
    (
        "Agenda 2030, ODS y marcos multilaterales (OCDE, CEPAL, Naciones Unidas)",
        "La Agenda 2030 y los Objetivos de Desarrollo Sostenible (Naciones Unidas, "
        "2015) sitúan la gestión pública territorial dentro de compromisos globales, "
        "en particular el ODS 16 (paz, justicia e instituciones sólidas), que exige "
        "instituciones eficaces, responsables, transparentes y con participación "
        "ciudadana real. Colombia adoptó formalmente esta agenda mediante el "
        "documento CONPES 3918 de 2018 (estrategia de implementación de los ODS en el "
        "país, sin artículos por ser documento de política, no ley), y es miembro "
        "pleno de la OCDE desde 2020, lo que la obliga a alinear su gestión pública "
        "con los estándares técnicos de esa organización. La CEPAL adapta esos "
        "estándares al contexto latinoamericano y sus asimetrías estructurales, y "
        "Naciones Unidas los articula en metas verificables — el marco de referencia "
        "final frente al cual se mide, en último término, el desempeño institucional "
        "que reporta este sistema.",
    ),
    (
        "Gobierno como Plataforma",
        "El Gobierno como Plataforma (O'Reilly, 2010) propone que el Estado deje de "
        "construir un sistema aislado por cada trámite y en su lugar ofrezca "
        "infraestructura digital compartida y reutilizable (identidad, autenticación, "
        "interoperabilidad) sobre la cual cualquier entidad pueda montar sus "
        "servicios, evitando duplicar esfuerzos. Colombia lo implementa de forma "
        "concreta mediante el Decreto 620 de 2020: el artículo 2.2.17.2.2.1 garantiza "
        "el acceso a los servicios ciudadanos digitales base a través del "
        "'Articulador' (la Agencia Nacional Digital), y el artículo 2.2.17.2.2.2 "
        "establece que el servicio de interoperabilidad del Estado se presta de forma "
        "exclusiva a través de ese Articulador — es decir, la plataforma común que "
        "describe la teoría ya es, en Colombia, una obligación normativa concreta, "
        "fundamentada además en el artículo 147 de la Ley 1955 de 2019 (Plan Nacional "
        "de Desarrollo 2018-2022).",
    ),
    (
        "Interoperabilidad de datos entre entidades",
        "La interoperabilidad —la capacidad de que distintos sistemas de información "
        "intercambien datos sin fricción, siguiendo el modelo de referencia de países "
        "como Estonia (X-Road)— es la condición técnica que hace posible el Gobierno "
        "como Plataforma y la Gobernanza de Datos. En Colombia queda regulada por el "
        "mismo Decreto 620 de 2020 (artículo 2.2.17.2.2.2, servicio de "
        "interoperabilidad exclusivo del Articulador) y por el Decreto 1078 de 2015, "
        "que define el 'marco de interoperabilidad' como la estructura común de "
        "principios, recomendaciones y directrices —políticas, legales, "
        "organizacionales, semánticas y técnicas— que orientan el intercambio de "
        "información entre entidades del Estado.",
    ),
    (
        "Co-creación y Design Thinking en política pública",
        "La co-creación y el Design Thinking proponen que la ciudadanía deje de ser "
        "receptora pasiva de un servicio público y pase a codiseñarlo junto con la "
        "entidad, mediante prototipado rápido y validación directa con el usuario "
        "final. En Colombia esta corriente está parcialmente respaldada: el Decreto "
        "767 de 2022 incorpora la 'Innovación Pública Digital' como uno de los "
        "elementos estructurales de la Política de Gobierno Digital, y bajo ese "
        "marco el Centro de Innovación Pública Digital de MinTIC opera la "
        "metodología CoCrearE. Es importante precisar, con honestidad metodológica, "
        "que CoCrearE es una metodología operativa del Centro, no una norma con "
        "artículos propios: su respaldo legal es el Decreto 767 de 2022 que crea el "
        "marco dentro del cual esa metodología funciona.",
    ),
    (
        "Administración Pública Conductual (Nudge)",
        "La Administración Pública Conductual, o enfoque Nudge (Thaler y Sunstein, "
        "2008, \"Nudge: Improving Decisions About Health, Wealth, and Happiness\"), "
        "propone rediseñar la forma en que se presentan las opciones al ciudadano "
        "(el 'arquitecto de decisiones') para facilitar mejores decisiones sin "
        "restringir la libertad de elegir, apoyándose en evidencia de las ciencias "
        "del comportamiento. Es el enfoque contemporáneo que ha tenido más desarrollo "
        "internacional (la Behavioural Insights Team del Reino Unido, la oficina de "
        "Cass Sunstein en la administración Obama en Estados Unidos) pero, siendo "
        "estrictos con la evidencia disponible, no se identificó una ley o CONPES "
        "colombiano que institucionalice de forma específica una unidad o política "
        "de ciencias del comportamiento a nivel nacional. Se incluye aquí como "
        "corriente teórica vigente y de vanguardia, no como corriente ya "
        "normativizada en Colombia — una oportunidad de política pública más que un "
        "hecho normativo consumado.",
    ),
    (
        "Gemelos Digitales de Territorio (Digital Twins)",
        "Un gemelo digital de territorio es una réplica virtual, alimentada con "
        "datos reales y actualizados, de una ciudad o región, que permite simular el "
        "efecto de una decisión de política pública (una obra, una reubicación de "
        "población, un cambio de uso del suelo) antes de ejecutarla. Es, de los cinco "
        "enfoques añadidos en esta ampliación, el más incipiente en Colombia: no se "
        "identificó un decreto, ley o CONPES que regule específicamente los gemelos "
        "digitales de territorio a la fecha de este informe; el CONPES 4144 de "
        "2025 menciona tecnologías emergentes de forma general dentro de la Política "
        "Nacional de Inteligencia Artificial, pero sin un desarrollo propio para "
        "esta técnica en particular. Se incluye por su relevancia para el debate "
        "actual de vanguardia en gestión territorial, dejando explícito que aún es "
        "una oportunidad de política pública pendiente de desarrollo normativo en el "
        "país.",
    ),
]

CONTEXTUALIZACION_ESQUEMA_INTEGRADOR = [
    "Estado Digital",
    "Transformación Digital",
    "Gobernanza Inteligente",
    "Inteligencia Artificial",
    "Gobierno Abierto",
    "Gobernanza de Datos",
    "Valor Público",
    "Administración Pública Basada en Evidencia",
    "Capacidades Estatales",
    "Resiliencia Institucional",
    "Agenda 2030 y Objetivos de Desarrollo Sostenible (ODS)",
]

CONTEXTUALIZACION_CIERRE = (
    "Este esquema integrador evidencia que la Administración Pública contemporánea "
    "no es una ruptura frente a las teorías clásicas, sino su evolución: el Estado "
    "Digital y la Transformación Digital habilitan una Gobernanza Inteligente "
    "apoyada en Inteligencia Artificial; el Gobierno Abierto y la Gobernanza de "
    "Datos fortalecen el Valor Público y una Administración Pública Basada en "
    "Evidencia; y todo ello construye Capacidades Estatales y Resiliencia "
    "Institucional, en línea con la Agenda 2030 y los ODS y los marcos de "
    "referencia de organismos internacionales como la OCDE, la CEPAL y las "
    "Naciones Unidas, que orientan a los países en la modernización de su gestión "
    "pública hacia la innovación, la transparencia y la generación de valor "
    "público."
)

CADENA_INTERPRETACION_TITULO = "Cadena de interpretación institucional del SIIEAP"

CADENA_INTERPRETACION_INTRO = (
    "El SIIEAP no se limita a contar recomendaciones de Función Pública: cada "
    "brecha detectada se interpreta siguiendo una cadena de análisis que va desde "
    "el dato oficial hasta el plan de tratamiento y su seguimiento posterior. Esta "
    "es la ruta que sigue el sistema para cada brecha priorizada de la entidad:"
)

CADENA_INTERPRETACION_PASOS = [
    "Pregunta FURAG",
    "Dimensión – Política – Índice MIPG",
    "Estándar metodológico esperado",
    "Resultado oficial MIPG",
    "Recomendación oficial de Función Pública",
    "Brecha de implementación identificada",
    "Interpretación desde las teorías y enfoques de la Administración Pública",
    "Paradigma administrativo predominante",
    "Modelo de gestión pública relacionado",
    "Capacidad estatal comprometida",
    "Valor público afectado",
    "Gobernanza comprometida",
    "Transformación digital involucrada",
    "Normatividad aplicable",
    "Riesgo institucional",
    "Impacto en: índice, política, dimensión, IDI, capacidades institucionales y valor público",
    "Plan de tratamiento",
    "Indicadores KPI e indicadores KRI",
    "Seguimiento",
    "Nueva medición FURAG",
]

CADENA_INTERPRETACION_CIERRE = (
    "Esta cadena evita el error de tratar todas las recomendaciones como "
    "equivalentes: lo que importa no es cuántas brechas tiene una entidad, sino "
    "qué tan lejos está cada una del estándar exigido, qué capacidad estatal y "
    "qué valor público compromete, y qué tan bien se cierra el ciclo hasta la "
    "siguiente medición FURAG. El análisis de riesgo que se presenta más "
    "adelante en este informe sigue esta misma lógica de distancia al estándar, "
    "no de simple conteo."
)

DISCLAIMER_INFORME = (
    "Este informe fue generado por el Sistema Integral de Diagnóstico del "
    "Desempeño Institucional (SIIEAP), a partir de datos reales del Índice de "
    "Desempeño Institucional (IDI-MIPG) de Función Pública, complementado con un "
    "análisis generado por inteligencia artificial (Claude, Anthropic) como punto "
    "de partida académico y metodológico. No constituye un dictamen oficial de "
    "Función Pública ni sustituye la validación técnica, jurídica y del líder de "
    "proceso de la entidad analizada. Se entrega con fines académicos, como "
    "evidencia de ejecución del sistema y como insumo de trabajo para la "
    "asignatura."
)

# ---------------------------------------------------------------------------
# Glosario y convenciones del informe
#
# Fuente de los términos OFICIALES: Glosario MIPG v7 (octubre de 2021),
# Departamento Administrativo de la Función Pública — verificado línea por
# línea contra el PDF oficial (2021-10-25_Glosario_mipg_v7.pdf) el 8 de
# agosto de 2026. Definiciones parafraseadas aquí, no citadas textualmente.
# Se listan aparte los términos de CONVENCIÓN PROPIA de SIIEAP, que NO
# existen en el glosario oficial y son exclusivos de este sistema, para que
# el lector nunca confunda unos con otros.
#
# Correcciones aplicadas en la verificación del 8 de agosto de 2026:
#   - MIPG: el glosario oficial v7 remite a Decreto 1083 de 2015 (Único
#     Reglamentario), no a Decreto 1499 de 2017 como decía la versión
#     anterior. Se citan ambos para no perder la referencia de creación.
#   - Grupo par: se añadió el mecanismo de quintiles, presente en el
#     glosario oficial y ausente en la versión anterior de este código.
#   - Valor público: se ajustó la redacción para no editorializar más allá
#     de lo que dice la fuente oficial.
#   - Se agregaron 6 términos oficiales adicionales que ya se usan en los
#     tres informes del SIIEAP pero no tenían entrada en este glosario:
#     Grupos de valor, Talento humano, Autocontrol, Rendición de cuentas,
#     Transparencia activa y Transparencia pasiva.
# ---------------------------------------------------------------------------

GLOSARIO_OFICIAL_MIPG = [
    (
        "MIPG (Modelo Integrado de Planeación y Gestión)",
        "Marco de referencia oficial del Estado colombiano para dirigir, "
        "planear, ejecutar, hacer seguimiento, evaluar y controlar la "
        "gestión de las entidades públicas, con integridad y calidad en el "
        "servicio, con el fin de generar resultados que atiendan los planes "
        "de desarrollo y resuelvan las necesidades de la ciudadanía. Creado "
        "por el Decreto 1499 de 2017, hoy compilado en el Decreto 1083 de "
        "2015 (Único Reglamentario del sector Función Pública), que es la "
        "norma que cita expresamente el glosario oficial vigente.",
    ),
    (
        "Índice de Desempeño Institucional (IDI)",
        "Cifra oficial, publicada por Función Pública, que resume qué tan "
        "orientada está una entidad hacia la eficacia, la eficiencia y la "
        "calidad de su gestión. Es siempre la cifra de referencia principal en "
        "este informe.",
    ),
    (
        "FURAG (Formulario Único de Reporte y Avance de Gestión)",
        "Cuestionario oficial en línea, aplicado anualmente por Función "
        "Pública, con el que se recolecta la información que luego produce el "
        "IDI y mide el avance de las políticas de MIPG y del Modelo Estándar "
        "de Control Interno (MECI).",
    ),
    (
        "MECI (Modelo Estándar de Control Interno)",
        "Estructura oficial para evaluar la estrategia, la gestión y los "
        "mecanismos de autoevaluación de una entidad, adaptable a las "
        "necesidades propias de cada una.",
    ),
    (
        "Dimensión",
        "Cada uno de los grandes bloques de políticas y prácticas de gestión "
        "en que se organiza MIPG (por ejemplo, Talento Humano, Direccionamiento "
        "Estratégico, Evaluación de Resultados). El IDI se calcula agregando "
        "el desempeño de la entidad en estas dimensiones.",
    ),
    (
        "Grupo par",
        "Agrupación oficial de entidades con características homogéneas que "
        "Función Pública usa para que los resultados del FURAG sean "
        "comparables entre entidades similares. Dentro de cada grupo par, los "
        "resultados se ordenan y se subagrupan en quintiles: cinco categorías "
        "de igual tamaño, cada una con el 20 % de las entidades del grupo.",
    ),
    (
        "Valor público",
        "Cambios sociales observables y medibles que el Estado produce como "
        "respuesta a las necesidades y demandas de la ciudadanía, legitimados "
        "democráticamente. Son, en últimas, los resultados que la gestión "
        "pública busca alcanzar.",
    ),
    (
        "ODS (Objetivos de Desarrollo Sostenible)",
        "Los 17 objetivos globales adoptados por la ONU en 2015 para poner fin "
        "a la pobreza, proteger el planeta y garantizar la prosperidad para "
        "2030; se usan aquí como marco de referencia para la prospectiva.",
    ),
    (
        "Grupos de valor",
        "Personas naturales (ciudadanos) o jurídicas —públicas o privadas— a "
        "quienes se dirigen los bienes y servicios de una entidad, y que "
        "reciben directamente los resultados de su gestión.",
    ),
    (
        "Talento humano",
        "El activo más importante de una entidad pública: todas las personas "
        "que, en el marco de los valores del servicio público, contribuyen "
        "con su trabajo al cumplimiento de la misión estatal y a responder "
        "las demandas de la ciudadanía.",
    ),
    (
        "Autocontrol",
        "Capacidad que debe desarrollar todo servidor público, sin importar "
        "su nivel jerárquico, para evaluar y controlar su propio trabajo, "
        "detectar desviaciones y corregirlas oportunamente en el ejercicio de "
        "sus funciones.",
    ),
    (
        "Rendición de cuentas",
        "Conjunto de normas, procedimientos y prácticas mediante los cuales "
        "las entidades públicas y sus servidores informan y explican a la "
        "ciudadanía, a la sociedad civil y a los organismos de control los "
        "resultados de su gestión. Es, en sí misma, una expresión de control "
        "social.",
    ),
    (
        "Transparencia activa",
        "Obligación de publicar proactivamente información de la entidad, "
        "sin que medie una solicitud, a través de los medios oficiales "
        "(sitio web, carteleras, etc.).",
    ),
    (
        "Transparencia pasiva",
        "Obligación de la entidad de gestionar y responder, dentro de los "
        "plazos legales, las solicitudes de información que presente la "
        "ciudadanía.",
    ),
]

GLOSARIO_CONVENCIONES_SIIEAP = [
    (
        "IDI-MIPG (como se usa en este informe)",
        "Forma abreviada, propia de este sistema, para referirse al Índice de "
        "Desempeño Institucional dentro del marco de MIPG. No es una sigla "
        "oficial de Función Pública, aunque el IDI que reporta sí lo es.",
    ),
    (
        "Brecha (detectada por SIIEAP)",
        "Convención EXCLUSIVA de este sistema: un índice MIPG de la entidad "
        "con puntaje por debajo de 60 puntos, umbral de alerta metodológico "
        "interno que SIIEAP usa para priorizar qué revisar primero. NO es una "
        "cifra que publique Función Pública, y no debe confundirse con el IDI "
        "oficial ni con las recomendaciones oficiales del banco consolidado.",
    ),
    (
        "ISVPT (Índice Sintético de Valor Público Territorial)",
        "Indicador propio de este informe, construido con metodología "
        "académica (normalización min-max y agregación simple, siguiendo "
        "directrices de la OCDE de 2008 para indicadores compuestos), que "
        "compara las 7 dimensiones del IDI-MIPG de la entidad dentro de su "
        "grupo par. No es un índice oficial de Función Pública; es un "
        "ejercicio complementario de este sistema.",
    ),
    (
        "Umbral de alerta metodológico interno (60 puntos)",
        "El punto de corte que SIIEAP usa internamente para marcar un índice "
        "como brecha. Se insiste en todo el informe en que la meta real y "
        "oficial de la gestión pública es siempre el 100%, nunca el 60%.",
    ),
    (
        "Recomendaciones oficiales FP",
        "A diferencia de las brechas, este dato SÍ es oficial: es el total de "
        "recomendaciones del banco consolidado del Departamento Administrativo "
        "de la Función Pública, emitidas específicamente para la entidad "
        "analizada.",
    ),
]


def _agregar_glosario_docx(doc, estilo_normal=None):
    """Inserta la sección 'Glosario y convenciones' en el documento Word,
    en formato de MATRIZ (tabla de dos columnas: Término | Definición),
    separando siempre los términos OFICIALES (Glosario MIPG v7, Función
    Pública) de las CONVENCIONES PROPIAS de SIIEAP, para que el lector nunca
    confunda un dato exclusivo de este informe con un dato oficial.
    """
    doc.add_heading("Glosario y convenciones de este informe", level=1)
    doc.add_paragraph(
        "Los términos técnicos usados a lo largo de este informe se explican "
        "aquí en dos matrices: primero los términos OFICIALES del Modelo "
        "Integrado de Planeación y Gestión, tal como los define el Glosario "
        "MIPG versión 7 (octubre de 2021) del Departamento Administrativo de "
        "la Función Pública; y luego las CONVENCIONES PROPIAS de este sistema "
        "(SIIEAP), que no existen en el glosario oficial y no deben "
        "confundirse con cifras que publique Función Pública."
    )

    def _tabla_glosario_docx(lista_terminos, color_encabezado, color_texto_encabezado=(0xFF, 0xFF, 0xFF)):
        tabla = doc.add_table(rows=1, cols=2)
        tabla.style = "Light Grid Accent 1"
        enc = tabla.rows[0].cells
        for celda, texto in zip(enc, ["Término", "Definición"]):
            celda.text = texto
            _sombrear_celda(celda, color_encabezado)
            for parrafo in celda.paragraphs:
                for run_enc in parrafo.runs:
                    run_enc.bold = True
                    run_enc.font.color.rgb = RGBColor(*color_texto_encabezado)
        for termino, definicion in lista_terminos:
            fila = tabla.add_row().cells
            fila[0].text = ""
            run_term = fila[0].paragraphs[0].add_run(termino)
            run_term.bold = True
            fila[1].text = definicion
        _ajustar_tabla_docx(tabla, anchos_cm=[4.5, 12.5], tamano_fuente_pt=9)
        _franjas_alternas_docx(tabla)
        doc.add_paragraph()

    doc.add_heading("Términos oficiales (Glosario MIPG v7, Función Pública)", level=2)
    _tabla_glosario_docx(GLOSARIO_OFICIAL_MIPG, COLOR_INSTITUCIONAL)

    doc.add_heading("Convenciones propias de SIIEAP (NO oficiales)", level=2)
    _tabla_glosario_docx(GLOSARIO_CONVENCIONES_SIIEAP, "B85C00")
    doc.add_page_break()


def _agregar_glosario_pdf(elementos, estilos, estilo_normal, estilo_h2):
    """Versión PDF (reportlab) de _agregar_glosario_docx: mismo formato de
    MATRIZ (tabla de dos columnas) y misma separación estricta entre
    términos oficiales (Glosario MIPG v7, Función Pública) y convenciones
    propias de SIIEAP."""
    estilo_celda = ParagraphStyle(
        "GlosarioCelda", parent=estilo_normal, fontSize=8.3, leading=10.5,
    )
    estilo_celda_termino = ParagraphStyle(
        "GlosarioTermino", parent=estilo_celda, fontName="Helvetica-Bold",
    )
    estilo_encabezado_tabla = ParagraphStyle(
        "GlosarioEncabezado", parent=estilo_celda, fontName="Helvetica-Bold",
        textColor=colors.white,
    )

    def _tabla_glosario_pdf(lista_terminos, color_hex_encabezado):
        filas = [[
            Paragraph("Término", estilo_encabezado_tabla),
            Paragraph("Definición", estilo_encabezado_tabla),
        ]]
        for termino, definicion in lista_terminos:
            filas.append([
                Paragraph(termino, estilo_celda_termino),
                Paragraph(definicion, estilo_celda),
            ])
        tabla = Table(filas, colWidths=[4.2 * cm, 12.8 * cm], repeatRows=1)
        estilo_tabla = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{color_hex_encabezado}")),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D0D0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]
        for i in range(1, len(filas)):
            if i % 2 == 0:
                estilo_tabla.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#F2F4F8")))
        tabla.setStyle(TableStyle(estilo_tabla))
        return tabla

    elementos.append(Paragraph("Glosario y convenciones de este informe", estilos["Heading1"]))
    elementos.append(Paragraph(
        "Los términos técnicos usados a lo largo de este informe se explican aquí en dos "
        "matrices: primero los términos OFICIALES del Modelo Integrado de Planeación y "
        "Gestión, tal como los define el Glosario MIPG versión 7 (octubre de 2021) del "
        "Departamento Administrativo de la Función Pública; y luego las CONVENCIONES "
        "PROPIAS de este sistema (SIIEAP), que no existen en el glosario oficial y no deben "
        "confundirse con cifras que publique Función Pública.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph("Términos oficiales (Glosario MIPG v7, Función Pública)", estilo_h2))
    elementos.append(Spacer(1, 4))
    elementos.append(_tabla_glosario_pdf(GLOSARIO_OFICIAL_MIPG, COLOR_INSTITUCIONAL))
    elementos.append(Spacer(1, 12))
    elementos.append(Paragraph("Convenciones propias de SIIEAP (NO oficiales)", estilo_h2))
    elementos.append(Spacer(1, 4))
    elementos.append(_tabla_glosario_pdf(GLOSARIO_CONVENCIONES_SIIEAP, "B85C00"))
    elementos.append(PageBreak())


# ---------------------------------------------------------------------------
# Anexo: Normativa vigente por política MIPG (las 19 políticas)
#
# Se agrega en los 3 informes (Técnico, Estudio de Caso, Ejecutivo), en
# Word y PDF, como sustento jurídico DIRECTO de dos cosas que pide la
# entidad: (a) que las brechas priorizadas por SIIEAP tengan como guía la
# normativa vigente de esa política concreta, y (b) que el plan de
# mejoramiento cite esa misma normativa como sustento de cada acción.
#
# Por eso la tabla se divide en dos bloques, siempre en este orden:
#   1. Políticas CON brecha detectada en la entidad (si hay diagnóstico
#      real disponible): aparecen primero y resaltadas, porque son las que
#      de verdad hay que usar para estructurar el plan de mejoramiento.
#   2. El resto de las 19 políticas, como anexo de referencia completo.
#
# Fuente y alcance: ver docstring de
# backend/base_conocimiento/normativa_politicas.py — no es una
# verificación exhaustiva artículo por artículo; prioriza los cambios
# normativos de 2022 en adelante.
# ---------------------------------------------------------------------------

def _politicas_con_brecha(diag=None) -> set[str]:
    """Códigos de política (POLxx) con al menos una brecha detectada para
    la entidad, según el diagnóstico real. Si no hay diagnóstico
    disponible, retorna un conjunto vacío (no se resalta nada)."""
    if diag is None or not getattr(diag, "brechas", None):
        return set()
    return {b.codigo_politica for b in diag.brechas}


def _agregar_normativa_politicas_docx(doc, diag=None):
    doc.add_heading("Normativa vigente por política MIPG", level=1)
    doc.add_paragraph(
        "Esta sección consolida, política por política, la normatividad que la "
        "entidad reportó como vigente para cada una de las 19 políticas de "
        "Gestión y Desempeño de MIPG. Es la referencia jurídica directa que "
        "este informe usa para priorizar acciones frente a las brechas "
        "detectadas y para sustentar el plan de mejoramiento. Las políticas con "
        "brecha detectada en esta entidad aparecen primero y resaltadas."
    )

    codigos_con_brecha = _politicas_con_brecha(diag)
    todas = todas_las_politicas_con_normativa()
    orden = sorted(
        todas.keys(),
        key=lambda c: (c not in codigos_con_brecha, int(c[3:])),
    )

    if codigos_con_brecha:
        doc.add_heading(
            "Políticas con brecha detectada en esta entidad (prioridad para el plan de mejoramiento)",
            level=2,
        )

    bloque_actual = "con_brecha" if codigos_con_brecha else None
    for codigo in orden:
        info = todas[codigo]
        es_con_brecha = codigo in codigos_con_brecha
        if codigos_con_brecha and bloque_actual == "con_brecha" and not es_con_brecha:
            doc.add_heading("Resto de políticas (anexo de referencia completo)", level=2)
            bloque_actual = "resto"

        color_encabezado = "C0392B" if es_con_brecha else COLOR_INSTITUCIONAL
        titulo = f"{codigo} — {info['nombre']}" + ("  ⚠ BRECHA DETECTADA" if es_con_brecha else "")
        doc.add_heading(titulo, level=3)

        tabla = doc.add_table(rows=1, cols=1)
        tabla.style = "Light Grid Accent 1"
        enc = tabla.rows[0].cells[0]
        enc.text = "Normativa"
        _sombrear_celda(enc, color_encabezado)
        for parrafo in enc.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for norma in info["normas"]:
            fila = tabla.add_row().cells[0]
            fila.text = norma
        _ajustar_tabla_docx(tabla, anchos_cm=[17.0], tamano_fuente_pt=8.5)
        doc.add_paragraph()

    doc.add_page_break()


def _agregar_normativa_politicas_pdf(elementos, estilos, estilo_normal, estilo_h2, diag=None):
    elementos.append(Paragraph("Normativa vigente por política MIPG", estilos["Heading1"]))
    elementos.append(Paragraph(
        "Esta sección consolida, política por política, la normatividad que la entidad "
        "reportó como vigente para cada una de las 19 políticas de Gestión y Desempeño de "
        "MIPG. Es la referencia jurídica directa que este informe usa para priorizar "
        "acciones frente a las brechas detectadas y para sustentar el plan de mejoramiento. "
        "Las políticas con brecha detectada en esta entidad aparecen primero y resaltadas.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 8))

    estilo_norma = ParagraphStyle("NormaCelda", parent=estilo_normal, fontSize=8.2, leading=10.2)

    codigos_con_brecha = _politicas_con_brecha(diag)
    todas = todas_las_politicas_con_normativa()
    orden = sorted(
        todas.keys(),
        key=lambda c: (c not in codigos_con_brecha, int(c[3:])),
    )

    if codigos_con_brecha:
        elementos.append(Paragraph(
            "Políticas con brecha detectada en esta entidad (prioridad para el plan de mejoramiento)",
            estilo_h2,
        ))
        elementos.append(Spacer(1, 4))

    bloque_actual = "con_brecha" if codigos_con_brecha else None
    for codigo in orden:
        info = todas[codigo]
        es_con_brecha = codigo in codigos_con_brecha
        if codigos_con_brecha and bloque_actual == "con_brecha" and not es_con_brecha:
            elementos.append(Paragraph("Resto de políticas (anexo de referencia completo)", estilo_h2))
            elementos.append(Spacer(1, 4))
            bloque_actual = "resto"

        color_hex = "C0392B" if es_con_brecha else COLOR_INSTITUCIONAL
        titulo = f"{codigo} — {info['nombre']}" + ("  ⚠ BRECHA DETECTADA" if es_con_brecha else "")
        elementos.append(Paragraph(titulo, estilo_h2))
        elementos.append(Spacer(1, 3))

        filas = [[Paragraph("Normativa", ParagraphStyle(
            "EncNorma", parent=estilo_norma, fontName="Helvetica-Bold", textColor=colors.white,
        ))]]
        for norma in info["normas"]:
            filas.append([Paragraph(norma, estilo_norma)])
        tabla = Table(filas, colWidths=[17.0 * cm], repeatRows=1)
        estilo_tabla = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{color_hex}")),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D0D0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]
        for i in range(1, len(filas)):
            if i % 2 == 0:
                estilo_tabla.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#F2F4F8")))
        tabla.setStyle(TableStyle(estilo_tabla))
        elementos.append(tabla)
        elementos.append(Spacer(1, 10))

    elementos.append(PageBreak())


# ---------------------------------------------------------------------------
# Anexo: Marco de descentralización 2026-2030
#
# Compartido por los TRES informes del SIIEAP (Técnico, Estudio de Caso,
# Ejecutivo). Resume, en formato de matriz, la Misión de Descentralización,
# el Acto Legislativo 3 de 2024 (VIGENTE) y el Proyecto de Ley de
# Competencias (EN TRÁMITE, texto puede cambiar), con la distribución del
# SGP hoy y la meta del 39,5 %, y el paralelo entre la Ley 617 de 2000, la
# Ley 715 de 2001 (SGP) y la Ley 2056 de 2020 (SGR — regalías).
#
# LEY_COMPETENCIAS_VIGENTE = False: bandera explícita de que el proyecto de
# ley AÚN NO es derecho vigente al momento de generar este informe. Debe
# actualizarse a True (y revisarse el contenido) el día en que el Congreso
# sancione la Ley de Competencias, o eliminarse el anexo si el proyecto se
# archiva.
# ---------------------------------------------------------------------------

LEY_COMPETENCIAS_VIGENTE = False

MARCO_DESCENTRALIZACION_INTRO = (
    "Este anexo incorpora, en su totalidad, el Capítulo Especial de Descentralización del "
    "SIIEAP: el marco de descentralización territorial que enmarca la labor de esta entidad, "
    "con independencia de quién ocupe la Presidencia o la alcaldía/gobernación de turno. Es "
    "resultado de un proceso institucional de más de cuatro años (Misión de Descentralización, "
    "2022-2024) y de actos del Congreso de la República. En todo el capítulo se distingue "
    "expresamente entre lo que YA ES norma vigente y lo que ES UN PROYECTO EN TRÁMITE."
)

MARCO_ENCUADRE_METODOLOGICO = (
    "Antes de entrar en materia: lo que aquí importa no es quién ocupa la Presidencia, sino "
    "que el régimen de competencias que se analiza es resultado de un proceso institucional "
    "que lleva más de cuatro años en construcción —la Misión de Descentralización (2022-2024), "
    "con participación del DNP, el Ministerio de Hacienda, el Ministerio del Interior, "
    "alcaldes, gobernadores, pueblos indígenas y organismos de cooperación internacional— y de "
    "actos del Congreso de la República aprobados en dos legislaturas sucesivas —el Acto "
    "Legislativo 3 de 2024 y, en trámite, el Proyecto de Ley de Competencias—. Ese proceso no "
    "depende de una persona ni de un gobierno en particular, y continuaría con independencia "
    "de quién gobierne."
)

MARCO_CONTINUIDAD_7AGOSTO = (
    "El 7 de agosto de 2026 tomó posesión como presidente de la República Abelardo de la "
    "Espriella. La transmisión de mando —realizada en Cali, primera vez fuera de Bogotá— es "
    "relevante para el SIIEAP no por el mandatario que asumió, sino porque CONFIRMA que la "
    "construcción institucional descrita sigue en pie: hubo juramento ante el Congreso, respeto "
    "al período constitucional, y un compromiso explícito de gobernar dentro del marco "
    "constitucional vigente, sin Asamblea Constituyente. Eso es, en los términos del artículo 1 "
    "de la Constitución, la manifestación misma del \"Estado social de derecho\" — el orden "
    "institucional funcionando — y no un mérito atribuible a una persona. El discurso se cita "
    "aquí únicamente como evidencia de que el nuevo gobierno se sitúa dentro del marco "
    "normativo ya construido, no como fuente de autoridad jurídica."
)

MARCO_ARTICULOS_CP = [
    ("Artículo 1, C.P.", "\"Colombia es un Estado social de derecho, organizado en forma de "
     "República unitaria, descentralizada, con autonomía de sus entidades territoriales, "
     "democrática, participativa y pluralista (...)\""),
    ("Artículo 2, C.P.", "\"Son fines esenciales del Estado: servir a la comunidad, promover la "
     "prosperidad general (...); facilitar la participación de todos en las decisiones que los "
     "afectan (...) y asegurar la convivencia pacífica y la vigencia de un orden justo.\""),
]

MARCO_MISION_EJES = [
    ("1. Competencias entre niveles de gobierno", "Definir y distribuir competencias entre Nación y entidades territoriales."),
    ("2. Fuentes y usos de los recursos para el desarrollo", "Proponer fuentes de financiación territorial articuladas con la nueva distribución de competencias."),
    ("3. Arquitectura institucional y modernización de la administración pública", "Revisar la estructura institucional del Estado para hacerla coherente con el nuevo reparto de funciones."),
    ("4. Estado abierto y participación ciudadana territorial", "Fortalecer la transparencia, la rendición de cuentas y la participación ciudadana territorial."),
    ("5. Descentralización y territorios indígenas", "Definir las acciones para el ordenamiento, la planeación y la institucionalidad de los territorios indígenas."),
]

MARCO_MISION_PROPUESTAS = (
    "El informe final (firmado en junio de 2024, publicado por el DNP el 5 de agosto de 2024 y "
    "copublicado por el PNUD Colombia el 6 de agosto de 2024) formuló 9 propuestas de reforma: "
    "1) nuevas categorías de entidades territoriales; 2) nueva Ley Orgánica de Ordenamiento "
    "Territorial; 3) política de arquitectura institucional; 4) reforma al Sistema General de "
    "Participaciones (SGP); 5) creación del Fondo de Convergencia Económica Territorial (FECET); "
    "6) modificaciones al Sistema General de Regalías (SGR); 7) reforma al Marco de "
    "Responsabilidad Fiscal Subnacional; 8) Estado abierto y participación ciudadana; y 9) "
    "caminos para la conformación de los territorios indígenas."
)

MARCO_COOPERACION = [
    ("PNUD", "Socio técnico principal: cofinanció y coejecutó la ruta de diálogo territorial (68 espacios, 2.200+ personas); copublicó el informe final."),
    ("USAID y Agencia Francesa de Desarrollo", "Cooperación técnica y financiera para los estudios de arquitectura institucional y Estado abierto."),
    ("OCDE", "Citada como fuente metodológica directa (2019); Colombia es miembro pleno desde 2020."),
    ("BID", "Citado como fuente (2019); financia además un programa de apoyo al cumplimiento de compromisos OCDE de Colombia, con misión técnica activa en julio de 2026."),
    ("OEA", "Citada como fuente conceptual (Secretaría General, 2008) sobre riesgos de asimetría de capacidades entre gobiernos locales."),
]

MARCO_COMPETENCIAS_NACION = [
    ("5", "Fortalecimiento institucional y cooperación técnica", "El plan de mejoramiento que ya genera el SIIEAP es, en la práctica, un insumo de fortalecimiento institucional."),
    ("6", "Indicadores para clasificar entidades territoriales", "El IDI-MIPG y sus 69 índices ya son un sistema de indicadores de capacidad institucional listo para esa clasificación."),
    ("8", "Monitoreo, seguimiento y control con alertas tempranas", "Las brechas priorizadas que detecta el SIIEAP son, de hecho, alertas tempranas de riesgo institucional."),
    ("9", "Participación ciudadana y Estado abierto", "Las políticas de Participación Ciudadana y Transparencia, ya evaluadas por el SIIEAP, son exactamente ese frente."),
    ("11", "Coherencia metodológica e integración de indicadores", "El banco de recomendaciones consolidado de Función Pública, que el SIIEAP ya integra, es un ejercicio de esa misma coherencia."),
    ("13", "Ajustes a la estructura de la administración pública", "La Política de Fortalecimiento Organizacional, ya evaluada por el SIIEAP, mide justo esa capacidad de ajuste."),
]

MARCO_MATRIZ_DIMENSIONES = [
    ("D1. Talento Humano", "Art. 9 (capacidades institucionales); art. 15 núm. 5", "Diagnóstico por índice con normativa de Ley 909/2004, Decreto 1083/2015 y Ley 1960/2019."),
    ("D2. Direccionamiento Estratégico", "Art. 180 (ruta de trabajo del Programa de Fortalecimiento)", "Diagnóstico de POL03 como insumo directo de esa ruta de trabajo."),
    ("D3. Gestión para Resultados", "Cap. III, Tít. II (cierre de brechas)", "El concepto \"brecha SIIEAP\" opera igual que el \"cierre de brechas\" que exige la ley en trámite."),
    ("D4. Evaluación de Resultados", "Art. 15 núm. 8 (alertas tempranas)", "POL14 es, hoy, una alerta temprana por entidad."),
    ("D5. Información y Comunicación", "Art. 186 (SUI-SGP)", "Base de datos estructurada, interoperable en el mediano plazo con el SUI-SGP."),
    ("D6. Gestión del Conocimiento", "Art. 15 núm. 11 (coherencia metodológica)", "Integración del banco de recomendaciones consolidado dentro de cada informe."),
    ("D7. Control Interno", "Art. 15 núm. 8; art. 181", "Diagnóstico de ambiente de control, riesgo y monitoreo, ya conectado a Ley 87/1993."),
]

MARCO_DESCENTRALIZACION_TIMELINE = [
    ("2019", "Ley 1962 de 2019 (Ley de Regiones)", "Ordena crear la Misión de Descentralización."),
    ("Dic. 2021", "Decreto 1665 de 2021", "Crea la Misión de Descentralización (5 componentes)."),
    ("Ago. 2024", "Informe final de la Misión", "9 propuestas de reforma; publicado por DNP y PNUD."),
    ("Dic. 2024", "Acto Legislativo 3 de 2024 — VIGENTE", "Reforma arts. 356-357 C.P.; incrementa el SGP hasta 39,5 % de los ICN."),
    ("Dic. 2025", "Proyecto de Ley Orgánica de Competencias — EN TRÁMITE", "Reglamenta el Acto Legislativo; su articulado puede cambiar."),
]

MARCO_SGP_DISTRIBUCION = [
    ("Educación", "58,5 %"),
    ("Salud", "24,5 %"),
    ("Agua Potable y Saneamiento Básico", "5,4 %"),
    ("Propósito General", "11,6 %"),
]

MARCO_SGP_CRECIMIENTO = [
    ("Cierre de brechas sociales, económicas E INSTITUCIONALES", "80 %"),
    ("Propósito general (70 % del 20 % restante)", "14 %"),
    ("Desarrollo económico (30 % del 20 % restante)", "6 %"),
]

MARCO_CATEGORIZACION_NUEVA = [
    ("1. Capacidades fiscales", "Ingresos propios y dependencia de transferencias"),
    ("2. Capacidades institucionales", "Exactamente lo que mide el IDI-MIPG y, por tanto, el SIIEAP"),
    ("3. Densidad poblacional", "Concentración de población en el territorio"),
    ("4. Conectividad territorial", "Vías, infraestructura y acceso físico al territorio"),
]

MARCO_PARALELO_LEYES = [
    ("Objeto", "Categorizar entidades territoriales (gasto público, sostenibilidad fiscal)", "Distribuir competencias y recursos del SGP", "Regular el SGR (regalías)"),
    ("Fundamento C.P.", "Arts. 302 y 320", "Arts. 356 y 357", "Arts. 360 y 361"),
    ("Última reforma", "Ley 2082 de 2021 (coexiste)", "Ley 1176 de 2007 (fijó 58,5/24,5/5,4/11,6 %)", "Reforma integral de 2020 (reemplazó Ley 1530/2012)"),
    ("¿La toca la reforma 2024-2027?", "NO se deroga; la nueva categorización SE SUMA (art. 9, parág. 4)", "SÍ — es su objeto central", "NO — el SGR queda fuera de esta reforma"),
]

MARCO_PLAN_TRES_NIVELES = [
    ("Estratégico", "Alta dirección nacional y territorial", "Priorizar, con base en las brechas del SIIEAP, qué entidades necesitan asunción más gradual de nuevas competencias; usar el histórico de informes como insumo para el informe técnico trienal del DNP al Congreso (art. 10, parág. 4)."),
    ("Táctico", "Equipos técnicos y de planeación, secretarías departamentales", "Convertir el plan de mejoramiento del SIIEAP en la \"ruta de trabajo\" formal del art. 180; articular el cronograma de cierre de brechas con la Estrategia de Monitoreo del SGP."),
    ("Operacional", "Servidores públicos responsables de cada índice/política", "Las recomendaciones que ya genera el SIIEAP por índice siguen siendo el nivel operativo — ahora con relevancia legal directa frente a la futura categorización territorial."),
]

MARCO_INVERSION_POR_DIMENSION = {
    "D1": ("Plan de capacitación (convenios ESAP/SENA), fortalecimiento de la carrera administrativa", "Propósito general"),
    "D2": ("Fortalecimiento de la planeación institucional y presupuestal, articulación con el Marco de Gasto de Mediano Plazo", "Cierre de brechas institucionales"),
    "D3": ("Conectividad, arquitectura empresarial, tablero de indicadores para decisiones basadas en datos, fortalecimiento del servicio al ciudadano", "Cierre de brechas institucionales"),
    "D4": ("Sistema de seguimiento y evaluación del desempeño institucional, tablero de control", "Cierre de brechas institucionales"),
    "D5": ("Infraestructura de gestión documental y archivo, fortalecimiento de registros administrativos", "Cierre de brechas institucionales"),
    "D6": ("Repositorio digital institucional, capacitación en gestión del conocimiento, analítica institucional básica", "Cierre de brechas institucionales"),
    "D7": ("Fortalecimiento del equipo de control interno, mapa de riesgos, software de seguimiento", "Cierre de brechas institucionales"),
}

MARCO_YOUTUBE_ESAP = [
    ("Foro Internacional MIPG — retos y oportunidades para la gestión pública territorial", "Escuela de Alto Gobierno ESAP + DAFP · 2 dic. 2025", "https://www.youtube.com/watch?v=L4SbFpWNN78"),
    ("Diálogos Territoriales de la Escuela de Alto Gobierno con la Gerencia Pública", "Escuela de Alto Gobierno ESAP", "https://www.youtube.com/watch?v=beRJf0qM880"),
    ("Playlist completa — Escuela de Alto Gobierno (60 videos)", "Foros, diálogos territoriales y capacitaciones", "https://www.youtube.com/playlist?list=PL3Wr6Pblz26PULiu-YVVhs2tWqVGmWRbu"),
    ("Canal oficial ESAP", "Para explorar todas las sesiones", "https://www.youtube.com/channel/UCsDXXFW16Dxe7QqXHz6_Yhg"),
]


def _filas_inversion_brechas(diag, top_n=5):
    """Construye, a partir del diagnóstico REAL de la entidad (diag), las filas
    de la tabla 'en qué invertir, brecha por brecha' — dinámica por entidad,
    no un ejemplo fijo. Toma las dimensiones con peor promedio (excluyendo las
    que no tienen información) y las cruza con MARCO_INVERSION_POR_DIMENSION.
    """
    filas = []
    if diag is None:
        return filas
    dims_validas = [d for d in diag.resultados_por_dimension if d.promedio is not None]
    dims_ordenadas = sorted(dims_validas, key=lambda d: d.promedio)[:top_n]
    for d in dims_ordenadas:
        inversion, canal = MARCO_INVERSION_POR_DIMENSION.get(
            d.codigo, ("Fortalecimiento institucional general de la dimensión", "Cierre de brechas institucionales")
        )
        etiqueta = f"{d.nombre} ({d.promedio:.2f})"
        filas.append((etiqueta, inversion, canal))
    return filas


def _agregar_marco_descentralizacion_docx(doc, diag=None, nombre_entidad=None):
    """Inserta el Capítulo Especial de Descentralización COMPLETO (no solo un
    resumen), compartido por los tres informes del SIIEAP. Personaliza la
    tabla de inversión por brecha con los datos reales de `diag`."""
    doc.add_page_break()
    doc.add_heading("Anexo: Capítulo Especial de Descentralización 2026-2030", level=1)
    doc.add_paragraph(MARCO_DESCENTRALIZACION_INTRO)

    # ---- 1. Contexto ----
    doc.add_heading("1. Contexto: continuidad institucional, no personalismo", level=2)
    doc.add_paragraph(MARCO_ENCUADRE_METODOLOGICO)
    doc.add_paragraph(MARCO_CONTINUIDAD_7AGOSTO)
    for etiqueta, texto in MARCO_ARTICULOS_CP:
        p = doc.add_paragraph()
        run_e = p.add_run(f"{etiqueta}: ")
        run_e.bold = True
        run_c = p.add_run(texto)
        run_c.italic = True

    # ---- 2. Misión de Descentralización ----
    doc.add_heading("2. La Misión de Descentralización (2022-2024)", level=2)
    tabla_ejes = doc.add_table(rows=1, cols=2)
    tabla_ejes.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_ejes.rows[0].cells, ["Eje", "Objetivo"]):
        celda.text = texto
        _sombrear_celda(celda, COLOR_INSTITUCIONAL)
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for eje, objetivo in MARCO_MISION_EJES:
        fila = tabla_ejes.add_row().cells
        run_e = fila[0].paragraphs[0].add_run(eje)
        run_e.bold = True
        fila[1].text = objetivo
    _ajustar_tabla_docx(tabla_ejes, anchos_cm=[7.0, 9.0], tamano_fuente_pt=8.7)
    _franjas_alternas_docx(tabla_ejes)
    doc.add_paragraph(MARCO_MISION_PROPUESTAS)

    doc.add_heading("Cooperación multilateral y bilateral", level=3)
    tabla_coop = doc.add_table(rows=1, cols=2)
    tabla_coop.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_coop.rows[0].cells, ["Organismo", "Rol documentado"]):
        celda.text = texto
        _sombrear_celda(celda, COLOR_INSTITUCIONAL)
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for org, rol in MARCO_COOPERACION:
        fila = tabla_coop.add_row().cells
        run_o = fila[0].paragraphs[0].add_run(org)
        run_o.bold = True
        fila[1].text = rol
    _ajustar_tabla_docx(tabla_coop, anchos_cm=[4.5, 11.5], tamano_fuente_pt=8.7)
    _franjas_alternas_docx(tabla_coop)

    # ---- 3. Acto Legislativo ----
    doc.add_heading("3. Acto Legislativo 3 de 2024 (norma constitucional VIGENTE)", level=2)
    doc.add_paragraph(
        "Modifica los arts. 356 y 357 C.P.; incrementa el SGP hasta 39,5 % de los ICN; "
        "ordena que la Ley de Competencias distribuya esos recursos según capacidad "
        "institucional de las entidades territoriales; mientras esa ley no se expida, el SGP "
        "se calcula con la fórmula de transición vigente."
    )

    # ---- 4. Ley de Competencias ----
    doc.add_heading("4. Proyecto de Ley de Competencias (EN TRÁMITE, no vigente)", level=2)
    doc.add_paragraph(
        "Objeto: fortalecer la autonomía y el desarrollo territorial mediante la asignación y "
        "distribución de competencias y recursos entre la Nación, las entidades territoriales y "
        "las beneficiarias del SGP, priorizando salud, educación y agua potable."
    )
    doc.add_heading("Competencias de la Nación que dialogan con el SIIEAP (art. 15)", level=3)
    tabla_comp = doc.add_table(rows=1, cols=3)
    tabla_comp.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_comp.rows[0].cells, ["Núm.", "Competencia", "Vínculo con el SIIEAP"]):
        celda.text = texto
        _sombrear_celda(celda, COLOR_INSTITUCIONAL)
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for num, comp, vinculo in MARCO_COMPETENCIAS_NACION:
        fila = tabla_comp.add_row().cells
        fila[0].text = num
        fila[1].text = comp
        fila[2].text = vinculo
    _ajustar_tabla_docx(tabla_comp, anchos_cm=[1.3, 6.7, 8.0], tamano_fuente_pt=8.2)
    _franjas_alternas_docx(tabla_comp)
    doc.add_paragraph(
        "El Título IV crea el Sistema de Autonomía y Descentralización Territorial, con el "
        "Consejo Superior de Autonomía y Descentralización (el DAFP como miembro CON VOZ Y "
        "VOTO), el Programa para el Fortalecimiento de la Autonomía y la Descentralización "
        "Territorial (art. 180) y el Programa de Adecuación Institucional del Orden Nacional "
        "(art. 181, con acompañamiento del DAFP). El Título V (art. 186) crea el Sistema Único "
        "de Información del SGP (SUI-SGP)."
    )

    # ---- 5. Lectura estratégica ----
    doc.add_heading("5. Por qué esto es el momento del SIIEAP", level=2)
    doc.add_paragraph(
        "El nuevo régimen no inventa un concepto nuevo de \"capacidad institucional\" — lo "
        "convierte en criterio legal de clasificación territorial, algo que el IDI-MIPG, y por "
        "tanto el SIIEAP, ya miden hoy con 7 dimensiones, 19 políticas y 69 índices."
    )

    # ---- 6. Matriz de correspondencia ----
    doc.add_heading("6. Matriz de correspondencia: 7 dimensiones IDI-MIPG", level=2)
    tabla_matriz = doc.add_table(rows=1, cols=3)
    tabla_matriz.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_matriz.rows[0].cells, ["Dimensión", "Disposición del nuevo marco", "Qué aporta el SIIEAP"]):
        celda.text = texto
        _sombrear_celda(celda, COLOR_INSTITUCIONAL)
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for dim, disp, aporte in MARCO_MATRIZ_DIMENSIONES:
        fila = tabla_matriz.add_row().cells
        run_d = fila[0].paragraphs[0].add_run(dim)
        run_d.bold = True
        fila[1].text = disp
        fila[2].text = aporte
    _ajustar_tabla_docx(tabla_matriz, anchos_cm=[4.0, 5.5, 6.5], tamano_fuente_pt=8.2)
    _franjas_alternas_docx(tabla_matriz)

    # ---- 7. Plan en 3 niveles ----
    doc.add_heading("7. Plan de implementación en tres niveles", level=2)
    for nivel, quien, accion in MARCO_PLAN_TRES_NIVELES:
        p = doc.add_paragraph()
        run_n = p.add_run(f"{nivel} ({quien}): ")
        run_n.bold = True
        p.add_run(accion)

    # ---- 8. Línea de tiempo ----
    doc.add_heading("8. Línea de tiempo", level=2)
    tabla_tl = doc.add_table(rows=1, cols=3)
    tabla_tl.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_tl.rows[0].cells, ["Fecha", "Hito", "Qué significa"]):
        celda.text = texto
        _sombrear_celda(celda, COLOR_INSTITUCIONAL)
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for fecha, hito, significado in MARCO_DESCENTRALIZACION_TIMELINE:
        fila = tabla_tl.add_row().cells
        fila[0].text = fecha
        run_h = fila[1].paragraphs[0].add_run(hito)
        run_h.bold = True
        fila[2].text = significado
    _ajustar_tabla_docx(tabla_tl, anchos_cm=[2.8, 5.2, 9.0], tamano_fuente_pt=8.7)
    _franjas_alternas_docx(tabla_tl)

    # ---- 9. SGP hoy y meta ----
    doc.add_heading("9. El SGP hoy y la meta del 39,5 %", level=2)
    doc.add_paragraph(
        "Distribución sectorial vigente del SGP (Ley 715 de 2001, sobre el 96 % que queda "
        "tras el 4 % de Asignaciones Especiales):"
    )
    tabla_sgp = doc.add_table(rows=1, cols=2)
    tabla_sgp.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_sgp.rows[0].cells, ["Sector", "% del SGP sectorial"]):
        celda.text = texto
        _sombrear_celda(celda, COLOR_INSTITUCIONAL)
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for sector, pct in MARCO_SGP_DISTRIBUCION:
        fila = tabla_sgp.add_row().cells
        fila[0].text = sector
        fila[1].text = pct
    _ajustar_tabla_docx(tabla_sgp, anchos_cm=[10.0, 5.0], tamano_fuente_pt=9)
    _franjas_alternas_docx(tabla_sgp)

    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run(
        "Meta: de 29,5 % a 39,5 % de los ICN, en 12 años desde el 1 de enero de 2027, en "
        "incrementos anuales iguales (~0,83 pp/año)."
    )
    run_meta.italic = True

    doc.add_heading("Cómo se reparte el crecimiento del SGP", level=3)
    tabla_crec = doc.add_table(rows=1, cols=2)
    tabla_crec.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_crec.rows[0].cells, ["Destino del incremento", "% del incremento"]):
        celda.text = texto
        _sombrear_celda(celda, "B9752A")
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for destino, pct in MARCO_SGP_CRECIMIENTO:
        fila = tabla_crec.add_row().cells
        fila[0].text = destino
        fila[1].text = pct
    _ajustar_tabla_docx(tabla_crec, anchos_cm=[11.0, 4.0], tamano_fuente_pt=9)
    _franjas_alternas_docx(tabla_crec)

    p_puente = doc.add_paragraph()
    run_puente_t = p_puente.add_run("El puente con el SIIEAP: ")
    run_puente_t.bold = True
    p_puente.add_run(
        "el 80 % del crecimiento del SGP se destina al cierre de brechas \"sociales, "
        "económicas E INSTITUCIONALES\" — esa palabra conecta, por mandato constitucional, ese "
        "recurso con exactamente lo que mide el IDI-MIPG y con las brechas que este informe ya "
        "diagnosticó."
    )

    # ---- 10. Categorización ----
    doc.add_heading("10. Categorización territorial: Ley 617 y el nuevo régimen", level=2)
    doc.add_paragraph(
        "La Ley 617 de 2000 usa 2 criterios (población e ingresos en SMLV) y clasifica al 92 % "
        "de los municipios en categorías 5ª y 6ª. El proyecto de Ley de Competencias (art. 9) "
        "propone 4 criterios adicionales, SIN derogar la Ley 617 — coexisten (art. 9, parág. 4):"
    )
    tabla_cat = doc.add_table(rows=1, cols=2)
    tabla_cat.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_cat.rows[0].cells, ["Criterio nuevo (desde 2027, en trámite)", "Qué mide"]):
        celda.text = texto
        _sombrear_celda(celda, COLOR_INSTITUCIONAL)
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for criterio, que_mide in MARCO_CATEGORIZACION_NUEVA:
        fila = tabla_cat.add_row().cells
        fila[0].text = criterio
        fila[1].text = que_mide
    _ajustar_tabla_docx(tabla_cat, anchos_cm=[6.0, 9.0], tamano_fuente_pt=9)
    _franjas_alternas_docx(tabla_cat)

    # ---- 11. Paralelo de las 3 leyes ----
    doc.add_heading("11. Paralelo: Ley 617/2000, Ley 715/2001 (SGP) y Ley 2056/2020 (SGR)", level=2)
    tabla_par = doc.add_table(rows=1, cols=4)
    tabla_par.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla_par.rows[0].cells, ["Aspecto", "Ley 617/2000", "Ley 715/2001 (SGP)", "Ley 2056/2020 (SGR)"]):
        celda.text = texto
        _sombrear_celda(celda, COLOR_INSTITUCIONAL)
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True
                run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for aspecto, l617, l715, l2056 in MARCO_PARALELO_LEYES:
        fila = tabla_par.add_row().cells
        run_a = fila[0].paragraphs[0].add_run(aspecto)
        run_a.bold = True
        fila[1].text = l617
        fila[2].text = l715
        fila[3].text = l2056
    _ajustar_tabla_docx(tabla_par, anchos_cm=[3.2, 3.9, 3.9, 4.0], tamano_fuente_pt=7.8)
    _franjas_alternas_docx(tabla_par)
    p_confusion = doc.add_paragraph()
    run_confusion = p_confusion.add_run(
        "La confusión más frecuente: creer que la reforma 2024-2027 también cambia las "
        "regalías. No es así — son dos sistemas y dos artículos constitucionales distintos; "
        "el SGR sigue rigiéndose, sin cambios por esta reforma, por la Ley 2056 de 2020 y el "
        "Decreto 1821 de 2020."
    )
    run_confusion.italic = True
    run_confusion.font.size = Pt(9.5)
    run_confusion.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ---- 12. Inversión brecha por brecha (DINÁMICO por entidad) ----
    nombre_mostrar = nombre_entidad or (diag.entidad if diag is not None else "esta entidad")
    doc.add_heading(f"12. En qué invertir esos recursos: brecha por brecha para {nombre_mostrar}", level=2)
    filas_inversion = _filas_inversion_brechas(diag)
    if filas_inversion:
        doc.add_paragraph(
            f"Con base en el diagnóstico REAL ya calculado en este informe para {nombre_mostrar}, "
            "así se traduciría la inversión del 80 % de cierre de brechas institucionales del SGP, "
            "empezando por las dimensiones con el promedio más bajo:"
        )
        tabla_inv = doc.add_table(rows=1, cols=3)
        tabla_inv.style = "Light Grid Accent 1"
        for celda, texto in zip(tabla_inv.rows[0].cells, ["Dimensión (puntaje real)", "Inversión recomendada", "Canal SGP"]):
            celda.text = texto
            _sombrear_celda(celda, COLOR_INSTITUCIONAL)
            for parrafo in celda.paragraphs:
                for run_enc in parrafo.runs:
                    run_enc.bold = True
                    run_enc.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for etiqueta, inversion, canal in filas_inversion:
            fila = tabla_inv.add_row().cells
            run_e = fila[0].paragraphs[0].add_run(etiqueta)
            run_e.bold = True
            fila[1].text = inversion
            fila[2].text = canal
        _ajustar_tabla_docx(tabla_inv, anchos_cm=[4.5, 7.0, 3.5], tamano_fuente_pt=8.5)
        _franjas_alternas_docx(tabla_inv)
    else:
        doc.add_paragraph(
            "Esta entidad no reportó dimensiones con información suficiente para construir la "
            "tabla dinámica de inversión — consulte el diagnóstico por dimensión de este mismo "
            "informe."
        )
    nota_inv = doc.add_paragraph()
    run_nota_inv = nota_inv.add_run(
        "Nota metodológica: esta tabla se genera automáticamente a partir del diagnóstico real "
        f"de {nombre_mostrar} en este informe — no es un ejemplo genérico."
    )
    run_nota_inv.italic = True
    run_nota_inv.font.size = Pt(9)
    run_nota_inv.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ---- 13. Protocolo PDT ----
    doc.add_heading(f"13. Cómo validar esto contra el PDT de {nombre_mostrar}", level=2)
    doc.add_paragraph(
        "El SIIEAP no reemplaza al Plan de Desarrollo Territorial (PDT) — lo alimenta. "
        "Protocolo antes de ejecutar cualquier inversión con estos recursos:"
    )
    for paso in [
        "1. Ubicar el PDT vigente de la entidad (Ley 152 de 1994), que cubre el periodo de gobierno de 4 años en curso.",
        "2. Revisar el Plan Plurianual de Inversiones del PDT: solo lo programado allí —o incorporado por modificación presupuestal— puede recibir estos recursos.",
        "3. Cruzar cada brecha priorizada por el SIIEAP con el programa o subprograma del PDT que más se le parezca por objeto de gasto.",
        "4. Donde no exista un programa correspondiente, tramitar su inclusión (ajuste al plan de acción anual) antes de comprometer recursos del incremento del SGP.",
    ]:
        doc.add_paragraph(paso, style="List Bullet")
    p_honestidad = doc.add_paragraph()
    run_honestidad = p_honestidad.add_run(
        f"Honestidad metodológica: el SIIEAP no tiene acceso al texto real del PDT de "
        f"{nombre_mostrar} — este es un protocolo de cruce, no una lectura ya hecha de un PDT "
        "real. Para aplicarlo, la entidad debe usar su propio PDT vigente."
    )
    run_honestidad.italic = True

    # ---- 14. ESAP YouTube ----
    doc.add_heading("14. ESAP en YouTube: sesiones para conectarse", level=2)
    for titulo, fuente, url in MARCO_YOUTUBE_ESAP:
        p = doc.add_paragraph()
        run_t = p.add_run(f"{titulo} — ")
        run_t.bold = True
        run_f = p.add_run(f"{fuente}. ")
        run_f.italic = True
        _agregar_hipervinculo_docx(p, url, url)

    # ---- Advertencia de vigencia final ----
    p_vigencia = doc.add_paragraph()
    run_vig_t = p_vigencia.add_run("Advertencia de vigencia: ")
    run_vig_t.bold = True
    run_vig_t.font.color.rgb = RGBColor(0xB8, 0x5C, 0x00)
    p_vigencia.add_run(
        "a la fecha de generación de este informe, el Proyecto de Ley Orgánica de Competencias "
        "NO es ley de la República — está en trámite en el Congreso y su articulado puede "
        "cambiar antes de la sanción presidencial, o el proyecto puede no llegar a aprobarse. "
        "Solo el Acto Legislativo 3 de 2024 es, hoy, norma constitucional vigente."
    )


def _agregar_hipervinculo_docx(paragraph, url, texto):
    """Inserta un hipervínculo real y clicable dentro de un párrafo de python-docx."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F3864")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = texto
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _agregar_marco_descentralizacion_pdf(elementos, estilos, estilo_normal, estilo_h2, diag=None, nombre_entidad=None):
    """Versión PDF (reportlab) de _agregar_marco_descentralizacion_docx — el
    Capítulo Especial de Descentralización COMPLETO, personalizado por entidad."""
    estilo_celda = ParagraphStyle("MarcoCelda", parent=estilo_normal, fontSize=8.2, leading=10.3)
    estilo_celda_b = ParagraphStyle("MarcoCeldaB", parent=estilo_celda, fontName="Helvetica-Bold")
    estilo_enc = ParagraphStyle("MarcoEnc", parent=estilo_celda, fontName="Helvetica-Bold", textColor=colors.white)
    estilo_h3 = ParagraphStyle("MarcoH3", parent=estilo_h2, fontSize=12)
    estilo_nota = ParagraphStyle("MarcoNota", parent=estilo_normal, fontSize=8.5, textColor=colors.HexColor("#666666"))
    estilo_vigencia = ParagraphStyle("MarcoVigencia", parent=estilo_normal, textColor=colors.HexColor("#B85C00"))
    estilo_link = ParagraphStyle("MarcoLink", parent=estilo_normal, textColor=colors.HexColor("#1F3864"))

    def _tabla_pdf(encabezados, filas_datos, anchos_cm, color_hex_enc=COLOR_INSTITUCIONAL, negrita_primera_col=False):
        filas = [[Paragraph(h, estilo_enc) for h in encabezados]]
        for fila_datos in filas_datos:
            fila_p = []
            for i, valor in enumerate(fila_datos):
                estilo_usar = estilo_celda_b if (negrita_primera_col and i == 0) else estilo_celda
                fila_p.append(Paragraph(str(valor), estilo_usar))
            filas.append(fila_p)
        tabla = Table(filas, colWidths=[a * cm for a in anchos_cm], repeatRows=1)
        estilo_tabla = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{color_hex_enc}")),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D0D0D0")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]
        for i in range(1, len(filas)):
            if i % 2 == 0:
                estilo_tabla.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor("#F2F4F8")))
        tabla.setStyle(TableStyle(estilo_tabla))
        return tabla

    elementos.append(PageBreak())
    elementos.append(Paragraph("Anexo: Capítulo Especial de Descentralización 2026-2030", estilos["Heading1"]))
    elementos.append(Paragraph(MARCO_DESCENTRALIZACION_INTRO, estilo_normal))
    elementos.append(Spacer(1, 8))

    elementos.append(Paragraph("1. Contexto: continuidad institucional, no personalismo", estilo_h2))
    elementos.append(Paragraph(MARCO_ENCUADRE_METODOLOGICO, estilo_normal))
    elementos.append(Paragraph(MARCO_CONTINUIDAD_7AGOSTO, estilo_normal))
    for etiqueta, texto in MARCO_ARTICULOS_CP:
        elementos.append(Paragraph(f"<b>{etiqueta}:</b> <i>{texto}</i>", estilo_normal))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("2. La Misión de Descentralización (2022-2024)", estilo_h2))
    elementos.append(Spacer(1, 4))
    elementos.append(_tabla_pdf(["Eje", "Objetivo"], list(MARCO_MISION_EJES), [7.0, 9.5], negrita_primera_col=True))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph(MARCO_MISION_PROPUESTAS, estilo_normal))
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph("Cooperación multilateral y bilateral", estilo_h3))
    elementos.append(Spacer(1, 4))
    elementos.append(_tabla_pdf(["Organismo", "Rol documentado"], list(MARCO_COOPERACION), [4.5, 12.0], negrita_primera_col=True))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("3. Acto Legislativo 3 de 2024 (norma constitucional VIGENTE)", estilo_h2))
    elementos.append(Paragraph(
        "Modifica los arts. 356 y 357 C.P.; incrementa el SGP hasta 39,5 % de los ICN; ordena "
        "que la Ley de Competencias distribuya esos recursos según capacidad institucional.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("4. Proyecto de Ley de Competencias (EN TRÁMITE, no vigente)", estilo_h2))
    elementos.append(Paragraph(
        "Objeto: fortalecer la autonomía y el desarrollo territorial mediante la asignación y "
        "distribución de competencias y recursos entre la Nación y las entidades territoriales.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph("Competencias de la Nación que dialogan con el SIIEAP (art. 15)", estilo_h3))
    elementos.append(Spacer(1, 4))
    filas_comp = [[n, c, v] for n, c, v in MARCO_COMPETENCIAS_NACION]
    elementos.append(_tabla_pdf(["Núm.", "Competencia", "Vínculo con el SIIEAP"], filas_comp, [1.3, 6.7, 8.5]))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph(
        "El Título IV crea el Sistema de Autonomía y Descentralización Territorial (DAFP con "
        "voz y voto en su Consejo Superior) y los programas de fortalecimiento institucional "
        "(arts. 180-181). El Título V (art. 186) crea el Sistema Único de Información del SGP.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("5. Por qué esto es el momento del SIIEAP", estilo_h2))
    elementos.append(Paragraph(
        "El nuevo régimen convierte \"capacidad institucional\" en criterio legal de "
        "clasificación territorial — exactamente lo que el IDI-MIPG ya mide.",
        estilo_normal,
    ))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("6. Matriz de correspondencia: 7 dimensiones IDI-MIPG", estilo_h2))
    elementos.append(Spacer(1, 4))
    filas_matriz = [[d, disp, ap] for d, disp, ap in MARCO_MATRIZ_DIMENSIONES]
    elementos.append(_tabla_pdf(["Dimensión", "Disposición del nuevo marco", "Qué aporta el SIIEAP"], filas_matriz, [3.8, 5.5, 7.2], negrita_primera_col=True))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("7. Plan de implementación en tres niveles", estilo_h2))
    for nivel, quien, accion in MARCO_PLAN_TRES_NIVELES:
        elementos.append(Paragraph(f"<b>{nivel} ({quien}):</b> {accion}", estilo_normal))
        elementos.append(Spacer(1, 4))
    elementos.append(Spacer(1, 8))

    elementos.append(Paragraph("8. Línea de tiempo", estilo_h2))
    elementos.append(Spacer(1, 4))
    filas_tl = [[f, h, s] for f, h, s in MARCO_DESCENTRALIZACION_TIMELINE]
    elementos.append(_tabla_pdf(["Fecha", "Hito", "Qué significa"], filas_tl, [2.6, 4.8, 9.1]))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("9. El SGP hoy y la meta del 39,5 %", estilo_h2))
    elementos.append(Paragraph(
        "Distribución sectorial vigente del SGP (Ley 715 de 2001):", estilo_normal,
    ))
    elementos.append(Spacer(1, 4))
    elementos.append(_tabla_pdf(["Sector", "% del SGP sectorial"], list(MARCO_SGP_DISTRIBUCION), [9.0, 4.0]))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph(
        "<i>Meta: de 29,5 % a 39,5 % de los ICN, en 12 años desde 2027 (~0,83 pp/año).</i>", estilo_normal,
    ))
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph("Cómo se reparte el crecimiento del SGP", estilo_h3))
    elementos.append(Spacer(1, 4))
    elementos.append(_tabla_pdf(["Destino del incremento", "% del incremento"], list(MARCO_SGP_CRECIMIENTO), [10.0, 3.0], color_hex_enc="B9752A"))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph(
        "<b>El puente con el SIIEAP:</b> el 80 % del crecimiento del SGP se destina al cierre "
        "de brechas \"sociales, económicas E INSTITUCIONALES\" — el mismo dato que mide el "
        "IDI-MIPG.", estilo_normal,
    ))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("10. Categorización territorial: Ley 617 y el nuevo régimen", estilo_h2))
    elementos.append(Paragraph(
        "La Ley 617 de 2000 usa 2 criterios y clasifica al 92 % de los municipios en categorías "
        "5ª y 6ª. El proyecto de Ley de Competencias (art. 9) propone 4 criterios adicionales, "
        "SIN derogar la Ley 617:", estilo_normal,
    ))
    elementos.append(Spacer(1, 4))
    elementos.append(_tabla_pdf(["Criterio nuevo (desde 2027, en trámite)", "Qué mide"], list(MARCO_CATEGORIZACION_NUEVA), [6.5, 9.5]))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("11. Paralelo: Ley 617/2000, Ley 715/2001 (SGP) y Ley 2056/2020 (SGR)", estilo_h2))
    elementos.append(Spacer(1, 4))
    filas_par = [[a, b, c, d] for a, b, c, d in MARCO_PARALELO_LEYES]
    elementos.append(_tabla_pdf(["Aspecto", "Ley 617/2000", "Ley 715/2001", "Ley 2056/2020"], filas_par, [3.4, 3.9, 3.9, 4.8], negrita_primera_col=True))
    elementos.append(Spacer(1, 6))
    elementos.append(Paragraph(
        "La confusión más frecuente: creer que la reforma también cambia las regalías. No es "
        "así — el SGR sigue rigiéndose por la Ley 2056 de 2020 y el Decreto 1821 de 2020.",
        estilo_nota,
    ))
    elementos.append(Spacer(1, 10))

    nombre_mostrar = nombre_entidad or (diag.entidad if diag is not None else "esta entidad")
    elementos.append(Paragraph(f"12. En qué invertir esos recursos: brecha por brecha para {nombre_mostrar}", estilo_h2))
    filas_inversion = _filas_inversion_brechas(diag)
    if filas_inversion:
        elementos.append(Paragraph(
            f"Con base en el diagnóstico REAL de {nombre_mostrar} en este informe, así se "
            "traduciría la inversión del 80 % de cierre de brechas institucionales del SGP, "
            "empezando por las dimensiones con el promedio más bajo:", estilo_normal,
        ))
        elementos.append(Spacer(1, 4))
        elementos.append(_tabla_pdf(["Dimensión (puntaje real)", "Inversión recomendada", "Canal SGP"], filas_inversion, [4.5, 7.5, 3.5], negrita_primera_col=True))
    else:
        elementos.append(Paragraph(
            "Esta entidad no reportó dimensiones con información suficiente para la tabla "
            "dinámica de inversión.", estilo_normal,
        ))
    elementos.append(Spacer(1, 4))
    elementos.append(Paragraph(
        f"<i>Nota metodológica: esta tabla se genera automáticamente a partir del diagnóstico "
        f"real de {nombre_mostrar} — no es un ejemplo genérico.</i>", estilo_nota,
    ))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph(f"13. Cómo validar esto contra el PDT de {nombre_mostrar}", estilo_h2))
    elementos.append(Paragraph(
        "El SIIEAP no reemplaza al Plan de Desarrollo Territorial (PDT) — lo alimenta. "
        "Protocolo antes de ejecutar cualquier inversión:", estilo_normal,
    ))
    for paso in [
        "1. Ubicar el PDT vigente de la entidad (Ley 152 de 1994).",
        "2. Revisar el Plan Plurianual de Inversiones del PDT.",
        "3. Cruzar cada brecha priorizada por el SIIEAP con el programa del PDT correspondiente.",
        "4. Donde no exista programa, tramitar su inclusión antes de comprometer recursos.",
    ]:
        elementos.append(Paragraph(paso, estilo_normal))
    elementos.append(Spacer(1, 4))
    elementos.append(Paragraph(
        f"<i>Honestidad metodológica: el SIIEAP no tiene acceso al texto real del PDT de "
        f"{nombre_mostrar} — este es un protocolo de cruce, no una lectura ya hecha.</i>",
        estilo_nota,
    ))
    elementos.append(Spacer(1, 10))

    elementos.append(Paragraph("14. ESAP en YouTube: sesiones para conectarse", estilo_h2))
    for titulo, fuente, url in MARCO_YOUTUBE_ESAP:
        elementos.append(Paragraph(
            f'<b>{titulo}</b> — <i>{fuente}</i>. <link href="{url}"><u>{url}</u></link>',
            estilo_link,
        ))
        elementos.append(Spacer(1, 4))

    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph(
        "<b>Advertencia de vigencia:</b> a la fecha de generación de este informe, el Proyecto "
        "de Ley Orgánica de Competencias NO es ley de la República — está en trámite y su "
        "articulado puede cambiar, o el proyecto puede no llegar a aprobarse. Solo el Acto "
        "Legislativo 3 de 2024 es, hoy, norma constitucional vigente.", estilo_vigencia,
    ))
    elementos.append(PageBreak())



def _ajustar_tabla_docx(tabla, anchos_cm, tamano_fuente_pt=9.5):
    """Fija anchos de columna reales (no solo sugeridos) y un tamaño de letra
    más pequeño para toda la tabla, evitando que el texto se desborde de las
    celdas — python-docx requiere fijar el ancho en CADA celda de CADA fila,
    no solo en la tabla, para que Word lo respete de forma consistente."""
    tabla.autofit = False
    tabla.allow_autofit = False
    for fila in tabla.rows:
        for celda, ancho in zip(fila.cells, anchos_cm):
            celda.width = Cm(ancho)
            for parrafo in celda.paragraphs:
                parrafo.paragraph_format.space_after = Pt(2)
                for run in parrafo.runs:
                    run.font.size = Pt(tamano_fuente_pt)


def _fecha_hoy_es():
    meses = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
        "agosto", "septiembre", "octubre", "noviembre", "diciembre",
    ]
    hoy = datetime.now()
    return f"{hoy.day} de {meses[hoy.month - 1]} de {hoy.year}"


_COLOR_HEX_POR_RIESGO = {
    "alta": "F5B7B1",
    "media": "FAD7A0",
    "baja": "A9DFBF",
}


# Mapeo determinístico política MIPG -> enfoque(s) contemporáneo(s) + norma.
# Esto NO depende de que la IA lo mencione: se calcula siempre por código a
# partir de las brechas reales de la entidad, garantizando que la conexión
# con los 15 enfoques contemporáneos sea visible en TODOS los informes.
POLITICA_A_ENFOQUE_Y_NORMA = {
    "gestión estratégica del talento humano": ("Capacidades Estatales", "Ley 909 de 2004 (Sistema de Empleo Público); Decreto 1083 de 2015, Libro 2, art. 2.2.4.5 (DUR Función Pública, compila las competencias laborales del derogado Decreto 2539/2005); Ley 1960 de 2019 (modifica la Ley 909, encargos y evaluación del desempeño)"),
    "integridad": ("Gobierno Abierto", "Ley 1712 de 2014; Ley 2013 de 2019 y Decreto 830 de 2021 (conflictos de interés); Decreto 648 de 2017 (adopta el Código de Integridad del Servicio Público); Ley 1474 de 2011 (Estatuto Anticorrupción)"),
    "planeación institucional": ("Administración Pública Basada en Evidencia", "Decreto 1499 de 2017, art. 2.2.22.3.2-3.3"),
    "gestión presupuestal y eficiencia del gasto": ("Capacidades Estatales", "Ley 610 de 2000, art. 1 (responsabilidad fiscal)"),
    "compras y contratación pública": ("Gobierno como Plataforma", "Decreto 620 de 2020, art. 2.2.17.2.2.1-2.2.2 (TVEC/Articulador)"),
    "fortalecimiento organizacional y simplificación de procesos": ("Capacidades Estatales", "Ley 617 de 2000; Ley 1454 de 2011"),
    "gobierno digital": ("Estado Digital / Transformación Digital del Estado", "Decreto 767 de 2022; Decreto 1263 de 2022"),
    "seguridad digital": ("Gobernanza de Datos", "CONPES 3995 de 2020 (Confianza y Seguridad Digital)"),
    "defensa jurídica": ("Capacidades Estatales", "Ley 87 de 1993 (control jurídico interno)"),
    "mejora normativa": ("Administración Pública Basada en Evidencia", "Decreto 1499 de 2017"),
    "servicio a las ciudadanías": ("Gobierno Abierto / Co-creación y Design Thinking", "Ley 1712 de 2014; Decreto 767 de 2022 (Innovación Pública Digital)"),
    "racionalización de trámites": ("Estado Digital", "Decreto 019 de 2012 (antitrámites); Decreto 767 de 2022"),
    "participación ciudadana en la gestión pública": ("Gobierno Abierto", "Ley 1712 de 2014; Colombia miembro OGP desde 2011"),
    "seguimiento y evaluación del desempeño institucional": ("Administración Pública Basada en Evidencia / Resiliencia Institucional", "Decreto 1499 de 2017; Ley 1523 de 2012"),
    "transparencia, acceso a la información y lucha contra la corrupción": ("Gobierno Abierto", "Ley 1712 de 2014, art. 1-3; ODS 16 (marco externo, ONU 2015)"),
    "gestión documental": ("Gobernanza de Datos", "Ley 594 de 2000 (Ley General de Archivos); Decreto 1389 de 2022"),
    "gestión de la información estadística": ("Gobernanza de Datos / Administración Pública Basada en Evidencia", "Decreto 1389 de 2022"),
    "gestión del conocimiento": ("Administración Pública Basada en Evidencia / Capacidades Estatales", "Decreto 1499 de 2017"),
    "control interno": ("Resiliencia Institucional / Capacidades Estatales", "Ley 87 de 1993, art. 1-2-9; Ley 1523 de 2012, art. 8"),
}


def _enfoque_y_norma_de_politica(nombre_politica: str):
    """Busca el enfoque contemporáneo y la norma asociada a una política MIPG,
    normalizando el nombre para tolerar variantes de mayúsculas/redacción."""
    clave = str(nombre_politica).strip().lower()
    clave = clave.replace("política de ", "").replace("política ", "")
    for politica_conocida, valor in POLITICA_A_ENFOQUE_Y_NORMA.items():
        if politica_conocida in clave or clave in politica_conocida:
            return valor
    return ("Capacidades Estatales", "Decreto 1499 de 2017 (MIPG)")


# ---------------------------------------------------------------------------
# Marco de responsabilidad de la Alta Dirección: artículo 6 de la
# Constitución y cruce política MIPG -> tipo de riesgo (jurídico, fiscal,
# administrativo o disciplinario) -> norma -> consecuencia posible.
# Uso exclusivo del Informe Ejecutivo para Alcaldes/Gobernadores, pero vive
# aquí (módulo compartido) por si otro informe necesita la misma lectura.
# ---------------------------------------------------------------------------

ARTICULO_6_CONSTITUCION_TITULO = "Artículo 6 de la Constitución Política de Colombia (1991)"

ARTICULO_6_CONSTITUCION_TEXTO = (
    "Los particulares sólo son responsables ante las autoridades por infringir "
    "la Constitución y las leyes. Los servidores públicos lo son por la misma "
    "causa y por omisión o extralimitación en el ejercicio de sus funciones."
)

ARTICULO_6_CONSTITUCION_EXPLICACION = (
    "Esta es la razón de fondo por la que cada brecha de este informe importa más "
    "allá del puntaje: a diferencia de un particular, que solo responde por infringir "
    "la ley, el servidor público de esta entidad responde también por OMISIÓN — es "
    "decir, por no hacer lo que la ley le exige — y por EXTRALIMITACIÓN — por hacer "
    "más de lo que la ley le permite. Una brecha del MIPG sostenida en el tiempo, sin "
    "plan de mejoramiento ni evidencia de gestión, es precisamente ese escenario de "
    "omisión que el artículo 6 sanciona, y puede derivar en responsabilidad jurídica, "
    "fiscal, administrativa y/o disciplinaria de quien tenía el deber legal de actuar."
)

# (tipo_de_riesgo, norma_principal, consecuencia_posible_breve)
POLITICA_A_RIESGO_ALTA_DIRECCION = {
    "gestión estratégica del talento humano": (
        "Administrativo y disciplinario",
        "Ley 909 de 2004 (empleo público); Decreto 1083 de 2015, Libro 2 (DUR Función Pública); Ley 1960 de 2019 (modifica la Ley 909, encargos y evaluación del desempeño); Ley 1952 de 2019 modif. Ley 2094 de 2021 (Código General Disciplinario)",
        "Faltas en la carrera administrativa y en la evaluación del desempeño pueden derivar en nulidad de actos de personal y en investigación disciplinaria contra el nominador.",
    ),
    "integridad": (
        "Disciplinario",
        "Decreto 648 de 2017 (Código de Integridad del Servicio Público); Ley 1952 de 2019 modif. Ley 2094 de 2021; Ley 2013 de 2019 y Decreto 830 de 2021 (conflictos de interés)",
        "Los conflictos de interés no declarados son falta disciplinaria autónoma, independiente de si hubo o no un daño patrimonial.",
    ),
    "planeación institucional": (
        "Administrativo",
        "Decreto 1499 de 2017, art. 2.2.22.3.2-3.3",
        "La ausencia de planeación institucional debilita la defensa de la entidad ante entes de control al no poder demostrar la debida diligencia de gestión.",
    ),
    "gestión presupuestal y eficiencia del gasto": (
        "Fiscal",
        "Ley 610 de 2000 (proceso de responsabilidad fiscal); Ley 617 de 2000 (límites de gasto); "
        "Acto Legislativo 05 de 2019, Ley 2056 de 2020 y Decreto 1821 de 2020 (Sistema General de "
        "Regalías, para entidades territoriales receptoras)",
        "El detrimento patrimonial derivado de una gestión presupuestal deficiente —incluida la "
        "ejecución irregular de recursos de regalías— puede abrir proceso de responsabilidad fiscal "
        "ante la Contraloría, con obligación de resarcir el daño con el patrimonio del responsable.",
    ),
    "compras y contratación pública": (
        "Fiscal y disciplinario",
        "Ley 80 de 1993 y Ley 1150 de 2007 (contratación estatal); Ley 1474 de 2011 (Estatuto Anticorrupción); Ley 610 de 2000",
        "Irregularidades contractuales son la causa más frecuente de procesos de responsabilidad fiscal y disciplinaria contra alcaldes y gobernadores en Colombia.",
    ),
    "fortalecimiento organizacional y simplificación de procesos": (
        "Administrativo",
        "Ley 617 de 2000; Ley 1454 de 2011 (ordenamiento territorial)",
        "Una estructura organizacional no ajustada a la capacidad fiscal real de la entidad es causal de observaciones de viabilidad fiscal ante el Ministerio de Hacienda.",
    ),
    "gobierno digital": (
        "Administrativo",
        "Decreto 767 de 2022; Decreto 1263 de 2022",
        "El incumplimiento de la política de Gobierno Digital es hoy objeto de seguimiento directo por parte del MinTIC y de Función Pública dentro del FURAG.",
    ),
    "seguridad digital": (
        "Administrativo y disciplinario",
        "CONPES 3995 de 2020; Ley 1952 de 2019 (deber de custodia de la información)",
        "La pérdida o filtración de información institucional por fallas de seguridad digital compromete disciplinariamente a quien tenía el deber de custodia.",
    ),
    "defensa jurídica": (
        "Jurídico y fiscal",
        "Ley 2213 de 2022 (procesos judiciales); Ley 678 de 2001 (acción de repetición); Ley 610 de 2000",
        "Una defensa jurídica débil incrementa el riesgo de condenas judiciales contra la entidad, que a su vez pueden derivar en acción de repetición contra el servidor responsable.",
    ),
    "mejora normativa": (
        "Administrativo",
        "Decreto 1499 de 2017",
        "Normas internas desactualizadas o contradictorias con la ley superior son causal de nulidad de actos administrativos ante lo contencioso administrativo.",
    ),
    "servicio a las ciudadanías": (
        "Disciplinario",
        "Ley 1755 de 2015 (derecho de petición); Ley 1952 de 2019",
        "No responder oportunamente las peticiones ciudadanas es falta disciplinaria expresa y puede derivar en acción de tutela contra la entidad.",
    ),
    "racionalización de trámites": (
        "Administrativo",
        "Decreto 019 de 2012 (antitrámites); Decreto 767 de 2022",
        "Exigir requisitos no autorizados por la ley es causal de responsabilidad disciplinaria del funcionario que los exige.",
    ),
    "participación ciudadana en la gestión pública": (
        "Disciplinario y administrativo",
        "Ley 1757 de 2015 (participación democrática)",
        "El incumplimiento de los espacios de participación obligatorios (rendición de cuentas, presupuesto participativo) es objeto de seguimiento por los entes de control.",
    ),
    "seguimiento y evaluación del desempeño institucional": (
        "Administrativo",
        "Decreto 1499 de 2017; Ley 1523 de 2012 (gestión del riesgo)",
        "La falta de seguimiento y evaluación impide demostrar la debida diligencia de la Alta Dirección ante un eventual proceso de responsabilidad.",
    ),
    "transparencia, acceso a la información y lucha contra la corrupción": (
        "Disciplinario y penal",
        "Ley 1712 de 2014 (transparencia); Ley 1474 de 2011 (Estatuto Anticorrupción); Código Penal, Título XV",
        "Las faltas en transparencia son, junto con la contratación, la puerta de entrada más común a investigaciones disciplinarias y, en los casos más graves, penales.",
    ),
    "gestión documental": (
        "Disciplinario",
        "Ley 594 de 2000 (Ley General de Archivos); Decreto 1389 de 2022",
        "La pérdida, alteración o destrucción de documentos públicos es falta disciplinaria gravísima según el Código General Disciplinario.",
    ),
    "gestión de la información estadística": (
        "Administrativo",
        "Decreto 1389 de 2022",
        "Reportar información estadística inexacta a Función Pública compromete la validez del FURAG y de los indicadores oficiales de la entidad.",
    ),
    "gestión del conocimiento": (
        "Administrativo",
        "Decreto 1499 de 2017",
        "La pérdida de conocimiento institucional por rotación de personal debilita la capacidad de respuesta de la entidad ante requerimientos de entes de control.",
    ),
    "control interno": (
        "Disciplinario y fiscal",
        "Ley 87 de 1993, art. 1-2-9; Ley 1523 de 2012, art. 8",
        "Un sistema de control interno débil o inexistente agrava la responsabilidad de la Alta Dirección, porque elimina la primera línea de defensa que debía advertir el riesgo a tiempo.",
    ),
}


def _riesgo_alta_direccion_de_politica(nombre_politica: str):
    """Análogo a _enfoque_y_norma_de_politica, pero para la lectura de
    riesgo legal/fiscal/administrativo/disciplinario dirigida a la Alta
    Dirección (alcalde/gobernador)."""
    clave = str(nombre_politica).strip().lower()
    clave = clave.replace("política de ", "").replace("política ", "")
    for politica_conocida, valor in POLITICA_A_RIESGO_ALTA_DIRECCION.items():
        if politica_conocida in clave or clave in politica_conocida:
            return valor
    return (
        "Administrativo",
        "Decreto 1499 de 2017 (MIPG)",
        "Toda brecha sostenida en el MIPG debilita la capacidad de la entidad para demostrar debida diligencia ante los entes de control.",
    )


def _color_hex_riesgo(nivel_riesgo) -> str | None:
    if not nivel_riesgo:
        return None
    return _COLOR_HEX_POR_RIESGO.get(str(nivel_riesgo).strip().lower())


def _sombrear_celda(celda, color_hex: str) -> None:
    """Colorea el fondo de una celda de tabla de python-docx (no hay API
    pública directa, se manipula el XML de la celda)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    propiedades_celda = celda._tc.get_or_add_tcPr()
    sombreado = OxmlElement("w:shd")
    sombreado.set(qn("w:val"), "clear")
    sombreado.set(qn("w:color"), "auto")
    sombreado.set(qn("w:fill"), color_hex)
    propiedades_celda.append(sombreado)


def _parsear_bloques_markdown(texto: str) -> list[tuple[str, object]]:
    """Convierte el markdown simple que produce motor_analisis_ia.py (encabezados
    ``##``/``###``, tablas con ``|``, listas con ``-``/``*`` y negrita ``**texto**``)
    en una lista de bloques estructurados que los renderizadores de docx y PDF
    pueden dibujar como encabezados, tablas y párrafos reales — en vez de volcar
    el texto crudo (con los símbolos ``#``, ``*`` y ``|`` literales) como hacía
    antes ``analisis_ia_texto.split("\\n")``.

    Devuelve una lista de tuplas ``(tipo, contenido)`` donde ``tipo`` es uno de
    ``"heading"`` (contenido = ``(nivel, texto)``), ``"table"`` (contenido =
    lista de filas, la primera es el encabezado), ``"bullet"`` (contenido =
    texto del ítem) o ``"parrafo"`` (contenido = texto plano).
    """
    lineas = texto.split("\n")
    bloques: list[tuple[str, object]] = []
    i = 0
    n = len(lineas)
    patron_separador_tabla = re.compile(r"^\|?[\s:\-|]+\|?$")
    patron_encabezado = re.compile(r"^(#{1,4})\s+(.*)$")
    patron_bullet = re.compile(r"^[-*]\s+(.*)$")

    while i < n:
        linea = lineas[i].strip()
        if not linea:
            i += 1
            continue

        m_enc = patron_encabezado.match(linea)
        if m_enc:
            nivel = len(m_enc.group(1))
            bloques.append(("heading", (nivel, m_enc.group(2).strip())))
            i += 1
            continue

        if linea.startswith("|"):
            filas_tabla: list[list[str]] = []
            while i < n and lineas[i].strip().startswith("|"):
                fila_cruda = lineas[i].strip()
                if "-" in fila_cruda and patron_separador_tabla.match(fila_cruda):
                    i += 1
                    continue
                celdas = [c.strip() for c in fila_cruda.strip("|").split("|")]
                filas_tabla.append(celdas)
                i += 1
            if filas_tabla:
                bloques.append(("table", filas_tabla))
            continue

        m_bullet = patron_bullet.match(linea)
        if m_bullet and not linea.startswith("**"):
            bloques.append(("bullet", m_bullet.group(1).strip()))
            i += 1
            continue

        bloques.append(("parrafo", linea))
        i += 1

    return bloques


_PATRON_NEGRITA = re.compile(r"(\*\*[^*]+\*\*)")


def _agregar_texto_con_negrita_docx(paragraph, texto, tamano_fuente_pt=None, color_rgb=None, negrita_base=False):
    """Agrega ``texto`` a un párrafo de python-docx, convirtiendo los tramos
    ``**así**`` en runs en negrita (en vez de dejar los asteriscos literales)."""
    partes = _PATRON_NEGRITA.split(texto)
    for parte in partes:
        if not parte:
            continue
        es_negrita = parte.startswith("**") and parte.endswith("**") and len(parte) > 4
        run = paragraph.add_run(parte[2:-2] if es_negrita else parte)
        run.bold = es_negrita or negrita_base
        if tamano_fuente_pt:
            run.font.size = Pt(tamano_fuente_pt)
        if color_rgb:
            run.font.color.rgb = color_rgb


def _negrita_markdown_a_html_pdf(texto: str) -> str:
    """Convierte ``**texto**`` a ``<b>texto</b>`` para que reportlab lo dibuje
    en negrita real dentro de un ``Paragraph`` (en vez del asterisco literal)."""
    return _PATRON_NEGRITA.sub(lambda m: f"<b>{m.group(0)[2:-2]}</b>", texto)


def _agregar_texto_markdown_docx(doc, texto: str, nivel_encabezado_base: int = 2) -> None:
    """Dibuja en ``doc`` el texto markdown generado por la IA como encabezados,
    tablas y párrafos con formato real, en lugar de párrafos con los símbolos
    ``##``, ``**`` y ``|`` sin interpretar."""
    for tipo, contenido in _parsear_bloques_markdown(texto):
        if tipo == "heading":
            nivel_md, texto_h = contenido
            nivel_doc = min(nivel_encabezado_base + (nivel_md - 1), 4)
            doc.add_heading(texto_h, level=max(nivel_doc, 1))
        elif tipo == "table":
            filas = contenido
            encabezado, *resto = filas
            n_cols = len(encabezado)
            if n_cols == 0:
                continue
            tabla = doc.add_table(rows=1, cols=n_cols)
            tabla.style = "Light Grid Accent 1"
            for idx_col in range(n_cols):
                celda = tabla.rows[0].cells[idx_col]
                p_celda = celda.paragraphs[0]
                _agregar_texto_con_negrita_docx(
                    p_celda, encabezado[idx_col], tamano_fuente_pt=9.5,
                    color_rgb=RGBColor(0xFF, 0xFF, 0xFF), negrita_base=True,
                )
                _sombrear_celda(celda, COLOR_INSTITUCIONAL)
            for fila_datos in resto:
                celdas_fila = tabla.add_row().cells
                for idx_col in range(n_cols):
                    texto_celda = fila_datos[idx_col] if idx_col < len(fila_datos) else ""
                    _agregar_texto_con_negrita_docx(celdas_fila[idx_col].paragraphs[0], texto_celda, tamano_fuente_pt=9)
            ancho_col = round(17.0 / n_cols, 2)
            _ajustar_tabla_docx(tabla, anchos_cm=[ancho_col] * n_cols, tamano_fuente_pt=9)
            doc.add_paragraph()
        elif tipo == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _agregar_texto_con_negrita_docx(p, contenido)
        else:
            p = doc.add_paragraph()
            _agregar_texto_con_negrita_docx(p, contenido)


def _texto_markdown_a_pdf_flowables(texto: str, estilo_normal, estilo_h2, estilo_h3=None):
    """Versión para reportlab de ``_agregar_texto_markdown_docx``: devuelve una
    lista de flowables (Paragraph/Table/Spacer) con encabezados, tablas y
    negrita reales en vez de los símbolos markdown sin interpretar."""
    if estilo_h3 is None:
        estilo_h3 = ParagraphStyle("H3SIIEAP_md", parent=estilo_h2, fontSize=11.5, spaceBefore=8)
    flowables = []
    for tipo, contenido in _parsear_bloques_markdown(texto):
        if tipo == "heading":
            nivel_md, texto_h = contenido
            estilo_usar = estilo_h2 if nivel_md <= 2 else estilo_h3
            texto_html = texto_h.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(_negrita_markdown_a_html_pdf(texto_html), estilo_usar))
        elif tipo == "table":
            filas = contenido
            encabezado, *resto = filas
            n_cols = len(encabezado)
            if n_cols == 0:
                continue
            def _celda(txt, encabezado_flag):
                txt_html = txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                txt_html = _negrita_markdown_a_html_pdf(txt_html)
                estilo_celda = _ESTILO_CELDA_TABLA_PDF_ENCABEZADO if encabezado_flag else _ESTILO_CELDA_TABLA_PDF
                return Paragraph(txt_html, estilo_celda)
            datos = [[_celda(c, True) for c in encabezado]]
            for fila_datos in resto:
                fila_completa = [fila_datos[idx] if idx < len(fila_datos) else "" for idx in range(n_cols)]
                datos.append([_celda(c, False) for c in fila_completa])
            ancho_col = (17 * cm) / n_cols
            tabla_pdf = Table(datos, hAlign="LEFT", colWidths=[ancho_col] * n_cols)
            tabla_pdf.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{COLOR_INSTITUCIONAL}")),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            flowables.append(tabla_pdf)
            flowables.append(Spacer(1, 6))
        elif tipo == "bullet":
            texto_html = contenido.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(f"•&nbsp;&nbsp;{_negrita_markdown_a_html_pdf(texto_html)}", estilo_normal))
            flowables.append(Spacer(1, 2))
        else:
            texto_html = contenido.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            flowables.append(Paragraph(_negrita_markdown_a_html_pdf(texto_html), estilo_normal))
            flowables.append(Spacer(1, 4))
    return flowables


# ---------------------------------------------------------------------------
# Generación del .docx
# ---------------------------------------------------------------------------

def generar_reporte_docx(nombre_entidad, diag, analisis_ia_texto, resultado_isvpt=None, resultado_360=None, idi_oficial=None, cruce_recomendaciones=None, total_recomendaciones_entidad=None, tipo_regimen_especial=None):
    """Devuelve un BytesIO con el informe técnico en formato Word (.docx).

    resultado_isvpt (opcional): un motor_isvpt.ResultadoISVPT ya calculado
    para el grupo de comparación de la entidad. Si se pasa, se agrega una
    sección con el Índice Sintético de Valor Público Territorial (novedad
    metodológica inspirada en el ISDEL de Vélez Tamayo et al., 2026).
    resultado_360 (opcional): un motor_analisis_360.ResultadoAnalisis360 ya
    calculado, usado para el Resumen Ejecutivo (comparación real contra el
    grupo par, igual que en un informe profesional de auditoría).
    idi_oficial (opcional pero MUY recomendado): el IDI oficial publicado
    por Función Pública para esta entidad, leído directamente del archivo
    oficial. Si se proporciona, este informe SIEMPRE lo usa como la cifra
    protagonista (portada, resumen ejecutivo); diag.idi_estimado (el
    cálculo interno de SIIEAP) se muestra únicamente como nota metodológica
    de verificación, nunca como reemplazo del dato oficial.
    """
    doc = Document()

    # Tipografía base más grande y legible (el informe se lee en pantalla y se imprime)
    estilo_normal_doc = doc.styles["Normal"]
    estilo_normal_doc.font.size = Pt(11.5)
    estilo_normal_doc.font.name = "Calibri"
    estilo_normal_doc.paragraph_format.space_after = Pt(8)
    estilo_normal_doc.paragraph_format.line_spacing = 1.15

    # El IDI OFICIAL de Función Pública es siempre la cifra protagonista.
    # Si por alguna razón no se cargó (ej. captura manual), se usa el cálculo
    # interno como única cifra disponible, dejándolo claro en el texto.
    idi_protagonista = idi_oficial if idi_oficial is not None else diag.idi_estimado
    hay_diferencia_idi = (
        idi_oficial is not None and diag.idi_estimado is not None
        and round(idi_oficial, 2) != round(diag.idi_estimado, 2)
    )
    nivel_riesgo_global = "ALTO" if (idi_protagonista or 0) < 40 else ("MEDIO" if (idi_protagonista or 0) < 70 else "BAJO")

    # Portada — logos institucionales + banner de color en vez de un título plano
    _agregar_logos_docx(doc)
    _agregar_banner_docx(
        doc,
        "Modelo de Conocimiento Institucional del Sistema de Inteligencia Artificial "
        "para la Evaluación Integral del Desempeño Institucional en Entidades Públicas (SIIEAP)",
        "Informe de Diagnóstico Institucional y Plan de Mejoramiento Prospectivo",
    )
    _agregar_nota_regimen_especial_docx(doc, tipo_regimen_especial)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nombre_entidad)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"Generado el {_fecha_hoy_es()} · Índice de Desempeño Institucional (IDI-MIPG), Decreto 1499 de 2017")

    p_autoria = doc.add_paragraph()
    p_autoria.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_autoria = p_autoria.add_run(
        "Docente: Norma Elizabeth Álvarez Grajales · Área del conocimiento "
        "Organizaciones Públicas y Gestión · Escuela Superior de Administración Pública (ESAP)"
    )
    run_autoria.italic = True
    run_autoria.font.size = Pt(9)

    # Total de recomendaciones OFICIALES de Función Pública para esta entidad.
    # Prioridad: el total real de recomendaciones cargadas para la entidad
    # (todas las que trae el consolidado/archivo, sin filtrar por brecha).
    # Solo si no se recibió ese dato, se cae al conteo derivado del cruce
    # (que subestima el total real, porque solo cuenta recomendaciones
    # vinculadas a brechas detectadas).
    total_recomendaciones = (
        total_recomendaciones_entidad
        if total_recomendaciones_entidad is not None
        else (sum(len(lista) for lista in cruce_recomendaciones.values()) if cruce_recomendaciones else None)
    )

    doc.add_paragraph()
    n_columnas_portada = 4 if total_recomendaciones is not None else 3
    tabla_portada = doc.add_table(rows=1, cols=n_columnas_portada)
    tabla_portada.style = "Light Grid Accent 1"
    tabla_portada.autofit = False
    tabla_portada.allow_autofit = False
    celdas_portada = tabla_portada.rows[0].cells
    etiqueta_idi = "IDI oficial (Función Pública)" if diag.aplica_mipg_integral else "Índice Control Interno (MECI)"
    celdas_portada[0].text = f"{etiqueta_idi}\n{idi_protagonista}"
    celdas_portada[1].text = f"Nivel de riesgo global\n{nivel_riesgo_global}"
    celdas_portada[2].text = f"Brechas detectadas (dato exclusivo de este informe)\n{len(diag.brechas)}"
    ancho_celda = Cm(4.2) if n_columnas_portada == 4 else Cm(5.5)
    if total_recomendaciones is not None:
        celdas_portada[3].text = f"Recomendaciones oficiales FP\n{total_recomendaciones}"
    for celda in celdas_portada:
        celda.width = ancho_celda
        for parrafo in celda.paragraphs:
            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run_c in parrafo.runs:
                run_c.bold = True
                run_c.font.size = Pt(13)

    doc.add_paragraph()
    p_explica = doc.add_paragraph()
    run_explica = p_explica.add_run(
        "⚠ Brechas detectadas: esta cifra es un cálculo EXCLUSIVO de este informe (SIIEAP) — "
        "NO es un dato que publique el Departamento Administrativo de la Función Pública. "
        "Corresponde al número de índices MIPG de esta entidad con puntaje por debajo de 60 "
        "puntos, umbral de alerta metodológico interno del sistema (la meta plena de la gestión "
        "pública es siempre el 100%, no 60)."
        + (
            f" ✔ Recomendaciones oficiales FP: este dato SÍ es oficial — corresponde al total de "
            f"recomendaciones del banco consolidado del Departamento Administrativo de la Función "
            f"Pública emitidas para esta entidad (no se limita a las brechas detectadas por este "
            f"informe)."
            if total_recomendaciones is not None else ""
        )
    )
    run_explica.italic = True
    run_explica.font.size = Pt(9)
    run_explica.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        p_grupo = doc.add_paragraph()
        p_grupo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_grupo.add_run(
            f"Grupo de comparación: {resultado_360.filtro_descripcion} "
            f"({resultado_360.n_entidades} entidades) · IDI promedio del grupo: {resultado_360.promedio_idi}"
        )
    if hay_diferencia_idi:
        p_nota_idi = doc.add_paragraph()
        p_nota_idi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nota_idi = p_nota_idi.add_run(
            f"(Nota: el cálculo interno de verificación de SIIEAP, revisando la suma/ponderación de los "
            f"índices reportados, arroja {diag.idi_estimado}. Esta es una cifra de validación metodológica "
            f"interna que NO reemplaza ni desvirtúa el IDI oficial de Función Pública, que es siempre el "
            f"que prevalece.)"
        )
        run_nota_idi.italic = True
        run_nota_idi.font.size = Pt(9)
        run_nota_idi.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    _agregar_tabla_contenido_docx(doc)
    _agregar_razon_de_ser_docx(doc, "tecnico")

    _agregar_glosario_docx(doc)
    _agregar_normativa_politicas_docx(doc, diag=diag)

    # 1. Resumen ejecutivo — cifras reales de ESTA entidad, arriba de todo
    _agregar_divisor_seccion_docx(doc, "1. Resumen ejecutivo", icono="📊")
    parrafos_resumen = [
        f"{nombre_entidad} obtuvo un Índice de Desempeño Institucional (IDI) oficial de {idi_protagonista} "
        f"sobre 100 en la vigencia analizada, con un nivel de riesgo global {nivel_riesgo_global}."
    ]
    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        brecha_grupo = round((idi_protagonista or 0) - resultado_360.promedio_idi, 2)
        comparativo = "por debajo" if brecha_grupo < 0 else "por encima"
        parrafos_resumen.append(
            f"Frente a su grupo de comparación ({resultado_360.filtro_descripcion}, "
            f"{resultado_360.n_entidades} entidades), la entidad se ubica {abs(brecha_grupo)} puntos "
            f"{comparativo} del promedio del grupo ({resultado_360.promedio_idi})."
        )
        if resultado_360.percentil_entidad_referencia is not None:
            parrafos_resumen.append(
                f"Esto la ubica en el percentil {resultado_360.percentil_entidad_referencia}% de su grupo."
            )
    dims_criticas = sorted(diag.resultados_por_dimension, key=lambda r: (r.promedio if r.promedio is not None else 0))[:3]
    if dims_criticas:
        nombres_criticas = ", ".join(f"{r.codigo} {r.nombre} ({r.promedio})" for r in dims_criticas)
        parrafos_resumen.append(f"Las dimensiones más críticas son: {nombres_criticas}.")
    dims_fuertes = [r for r in diag.resultados_por_dimension if (r.promedio or 0) >= 60]
    if dims_fuertes:
        nombres_fuertes = ", ".join(f"{r.codigo} {r.nombre} ({r.promedio})" for r in dims_fuertes)
        parrafos_resumen.append(f"Como fortaleza relativa, se destaca(n): {nombres_fuertes}.")
    else:
        parrafos_resumen.append(
            "Ninguna dimensión alcanza el umbral de 60 puntos: el patrón de brechas simultáneas en "
            "múltiples dimensiones sugiere una debilidad estructural de capacidad institucional "
            "(Oszlak), más que fallas puntuales y aisladas de gestión."
        )
    if hay_diferencia_idi:
        parrafos_resumen.append(
            f"Nota metodológica: Función Pública reporta un IDI oficial de {idi_oficial} para esta "
            f"entidad, cifra que este informe usa siempre como referencia principal. El sistema SIIEAP, "
            f"al revisar internamente la suma y ponderación de los índices del archivo oficial, calcula "
            f"además un valor de verificación de {diag.idi_estimado}; esta diferencia es un insumo para "
            f"depurar la metodología interna del sistema, y en ningún caso reemplaza, desvirtúa ni "
            f"contradice la cifra oficial de Función Pública."
        )
    parrafos_resumen.append(
        f"Este informe (SIIEAP) detectó, con metodología propia y exclusiva del sistema — no con "
        f"una cifra publicada por la Función Pública —, {len(diag.brechas)} brechas de "
        "implementación por debajo de 60 puntos (el umbral de alerta metodológico interno que usa "
        "este sistema para priorizar el análisis; la meta plena de la gestión pública, se insiste, "
        "es el 100% de cumplimiento, no 60 puntos), todas desarrolladas en detalle en este informe "
        "— sin selección ni recorte."
    )
    for parrafo_r in parrafos_resumen:
        doc.add_paragraph(parrafo_r)

    # Contextualización académica (bloque fijo)
    _agregar_divisor_seccion_docx(doc, CONTEXTUALIZACION_TITULO, icono="📚")
    doc.add_paragraph(CONTEXTUALIZACION_INTRO)

    doc.add_heading("Arco teórico: de los griegos a la Administración Pública contemporánea", level=2)
    for titulo_hito, texto_hito in CONTEXTUALIZACION_ARCO_HISTORICO:
        p = doc.add_paragraph()
        run_hito = p.add_run(f"{titulo_hito}. ")
        run_hito.bold = True
        p.add_run(texto_hito)

    doc.add_heading(CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_TITULO, level=2)
    doc.add_paragraph(CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_INTRO)
    for titulo_teoria, texto_teoria in CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS:
        p_teoria = doc.add_paragraph()
        run_titulo_teoria = p_teoria.add_run(f"{titulo_teoria}. ")
        run_titulo_teoria.bold = True
        p_teoria.add_run(texto_teoria)

    doc.add_heading("Esquema integrador de la Administración Pública contemporánea", level=2)
    p_esquema = doc.add_paragraph()
    run_esquema = p_esquema.add_run(" → ".join(CONTEXTUALIZACION_ESQUEMA_INTEGRADOR))
    run_esquema.italic = True
    doc.add_paragraph(CONTEXTUALIZACION_CIERRE)

    doc.add_page_break()

    # Cadena de interpretación institucional (bloque fijo)
    _agregar_divisor_seccion_docx(doc, CADENA_INTERPRETACION_TITULO, icono="🔗")
    doc.add_paragraph(CADENA_INTERPRETACION_INTRO)
    for paso in CADENA_INTERPRETACION_PASOS:
        doc.add_paragraph(paso, style="List Bullet")
    doc.add_paragraph()
    p_cierre_cadena = doc.add_paragraph()
    run_cierre_cadena = p_cierre_cadena.add_run(CADENA_INTERPRETACION_CIERRE)
    run_cierre_cadena.italic = True

    doc.add_page_break()

    # Resultado real del diagnóstico
    _agregar_divisor_seccion_docx(doc, "Resultado del diagnóstico institucional (datos reales)", icono="📈")
    p_idi = doc.add_paragraph()
    run_idi = p_idi.add_run(f"IDI estimado: {diag.idi_estimado}")
    run_idi.bold = True
    run_idi.font.size = Pt(14)

    tabla = doc.add_table(rows=1, cols=4)
    tabla.style = "Light Grid Accent 1"
    encabezados = tabla.rows[0].cells
    for celda, texto in zip(encabezados, ["Dimensión", "Promedio", "Riesgo", "Índices"]):
        celda.text = texto
        for parrafo in celda.paragraphs:
            for run_enc in parrafo.runs:
                run_enc.bold = True

    for r in diag.resultados_por_dimension:
        fila = tabla.add_row().cells
        fila[0].text = f"{r.codigo} {r.nombre}"
        fila[1].text = str(r.promedio)
        fila[2].text = str(r.nivel_riesgo)
        fila[3].text = f"{r.n_indices_evaluados}/{r.n_indices_esperados}"
        color_fondo = _color_hex_riesgo(r.nivel_riesgo)
        if color_fondo:
            _sombrear_celda(fila[2], color_fondo)
        color_quintil = _color_hex_quintil_mipg(r.promedio)
        if color_quintil:
            _sombrear_celda(fila[1], color_quintil)
    _ajustar_tabla_docx(tabla, anchos_cm=[9.5, 2.3, 2.3, 2.0])

    doc.add_paragraph()
    try:
        buffer_grafica_dim = generar_grafica_dimensiones(diag)
        doc.add_picture(buffer_grafica_dim, width=Inches(6.2))
        parrafo_imagen = doc.paragraphs[-1]
        parrafo_imagen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass  # si falla la gráfica, el informe sigue sin ella

    doc.add_heading("Brechas priorizadas (todas las detectadas)", level=2)
    if diag.brechas:
        try:
            buffer_grafica_brechas = generar_grafica_brechas(diag)
            if buffer_grafica_brechas:
                doc.add_picture(buffer_grafica_brechas, width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
        except Exception:
            pass
        for b in diag.brechas:
            doc.add_paragraph(
                f"{b.codigo_indice} ({b.puntaje}) — {b.nombre_indice} · {b.politica}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("No se detectaron brechas por debajo del umbral con los datos disponibles.")

    # Tabla de conexión política -> enfoque contemporáneo -> norma (garantizada por código)
    if diag.brechas:
        doc.add_page_break()
        doc.add_heading("Conexión de las brechas con los enfoques contemporáneos de la Administración Pública", level=2)
        doc.add_paragraph(
            "Cada política con brechas se conecta aquí, de forma sistemática, con el enfoque "
            "contemporáneo que mejor la explica (de los 15 desarrollados en la contextualización "
            "de este informe) y con la norma colombiana que lo respalda."
        )
        politicas_con_brecha = []
        vistas = set()
        for b in diag.brechas:
            if b.politica not in vistas:
                vistas.add(b.politica)
                politicas_con_brecha.append(b.politica)
        tabla_enfoques = doc.add_table(rows=1, cols=3)
        tabla_enfoques.style = "Light Grid Accent 1"
        enc_ef = tabla_enfoques.rows[0].cells
        for celda, texto in zip(enc_ef, ["Política con brecha", "Enfoque contemporáneo", "Norma"]):
            celda.text = texto
            for parrafo in celda.paragraphs:
                for run_enc in parrafo.runs:
                    run_enc.bold = True
        for politica in politicas_con_brecha:
            enfoque, norma = _enfoque_y_norma_de_politica(politica)
            fila = tabla_enfoques.add_row().cells
            fila[0].text = politica
            fila[1].text = enfoque
            fila[2].text = norma
        _ajustar_tabla_docx(tabla_enfoques, anchos_cm=[5.0, 4.5, 6.0], tamano_fuente_pt=8.5)
        _franjas_alternas_docx(tabla_enfoques)

    if resultado_isvpt is not None and resultado_isvpt.isvpt_entidad_referencia is not None:
        doc.add_page_break()
        _agregar_divisor_seccion_docx(doc, "El Termómetro del Valor Público: Índice Sintético de Valor Público Territorial (ISVPT)", icono="🎯")
        doc.add_paragraph(
            "Como complemento al IDI oficial, este informe incorpora un ejercicio de índice "
            "sintético construido con la metodología académica validada por Vélez Tamayo, "
            "Ortiz-Muñoz y Cardona Montoya (2026) para el Índice Sintético de Desarrollo "
            "Económico Local (ISDEL) — normalización min-max y agregación aritmética simple, "
            "siguiendo las directrices de la OCDE (2008) para indicadores compuestos — "
            "aplicada aquí a las 7 dimensiones reales del IDI-MIPG dentro del grupo de "
            "comparación de esta entidad."
        )
        p_isvpt = doc.add_paragraph()
        run_isvpt = p_isvpt.add_run(
            f"ISVPT de {nombre_entidad}: {resultado_isvpt.isvpt_entidad_referencia} "
            f"(posición {resultado_isvpt.posicion_entidad_referencia} de "
            f"{resultado_isvpt.n_entidades} en su grupo de comparación)."
        )
        run_isvpt.bold = True

        if resultado_isvpt.subindices_por_dimension_entidad:
            doc.add_heading("Subíndices normalizados por dimensión (0 = el más rezagado del grupo, 1 = el mejor)", level=2)
            tabla_isvpt = doc.add_table(rows=1, cols=2)
            tabla_isvpt.style = "Light Grid Accent 1"
            enc = tabla_isvpt.rows[0].cells
            enc[0].text, enc[1].text = "Dimensión", "Subíndice normalizado"
            for celda in enc:
                for parrafo in celda.paragraphs:
                    for run_enc in parrafo.runs:
                        run_enc.bold = True
            for dim, valor in resultado_isvpt.subindices_por_dimension_entidad.items():
                fila = tabla_isvpt.add_row().cells
                fila[0].text = dim
                fila[1].text = str(valor)

        doc.add_paragraph()
        p_nota_isvpt = doc.add_paragraph()
        run_nota_isvpt = p_nota_isvpt.add_run(resultado_isvpt.nota_metodologica)
        run_nota_isvpt.italic = True
        run_nota_isvpt.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if resultado_isvpt.nota_valor_extremo:
            p_nota_extremo = doc.add_paragraph()
            p_nota_extremo.paragraph_format.space_before = Pt(6)
            run_nota_extremo = p_nota_extremo.add_run(resultado_isvpt.nota_valor_extremo)
            run_nota_extremo.bold = True
            run_nota_extremo.font.color.rgb = RGBColor(0xC0, 0x50, 0x00)

    doc.add_page_break()

    # Análisis integral IA
    _agregar_divisor_seccion_docx(doc, "Análisis integral, señales de riesgo multinivel y Plan de Mejoramiento Prospectivo (generado por IA a partir de los datos reales)", icono="⚠️")
    p_cita_riesgo = doc.add_paragraph()
    run_cita_riesgo = p_cita_riesgo.add_run(
        "La lectura de riesgo institucional que sigue retoma la metodología de gestión integrada "
        "de riesgos (ISO 31000) documentada por Jurado-Zambrano y Villanueva (2021, ESAP) para el "
        "sector público colombiano, aplicada aquí a los datos reales de esta entidad."
    )
    run_cita_riesgo.italic = True
    run_cita_riesgo.font.size = Pt(9.5)
    _agregar_texto_markdown_docx(doc, analisis_ia_texto)

    doc.add_page_break()

    # Disclaimer
    doc.add_heading("Nota metodológica", level=2)
    p_disc = doc.add_paragraph()
    run_disc = p_disc.add_run(DISCLAIMER_INFORME)
    run_disc.italic = True
    run_disc.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _agregar_marco_descentralizacion_docx(doc, diag=diag, nombre_entidad=nombre_entidad)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# Generación del PDF (reportlab)
# ---------------------------------------------------------------------------

def generar_reporte_pdf(nombre_entidad, diag, analisis_ia_texto, resultado_isvpt=None, resultado_360=None, idi_oficial=None, cruce_recomendaciones=None, total_recomendaciones_entidad=None, tipo_regimen_especial=None):
    """Devuelve un BytesIO con el informe técnico en formato PDF.

    resultado_isvpt (opcional): ver docstring de generar_reporte_docx.
    resultado_360 (opcional): ver docstring de generar_reporte_docx.
    idi_oficial (opcional pero MUY recomendado): ver docstring de generar_reporte_docx.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=LETTER,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )
    estilos = getSampleStyleSheet()
    estilo_titulo = ParagraphStyle("TituloSIIEAP", parent=estilos["Title"], fontSize=20)
    estilo_h2 = ParagraphStyle("H2SIIEAP", parent=estilos["Heading2"], spaceBefore=14, fontSize=14)
    estilo_normal = ParagraphStyle("NormalSIIEAP", parent=estilos["Normal"], fontSize=11, leading=15)
    estilo_cursiva = ParagraphStyle("Cursiva", parent=estilo_normal, fontName="Helvetica-Oblique", textColor=colors.grey)

    idi_protagonista = idi_oficial if idi_oficial is not None else diag.idi_estimado
    hay_diferencia_idi = (
        idi_oficial is not None and diag.idi_estimado is not None
        and round(idi_oficial, 2) != round(diag.idi_estimado, 2)
    )
    nivel_riesgo_global = "ALTO" if (idi_protagonista or 0) < 40 else ("MEDIO" if (idi_protagonista or 0) < 70 else "BAJO")

    elementos = []

    # Portada — logos institucionales + banner de color en vez de un título plano
    elementos.extend(_logos_pdf_flowables())
    elementos.extend(_banner_portada_pdf_flowables(
        "Modelo de Conocimiento Institucional del Sistema de Inteligencia Artificial "
        "para la Evaluación Integral del Desempeño Institucional en Entidades Públicas (SIIEAP)",
        "Informe de Diagnóstico Institucional y Plan de Mejoramiento Prospectivo",
    ))
    elementos.extend(_nota_regimen_especial_pdf_flowables(tipo_regimen_especial))
    elementos.append(Paragraph(f"<b>{nombre_entidad}</b>", ParagraphStyle("EntidadSIIEAP", parent=estilos["Heading2"], textColor=colors.HexColor("#1F3864"))))
    elementos.append(Paragraph(f"Generado el {_fecha_hoy_es()} · Índice de Desempeño Institucional (IDI-MIPG), Decreto 1499 de 2017", estilo_normal))
    elementos.append(Paragraph(
        "Docente: Norma Elizabeth Álvarez Grajales · Área del conocimiento "
        "Organizaciones Públicas y Gestión · Escuela Superior de Administración Pública (ESAP)",
        ParagraphStyle("AutoriaSIIEAP", parent=estilo_normal, fontSize=9, fontName="Helvetica-Oblique"),
    ))
    elementos.append(Spacer(1, 14))

    # Igual que en la versión docx: se prioriza el total real de recomendaciones
    # de la entidad (todas las cargadas, sin filtrar por brecha) sobre el conteo
    # derivado del cruce, que subestima el total oficial.
    total_recomendaciones = (
        total_recomendaciones_entidad
        if total_recomendaciones_entidad is not None
        else (sum(len(lista) for lista in cruce_recomendaciones.values()) if cruce_recomendaciones else None)
    )

    etiqueta_idi_pdf = "IDI oficial (Función Pública)" if diag.aplica_mipg_integral else "Índice Control Interno (MECI)"
    if total_recomendaciones is not None:
        datos_tabla_portada = [
            [etiqueta_idi_pdf, "Nivel de riesgo global", "Brechas detectadas\n(dato exclusivo de este informe)", "Recomendaciones oficiales FP"],
            [str(idi_protagonista), nivel_riesgo_global, str(len(diag.brechas)), str(total_recomendaciones)],
        ]
        anchos_portada = [4 * cm, 4 * cm, 4 * cm, 4.5 * cm]
    else:
        datos_tabla_portada = [
            [etiqueta_idi_pdf, "Nivel de riesgo global", "Brechas detectadas\n(dato exclusivo de este informe)"],
            [str(idi_protagonista), nivel_riesgo_global, str(len(diag.brechas))],
        ]
        anchos_portada = [5 * cm, 5 * cm, 5 * cm]
    tabla_portada = Table(datos_tabla_portada, hAlign="CENTER", colWidths=anchos_portada)
    tabla_portada.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("FONTSIZE", (0, 1), (-1, 1), 14),
        ("FONTNAME", (0, 1), (-1, 1), "Helvetica-Bold"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elementos.append(tabla_portada)
    elementos.append(Spacer(1, 6))
    texto_explica_portada = (
        "⚠ Brechas detectadas: cifra EXCLUSIVA de este informe (SIIEAP) — NO es un dato que "
        "publique el Departamento Administrativo de la Función Pública. Corresponde al número de "
        "índices MIPG de esta entidad con puntaje por debajo de 60 puntos, umbral de alerta "
        "metodológico interno del sistema (la meta plena de la gestión pública es siempre el "
        "100%, no 60)."
    )
    if total_recomendaciones is not None:
        texto_explica_portada += (
            " ✔ Recomendaciones oficiales FP: este dato SÍ es oficial — corresponde al total de "
            "recomendaciones del banco consolidado del Departamento Administrativo de la Función "
            "Pública emitidas para esta entidad (no se limita a las brechas detectadas por este "
            "informe)."
        )
    elementos.append(Paragraph(texto_explica_portada, estilo_cursiva))
    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        elementos.append(Spacer(1, 10))
        elementos.append(Paragraph(
            f"Grupo de comparación: {resultado_360.filtro_descripcion} "
            f"({resultado_360.n_entidades} entidades) · IDI promedio del grupo: {resultado_360.promedio_idi}",
            estilo_cursiva,
        ))
    if hay_diferencia_idi:
        elementos.append(Spacer(1, 8))
        elementos.append(Paragraph(
            f"(Nota: el cálculo interno de verificación de SIIEAP, revisando la suma/ponderación de los "
            f"índices reportados, arroja {diag.idi_estimado}. Esta es una cifra de validación metodológica "
            f"interna que NO reemplaza ni desvirtúa el IDI oficial de Función Pública, que es siempre el "
            f"que prevalece.)",
            ParagraphStyle("NotaIDI", parent=estilo_cursiva, fontSize=8),
        ))
    elementos.append(PageBreak())

    elementos.extend(_toc_pdf_flowables([
        "Contextualización de la Administración Pública Contemporánea",
        "Diagnóstico institucional por dimensión",
        "Índice Sintético de Valor Público Territorial (ISVPT)",
        "Análisis integral generado por IA",
        "Plan de mejoramiento prospectivo",
        "Glosario y notas de trazabilidad",
    ]))
    elementos.extend(_razon_de_ser_pdf_flowables("tecnico"))

    _agregar_glosario_pdf(elementos, estilos, estilo_normal, estilo_h2)
    _agregar_normativa_politicas_pdf(elementos, estilos, estilo_normal, estilo_h2, diag=diag)

    # 1. Resumen ejecutivo — cifras reales de ESTA entidad, arriba de todo
    elementos.extend(_divisor_seccion_pdf("1. Resumen ejecutivo", icono="📊"))
    parrafos_resumen = [
        f"{nombre_entidad} obtuvo un Índice de Desempeño Institucional (IDI) oficial de {idi_protagonista} "
        f"sobre 100 en la vigencia analizada, con un nivel de riesgo global {nivel_riesgo_global}."
    ]
    if resultado_360 is not None and resultado_360.promedio_idi is not None:
        brecha_grupo = round((idi_protagonista or 0) - resultado_360.promedio_idi, 2)
        comparativo = "por debajo" if brecha_grupo < 0 else "por encima"
        parrafos_resumen.append(
            f"Frente a su grupo de comparación ({resultado_360.filtro_descripcion}, "
            f"{resultado_360.n_entidades} entidades), la entidad se ubica {abs(brecha_grupo)} puntos "
            f"{comparativo} del promedio del grupo ({resultado_360.promedio_idi})."
        )
        if resultado_360.percentil_entidad_referencia is not None:
            parrafos_resumen.append(f"Esto la ubica en el percentil {resultado_360.percentil_entidad_referencia}% de su grupo.")
    dims_criticas = sorted(diag.resultados_por_dimension, key=lambda r: (r.promedio if r.promedio is not None else 0))[:3]
    if dims_criticas:
        nombres_criticas = ", ".join(f"{r.codigo} {r.nombre} ({r.promedio})" for r in dims_criticas)
        parrafos_resumen.append(f"Las dimensiones más críticas son: {nombres_criticas}.")
    dims_fuertes = [r for r in diag.resultados_por_dimension if (r.promedio or 0) >= 60]
    if dims_fuertes:
        nombres_fuertes = ", ".join(f"{r.codigo} {r.nombre} ({r.promedio})" for r in dims_fuertes)
        parrafos_resumen.append(f"Como fortaleza relativa, se destaca(n): {nombres_fuertes}.")
    else:
        parrafos_resumen.append(
            "Ninguna dimensión alcanza el umbral de 60 puntos: el patrón de brechas simultáneas en "
            "múltiples dimensiones sugiere una debilidad estructural de capacidad institucional "
            "(Oszlak), más que fallas puntuales y aisladas de gestión."
        )
    if hay_diferencia_idi:
        parrafos_resumen.append(
            f"Nota metodológica: Función Pública reporta un IDI oficial de {idi_oficial} para esta "
            f"entidad, cifra que este informe usa siempre como referencia principal. El sistema SIIEAP, "
            f"al revisar internamente la suma y ponderación de los índices del archivo oficial, calcula "
            f"además un valor de verificación de {diag.idi_estimado}; esta diferencia es un insumo para "
            f"depurar la metodología interna del sistema, y en ningún caso reemplaza, desvirtúa ni "
            f"contradice la cifra oficial de Función Pública."
        )
    parrafos_resumen.append(
        f"Este informe (SIIEAP) detectó, con metodología propia y exclusiva del sistema — no con "
        f"una cifra publicada por la Función Pública —, {len(diag.brechas)} brechas de "
        "implementación por debajo de 60 puntos (el umbral de alerta metodológico interno que usa "
        "este sistema para priorizar el análisis; la meta plena de la gestión pública, se insiste, "
        "es el 100% de cumplimiento, no 60 puntos), todas desarrolladas en detalle en este informe "
        "— sin selección ni recorte."
    )
    for parrafo_r in parrafos_resumen:
        elementos.append(Paragraph(parrafo_r, estilo_normal))
        elementos.append(Spacer(1, 4))
    elementos.append(PageBreak())

    # Contextualización
    elementos.extend(_divisor_seccion_pdf(CONTEXTUALIZACION_TITULO, icono="📚"))
    elementos.append(Paragraph(CONTEXTUALIZACION_INTRO, estilo_normal))
    elementos.append(Paragraph("Arco teórico: de los griegos a la Administración Pública contemporánea", estilo_h2))
    for titulo_hito, texto_hito in CONTEXTUALIZACION_ARCO_HISTORICO:
        elementos.append(Paragraph(f"<b>{titulo_hito}.</b> {texto_hito}", estilo_normal))
        elementos.append(Spacer(1, 6))

    elementos.append(Paragraph(CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_TITULO, estilo_h2))
    elementos.append(Paragraph(CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS_INTRO, estilo_normal))
    elementos.append(Spacer(1, 4))
    for titulo_teoria, texto_teoria in CONTEXTUALIZACION_TEORIAS_CONTEMPORANEAS:
        elementos.append(Paragraph(f"<b>{titulo_teoria}.</b> {texto_teoria}", estilo_normal))
        elementos.append(Spacer(1, 6))

    elementos.append(Paragraph("Esquema integrador de la Administración Pública contemporánea", estilo_h2))
    elementos.append(Paragraph("<i>" + " → ".join(CONTEXTUALIZACION_ESQUEMA_INTEGRADOR) + "</i>", estilo_normal))
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph(CONTEXTUALIZACION_CIERRE, estilo_normal))
    elementos.append(PageBreak())

    # Cadena de interpretación institucional (bloque fijo)
    elementos.extend(_divisor_seccion_pdf(CADENA_INTERPRETACION_TITULO, icono="🔗"))
    elementos.append(Paragraph(CADENA_INTERPRETACION_INTRO, estilo_normal))
    elementos.append(Spacer(1, 6))
    for paso in CADENA_INTERPRETACION_PASOS:
        elementos.append(Paragraph(f"• {paso}", estilo_normal))
    elementos.append(Spacer(1, 8))
    elementos.append(Paragraph(CADENA_INTERPRETACION_CIERRE, estilo_cursiva))
    elementos.append(PageBreak())

    # Resultado real del diagnóstico
    elementos.extend(_divisor_seccion_pdf("Resultado del diagnóstico institucional (datos reales)", icono="📈"))
    elementos.append(Paragraph(f"<b>IDI estimado: {diag.idi_estimado}</b>", ParagraphStyle("IDIGrande", parent=estilo_normal, fontSize=15)))
    elementos.append(Spacer(1, 8))

    datos_tabla = [[_celda_pdf("Dimensión", encabezado=True, tamano_fuente=9), "Promedio", "Riesgo", "Índices"]]
    for r in diag.resultados_por_dimension:
        datos_tabla.append([
            _celda_pdf(f"{r.codigo} {r.nombre}", tamano_fuente=9), str(r.promedio), str(r.nivel_riesgo),
            f"{r.n_indices_evaluados}/{r.n_indices_esperados}",
        ])
    tabla = Table(datos_tabla, hAlign="LEFT", colWidths=[7 * cm, 2.5 * cm, 2.5 * cm, 2.5 * cm])
    estilo_tabla_dim = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]
    _COLOR_PDF_POR_RIESGO = {
        "alta": colors.HexColor("#F5B7B1"),
        "media": colors.HexColor("#FAD7A0"),
        "baja": colors.HexColor("#A9DFBF"),
    }
    for indice_fila, r in enumerate(diag.resultados_por_dimension, start=1):
        color_riesgo = _COLOR_PDF_POR_RIESGO.get(str(r.nivel_riesgo).strip().lower())
        if color_riesgo:
            estilo_tabla_dim.append(("BACKGROUND", (2, indice_fila), (2, indice_fila), color_riesgo))
        color_quintil_hex = _color_hex_quintil_mipg(r.promedio)
        if color_quintil_hex:
            estilo_tabla_dim.append(("BACKGROUND", (1, indice_fila), (1, indice_fila), colors.HexColor(f"#{color_quintil_hex}")))
    tabla.setStyle(TableStyle(estilo_tabla_dim))
    elementos.append(tabla)
    elementos.append(Spacer(1, 10))

    try:
        buffer_grafica_dim = generar_grafica_dimensiones(diag)
        elementos.append(Image(buffer_grafica_dim, width=16 * cm, height=16 * cm * 0.5))
        elementos.append(Spacer(1, 10))
    except Exception:
        pass

    elementos.append(Paragraph("Brechas priorizadas (todas las detectadas)", estilo_h2))
    if diag.brechas:
        try:
            buffer_grafica_brechas = generar_grafica_brechas(diag)
            if buffer_grafica_brechas:
                elementos.append(Image(buffer_grafica_brechas, width=16 * cm, height=16 * cm * 0.55))
                elementos.append(Spacer(1, 10))
        except Exception:
            pass
        for b in diag.brechas:
            elementos.append(Paragraph(
                f"• {b.codigo_indice} ({b.puntaje}) — {b.nombre_indice} · {b.politica}", estilo_normal,
            ))
    else:
        elementos.append(Paragraph("No se detectaron brechas por debajo del umbral con los datos disponibles.", estilo_normal))
    elementos.append(PageBreak())

    # Tabla de conexión política -> enfoque contemporáneo -> norma (garantizada por código)
    if diag.brechas:
        elementos.append(Paragraph("Conexión de las brechas con los enfoques contemporáneos de la Administración Pública", estilo_h2))
        elementos.append(Paragraph(
            "Cada política con brechas se conecta aquí, de forma sistemática, con el enfoque "
            "contemporáneo que mejor la explica y con la norma colombiana que lo respalda.",
            estilo_normal,
        ))
        elementos.append(Spacer(1, 6))
        politicas_con_brecha = []
        vistas = set()
        for b in diag.brechas:
            if b.politica not in vistas:
                vistas.add(b.politica)
                politicas_con_brecha.append(b.politica)
        datos_tabla_enfoques = [[_celda_pdf("Política con brecha", encabezado=True), _celda_pdf("Enfoque contemporáneo", encabezado=True), _celda_pdf("Norma", encabezado=True)]]
        for politica in politicas_con_brecha:
            enfoque, norma = _enfoque_y_norma_de_politica(politica)
            datos_tabla_enfoques.append([_celda_pdf(politica), _celda_pdf(enfoque), _celda_pdf(norma)])
        tabla_enfoques = Table(datos_tabla_enfoques, hAlign="LEFT", colWidths=[5 * cm, 5 * cm, 6 * cm])
        tabla_enfoques.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F4F8")]),
        ]))
        elementos.append(tabla_enfoques)
        elementos.append(PageBreak())

    if resultado_isvpt is not None and resultado_isvpt.isvpt_entidad_referencia is not None:
        elementos.extend(_divisor_seccion_pdf("El Termómetro del Valor Público: Índice Sintético de Valor Público Territorial (ISVPT)", icono="🎯"))
        elementos.append(Paragraph(
            "Como complemento al IDI oficial, este informe incorpora un ejercicio de índice "
            "sintético construido con la metodología académica validada por Vélez Tamayo, "
            "Ortiz-Muñoz y Cardona Montoya (2026) para el ISDEL — normalización min-max y "
            "agregación aritmética simple, siguiendo las directrices de la OCDE (2008) para "
            "indicadores compuestos — aplicada aquí a las 7 dimensiones reales del IDI-MIPG "
            "dentro del grupo de comparación de esta entidad.",
            estilo_normal,
        ))
        elementos.append(Spacer(1, 6))
        elementos.append(Paragraph(
            f"<b>ISVPT de {nombre_entidad}: {resultado_isvpt.isvpt_entidad_referencia} "
            f"(posición {resultado_isvpt.posicion_entidad_referencia} de "
            f"{resultado_isvpt.n_entidades} en su grupo de comparación).</b>",
            estilo_normal,
        ))
        elementos.append(Spacer(1, 8))

        if resultado_isvpt.subindices_por_dimension_entidad:
            elementos.append(Paragraph("Subíndices normalizados por dimensión (0 = más rezagado del grupo, 1 = mejor del grupo)", estilo_h2))
            datos_tabla_isvpt = [["Dimensión", "Subíndice normalizado"]]
            for dim, valor in resultado_isvpt.subindices_por_dimension_entidad.items():
                datos_tabla_isvpt.append([dim, str(valor)])
            tabla_isvpt = Table(datos_tabla_isvpt, hAlign="LEFT", colWidths=[10 * cm, 4 * cm])
            tabla_isvpt.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            elementos.append(tabla_isvpt)
            elementos.append(Spacer(1, 8))

        elementos.append(Paragraph(resultado_isvpt.nota_metodologica, estilo_cursiva))

        if resultado_isvpt.nota_valor_extremo:
            estilo_nota_extremo = ParagraphStyle(
                "NotaExtremoISVPT", parent=estilo_normal,
                textColor=colors.HexColor("#C05000"), fontName="Helvetica-Bold",
                spaceBefore=6,
            )
            elementos.append(Paragraph(resultado_isvpt.nota_valor_extremo, estilo_nota_extremo))

        elementos.append(PageBreak())

    # Análisis integral IA
    elementos.extend(_divisor_seccion_pdf("Análisis integral, señales de riesgo multinivel y Plan de Mejoramiento Prospectivo (generado por IA a partir de los datos reales)", icono="⚠️"))
    elementos.append(Paragraph(
        "La lectura de riesgo institucional que sigue retoma la metodología de gestión integrada de "
        "riesgos (ISO 31000) documentada por Jurado-Zambrano y Villanueva (2021, ESAP) para el sector "
        "público colombiano, aplicada aquí a los datos reales de esta entidad.",
        estilo_cursiva,
    ))
    elementos.append(Spacer(1, 6))
    elementos.extend(_texto_markdown_a_pdf_flowables(analisis_ia_texto, estilo_normal, estilo_h2))
    elementos.append(PageBreak())

    # Disclaimer
    elementos.append(Paragraph("Nota metodológica", estilos["Heading2"]))
    elementos.append(Paragraph(DISCLAIMER_INFORME, estilo_cursiva))

    _agregar_marco_descentralizacion_pdf(elementos, estilos, estilo_normal, estilo_h2, diag=diag, nombre_entidad=nombre_entidad)

    def _pie_de_pagina(canvas_pdf, doc_pdf):
        canvas_pdf.saveState()
        canvas_pdf.setFont("Helvetica", 8)
        canvas_pdf.setFillColor(colors.grey)
        canvas_pdf.drawString(2 * cm, 1.3 * cm, f"SIIEAP — {nombre_entidad}")
        canvas_pdf.drawRightString(LETTER[0] - 2 * cm, 1.3 * cm, f"Página {doc_pdf.page}")
        canvas_pdf.restoreState()

    doc.build(elementos, onFirstPage=_pie_de_pagina, onLaterPages=_pie_de_pagina)
    buffer.seek(0)
    return buffer
