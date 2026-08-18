# -*- coding: utf-8 -*-
"""
Contenido de referencia extraído y parafraseado de:
- Guía para la Gestión Estratégica y el Fortalecimiento de la Auditoría
  Interna en Entidades Públicas, V5, julio 2026 (Función Pública / IIA Colombia)
- Guía de Gestión Integral del Riesgo V7, septiembre 2025 (Función Pública)
- Programa de Transparencia y Ética Pública 2025 (Función Pública)
- Bienvenida a auditores de control interno (roles OCI, Decreto 648 de 2017)

Todo el contenido aquí es metodología GENÉRICA de estas guías. Los datos
específicos de cada entidad (puntajes, brechas, recomendaciones) se inyectan
aparte, siempre desde Territorio/Nación y el banco de recomendaciones —
nunca se inventan aquí.
"""

# --- Escalas oficiales de la Guía de Gestión Integral del Riesgo V7 ---
ESCALA_PROBABILIDAD = [
    ("Muy Baja", "20%"),
    ("Baja", "40%"),
    ("Media", "60%"),
    ("Alta", "80%"),
    ("Muy Alta", "100%"),
]

ESCALA_IMPACTO = [
    ("Leve", "20%"),
    ("Menor", "40%"),
    ("Moderado", "60%"),
    ("Mayor", "80%"),
    ("Catastrófico", "100%"),
]

NOTA_METODOLOGICA_RIESGO = (
    "Nota metodológica: la Guía de Gestión Integral del Riesgo V7 de Función Pública "
    "calcula la probabilidad a partir de la frecuencia real con la que se ejecuta cada "
    "actividad (veces por año), dato operativo interno que no está disponible en las "
    "fuentes públicas usadas para este informe (Territorio/Nación). Por eso, para fines "
    "de priorización, este informe deriva el nivel de riesgo inherente por dimensión "
    "directamente del puntaje oficial FP (a menor puntaje, mayor riesgo), sin sustituir el "
    "análisis de frecuencia y controles que la entidad debe desarrollar internamente "
    "siguiendo la metodología completa de la Guía."
)

CRITICIDAD_NIVELES = ["Muy baja", "Baja", "Media", "Alta", "Muy Alta"]


def nivel_criticidad_desde_puntaje(puntaje: float) -> str:
    """Deriva un nivel de criticidad (escala oficial de la Guía de Auditoría V5,
    Tabla de criticidad de hallazgos) a partir del puntaje oficial del índice.
    Puntajes más bajos = mayor criticidad. Regla explícita y trazable:
    <20 Muy Alta, 20-<35 Alta, 35-<50 Media, 50-<60 Baja, >=60 Muy baja (sin brecha)."""
    if puntaje is None:
        return "Sin información"
    if puntaje < 20:
        return "Muy Alta"
    if puntaje < 35:
        return "Alta"
    if puntaje < 50:
        return "Media"
    if puntaje < 60:
        return "Baja"
    return "Muy baja"


def nivel_riesgo_inherente_desde_puntaje(puntaje: float) -> str:
    """Nivel de riesgo inherente (matriz de calor, Guía de Riesgo V7) derivado
    del puntaje oficial de la dimensión, según la nota metodológica anterior."""
    if puntaje is None:
        return "Sin información"
    if puntaje < 40:
        return "Extremo"
    if puntaje < 60:
        return "Alto"
    if puntaje < 80:
        return "Moderado"
    return "Bajo"


ROLES_OCI = [
    ("Liderazgo estratégico",
     "Soporte estratégico para la toma de decisiones del nominador y del representante "
     "legal, agregando valor de manera independiente mediante informes, manejo de "
     "información estratégica y alertas oportunas."),
    ("Enfoque hacia la prevención",
     "Asesoría permanente y formulación de recomendaciones con alcance preventivo, "
     "fomentando la cultura del control para la toma de decisiones oportunas."),
    ("Evaluación de la gestión del riesgo",
     "Aseguramiento objetivo a la Alta Dirección sobre el diseño y la efectividad de las "
     "actividades de administración del riesgo de la entidad."),
    ("Evaluación y seguimiento",
     "Evaluación planeada, documentada y sistemática de metas, resultados, políticas, "
     "planes, programas, procesos, indicadores y riesgos institucionales."),
    ("Relación con entes externos de control",
     "Puente entre los entes externos de control (Contraloría, Procuraduría, etc.) y la "
     "entidad, facilitando el flujo de información con dichos organismos."),
]

ESTRUCTURA_HALLAZGO_EXPLICACION = (
    "Conforme a la Norma 14.2 (NOGAI™, The IIA Global, 2024) adoptada por la Guía de "
    "Auditoría Interna V5 de Función Pública, un hallazgo de auditoría se documenta "
    "identificando la diferencia entre el CRITERIO (la norma o estándar esperado) y la "
    "CONDICIÓN (la situación real observada), analizando su CAUSA raíz y su EFECTO o "
    "impacto sobre la gestión institucional."
)

MARCO_LEGAL_FORMATO_PLAN_MEJORAMIENTO = (
    "El Departamento Administrativo de la Función Pública (DAFP) no publica un único "
    "formato de Excel o Word de uso obligatorio idéntico para todas las entidades del "
    "país; cada entidad adapta su plantilla a su propio Sistema de Control Interno. Sin "
    "embargo, el contenido mínimo exigido es consistente en toda la administración "
    "pública colombiana, con fundamento en: la Ley 87 de 1993 (Sistema de Control "
    "Interno), el Decreto 1499 de 2017 (MIPG/MECI, Dimensión 7 — Control Interno), la "
    "Circular 100-003 de 2010 del DAFP (orientaciones para el manejo de planes de "
    "mejoramiento) y la Norma 15.1 de comunicación final del trabajo de auditoría "
    "(NOGAI™, The IIA Global, 2024), que exige que todo informe de auditoría incluya el "
    "plan de mejoramiento correspondiente a cada hallazgo, con responsable y fecha de "
    "implementación definidos por la propia entidad."
)

COLUMNAS_PLAN_MEJORAMIENTO_OFICIAL = [
    'N°', 'Hallazgo', 'Causa', 'Acción de mejora', 'Meta', 'Indicador de seguimiento',
    'Responsable', 'Fecha de inicio', 'Fecha de terminación', 'Estado',
]

CRITERIOS_CIERRE_PLAN_MEJORAMIENTO = (
    "Un hallazgo se cierra cuando la acción de mejora cumple dos criterios: EFICACIA "
    "(se ejecutaron las actividades planificadas en el plazo previsto) y EFECTIVIDAD "
    "(la acción realmente eliminó la causa raíz que originó el hallazgo, no solo su "
    "síntoma). La verificación de ambos criterios corresponde a la Oficina de Control "
    "Interno de la entidad, con base en las evidencias que aporte el responsable de "
    "cada acción."
)


CARACTERISTICAS_PLAN_MEJORAMIENTO = (
    "La Norma 15.1 exige que el informe de auditoría no se limite a presentar el "
    "hallazgo, sino que incluya también el plan de mejoramiento correspondiente, con "
    "responsable y plazo de implementación definidos por la propia entidad."
)


MARCO_TEORICO_INSTITUCIONAL = (
    "El Modelo Integrado de Planeación y Gestión (MIPG, Decreto 1499 de 2017) es la "
    "traducción institucional colombiana de tres corrientes de la administración pública: "
    "conserva del institucionalismo la atención a reglas, rutinas y estructura "
    "organizacional; toma de la Nueva Gestión Pública (NGP) la medición sistemática del "
    "desempeño a través del FURAG y el IDI; y adopta de la gobernanza y el valor público "
    "la idea de que la meta final es generar resultados verificables para la ciudadanía, "
    "no solo cumplir un formulario. Esta lectura teórica es general para toda entidad que "
    "reporte MIPG y no depende de datos específicos de una entidad en particular."
)

GESTION_PREVENTIVA_RIESGO_FISCAL = (
    "La Guía de Gestión Integral del Riesgo V7 de Función Pública dedica su Capítulo IV a "
    "la Gestión Preventiva de Riesgos Fiscales, con fundamento en la Ley 610 de 2000 "
    "(artículos 3 y 6) y el Acto Legislativo 04 de 2019, que modificó los artículos 267 y "
    "268 de la Constitución para dar al control fiscal un enfoque preventivo. El riesgo "
    "fiscal se define como un evento potencial con efecto dañoso sobre recursos, bienes o "
    "intereses patrimoniales públicos — distinto del daño patrimonial, que es la "
    "afectación real y ya consumada. El Decreto 403 de 2020 reglamenta este enfoque "
    "preventivo, y el Control Fiscal Interno (CFI) — que hace parte del Sistema de Control "
    "Interno de la entidad — es el primer nivel de esta vigilancia, evaluado por la "
    "Contraloría respectiva como parte del fenecimiento de la cuenta."
)

VALOR_PROBATORIO_INFORMES_CI = (
    "El artículo 9 de la Ley 1474 de 2011 dispone que los informes de los funcionarios de "
    "control interno tendrán valor probatorio en los procesos disciplinarios, "
    "administrativos, judiciales y fiscales, cuando las autoridades competentes así lo "
    "soliciten. Esto significa que las brechas documentadas en este informe no son solo "
    "una alerta de gestión: si la entidad no las atiende, pueden convertirse en evidencia "
    "formal ante Procuraduría, Contraloría u otras autoridades de control. Este informe NO "
    "determina responsabilidad disciplinaria, fiscal o penal alguna — esa competencia es "
    "exclusiva de las autoridades de control respectivas; su función es documentar el "
    "estado real de cada índice para que la propia entidad actúe primero, de forma "
    "preventiva."
)


def declaracion_ptep_adaptada(nombre_entidad: str, diag=None, cruce_recomendaciones=None) -> str:
    """Retorna una lista de (subtítulo, texto) con el desarrollo completo de la
    articulación con el Programa de Transparencia y Ética Pública (PTEP),
    parafraseado (no copiado textual, por derechos de autor) del documento
    'Programa de Transparencia y Ética Pública 2025' de Función Pública, y
    adaptado con los datos reales de la entidad cuando están disponibles
    (índice POL15 — Transparencia, índice POL02 — Integridad, y las
    recomendaciones oficiales de esas dos políticas si hay brecha)."""
    secciones = []

    secciones.append((
        'Marco normativo y alcance',
        'El Decreto 1122 de 2024 exige que toda entidad pública formule un Programa de '
        'Transparencia y Ética Pública (PTEP). El PTEP 2025-2028 de Función Pública se '
        'desarrolla en tres fases: una etapa inicial de construcción a cuatro años, '
        'seguida de implementación y monitoreo anuales, y finalmente evaluación — con '
        'participación activa de funcionarios, contratistas y proveedores en cada etapa. '
        'Cada entidad debe adaptar este ciclo a su propia realidad institucional; este '
        'informe usa el documento de Función Pública como referencia de buenas prácticas, '
        'no como un formulario que la entidad deba copiar literalmente.'
    ))

    secciones.append((
        'Compromisos institucionales que exige el Programa',
        'El Programa exige seis compromisos institucionales concretos: cumplimiento '
        'normativo de las políticas y procedimientos del PTEP; actuación ética de los '
        'servidores con honestidad e integridad; fomento activo de la transparencia y la '
        'rendición de cuentas; implementación y seguimiento efectivo de cada componente; '
        'canales de denuncia de irregularidades que protejan la confidencialidad de quien '
        'denuncia; y — donde aplique — debida diligencia en el marco del Sistema de '
        'Administración del Riesgo de Lavado de Activos y Financiación del Terrorismo '
        '(SARLAFT).'
    ))

    secciones.append((
        'Ciclo de gestión del Programa (Formulación → Publicación → Ejecución)',
        'El ciclo tiene 6 fases: Formulación (diseño de estrategias contra la corrupción); '
        'Validación (publicación de la versión inicial en el sitio web institucional por '
        '15 días calendario, para recibir observaciones ciudadanas); Consolidación '
        '(documento preliminar); Aprobación (por el Comité Institucional de Gestión y '
        'Desempeño); Publicación (antes del 31 de enero de cada vigencia, en el botón de '
        'Transparencia y Acceso a la Información Pública del sitio web); y Ejecución, a '
        'cargo de los líderes de cada proceso. Las modificaciones al Programa se reciben '
        'hasta el 15 de junio (corte de mitad de año) o el 30 de noviembre (corte de fin '
        'de año), y se socializan públicamente durante 5 días hábiles.'
    ))

    secciones.append((
        'Roles y responsables (Monitoreo, Administración, Supervisión, Auditoría)',
        'El Programa define 4 roles diferenciados: Monitoreo, a cargo de los líderes de '
        'cada proceso, quienes diseñan y ejecutan las actividades de su componente; '
        'Administración, a cargo de la dependencia que haga las veces de Oficina de '
        'Relación con la Ciudadanía (consolida el Programa completo y lo presenta al '
        'Comité); Supervisión, a cargo de la Alta Dirección a través del Comité '
        'Institucional de Gestión y Desempeño (aprueba el Programa y sus modificaciones); '
        'y Auditoría, evaluación y mejora, responsabilidad exclusiva de la Oficina de '
        'Control Interno, con seguimientos semestrales (cortes a 30 de junio y 31 de '
        'diciembre) y un informe consolidado publicado en la sección de Informes '
        'Institucionales del sitio web.'
    ))

    secciones.append((
        'Reportes y su relación con el FURAG',
        'El seguimiento se hace mediante reportes semestrales: los líderes de proceso '
        'envían su reporte antes del quinto día hábil siguiente al corte semestral; la '
        'Oficina de Control Interno tiene 30 días calendario desde el corte para publicar '
        'su informe de seguimiento. La evaluación anual de los componentes transversal y '
        'programático del PTEP se hace a través del mismo Formulario Único de Reporte de '
        'Avances de la Gestión (FURAG) que produce el IDI de este informe — es decir, el '
        'desempeño del PTEP de la entidad ya está siendo medido, año a año, por la misma '
        'fuente oficial que sustenta todas las demás cifras de este documento.'
    ))

    secciones.append((
        'Formación y comunicación',
        'El Programa exige actividades formativas a través del Plan Institucional de '
        'Capacitación y del proceso de Inducción y Reinducción, coordinadas entre el área '
        'de Talento Humano y quien administre el PTEP, con apoyo de Comunicaciones — '
        'incluyendo formación específica sobre el Código de Integridad y el manejo de '
        'conflictos de interés. La estrategia de comunicación debe divulgar activamente, '
        'en cada ciclo, tanto la publicación inicial y sus modificaciones como los '
        'informes de seguimiento, a través del sitio web institucional y redes sociales.'
    ))

    if diag is not None:
        pol15 = next((r for r in diag.brechas if r.codigo_politica == 'POL15'), None)
        pol02 = next((r for r in diag.brechas if r.codigo_politica == 'POL02'), None)
        cruce = cruce_recomendaciones or {}
        recs_transparencia = []
        for b in diag.brechas:
            if b.codigo_politica in ('POL15', 'POL02'):
                recs_transparencia.extend(cruce.get(b.codigo_indice, [])[:1])
        texto_estado = (
            f'{nombre_entidad} tiene {"una brecha detectada" if (pol15 or pol02) else "sin brecha detectada en este informe"} '
            f'en los índices directamente relacionados con este Programa '
            f'(POL15 — Transparencia, Acceso a la Información y lucha contra la Corrupción'
            f'{", y POL02 — Integridad" if pol02 else ""}). '
        )
        if pol15:
            texto_estado += f'El índice POL15 tiene un puntaje oficial de {pol15.puntaje} sobre 100, por debajo del umbral de 60 puntos. '
        if pol02:
            texto_estado += f'El índice POL02 (Integridad) tiene un puntaje oficial de {pol02.puntaje} sobre 100. '
        if recs_transparencia:
            texto_estado += (
                'Entre las recomendaciones oficiales de Función Pública específicas para esta entidad '
                'en estas políticas se incluye: "' + _limpiar(recs_transparencia[0])[:220] + '..." '
                '(ver el listado completo en la sección de recomendaciones de este informe).'
            )
        secciones.append(('Estado real de esta entidad frente al Programa', texto_estado))
    else:
        secciones.append((
            'Estado real de esta entidad frente al Programa',
            f'{nombre_entidad} debería contar con este Programa formulado, publicado y con '
            f'seguimiento semestral a cargo de su Oficina de Control Interno — su índice POL15 '
            f'(Transparencia, Acceso a la Información y lucha contra la Corrupción) en este '
            f'informe es el mejor proxy oficial disponible del estado de esa política.'
        ))

    return secciones


# ============================================================================
# Constructor reutilizable de las secciones nuevas (agosto 2026): Análisis de
# riesgos + Capítulo de auditoría (hallazgos + plan de mejoramiento) + PTEP.
# Usa datos 100% reales del propio diagnóstico (diag.brechas,
# diag.resultados_por_dimension, cruce_recomendaciones) — nada se inventa aquí.
# ============================================================================

def agregar_capitulo_riesgos_y_auditoria_docx(doc, diag, cruce_recomendaciones, nombre_entidad,
                                               numero_seccion_riesgos=4, numero_seccion_auditoria=5,
                                               numero_seccion_ptep=6, nivel_encabezado=1,
                                               incluir_roles_oci=True, max_hallazgos=20,
                                               recomendaciones_completas=None):
    """Agrega al final de un Document (python-docx) ya construido las 3
    secciones nuevas. Se importa aquí (no arriba del módulo) para evitar
    dependencia circular con generador_informe.py.

    recomendaciones_completas (opcional): list[Recomendacion] SIN filtrar
    por brecha — la lista íntegra que Función Pública entregó a la entidad.
    Si se pasa, el Plan de Mejoramiento (5.3) incluye el 100% de las
    recomendaciones oficiales, no solo las asociadas a un hallazgo."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _sombrear(celda, color_hex):
        tcPr = celda._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def _fnum(v):
        if v is None:
            return 'S/D'
        try:
            v = float(v)
            return str(int(v)) if v == int(v) else f'{v:.2f}'
        except (TypeError, ValueError):
            return str(v)

    # --- Sección: Análisis de riesgos institucionales ---
    doc.add_heading(f'{numero_seccion_riesgos}. Análisis de riesgos institucionales', level=nivel_encabezado)

    doc.add_heading(f'{numero_seccion_riesgos}.1 Marco teórico institucional', level=nivel_encabezado + 1)
    doc.add_paragraph(MARCO_TEORICO_INSTITUCIONAL).runs[0].font.size = Pt(9)

    doc.add_heading(f'{numero_seccion_riesgos}.2 Valoración de riesgo por dimensión', level=nivel_encabezado + 1)
    p = doc.add_paragraph()
    r = p.add_run(NOTA_METODOLOGICA_RIESGO)
    r.italic = True
    r.font.size = Pt(9)

    doc.add_paragraph(
        'Escala de probabilidad (Guía de Gestión Integral del Riesgo V7): ' +
        '; '.join(f'{n} ({pct})' for n, pct in ESCALA_PROBABILIDAD) + '.'
    ).runs[0].font.size = Pt(9)
    doc.add_paragraph(
        'Escala de impacto: ' + '; '.join(f'{n} ({pct})' for n, pct in ESCALA_IMPACTO) + '.'
    ).runs[0].font.size = Pt(9)

    if diag.aplica_mipg_integral:
        t = doc.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        for i, h in enumerate(['Dimensión', 'Puntaje oficial', 'Nivel de riesgo inherente']):
            t.rows[0].cells[i].text = h
            for run in t.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
            _sombrear(t.rows[0].cells[i], 'DDEBF7')
        for res_dim in diag.resultados_por_dimension:
            valor = res_dim.promedio_oficial if res_dim.promedio_oficial is not None else res_dim.promedio
            row = t.add_row().cells
            row[0].text = f'{res_dim.codigo} {res_dim.nombre}'
            row[1].text = _fnum(valor)
            row[2].text = nivel_riesgo_inherente_desde_puntaje(valor)
            for c in row:
                for run in c.paragraphs[0].runs:
                    run.font.size = Pt(9)

    doc.add_heading(f'{numero_seccion_riesgos}.3 Gestión preventiva de riesgos fiscales', level=nivel_encabezado + 1)
    doc.add_paragraph(GESTION_PREVENTIVA_RIESGO_FISCAL).runs[0].font.size = Pt(9)

    doc.add_heading(f'{numero_seccion_riesgos}.4 Valor probatorio de los hallazgos de este informe', level=nivel_encabezado + 1)
    doc.add_paragraph(VALOR_PROBATORIO_INFORMES_CI).runs[0].font.size = Pt(9)

    # --- Sección: Capítulo de auditoría ---
    doc.add_heading(f'{numero_seccion_auditoria}. Capítulo de auditoría: hallazgos y plan de mejoramiento', level=nivel_encabezado)
    p2 = doc.add_paragraph(ESTRUCTURA_HALLAZGO_EXPLICACION)
    p2.runs[0].font.size = Pt(9)
    p3 = doc.add_paragraph()
    r3 = p3.add_run(CARACTERISTICAS_PLAN_MEJORAMIENTO)
    r3.italic = True
    r3.font.size = Pt(9)

    if incluir_roles_oci:
        doc.add_heading(f'{numero_seccion_auditoria}.1 Roles de la Oficina de Control Interno (Decreto 648 de 2017)', level=nivel_encabezado + 1)
        for nombre_rol, desc_rol in ROLES_OCI:
            p4 = doc.add_paragraph()
            r4a = p4.add_run(f'{nombre_rol}: ')
            r4a.bold = True
            r4a.font.size = Pt(9)
            r4b = p4.add_run(desc_rol)
            r4b.font.size = Pt(9)

    doc.add_heading(f'{numero_seccion_auditoria}.2 Hallazgos de auditoría (criterio – condición – criticidad)', level=nivel_encabezado + 1)
    brechas_ordenadas = sorted(diag.brechas, key=lambda b: b.puntaje)[:max_hallazgos]
    ht = doc.add_table(rows=1, cols=4)
    ht.style = 'Table Grid'
    for i, h in enumerate(['Índice', 'Criterio (esperado)', 'Condición (puntaje real)', 'Criticidad']):
        ht.rows[0].cells[i].text = h
        for run in ht.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        _sombrear(ht.rows[0].cells[i], 'DDEBF7')
    for b in brechas_ordenadas:
        row = ht.add_row().cells
        row[0].text = f'{b.codigo_indice} {b.nombre_indice}'
        row[1].text = 'Desempeño adecuado según modelo MIPG (≥60 pts.)'
        row[2].text = f'{_fnum(b.puntaje)} puntos'
        row[3].text = nivel_criticidad_desde_puntaje(b.puntaje)
        for c in row:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.add_heading(f'{numero_seccion_auditoria}.3 Plan de mejoramiento — propuesta con la totalidad de recomendaciones oficiales de Función Pública', level=nivel_encabezado + 1)
    plazos = {'Muy Alta': 'Corto plazo (0-3 meses)', 'Alta': 'Corto plazo (0-6 meses)',
              'Media': 'Mediano plazo (6-12 meses)', 'Baja': 'Largo plazo (12-18 meses)',
              'Muy baja': 'Largo plazo (12-18 meses)'}

    # Mapa código de política -> criticidad más alta detectada entre sus brechas (si tiene)
    criticidad_por_politica = {}
    for b in diag.brechas:
        crit = nivel_criticidad_desde_puntaje(b.puntaje)
        orden = {'Muy Alta': 5, 'Alta': 4, 'Media': 3, 'Baja': 2, 'Muy baja': 1}
        actual = criticidad_por_politica.get(b.codigo_politica)
        if actual is None or orden.get(crit, 0) > orden.get(actual, 0):
            criticidad_por_politica[b.codigo_politica] = crit

    if recomendaciones_completas:
        p5 = doc.add_paragraph()
        r5 = p5.add_run(
            f'Total de recomendaciones oficiales de Función Pública para esta entidad: '
            f'{len(recomendaciones_completas)}. Se presentan TODAS, agrupadas por política — no '
            f'solo las asociadas a un hallazgo priorizado — porque el plan de mejoramiento debe '
            f'cubrir la totalidad de lo que Función Pública entregó a la entidad. El plazo sugerido '
            f'depende de si esa política tiene o no una brecha detectada (y su criticidad); las '
            f'políticas sin brecha se marcan como mantenimiento.'
        )
        r5.italic = True
        r5.font.size = Pt(9)

        agrupado = {}
        for rec in recomendaciones_completas:
            agrupado.setdefault((rec.codigo_politica, rec.politica_nombre), []).append(_limpiar(rec.texto))

        pt2 = doc.add_table(rows=1, cols=3)
        pt2.style = 'Table Grid'
        for i, h in enumerate(['Política', 'Recomendación oficial FP (acción de mejora)', 'Plazo sugerido']):
            pt2.rows[0].cells[i].text = h
            for run in pt2.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
            _sombrear(pt2.rows[0].cells[i], 'DDEBF7')

        for (cod_pol, nombre_pol), textos in agrupado.items():
            crit = criticidad_por_politica.get(cod_pol)
            plazo_txt = plazos.get(crit, 'Mediano plazo (mantenimiento — sin brecha detectada)') if crit else 'Mantenimiento (política sin brecha detectada)'
            for texto in textos:
                row = pt2.add_row().cells
                row[0].text = f'{cod_pol} — {nombre_pol}'
                row[1].text = texto
                row[2].text = plazo_txt
                for c in row:
                    for run in c.paragraphs[0].runs:
                        run.font.size = Pt(9)
    else:
        # Respaldo (sin lista completa disponible): solo lo asociado a hallazgos priorizados.
        p5 = doc.add_paragraph()
        r5 = p5.add_run('Cada hallazgo se asocia con las recomendaciones oficiales de Función Pública ya cruzadas por el sistema para ese índice. El plazo sugerido es una guía de priorización, no un plazo oficial.')
        r5.italic = True
        r5.font.size = Pt(9)
        pt2 = doc.add_table(rows=1, cols=4)
        pt2.style = 'Table Grid'
        for i, h in enumerate(['Hallazgo', 'Política', 'Acción de mejora (recomendación oficial FP)', 'Plazo sugerido']):
            pt2.rows[0].cells[i].text = h
            for run in pt2.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
            _sombrear(pt2.rows[0].cells[i], 'DDEBF7')
        cruce = cruce_recomendaciones or {}
        for b in brechas_ordenadas[:15]:
            acciones = cruce.get(b.codigo_indice, [])
            accion_texto = _limpiar(acciones[0]) if acciones else 'Ver recomendaciones oficiales completas en la sección de recomendaciones.'
            criticidad = nivel_criticidad_desde_puntaje(b.puntaje)
            row = pt2.add_row().cells
            row[0].text = f'{b.codigo_indice} {b.nombre_indice}'
            row[1].text = b.politica
            row[2].text = accion_texto
            row[3].text = plazos.get(criticidad, 'Por definir')
            for c in row:
                for run in c.paragraphs[0].runs:
                    run.font.size = Pt(9)

    # --- Sección: Articulación con el PTEP ---
    doc.add_heading(f'{numero_seccion_ptep}. Articulación con el Programa de Transparencia y Ética Pública', level=nivel_encabezado)
    p_intro = doc.add_paragraph()
    r_intro = p_intro.add_run(
        'Desarrollo basado en el "Programa de Transparencia y Ética Pública 2025" de '
        'Función Pública (Versión 01, agosto de 2025), parafraseado y adaptado con los '
        'datos reales de esta entidad — no es una copia literal del documento fuente.'
    )
    r_intro.italic = True
    r_intro.font.size = Pt(9)
    for subtitulo, texto in declaracion_ptep_adaptada(nombre_entidad, diag=diag, cruce_recomendaciones=cruce_recomendaciones):
        doc.add_heading(subtitulo, level=nivel_encabezado + 1)
        doc.add_paragraph(texto).runs[0].font.size = Pt(9)


def agregar_capitulo_plan_mejoramiento_formal_docx(doc, diag, cruce_recomendaciones, nombre_entidad,
                                                     numero_seccion=7, nivel_encabezado=1):
    """Capítulo especial: Plan de Mejoramiento en el formato con las columnas
    mínimas exigidas de forma consistente en la administración pública
    colombiana (ver MARCO_LEGAL_FORMATO_PLAN_MEJORAMIENTO). Se importa aquí
    para evitar dependencia circular."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def _sombrear(celda, color_hex):
        tcPr = celda._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def _fnum(v):
        if v is None:
            return 'S/D'
        try:
            v = float(v)
            return str(int(v)) if v == int(v) else f'{v:.2f}'
        except (TypeError, ValueError):
            return str(v)

    RESPONSABLE_POR_PALABRA_CLAVE = [
        ('talento humano', 'Área de Talento Humano'),
        ('integridad', 'Comité Institucional de Gestión y Desempeño'),
        ('planeación', 'Oficina Asesora de Planeación'),
        ('presupuestal', 'Área Financiera / Presupuesto'),
        ('contratación', 'Área de Contratación / Jurídica'),
        ('compras', 'Área de Contratación / Jurídica'),
        ('fortalecimiento organizacional', 'Oficina Asesora de Planeación'),
        ('gobierno digital', 'Área de Tecnologías de la Información'),
        ('seguridad digital', 'Área de Tecnologías de la Información'),
        ('defensa jurídica', 'Oficina Jurídica'),
        ('mejora normativa', 'Oficina Jurídica'),
        ('servicio a las ciudadanías', 'Área de Atención al Ciudadano'),
        ('trámites', 'Área de Atención al Ciudadano'),
        ('participación ciudadana', 'Oficina de Comunicaciones / Participación'),
        ('seguimiento y evaluación', 'Oficina Asesora de Planeación'),
        ('transparencia', 'Oficina de Control Interno / Comunicaciones'),
        ('gestión documental', 'Área de Gestión Documental / Archivo'),
        ('información estadística', 'Oficina Asesora de Planeación'),
        ('conocimiento', 'Oficina Asesora de Planeación'),
        ('control interno', 'Oficina de Control Interno'),
    ]

    def _responsable_sugerido(nombre_politica):
        n = (nombre_politica or '').lower()
        for clave, resp in RESPONSABLE_POR_PALABRA_CLAVE:
            if clave in n:
                return resp
        return 'Líder del proceso responsable de la política'

    plazos_dias = {'Muy Alta': 90, 'Alta': 180, 'Media': 365, 'Baja': 545, 'Muy baja': 545}

    doc.add_heading(f'{numero_seccion}. Capítulo especial: Plan de Mejoramiento', level=nivel_encabezado)
    doc.add_paragraph(MARCO_LEGAL_FORMATO_PLAN_MEJORAMIENTO).runs[0].font.size = Pt(9)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(CRITERIOS_CIERRE_PLAN_MEJORAMIENTO)
    r2.italic = True
    r2.font.size = Pt(9)
    p3 = doc.add_paragraph()
    r3 = p3.add_run(
        'Las columnas "Fecha de inicio", "Fecha de terminación" y "Responsable" son PROPUESTAS '
        'de priorización de este informe, con base en la criticidad de cada hallazgo — la entidad '
        'debe validarlas, ajustarlas y formalizarlas con su Comité Institucional de Coordinación de '
        'Control Interno antes de considerarlas el plan de mejoramiento oficial. La columna "Causa" '
        'requiere análisis de causa raíz por parte de la entidad; este informe no cuenta con '
        'información operativa interna para determinarla con precisión.'
    )
    r3.italic = True
    r3.font.size = Pt(9)

    from datetime import date, timedelta
    hoy = date.today()

    t = doc.add_table(rows=1, cols=10)
    t.style = 'Table Grid'
    t.autofit = False
    t.allow_autofit = False
    for i, h in enumerate(COLUMNAS_PLAN_MEJORAMIENTO_OFICIAL):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
        _sombrear(t.rows[0].cells[i], 'DDEBF7')

    anchos_cm = [0.6, 2.5, 1.6, 2.8, 1.6, 1.6, 1.6, 1.3, 1.3, 1.0]
    from docx.shared import Cm
    for idx_col, ancho in enumerate(anchos_cm):
        t.columns[idx_col].width = Cm(ancho)
        for fila_tabla in t.rows:
            fila_tabla.cells[idx_col].width = Cm(ancho)

    cruce = cruce_recomendaciones or {}
    brechas_ordenadas = sorted(diag.brechas, key=lambda b: b.puntaje)
    for n, b in enumerate(brechas_ordenadas, start=1):
        criticidad = nivel_criticidad_desde_puntaje(b.puntaje)
        acciones = cruce.get(b.codigo_indice, [])
        accion_texto = _limpiar(acciones[0]) if acciones else 'Ver banco de recomendaciones oficiales de Función Pública para esta política.'
        dias = plazos_dias.get(criticidad, 365)
        fecha_fin = hoy + timedelta(days=dias)
        row = t.add_row().cells
        row[0].text = str(n)
        row[1].text = f'{b.codigo_indice} — {b.nombre_indice}'
        row[2].text = 'A determinar por la entidad (análisis de causa raíz)'
        row[3].text = accion_texto
        row[4].text = f'Alcanzar un puntaje ≥60/100 en {b.codigo_indice}'
        row[5].text = f'Puntaje oficial de {b.codigo_indice} en el próximo reporte FURAG'
        row[6].text = _responsable_sugerido(b.politica)
        row[7].text = hoy.strftime('%d/%m/%Y')
        row[8].text = fecha_fin.strftime('%d/%m/%Y')
        row[9].text = 'Abierto'
        for idx_c, c in enumerate(row):
            c.width = Cm(anchos_cm[idx_c])
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(8)
    for idx_col, ancho in enumerate(anchos_cm):
        t.columns[idx_col].width = Cm(ancho)


def _limpiar(texto):
    """Colapsa líneas en blanco repetidas para que las celdas de tabla no
    queden con huecos vacíos innecesarios. No recorta ni modifica contenido."""
    if not texto:
        return texto
    lineas = [l.strip() for l in str(texto).split('\n')]
    salida = []
    vacia_antes = False
    for l in lineas:
        if l == '':
            if not vacia_antes:
                salida.append('')
            vacia_antes = True
        else:
            salida.append(l)
            vacia_antes = False
    return '\n'.join(salida).strip()
